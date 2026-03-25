"""
RAGnarok — FastAPI Backend (OpenSearch edition)
Embeddings via Ollama/OpenAI-compatible endpoint.
Vector storage and kNN search via opensearch-py.
"""

from __future__ import annotations

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from pydantic import AliasChoices, BaseModel, ConfigDict, Field
from contextlib import asynccontextmanager, suppress
import httpx
import ipaddress
import json
import re
import uuid
import os
import asyncio
import fnmatch
import hashlib
import io
import csv
import shutil
import math
import unicodedata
from datetime import datetime, timedelta, timezone
from typing import Any, Optional
from pathlib import Path
from zoneinfo import ZoneInfo, ZoneInfoNotFoundError
from opensearchpy import OpenSearch
from urllib.parse import urlparse
from typing_extensions import Literal, TypedDict
from mcp import ClientSession
from mcp.client.sse import sse_client
from mcp.client.streamable_http import streamable_http_client
from langchain_core.prompts import ChatPromptTemplate
from langgraph.graph import END, START, StateGraph

try:
    from sklearn.ensemble import GradientBoostingClassifier, GradientBoostingRegressor, RandomForestClassifier, RandomForestRegressor
    from sklearn.feature_extraction import DictVectorizer
    from sklearn.impute import SimpleImputer
    from sklearn.linear_model import LinearRegression, LogisticRegression
    from sklearn.metrics import (
        accuracy_score,
        f1_score,
        mean_absolute_error,
        mean_squared_error,
        precision_score,
        r2_score,
        recall_score,
    )
    from sklearn.model_selection import train_test_split
    HAVE_SKLEARN = True
except Exception:
    HAVE_SKLEARN = False

try:
    from xgboost import XGBClassifier, XGBRegressor
    HAVE_XGBOOST = True
except Exception:
    HAVE_XGBOOST = False


@asynccontextmanager
async def ragnarok_lifespan(app: FastAPI):
    await read_db_state()
    stop_event = asyncio.Event()
    app.state.planning_scheduler_stop = stop_event
    app.state.planning_scheduler_task = asyncio.create_task(planning_scheduler_loop(stop_event))
    try:
        yield
    finally:
        stop_event = getattr(app.state, "planning_scheduler_stop", None)
        task = getattr(app.state, "planning_scheduler_task", None)
        if stop_event is not None:
            stop_event.set()
        if task is not None:
            try:
                await task
            except Exception:
                pass


app = FastAPI(title="RAGnarok API", lifespan=ragnarok_lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class UpstreamServiceError(Exception):
    def __init__(self, detail: str, status_code: int = 400):
        super().__init__(detail)
        self.detail = detail
        self.status_code = status_code


TABLE_OUTPUT_GUIDANCE = (
    "If the user explicitly asks for a table, rows/columns, a matrix, a grid, a schema list, or a tabular preview, "
    "return the relevant structured result as a valid Markdown table whenever the data is naturally tabular. "
    "Keep the table readable and compact, prefer the most relevant columns, and clearly mention when rows or columns are truncated. "
    "Never invent missing cells or pretend to have columns you did not actually retrieve."
)


def _with_table_output_guidance(base_prompt: str) -> str:
    text = str(base_prompt or "").strip()
    if not text:
        return TABLE_OUTPUT_GUIDANCE
    if TABLE_OUTPUT_GUIDANCE in text:
        return text
    return f"{text}\n\n{TABLE_OUTPUT_GUIDANCE}"


# ── Agent Observability Log Bus ───────────────────────────────────────────────

import collections
import time as _time
from fastapi import Request
from fastapi.responses import StreamingResponse

_LOG_BUS_MAX_AGE_S = 600       # 10 minutes
_LOG_BUS_MAX_EVENTS = 500      # max retained events
_log_buffer: collections.deque = collections.deque(maxlen=_LOG_BUS_MAX_EVENTS)
_log_subscribers: list[asyncio.Queue] = []


@app.exception_handler(UpstreamServiceError)
async def handle_upstream_service_error(_request: Request, exc: UpstreamServiceError):
    return JSONResponse(status_code=exc.status_code, content={"detail": exc.detail})

def _emit_log(kind: str, agent: str, message: str, data: dict | None = None) -> None:
    """
    kind: "info" | "decision" | "tool_call" | "sql" | "llm" | "error" | "success" | "warning"
    agent: e.g. "manager", "clickhouse", "rag", "mcp", "planner", etc.
    """
    event = {
        "id": uuid.uuid4().hex[:8],
        "ts": _utc_now_iso(),
        "ts_epoch": _time.time(),
        "kind": kind,
        "agent": agent,
        "message": message,
        "data": data or {},
    }
    _log_buffer.append(event)
    for queue in _log_subscribers:
        try:
            queue.put_nowait(event)
        except asyncio.QueueFull:
            pass

async def _get_live_logs() -> list[dict]:
    cutoff = _time.time() - _LOG_BUS_MAX_AGE_S
    return [e for e in _log_buffer if e["ts_epoch"] >= cutoff]


@app.get("/api/logs")
async def get_logs():
    return {"logs": await _get_live_logs()}


@app.delete("/api/logs/clear")
async def clear_logs():
    _log_buffer.clear()
    return {"ok": True}


@app.get("/api/logs/stream")
async def stream_logs(request: Request):
    async def generator():
        history = await _get_live_logs()
        for event in history:
            yield f"data: {json.dumps(event)}\n\n"
            
        queue: asyncio.Queue = asyncio.Queue(maxsize=200)
        _log_subscribers.append(queue)
        try:
            while True:
                if await request.is_disconnected():
                    break
                try:
                    event = await asyncio.wait_for(queue.get(), timeout=15.0)
                    yield f"data: {json.dumps(event)}\n\n"
                except asyncio.TimeoutError:
                    yield ": ping\n\n"
        finally:
            try:
                _log_subscribers.remove(queue)
            except ValueError:
                pass

    return StreamingResponse(
        generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive"
        }
    )


# ── Persistent app database ───────────────────────────────────────────────────

DB_PATH = Path(__file__).parent / "DB.json"
DB_LOCK = asyncio.Lock()
LEGACY_DB_USER_ID = "__legacy__"

DEFAULT_APP_CONFIG = {
    "provider": "ollama",
    "baseUrl": "http://localhost:11434",
    "apiKey": "",
    "model": "llama3",
    "systemPrompt": (
        "You are a helpful, smart, and concise AI assistant. Present non-JSON answers "
        "in polished Markdown with clear sections, concise bullets, tasteful **bold** "
        "emphasis, and tables when they help. Safe semantic HTML fragments such as "
        "<section>, <article>, <details>, <summary>, <table>, <ul>, <ol>, and "
        "<blockquote> are allowed when they genuinely improve the layout. Never return "
        "a full HTML document, CSS, or JavaScript. Put the main answer first. If you add "
        "technical details, SQL, raw previews, reasoning notes, or appendices, place them "
        "after the main answer and preferably inside <details><summary>Expand details</summary>"
        "...</details> blocks. When offering choices or clarification "
        "options, always use markdown task lists (- [ ] Option) so the UI can present "
        "clickable replies. "
        "If the user explicitly asks for a table, rows/columns, a matrix, a grid, a schema list, or a tabular preview, "
        "return the relevant structured result as a valid Markdown table whenever the data is naturally tabular."
    ),
    "managerUseRagFunctionalContext": False,
    "disableSslVerification": False,
    "elasticsearchUrl": "http://localhost:9200",
    "elasticsearchIndex": "rag_documents",
    "elasticsearchUsername": "",
    "elasticsearchPassword": "",
    "embeddingBaseUrl": "http://localhost:11434/v1",
    "embeddingApiKey": "",
    "embeddingModel": "nomic-embed-text",
    "embeddingVerifySsl": True,
    "chunkSize": 512,
    "chunkOverlap": 50,
    "knnNeighbors": 50,
    "mcpTools": [
        {"id": "mcp_1", "label": "MCP Tool 1", "url": ""},
        {"id": "mcp_2", "label": "MCP Tool 2", "url": ""},
    ],
    "documentationUrl": "",
    "agenticDataVizUrl": "",
    "portalApps": [],
    "customAgents": [],
    "settingsAccessPassword": "MM@2026",
    "clickhouseHost": "localhost",
    "clickhousePort": 8123,
    "clickhouseDatabase": "default",
    "clickhouseUsername": "default",
    "clickhousePassword": "",
    "clickhouseSecure": False,
    "clickhouseVerifySsl": True,
    "clickhouseHttpPath": "",
    "clickhouseQueryLimit": 200,
    "oracleConnections": [
        {
            "id": "oracle_default",
            "label": "Default Oracle",
            "host": "localhost",
            "port": 1521,
            "serviceName": "",
            "sid": "",
            "dsn": "",
            "username": "",
            "password": "",
        }
    ],
    "oracleAnalystConfig": {
        "connectionId": "oracle_default",
        "rowLimit": 1000,
        "maxRetries": 3,
        "maxIterations": 8,
        "toolkitId": "",
        "systemPrompt": (
            "You are the Oracle SQL agent. Reply in English. Use Oracle tools before making assumptions, "
            "generate optimized Oracle SQL with explicit columns, and present final user-facing answers "
            "in polished Markdown with clear sections, concise bullets, and tasteful emphasis. "
            "Safe semantic HTML fragments are allowed when they improve readability. "
            "When the user explicitly asks for tabular output, include a readable Markdown table."
        ),
    },
    "fileManagerConfig": {
        "basePath": "",
        "maxIterations": 10,
        "systemPrompt": (
            "You are the File Management agent. Reply in English by default. Use "
            "filesystem tools instead of guessing, keep answers short and factual, "
            "ask for confirmation before destructive or overwrite actions, and present "
            "final user-facing answers in polished Markdown with concise structure and "
            "tasteful emphasis. When the user explicitly asks for tabular output, include a Markdown table whenever the result is tabular."
        ),
    },
}

DEFAULT_PREFERENCES = {
    "darkMode": False,
    "currentConversationId": None,
    "workflow": "LLM",
    "agentRole": "manager",
    "selectedMcpToolId": "",
    "selectedCustomAgentId": "",
    "page": "landing",
}

AGENT_ROLES = {"manager", "clickhouse_query", "file_management", "pdf_creator", "oracle_analyst", "data_analyst", "auto_ml", "data_cleaner", "anonymizer", "custom_agent"}
PLANNER_AGENT_ROLES = {"manager", "clickhouse_query", "file_management"}
PLANNER_TRIGGER_KINDS = {
    "once",
    "daily",
    "weekly",
    "interval",
    "clickhouse_watch",
    "file_watch",
}
PLANNER_WEEKDAYS = ["mon", "tue", "wed", "thu", "fri", "sat", "sun"]
PLANNER_WATCH_MODES = {"returns_rows", "count_increases", "result_changes"}
PLANNER_MAX_RUNS = 60
PLANNER_MAX_KNOWN_FILES = 2000
PLANNER_LOOP_INTERVAL_SECONDS = 20
CHAT_MEMORY_MIN_STEPS = 5
CHAT_MEMORY_MAX_STEPS = 10
DATA_QUALITY_MAX_GUIDED_TABLES = 20
DATA_QUALITY_MAX_SAMPLE_ROWS = 2_000_000
DATA_QUALITY_DEFAULT_SAMPLE_SIZE = 50_000
DATA_QUALITY_STRING_SENTINELS = ["n/a", "na", "null", "none", "unknown", "-1", "9999", "99999"]
DATA_ANALYST_DEFAULT_MAX_STEPS = 10
DATA_ANALYST_MAX_STEPS = 10
DATA_ANALYST_MAX_RESULT_ROWS = 2000
DATA_ANALYST_RESULT_PREVIEW_ROWS = 10
DATA_ANALYST_TABLE_OPTION_LIMIT = 8
DATA_ANALYST_MAX_KNOWLEDGE_RESULTS = 4
DATA_ANALYST_REVIEW_INTERVAL = 1
DATA_ANALYST_FORBIDDEN_DATE_FUNCTIONS = (
    "tostartofmonth",
    "tostartofweek",
    "tostartofyear",
    "tostartofquarter",
    "datediff",
    "adddays",
    "addmonths",
    "addyears",
    "subtractdays",
    "subtractmonths",
    "subtractyears",
    "year(",
    "month(",
    "day(",
    "today(",
    "yesterday(",
    "todate(",
    "todatetime(",
)


def _default_planning_state() -> dict:
    return {
        "plans": [],
        "runs": [],
    }


def _default_planning_trigger() -> dict:
    return {
        "kind": "daily",
        "timezone": "UTC",
        "oneTimeAt": "",
        "timeOfDay": "09:00",
        "weekdays": ["mon"],
        "intervalMinutes": 60,
        "pollMinutes": 5,
        "watchSql": "",
        "watchMode": "result_changes",
        "directory": "",
        "pattern": "*",
        "recursive": False,
    }


def _default_planning_runtime() -> dict:
    return {
        "lastCheckedAt": None,
        "lastSeenFingerprint": "",
        "lastSeenMetric": None,
        "knownFiles": [],
    }


def _normalize_planning_trigger(trigger: Optional[dict]) -> dict:
    normalized = _default_planning_trigger()
    if not isinstance(trigger, dict):
        return normalized

    kind = trigger.get("kind")
    normalized["kind"] = kind if kind in PLANNER_TRIGGER_KINDS else normalized["kind"]
    timezone_name = str(trigger.get("timezone") or normalized["timezone"]).strip() or normalized["timezone"]
    normalized["timezone"] = timezone_name
    normalized["oneTimeAt"] = str(trigger.get("oneTimeAt") or "").strip()
    normalized["timeOfDay"] = str(trigger.get("timeOfDay") or normalized["timeOfDay"]).strip() or normalized["timeOfDay"]
    weekdays = trigger.get("weekdays")
    if isinstance(weekdays, list):
        normalized["weekdays"] = [
            day for day in weekdays
            if isinstance(day, str) and day.lower() in PLANNER_WEEKDAYS
        ] or normalized["weekdays"]
    interval_minutes = trigger.get("intervalMinutes")
    if isinstance(interval_minutes, (int, float)):
        normalized["intervalMinutes"] = max(1, int(interval_minutes))
    poll_minutes = trigger.get("pollMinutes")
    if isinstance(poll_minutes, (int, float)):
        normalized["pollMinutes"] = max(1, int(poll_minutes))
    normalized["watchSql"] = str(trigger.get("watchSql") or "").strip()
    watch_mode = trigger.get("watchMode")
    normalized["watchMode"] = watch_mode if watch_mode in PLANNER_WATCH_MODES else normalized["watchMode"]
    normalized["directory"] = str(trigger.get("directory") or "").strip()
    normalized["pattern"] = str(trigger.get("pattern") or normalized["pattern"]).strip() or normalized["pattern"]
    normalized["recursive"] = bool(trigger.get("recursive", normalized["recursive"]))
    return normalized


def _normalize_planning_plan(plan: Optional[dict]) -> dict:
    normalized = {
        "id": uuid.uuid4().hex,
        "name": "",
        "prompt": "",
        "agents": [],
        "status": "active",
        "trigger": _default_planning_trigger(),
        "createdAt": "",
        "updatedAt": "",
        "nextRunAt": None,
        "lastRunAt": None,
        "lastStatus": None,
        "lastSummary": "",
        "runtime": _default_planning_runtime(),
    }
    if not isinstance(plan, dict):
        return normalized

    normalized["id"] = str(plan.get("id") or normalized["id"])
    normalized["name"] = str(plan.get("name") or "").strip()
    normalized["prompt"] = str(plan.get("prompt") or plan.get("objective") or "").strip()
    agents = plan.get("agents")
    if isinstance(agents, list):
        normalized["agents"] = [
            agent for agent in agents
            if isinstance(agent, str) and agent in PLANNER_AGENT_ROLES
        ]
    status = plan.get("status")
    normalized["status"] = status if status in {"active", "paused"} else "active"
    normalized["trigger"] = _normalize_planning_trigger(plan.get("trigger"))
    normalized["createdAt"] = str(plan.get("createdAt") or "").strip()
    normalized["updatedAt"] = str(plan.get("updatedAt") or "").strip()
    normalized["nextRunAt"] = plan.get("nextRunAt") if isinstance(plan.get("nextRunAt"), str) or plan.get("nextRunAt") is None else None
    normalized["lastRunAt"] = plan.get("lastRunAt") if isinstance(plan.get("lastRunAt"), str) or plan.get("lastRunAt") is None else None
    normalized["lastStatus"] = plan.get("lastStatus") if plan.get("lastStatus") in {"running", "success", "error", None} else None
    normalized["lastSummary"] = str(plan.get("lastSummary") or "").strip()

    runtime = _default_planning_runtime()
    incoming_runtime = plan.get("runtime")
    if isinstance(incoming_runtime, dict):
        runtime["lastCheckedAt"] = incoming_runtime.get("lastCheckedAt") if isinstance(incoming_runtime.get("lastCheckedAt"), str) or incoming_runtime.get("lastCheckedAt") is None else None
        runtime["lastSeenFingerprint"] = str(incoming_runtime.get("lastSeenFingerprint") or "").strip()
        last_seen_metric = incoming_runtime.get("lastSeenMetric")
        if isinstance(last_seen_metric, (int, float)) or last_seen_metric is None:
            runtime["lastSeenMetric"] = last_seen_metric
        known_files = incoming_runtime.get("knownFiles")
        if isinstance(known_files, list):
            runtime["knownFiles"] = [
                str(path) for path in known_files
                if isinstance(path, str)
            ][:PLANNER_MAX_KNOWN_FILES]
    normalized["runtime"] = runtime
    return normalized


def _normalize_planning_run(run: Optional[dict]) -> dict:
    normalized = {
        "id": uuid.uuid4().hex,
        "planId": "",
        "planName": "Unnamed plan",
        "triggerKind": "manual",
        "triggerLabel": "",
        "startedAt": "",
        "finishedAt": None,
        "status": "running",
        "summary": "",
        "outputs": [],
    }
    if not isinstance(run, dict):
        return normalized

    normalized["id"] = str(run.get("id") or normalized["id"])
    normalized["planId"] = str(run.get("planId") or "").strip()
    normalized["planName"] = str(run.get("planName") or normalized["planName"]).strip()
    trigger_kind = str(run.get("triggerKind") or "manual").strip()
    normalized["triggerKind"] = (
        trigger_kind if trigger_kind in PLANNER_TRIGGER_KINDS or trigger_kind == "manual" else "manual"
    )
    normalized["triggerLabel"] = str(run.get("triggerLabel") or "").strip()
    normalized["startedAt"] = str(run.get("startedAt") or "").strip()
    normalized["finishedAt"] = run.get("finishedAt") if isinstance(run.get("finishedAt"), str) or run.get("finishedAt") is None else None
    status = run.get("status")
    normalized["status"] = status if status in {"running", "success", "error"} else "running"
    normalized["summary"] = str(run.get("summary") or "").strip()
    outputs = run.get("outputs")
    if isinstance(outputs, list):
        normalized["outputs"] = [
            {
                "agent": str(item.get("agent") or "").strip(),
                "status": item.get("status") if item.get("status") in {"success", "error"} else "success",
                "content": str(item.get("content") or "").strip(),
            }
            for item in outputs
            if isinstance(item, dict)
        ]
    return normalized


def _normalize_planning_state(payload: Optional[dict]) -> dict:
    state = _default_planning_state()
    if not isinstance(payload, dict):
        return state

    plans = payload.get("plans")
    if isinstance(plans, list):
        state["plans"] = [_normalize_planning_plan(plan) for plan in plans]

    runs = payload.get("runs")
    if isinstance(runs, list):
        state["runs"] = [_normalize_planning_run(run) for run in runs[:PLANNER_MAX_RUNS]]

    return state


def _utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _normalize_db_user_id(raw_user_id: Optional[str]) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_-]", "", str(raw_user_id or "").strip())
    return cleaned[:64] or "anonymous"


def _default_user_db_state() -> dict:
    return {
        "conversations": [],
        "preferences": json.loads(json.dumps(DEFAULT_PREFERENCES)),
        "updatedAt": _utc_now_iso(),
    }


def _normalize_user_db_state(payload: Optional[dict]) -> dict:
    state = _default_user_db_state()
    if not isinstance(payload, dict):
        return state

    incoming_conversations = payload.get("conversations")
    if isinstance(incoming_conversations, list):
        state["conversations"] = incoming_conversations

    incoming_preferences = payload.get("preferences")
    if isinstance(incoming_preferences, dict):
        state["preferences"].update(incoming_preferences)
    if state["preferences"].get("agentRole") not in AGENT_ROLES:
        state["preferences"]["agentRole"] = "manager"

    incoming_updated_at = payload.get("updatedAt")
    if isinstance(incoming_updated_at, str) and incoming_updated_at:
        state["updatedAt"] = incoming_updated_at

    return state


def _db_state_for_user(full_state: dict, user_id: str) -> dict:
    users = full_state.get("users") or {}
    user_bucket = _normalize_user_db_state(users.get(user_id))
    return {
        "schemaVersion": full_state.get("schemaVersion", 2),
        "updatedAt": full_state.get("updatedAt") or _utc_now_iso(),
        "config": full_state.get("config") or json.loads(json.dumps(DEFAULT_APP_CONFIG)),
        "conversations": user_bucket.get("conversations") or [],
        "preferences": user_bucket.get("preferences") or json.loads(json.dumps(DEFAULT_PREFERENCES)),
        "planning": full_state.get("planning") or _default_planning_state(),
    }


def _ensure_user_db_state(full_state: dict, user_id: str) -> tuple[dict, bool]:
    normalized = _normalize_db_state(full_state)
    users = normalized.setdefault("users", {})
    changed = False

    if user_id not in users:
        legacy_bucket = users.get(LEGACY_DB_USER_ID)
        if legacy_bucket:
            users[user_id] = _normalize_user_db_state(legacy_bucket)
            users.pop(LEGACY_DB_USER_ID, None)
            changed = True
        else:
            users[user_id] = _default_user_db_state()
            changed = True

    users[user_id] = _normalize_user_db_state(users.get(user_id))
    return normalized, changed


def _default_db_state() -> dict:
    return {
        "schemaVersion": 2,
        "updatedAt": _utc_now_iso(),
        "config": json.loads(json.dumps(DEFAULT_APP_CONFIG)),
        "users": {},
        "planning": _default_planning_state(),
    }


def _normalize_db_state(payload: Optional[dict]) -> dict:
    # Normalize and backfill persisted state on every read/write so DB.json can
    # evolve safely even when older snapshots miss newer configuration blocks.
    state = _default_db_state()
    if not isinstance(payload, dict):
        return state

    incoming_config = payload.get("config")
    if isinstance(incoming_config, dict):
        state["config"].update(incoming_config)
        incoming_file_manager = incoming_config.get("fileManagerConfig")
        if isinstance(incoming_file_manager, dict):
            state["config"]["fileManagerConfig"] = {
                **DEFAULT_APP_CONFIG["fileManagerConfig"],
                **incoming_file_manager,
            }
        incoming_portal_apps = incoming_config.get("portalApps")
        if isinstance(incoming_portal_apps, list):
            state["config"]["portalApps"] = [
                {
                    "id": str(item.get("id") or f"portal_app_{index + 1}").strip() or f"portal_app_{index + 1}",
                    "name": str(item.get("name") or "").strip(),
                    "url": str(item.get("url") or "").strip(),
                    "description": str(item.get("description") or "").strip(),
                }
                for index, item in enumerate(incoming_portal_apps)
                if isinstance(item, dict)
            ]
        incoming_custom_agents = incoming_config.get("customAgents")
        if isinstance(incoming_custom_agents, list):
            state["config"]["customAgents"] = [
                {
                    "id": str(item.get("id") or f"custom_agent_{index + 1}").strip() or f"custom_agent_{index + 1}",
                    "title": str(item.get("title") or f"Custom Agent {index + 1}").strip() or f"Custom Agent {index + 1}",
                    "description": str(item.get("description") or "").strip(),
                    "pythonCode": str(item.get("pythonCode") or item.get("python_code") or "").strip(),
                    "systemPrompt": str(item.get("systemPrompt") or item.get("system_prompt") or "").strip(),
                    "managerRoutingHint": str(item.get("managerRoutingHint") or item.get("manager_routing_hint") or "").strip(),
                    "status": str(item.get("status") or "draft").strip() or "draft",
                    "statusMessage": str(item.get("statusMessage") or item.get("status_message") or "").strip(),
                    "enabled": bool(item.get("enabled")),
                    "badgeColor": str(item.get("badgeColor") or item.get("badge_color") or "zinc").strip() or "zinc",
                }
                for index, item in enumerate(incoming_custom_agents)
                if isinstance(item, dict)
            ]
        incoming_oracle_agent = incoming_config.get("oracleAnalystConfig")
        if isinstance(incoming_oracle_agent, dict):
            state["config"]["oracleAnalystConfig"] = {
                **DEFAULT_APP_CONFIG["oracleAnalystConfig"],
                **incoming_oracle_agent,
            }
        incoming_oracle_connections = incoming_config.get("oracleConnections")
        if isinstance(incoming_oracle_connections, list) and incoming_oracle_connections:
            state["config"]["oracleConnections"] = [
                {
                    **DEFAULT_APP_CONFIG["oracleConnections"][0],
                    **connection,
                }
                for connection in incoming_oracle_connections
                if isinstance(connection, dict)
            ] or json.loads(json.dumps(DEFAULT_APP_CONFIG["oracleConnections"]))

    normalized_users: dict[str, dict] = {}
    incoming_users = payload.get("users")
    if isinstance(incoming_users, dict):
        for raw_user_id, raw_user_state in incoming_users.items():
            normalized_users[_normalize_db_user_id(raw_user_id)] = _normalize_user_db_state(raw_user_state)
    else:
        # Backward compatibility: migrate the historical single-user payload to
        # a temporary legacy bucket that can be adopted by the first client.
        legacy_conversations = payload.get("conversations")
        legacy_preferences = payload.get("preferences")
        if isinstance(legacy_conversations, list) or isinstance(legacy_preferences, dict):
            normalized_users[LEGACY_DB_USER_ID] = _normalize_user_db_state(
                {
                    "conversations": legacy_conversations if isinstance(legacy_conversations, list) else [],
                    "preferences": legacy_preferences if isinstance(legacy_preferences, dict) else {},
                    "updatedAt": payload.get("updatedAt") if isinstance(payload.get("updatedAt"), str) else _utc_now_iso(),
                }
            )
    state["users"] = normalized_users

    state["planning"] = _normalize_planning_state(payload.get("planning"))

    incoming_updated_at = payload.get("updatedAt")
    if isinstance(incoming_updated_at, str) and incoming_updated_at:
        state["updatedAt"] = incoming_updated_at

    incoming_schema_version = payload.get("schemaVersion")
    if isinstance(incoming_schema_version, int):
        state["schemaVersion"] = incoming_schema_version

    return state


def _write_db_state_sync(state: dict) -> dict:
    # Always rewrite through a temporary file, then atomically replace the main
    # DB file. This avoids partial writes if the process is interrupted.
    normalized = _normalize_db_state(state)
    normalized["updatedAt"] = _utc_now_iso()

    DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    temp_path = DB_PATH.with_suffix(".json.tmp")
    temp_path.write_text(
        json.dumps(normalized, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    temp_path.replace(DB_PATH)
    return normalized


def _read_db_state_sync() -> dict:
    # Treat DB.json as self-healing storage: malformed or outdated payloads are
    # normalized immediately so the backend can keep operating predictably.
    if not DB_PATH.exists():
        return _write_db_state_sync(_default_db_state())

    try:
        raw = json.loads(DB_PATH.read_text(encoding="utf-8"))
    except Exception:
        return _write_db_state_sync(_default_db_state())

    normalized = _normalize_db_state(raw)
    if normalized != raw:
        return _write_db_state_sync(normalized)
    return normalized


async def read_db_state() -> dict:
    async with DB_LOCK:
        return await asyncio.to_thread(_read_db_state_sync)


async def write_db_state(payload: dict) -> dict:
    async with DB_LOCK:
        return await asyncio.to_thread(_write_db_state_sync, payload)


# ── OpenSearch client factory ─────────────────────────────────────────────────

def get_os_client(url: str, username: str = None, password: str = None) -> OpenSearch:
    parsed = urlparse(url)
    use_ssl = parsed.scheme == "https"
    host = parsed.hostname or "localhost"
    port = parsed.port or (443 if use_ssl else 9200)
    auth = (username, password) if username and password else None
    return OpenSearch(
        hosts=[{"host": host, "port": port}],
        http_auth=auth,
        use_ssl=use_ssl,
        verify_certs=False,
        ssl_show_warn=False,
    )


# ── ClickHouse helpers ────────────────────────────────────────────────────────

def _is_local_service_host(host: Optional[str]) -> bool:
    if not host:
        return False
    cleaned = str(host).strip().strip("[]").lower()
    if cleaned == "localhost":
        return True
    try:
        ip = ipaddress.ip_address(cleaned)
    except ValueError:
        return False
    return ip.is_loopback or ip.is_unspecified


def _normalize_local_service_host(host: Optional[str]) -> str:
    cleaned = str(host or "").strip().strip("[]")
    if not cleaned:
        return "127.0.0.1"
    return "127.0.0.1" if _is_local_service_host(cleaned) else cleaned


def _normalize_local_service_url(url: str) -> str:
    raw = (url or "").strip()
    if not raw:
        return raw
    parsed = urlparse(raw)
    if not parsed.scheme or not parsed.netloc or not parsed.hostname:
        return raw.rstrip("/")
    if not _is_local_service_host(parsed.hostname):
        return raw.rstrip("/")

    normalized_host = _normalize_local_service_host(parsed.hostname)
    auth = ""
    if parsed.username:
        auth = parsed.username
        if parsed.password:
            auth += f":{parsed.password}"
        auth += "@"
    port = f":{parsed.port}" if parsed.port else ""
    normalized = parsed._replace(netloc=f"{auth}{normalized_host}{port}")
    return normalized.geturl().rstrip("/")


def _httpx_async_client_kwargs(
    target: str,
    *,
    timeout: float,
    verify: Optional[bool] = None,
) -> dict[str, Any]:
    kwargs: dict[str, Any] = {"timeout": timeout}
    if verify is not None:
        kwargs["verify"] = verify

    hostname: Optional[str]
    if "://" in str(target or ""):
        hostname = urlparse(str(target)).hostname
    else:
        hostname = str(target or "").strip()

    if _is_local_service_host(hostname):
        # Ignore OS-level proxy variables for loopback calls. This avoids a
        # class of Windows issues where local backend-to-local-service traffic
        # is unexpectedly routed through a proxy and every agent fails.
        kwargs["trust_env"] = False

    return kwargs


def _ssl_verification_disabled(payload: Optional[dict[str, Any]]) -> bool:
    if not isinstance(payload, dict):
        return False
    raw = payload.get("disableSslVerification", payload.get("disable_ssl_verification", False))
    return bool(raw)


def _effective_verify_ssl(verify_ssl: bool, disable_ssl_verification: bool) -> bool:
    return False if disable_ssl_verification else bool(verify_ssl)


def _build_mcp_http_client_factory(
    target: str,
    *,
    disable_ssl_verification: bool = False,
):
    normalized_target = _normalize_local_service_url(target)
    hostname = urlparse(normalized_target).hostname

    def factory(
        headers: dict[str, Any] | None = None,
        timeout: httpx.Timeout | None = None,
        auth: httpx.Auth | None = None,
        **extra_kwargs: Any,
    ) -> httpx.AsyncClient:
        # This factory is only used by the MCP SDK SSE transport. It keeps the
        # client setup centralized so SSL and Windows-local-proxy behavior stay
        # consistent with the rest of the backend.
        kwargs: dict[str, Any] = {
            "follow_redirects": True,
            "timeout": timeout or httpx.Timeout(30.0, read=300.0),
        }
        if headers is not None:
            kwargs["headers"] = headers
        if auth is not None:
            kwargs["auth"] = auth
        if disable_ssl_verification:
            kwargs["verify"] = False
        if _is_local_service_host(hostname):
            kwargs["trust_env"] = False
        kwargs.update(extra_kwargs)
        return httpx.AsyncClient(**kwargs)

    return factory


def _mcp_default_headers() -> dict[str, str]:
    return {
        "Accept": "application/json, text/event-stream",
    }


def _build_mcp_async_client(
    target: str,
    *,
    disable_ssl_verification: bool = False,
    headers: Optional[dict[str, Any]] = None,
) -> httpx.AsyncClient:
    normalized_target = _normalize_local_service_url(target)
    hostname = urlparse(normalized_target).hostname
    kwargs: dict[str, Any] = {
        "headers": headers or _mcp_default_headers(),
        "follow_redirects": True,
        "timeout": httpx.Timeout(30.0, read=300.0),
    }
    if disable_ssl_verification:
        kwargs["verify"] = False
    if _is_local_service_host(hostname):
        kwargs["trust_env"] = False
    return httpx.AsyncClient(
        **kwargs,
    )


def _collect_exception_messages(exc: BaseException) -> list[str]:
    messages: list[str] = []
    seen: set[int] = set()

    def visit(current: BaseException) -> None:
        current_id = id(current)
        if current_id in seen:
            return
        seen.add(current_id)

        nested = getattr(current, "exceptions", None)
        if isinstance(nested, (list, tuple)) and nested:
            for item in nested:
                if isinstance(item, BaseException):
                    visit(item)
        else:
            text = str(current).strip()
            response = getattr(current, "response", None)
            if response is not None:
                status_code = getattr(response, "status_code", None)
                body = ""
                try:
                    body = str(getattr(response, "text", "") or "").strip()
                except Exception:
                    body = ""
                if len(body) > 280:
                    body = body[:280].rstrip() + "..."
                if status_code and body:
                    text = f"HTTP {status_code}: {body}"
                elif status_code:
                    text = f"HTTP {status_code}"
            if not text:
                text = current.__class__.__name__
            if text not in messages:
                messages.append(text)

        cause = getattr(current, "__cause__", None)
        if isinstance(cause, BaseException):
            visit(cause)
        context = getattr(current, "__context__", None)
        if isinstance(context, BaseException):
            visit(context)

    visit(exc)
    return messages


def _format_mcp_exception(exc: BaseException, target: str) -> str:
    normalized_target = _normalize_local_service_url(target)
    parsed = urlparse(normalized_target)
    path = parsed.path or "/"
    messages = _collect_exception_messages(exc)
    detail = " | ".join(messages[:4]) if messages else str(exc)
    if any("406" in message for message in messages):
        detail += (
            f" | The MCP server rejected the request on `{path}`. "
            "Check whether this endpoint expects streamable HTTP versus SSE, and whether it accepts `application/json, text/event-stream`."
        )
    return detail


@asynccontextmanager
async def _mcp_client_session(
    target: str,
    *,
    disable_ssl_verification: bool = False,
):
    # Prefer the official Python MCP SDK end-to-end:
    # - `/sse` endpoints use the SSE transport
    # - all other endpoints use streamable HTTP
    # This keeps the backend aligned with MCP reference clients and avoids
    # transport quirks from older wrappers.
    normalized_target = _normalize_local_service_url(target)
    parsed = urlparse(normalized_target)
    normalized_path = (parsed.path or "").rstrip("/").lower()

    if normalized_path.endswith("/sse") or normalized_path == "/sse":
        transport_factory = _build_mcp_http_client_factory(
            normalized_target,
            disable_ssl_verification=disable_ssl_verification,
        )
        async with sse_client(
            normalized_target,
            headers=_mcp_default_headers(),
            httpx_client_factory=transport_factory,
        ) as (read_stream, write_stream):
            async with ClientSession(read_stream, write_stream) as session:
                await session.initialize()
                yield session
        return

    async with _build_mcp_async_client(
        normalized_target,
        disable_ssl_verification=disable_ssl_verification,
        headers=_mcp_default_headers(),
    ) as http_client:
        async with streamable_http_client(
            url=normalized_target,
            http_client=http_client,
        ) as (read_stream, write_stream, _get_session_id):
            async with ClientSession(read_stream, write_stream) as session:
                await session.initialize()
                yield session


def _stringify_mcp_content_block(block: Any) -> str:
    text_value = getattr(block, "text", None)
    if isinstance(text_value, str) and text_value.strip():
        return text_value.strip()

    resource = getattr(block, "resource", None)
    if resource is not None:
        resource_text = getattr(resource, "text", None)
        if isinstance(resource_text, str) and resource_text.strip():
            return resource_text.strip()
        resource_uri = getattr(resource, "uri", None)
        if resource_uri:
            return str(resource_uri)
        if hasattr(resource, "model_dump"):
            return json.dumps(resource.model_dump(), ensure_ascii=False, indent=2)

    uri_value = getattr(block, "uri", None)
    if uri_value:
        name_value = getattr(block, "name", None)
        return f"{name_value}: {uri_value}" if name_value else str(uri_value)

    if hasattr(block, "model_dump"):
        return json.dumps(block.model_dump(), ensure_ascii=False, indent=2)

    return str(block).strip()


def _format_mcp_tool_result(result: Any) -> str:
    rendered_parts: list[str] = []
    content_blocks = getattr(result, "content", None) or []
    for block in content_blocks:
        rendered = _stringify_mcp_content_block(block)
        if rendered:
            rendered_parts.append(rendered)

    structured = getattr(result, "structuredContent", None)
    if structured:
        rendered_parts.append(json.dumps(structured, ensure_ascii=False, indent=2))

    output = "\n".join(part for part in rendered_parts if part).strip()
    if getattr(result, "isError", False):
        return f"[Tool error] {output or 'The MCP server returned an error.'}"
    return output

def quote_clickhouse_literal(value: str) -> str:
    escaped = value.replace("\\", "\\\\").replace("'", "\\'")
    return f"'{escaped}'"


def clean_sql_text(sql: str) -> str:
    text = (sql or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z0-9_-]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text)
    return text.strip().rstrip(";").strip()


def is_safe_read_only_sql(sql: str) -> bool:
    cleaned = clean_sql_text(sql).lower()
    if not cleaned:
        return False
    if ";" in cleaned:
        return False
    if not (
        cleaned.startswith("select")
        or cleaned.startswith("with")
        or cleaned.startswith("show")
        or cleaned.startswith("describe")
        or cleaned.startswith("desc")
        or cleaned.startswith("exists")
        or cleaned.startswith("explain")
    ):
        return False
    forbidden = [
        "insert", "update", "delete", "alter", "drop", "truncate", "create",
        "grant", "revoke", "rename", "optimize", "attach", "detach",
    ]
    return not any(re.search(rf"\b{keyword}\b", cleaned) for keyword in forbidden)


def enforce_query_limit(sql: str, limit: int) -> str:
    cleaned = clean_sql_text(sql)
    safe_limit = max(1, min(limit, 1000))
    if re.search(r"\blimit\s+\d+", cleaned, re.IGNORECASE):
        return cleaned
    return f"{cleaned}\nLIMIT {safe_limit}"


def extract_json_object(text: str) -> dict[str, Any]:
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if not match:
        return {}
    try:
        return json.loads(match.group())
    except Exception:
        return {}


def normalize_choice(text: str) -> str:
    cleaned = re.sub(r"^\s*i choose:\s*", "", text or "", flags=re.IGNORECASE)
    return cleaned.strip().strip("`").strip()


def normalize_intent_text(text: str) -> str:
    normalized = normalize_choice(text).lower()
    if not normalized:
        return ""
    without_accents = "".join(
        char
        for char in unicodedata.normalize("NFKD", normalized)
        if not unicodedata.combining(char)
    )
    return re.sub(r"\s+", " ", without_accents).strip()


def _normalized_history_messages(
    history: list[dict[str, Any]],
    current_message: Optional[str] = None,
    max_steps: int = CHAT_MEMORY_MAX_STEPS,
) -> list[dict[str, str]]:
    safe_max_steps = max(CHAT_MEMORY_MIN_STEPS, max_steps)
    normalized: list[dict[str, str]] = []
    for item in history:
        role = str(item.get("role") or "user")
        if role not in {"user", "assistant", "system"}:
            continue
        content = str(item.get("content") or "").strip()
        if not content:
            continue
        normalized.append(
            {
                "role": role,
                "content": _truncate_text_preview(content, 1200),
            }
        )

    if current_message:
        current_normalized = normalize_choice(current_message).lower()
        while normalized:
            last = normalized[-1]
            if last["role"] != "user":
                break
            if normalize_choice(last["content"]).lower() != current_normalized:
                break
            normalized.pop()

    return normalized[-safe_max_steps:]


def _conversation_memory_markdown(
    history: list[dict[str, Any]],
    current_message: Optional[str] = None,
    max_steps: int = CHAT_MEMORY_MAX_STEPS,
) -> str:
    trimmed = _normalized_history_messages(history, current_message=current_message, max_steps=max_steps)
    if not trimmed:
        return "No recent memory."
    lines = []
    for item in trimmed:
        label = "User" if item["role"] == "user" else "Assistant"
        lines.append(f"- {label}: {item['content']}")
    return "\n".join(lines)


def resolve_user_choice(user_text: str, options: list[str]) -> Optional[str]:
    if not options:
        return None
    normalized = normalize_choice(user_text).lower()
    if not normalized:
        return None
    for option in options:
        option_lower = option.lower()
        if normalized == option_lower:
            return option
        explicit_patterns = [
            rf"^(choose|use|pick)\s+{re.escape(option_lower)}$",
            rf"^(the\s+)?{re.escape(option_lower)}\s+table$",
            rf"^table\s+{re.escape(option_lower)}$",
        ]
        if any(re.fullmatch(pattern, normalized) for pattern in explicit_patterns):
            return option
    return None


def clickhouse_url(config: "ClickHouseConfig") -> str:
    scheme = "https" if config.secure else "http"
    host = _normalize_local_service_host(config.host)
    base = f"{scheme}://{host}:{config.port}"
    path = config.http_path.strip("/")
    return f"{base}/{path}" if path else base


async def execute_clickhouse_sql(
    config: "ClickHouseConfig",
    sql: str,
    readonly: bool = True,
    json_format: bool = True,
    max_result_rows_override: Optional[int] = None,
) -> dict[str, Any]:
    query = clean_sql_text(sql)
    if not query:
        raise HTTPException(status_code=400, detail="Empty ClickHouse query.")

    if readonly and not is_safe_read_only_sql(query):
        raise HTTPException(status_code=400, detail="Only read-only SELECT queries are allowed.")

    suffix = " FORMAT JSON" if json_format and "format json" not in query.lower() else ""
    final_query = query + suffix
    params = {
        "database": config.database,
        "readonly": 1 if readonly else 0,
        "wait_end_of_query": 1,
        "result_overflow_mode": "break",
        "max_result_rows": max(
            1,
            min(
                int(max_result_rows_override if max_result_rows_override is not None else config.query_limit),
                10_000,
            ),
        ),
    }
    auth = (config.username, config.password) if config.username else None

    endpoint = clickhouse_url(config)
    try:
        async with httpx.AsyncClient(
            **_httpx_async_client_kwargs(endpoint, timeout=60.0, verify=config.verify_ssl)
        ) as client:
            response = await client.post(
                endpoint,
                params=params,
                content=final_query.encode("utf-8"),
                auth=auth,
                headers={"Content-Type": "text/plain; charset=utf-8"},
            )
            response.raise_for_status()
            if json_format:
                return response.json()
            return {"raw": response.text}
    except httpx.HTTPError as exc:
        raise HTTPException(
            status_code=400,
            detail=(
                f"ClickHouse connection error for `{endpoint}`: {exc}. "
                "Check the host/port and avoid using `0.0.0.0` as a client URL."
            ),
        ) from exc


async def list_clickhouse_tables(config: "ClickHouseConfig") -> list[str]:
    result = await execute_clickhouse_sql(
        config,
        (
            "SELECT name FROM system.tables "
            f"WHERE database = {quote_clickhouse_literal(config.database)} "
            "ORDER BY name"
        ),
    )
    return [row.get("name", "") for row in result.get("data", []) if row.get("name")]


async def describe_clickhouse_table(
    config: "ClickHouseConfig",
    table_name: str,
) -> list[dict[str, str]]:
    result = await execute_clickhouse_sql(
        config,
        (
            "SELECT name, type, default_kind, default_expression "
            "FROM system.columns "
            f"WHERE database = {quote_clickhouse_literal(config.database)} "
            f"AND table = {quote_clickhouse_literal(table_name)} "
            "ORDER BY position"
        ),
    )
    return [
        {
            "name": row.get("name", ""),
            "type": row.get("type", ""),
            "default_kind": row.get("default_kind", "") or "",
            "default_expression": row.get("default_expression", "") or "",
        }
        for row in result.get("data", [])
        if row.get("name")
    ]


def find_date_columns(schema: list[dict[str, str]]) -> list[str]:
    date_like = []
    for column in schema:
        col_type = column.get("type", "").lower()
        if any(token in col_type for token in ["date", "time"]):
            date_like.append(column.get("name", ""))
    return [name for name in date_like if name]


def _clickhouse_schema_category(type_name: str) -> str:
    lowered = (type_name or "").lower()
    if any(token in lowered for token in ["date", "time"]):
        return "date"
    if any(token in lowered for token in ["int", "float", "decimal", "numeric", "double", "real"]):
        return "numeric"
    if any(token in lowered for token in ["string", "fixedstring", "uuid", "enum", "ipv", "bool"]):
        return "string"
    return "other"


def match_schema_columns(candidates: list[str], schema: list[dict[str, str]]) -> list[str]:
    lookup = {
        column.get("name", "").lower(): column.get("name", "")
        for column in schema
        if column.get("name")
    }
    matched: list[str] = []
    for candidate in candidates:
        name = lookup.get(candidate.lower())
        if name and name not in matched:
            matched.append(name)
    return matched


def match_available_options(candidates: list[str], options: list[str]) -> list[str]:
    lookup = {
        option.lower(): option
        for option in options
        if option
    }
    matched: list[str] = []
    for candidate in candidates:
        option = lookup.get((candidate or "").lower())
        if option and option not in matched:
            matched.append(option)
    return matched


def quote_clickhouse_identifier(name: str) -> str:
    escaped = (name or "").replace("`", "``")
    return f"`{escaped}`"


def classify_clickhouse_column_type(type_name: str) -> str:
    lowered = (type_name or "").lower()
    if any(token in lowered for token in ["date", "time"]):
        return "date"
    if any(token in lowered for token in ["int", "float", "decimal", "double"]):
        return "numeric"
    if any(token in lowered for token in ["string", "fixedstring", "uuid", "enum"]):
        return "string"
    return "other"


def is_clickhouse_schema_request(user_message: str) -> bool:
    normalized = normalize_intent_text(user_message)
    if not normalized:
        return False

    direct_patterns = [
        "list the columns",
        "list columns",
        "list all columns",
        "list the fields",
        "list fields",
        "show columns",
        "show the columns",
        "show fields",
        "show the fields",
        "what are the columns",
        "what are the fields",
        "table schema",
        "show schema",
        "describe table",
        "describe the table",
        "liste des champs",
        "liste les champs",
        "liste des colonnes",
        "liste les colonnes",
        "donne moi la liste des champs",
        "donne moi les champs",
        "donne moi la liste des colonnes",
        "quels sont les champs",
        "quelles sont les colonnes",
        "schema de la table",
        "structure de la table",
    ]
    if any(pattern in normalized for pattern in direct_patterns):
        return True

    has_column_word = any(token in normalized for token in [" champ", " champs", " colonne", " colonnes", " field", " fields", " column", " columns"])
    has_listing_verb = any(token in normalized for token in ["liste", "donne", "montre", "show", "list", "what are", "describe", "schema", "structure"])
    has_row_request = any(token in normalized for token in [" ligne", " lignes", " row", " rows", " sample", " preview", " tableau"])
    return has_column_word and has_listing_verb and not has_row_request


def is_clickhouse_table_list_request(user_message: str) -> bool:
    normalized = normalize_intent_text(user_message)
    if not normalized:
        return False

    direct_patterns = [
        "list tables",
        "list the tables",
        "show tables",
        "show the tables",
        "what tables",
        "what are the tables",
        "available tables",
        "table list",
        "liste des tables",
        "liste les tables",
        "montre les tables",
        "affiche les tables",
        "quelles sont les tables",
        "donne moi la liste des tables",
        "donne moi les tables",
    ]
    if any(pattern in normalized for pattern in direct_patterns):
        return True

    has_table_word = any(token in normalized for token in [" table", " tables"])
    has_listing_verb = any(token in normalized for token in ["list", "show", "what are", "available", "liste", "montre", "affiche", "quelles sont", "donne"])
    has_schema_word = any(token in normalized for token in ["column", "field", "schema", "colonne", "champ", "structure"])
    return has_table_word and has_listing_verb and not has_schema_word


def is_clickhouse_sample_rows_request(user_message: str) -> bool:
    normalized = normalize_intent_text(user_message)
    if not normalized:
        return False

    if "[resolved intent: row count]" in normalized:
        return False
    if "[resolved intent: sample rows]" in normalized:
        return True

    row_tokens = [
        " rows", " row", " lignes", " ligne", "sample", "preview",
        "first rows", "first row", "top rows", "top row",
        "premieres lignes", "premiere ligne", "quelques lignes",
    ]
    asks_for_rows = any(token in normalized for token in row_tokens) or bool(
        re.search(r"\b\d{1,4}\s*(rows?|lines?|lignes?)\b", normalized)
    )
    if not asks_for_rows:
        return False

    sample_verbs = ["show", "give", "display", "return", "donne", "montre", "affiche"]
    if not any(token in normalized for token in sample_verbs):
        return False

    return True


def is_clickhouse_row_count_request(user_message: str) -> bool:
    normalized = normalize_intent_text(user_message)
    if not normalized:
        return False

    if "[resolved intent: sample rows]" in normalized:
        return False
    if "[resolved intent: row count]" in normalized:
        return True

    direct_patterns = [
        "number of rows",
        "row count",
        "count the rows",
        "count rows",
        "how many rows",
        "how many records",
        "how many lines",
        "combien de lignes",
        "combien d enregistrements",
        "nombre de lignes",
        "nombre d enregistrements",
        "count(*)",
    ]
    if any(pattern in normalized for pattern in direct_patterns):
        return True

    has_count_word = any(
        token in normalized
        for token in [" count", "count ", "combien", "nombre", "how many", "total"]
    )
    has_row_word = any(
        token in normalized
        for token in [" row", " rows", " line", " lines", " ligne", " lignes", " record", " records", " enregistrement", " enregistrements"]
    )
    return has_count_word and has_row_word


def is_clickhouse_row_request_ambiguous(user_message: str) -> bool:
    normalized = normalize_intent_text(user_message)
    if not normalized:
        return False

    if "[resolved intent: row count]" in normalized or "[resolved intent: sample rows]" in normalized:
        return False

    has_row_word = any(
        token in normalized
        for token in [" row", " rows", " line", " lines", " ligne", " lignes", " record", " records", " enregistrement", " enregistrements"]
    )
    has_listing_verb = any(
        token in normalized
        for token in ["show", "give", "display", "return", "donne", "montre", "affiche", "liste", "list"]
    )
    return has_row_word and has_listing_verb and is_clickhouse_row_count_request(user_message) and is_clickhouse_sample_rows_request(user_message)


def extract_clickhouse_requested_row_limit(user_message: str, default: int = 10, max_limit: int = 100) -> int:
    normalized = normalize_intent_text(user_message)
    if not normalized:
        return default

    patterns = [
        r"\b(\d{1,4})\s*(rows?|lines?|lignes?)\b",
        r"\b(?:top|first|limit)\s+(\d{1,4})\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, normalized)
        if match:
            try:
                return max(1, min(int(match.group(1)), max_limit))
            except Exception:
                break
    return max(1, min(default, max_limit))


def _clickhouse_markdown_cell(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    text = str(value)
    text = text.replace("|", "\\|").replace("\n", "<br />")
    return text


def build_clickhouse_markdown_table(rows: list[dict[str, Any]], headers: Optional[list[str]] = None, max_rows: int = 20) -> str:
    preview_rows = rows[: max(1, max_rows)]
    if not preview_rows:
        return ""
    if not headers:
        headers = [str(key) for key in preview_rows[0].keys()]
    if not headers:
        return ""

    header_line = "| " + " | ".join(headers) + " |"
    divider_line = "| " + " | ".join(["---"] * len(headers)) + " |"
    body = [
        "| " + " | ".join(_clickhouse_markdown_cell(row.get(header)) for header in headers) + " |"
        for row in preview_rows
    ]
    table = "\n".join([header_line, divider_line, *body])
    if len(rows) > len(preview_rows):
        table += f"\n\n_Showing the first {len(preview_rows)} rows out of {len(rows)}._"
    return table


def _user_explicitly_requests_table(text: str) -> bool:
    normalized = str(text or "").lower()
    return any(
        token in normalized
        for token in [
            "table",
            "tabular",
            "markdown table",
            "show rows",
            "show row",
            "schema list",
            "as a grid",
            "as grid",
            "as matrix",
            "under the form of a table",
            "sous la forme d'un tableau",
            "sous forme de tableau",
            "dans un tableau",
            "liste des champs",
            "liste des colonnes",
            "rows and columns",
        ]
    )


def build_clickhouse_schema_section(table_name: str, schema: list[dict[str, str]]) -> str:
    schema_rows = [
        {
            "Column": column.get("name", ""),
            "Type": column.get("type", ""),
            "Default": column.get("default_expression", "") or column.get("default_kind", "") or "",
        }
        for column in schema
    ]
    table = build_clickhouse_markdown_table(schema_rows, headers=["Column", "Type", "Default"], max_rows=200)
    intro = f"## Table Schema\nHere is the column list for `{table_name}`."
    return f"{intro}\n\n{table}" if table else intro


# ── Oracle helpers ────────────────────────────────────────────────────────────

ORACLE_REACT_MAX_ITERATIONS = 8
ORACLE_MAX_ROW_LIMIT = 50_000
ORACLE_DEFAULT_ROW_LIMIT = 1_000
ORACLE_TABLE_PREVIEW_LIMIT = 40
ORACLE_RESULT_PREVIEW_ROWS = 12
ORACLE_SYSTEM_OWNERS = {
    "ANONYMOUS", "APPQOSSYS", "AUDSYS", "CTXSYS", "DBSNMP", "DIP", "DMSYS",
    "GGSYS", "GSMADMIN_INTERNAL", "GSMCATUSER", "LBACSYS", "MDSYS", "OJVMSYS",
    "OLAPSYS", "ORDDATA", "ORDPLUGINS", "ORDSYS", "OUTLN", "REMOTE_SCHEDULER_AGENT",
    "SI_INFORMTN_SCHEMA", "SYS", "SYSTEM", "WMSYS", "XDB",
}
ORACLE_FORBIDDEN_FILTER_KEYWORDS = [
    "drop", "delete", "insert", "update", "create", "alter", "exec", "execute",
    "merge", "grant", "revoke", "commit", "rollback", "union", "sleep", "dbms_lock",
]
ORACLE_REACT_TOOL_NAMES = {"list_tables", "get_schema", "check_query", "execute_query"}


def _import_oracledb():
    try:
        import oracledb
        return oracledb
    except ImportError as exc:
        raise ValueError("The optional dependency `oracledb` is required for Oracle features.") from exc


def quote_oracle_literal(value: str) -> str:
    escaped = str(value or "").replace("'", "''")
    return f"'{escaped}'"


def _normalize_oracle_identifier_part(value: str) -> str:
    text = str(value or "").strip().strip('"')
    return text.upper()


def parse_oracle_table_reference(table_ref: str) -> tuple[Optional[str], str]:
    text = str(table_ref or "").strip()
    if not text:
        raise ValueError("An Oracle table name is required.")
    if "." in text:
        owner, table_name = text.split(".", 1)
        return _normalize_oracle_identifier_part(owner), _normalize_oracle_identifier_part(table_name)
    return None, _normalize_oracle_identifier_part(text)


def quote_oracle_identifier(name: str) -> str:
    normalized = _normalize_oracle_identifier_part(name)
    if re.fullmatch(r"[A-Z][A-Z0-9_$#]*", normalized):
        return normalized
    escaped = str(name or "").replace('"', '""')
    return f'"{escaped}"'


def quote_oracle_table_reference(table_ref: str) -> str:
    owner, table_name = parse_oracle_table_reference(table_ref)
    if owner:
        return f"{quote_oracle_identifier(owner)}.{quote_oracle_identifier(table_name)}"
    return quote_oracle_identifier(table_name)


def is_safe_read_only_oracle_sql(sql: str) -> bool:
    cleaned = clean_sql_text(sql).lower()
    if not cleaned or ";" in cleaned:
        return False
    if not (cleaned.startswith("select") or cleaned.startswith("with")):
        return False
    forbidden = [
        "insert", "update", "delete", "merge", "alter", "drop", "truncate", "create",
        "grant", "revoke", "comment", "call", "execute immediate",
    ]
    return not any(re.search(rf"\b{keyword}\b", cleaned) for keyword in forbidden)


def enforce_oracle_row_limit(sql: str, row_limit: int) -> str:
    cleaned = clean_sql_text(sql)
    safe_limit = max(1, min(int(row_limit or ORACLE_DEFAULT_ROW_LIMIT), ORACLE_MAX_ROW_LIMIT))
    lowered = cleaned.lower()
    if re.search(r"\bfetch\s+first\s+\d+\s+rows\s+only\b", lowered):
        return cleaned
    if re.search(r"\brownum\s*<=\s*\d+\b", lowered):
        return cleaned
    return f"{cleaned}\nFETCH FIRST {safe_limit} ROWS ONLY"


def _serialize_oracle_value(value: Any) -> Any:
    if value is None:
        return None
    if isinstance(value, (str, int, float, bool)):
        return value
    if isinstance(value, datetime):
        return value.isoformat()
    if hasattr(value, "isoformat"):
        try:
            return value.isoformat()
        except Exception:
            pass
    if isinstance(value, bytes):
        return value.hex()
    return str(value)


def _oracle_make_dsn(connection: OracleConnectionConfig) -> str:
    if connection.dsn.strip():
        return connection.dsn.strip()
    if not connection.host.strip():
        raise ValueError("Oracle host is required unless a custom DSN is provided.")
    if connection.service_name.strip():
        oracledb = _import_oracledb()
        return oracledb.makedsn(connection.host.strip(), int(connection.port or 1521), service_name=connection.service_name.strip())
    if connection.sid.strip():
        oracledb = _import_oracledb()
        return oracledb.makedsn(connection.host.strip(), int(connection.port or 1521), sid=connection.sid.strip())
    raise ValueError("Oracle connection requires either a service name, a SID, or a custom DSN.")


def _oracle_connect_sync(connection: OracleConnectionConfig):
    oracledb = _import_oracledb()
    return oracledb.connect(
        user=connection.username.strip(),
        password=connection.password,
        dsn=_oracle_make_dsn(connection),
    )


def _oracle_list_tables_sync(connection: OracleConnectionConfig) -> list[str]:
    sql = """
        SELECT owner || '.' || table_name AS table_name
        FROM all_tables
        WHERE owner NOT IN ({owners})
        ORDER BY owner, table_name
        FETCH FIRST 2000 ROWS ONLY
    """.strip().format(
        owners=", ".join(quote_oracle_literal(owner) for owner in sorted(ORACLE_SYSTEM_OWNERS))
    )
    with _oracle_connect_sync(connection) as conn:
        with conn.cursor() as cursor:
            cursor.execute(sql)
            return [str(row[0]) for row in cursor.fetchall() if row and row[0]]


def _oracle_get_schema_sync(
    connection: OracleConnectionConfig,
    table_name: str,
    columns_filter: str = "",
) -> list[dict[str, Any]]:
    owner, parsed_table = parse_oracle_table_reference(table_name)
    filters = [
        "table_name = :table_name",
    ]
    params: dict[str, Any] = {"table_name": parsed_table}
    if owner:
        filters.append("owner = :owner")
        params["owner"] = owner

    filtered_columns = [
        _normalize_oracle_identifier_part(column)
        for column in re.split(r"[\s,]+", columns_filter or "")
        if str(column).strip()
    ]
    if filtered_columns:
        binds = []
        for index, column_name in enumerate(filtered_columns):
            key = f"col_{index}"
            binds.append(f":{key}")
            params[key] = column_name
        filters.append(f"column_name IN ({', '.join(binds)})")

    sql = f"""
        SELECT column_name,
               data_type ||
               CASE
                 WHEN data_precision IS NOT NULL AND data_scale IS NOT NULL THEN '(' || data_precision || ',' || data_scale || ')'
                 WHEN data_precision IS NOT NULL THEN '(' || data_precision || ')'
                 WHEN char_col_decl_length IS NOT NULL THEN '(' || char_col_decl_length || ')'
                 ELSE ''
               END AS column_type,
               nullable
        FROM all_tab_columns
        WHERE {' AND '.join(filters)}
        ORDER BY column_id
    """.strip()

    with _oracle_connect_sync(connection) as conn:
        with conn.cursor() as cursor:
            cursor.execute(sql, params)
            rows = cursor.fetchall()
    return [
        {
            "name": str(row[0]),
            "type": str(row[1] or ""),
            "nullable": str(row[2] or "").upper() == "Y",
        }
        for row in rows
    ]


def _oracle_check_query_sync(connection: OracleConnectionConfig, sql: str) -> dict[str, Any]:
    query = clean_sql_text(sql)
    if not is_safe_read_only_oracle_sql(query):
        raise ValueError("Only read-only Oracle SELECT queries are allowed.")

    with _oracle_connect_sync(connection) as conn:
        with conn.cursor() as cursor:
            try:
                cursor.execute(f"EXPLAIN PLAN FOR {query}")
                return {"valid": True, "mode": "explain"}
            except Exception as explain_exc:
                try:
                    wrapped = f"SELECT * FROM ({query}) WHERE 1 = 0"
                    cursor.execute(wrapped)
                    return {"valid": True, "mode": "parse_only", "warning": str(explain_exc)}
                except Exception as parse_exc:
                    raise ValueError(str(parse_exc)) from parse_exc


def _oracle_execute_query_sync(
    connection: OracleConnectionConfig,
    sql: str,
    row_limit: int,
) -> dict[str, Any]:
    query = clean_sql_text(sql)
    if not is_safe_read_only_oracle_sql(query):
        raise ValueError("Only read-only Oracle SELECT queries are allowed.")
    final_query = enforce_oracle_row_limit(query, row_limit)

    with _oracle_connect_sync(connection) as conn:
        with conn.cursor() as cursor:
            cursor.execute(final_query)
            description = cursor.description or []
            rows = cursor.fetchmany(max(1, min(int(row_limit or ORACLE_DEFAULT_ROW_LIMIT), ORACLE_MAX_ROW_LIMIT)))

    columns = [
        {
            "name": str(column[0]),
            "type": str(getattr(column[1], "__name__", column[1]) if len(column) > 1 else ""),
        }
        for column in description
    ]
    headers = [column["name"] for column in columns]
    data = [
        {
            header: _serialize_oracle_value(value)
            for header, value in zip(headers, row)
        }
        for row in rows
    ]
    return {"sql": final_query, "columns": columns, "rows": data, "row_count": len(data)}


async def list_oracle_tables(connection: OracleConnectionConfig) -> list[str]:
    return await asyncio.to_thread(_oracle_list_tables_sync, connection)


async def get_oracle_schema(
    connection: OracleConnectionConfig,
    table_name: str,
    columns_filter: str = "",
) -> list[dict[str, Any]]:
    return await asyncio.to_thread(_oracle_get_schema_sync, connection, table_name, columns_filter)


async def check_oracle_query(connection: OracleConnectionConfig, sql: str) -> dict[str, Any]:
    return await asyncio.to_thread(_oracle_check_query_sync, connection, sql)


async def execute_oracle_query(
    connection: OracleConnectionConfig,
    sql: str,
    row_limit: int,
) -> dict[str, Any]:
    return await asyncio.to_thread(_oracle_execute_query_sync, connection, sql, row_limit)


def _oracle_test_connection_sync(connection: OracleConnectionConfig) -> dict[str, Any]:
    with _oracle_connect_sync(connection) as conn:
        with conn.cursor() as cursor:
            cursor.execute(
                """
                SELECT
                    SYS_CONTEXT('USERENV', 'CURRENT_SCHEMA') AS current_schema,
                    SYS_CONTEXT('USERENV', 'SESSION_USER') AS session_user
                FROM dual
                """
            )
            row = cursor.fetchone() or ("", "")
    tables = _oracle_list_tables_sync(connection)
    return {
        "connection_id": connection.id,
        "label": connection.label or connection.id,
        "current_schema": str(row[0] or "").strip(),
        "session_user": str(row[1] or "").strip(),
        "table_count": len(tables),
        "tables": tables[:ORACLE_TABLE_PREVIEW_LIMIT],
    }


async def test_oracle_connection(connection: OracleConnectionConfig) -> dict[str, Any]:
    return await asyncio.to_thread(_oracle_test_connection_sync, connection)


def _default_data_quality_state() -> dict[str, Any]:
    return {
        "stage": "idle",
        "table": None,
        "columns": [],
        "sample_size": DATA_QUALITY_DEFAULT_SAMPLE_SIZE,
        "row_filter": "",
        "time_column": None,
        "db_type": "clickhouse",
        "schema_info": [],
        "column_stats": {},
        "volumetric_stats": None,
        "llm_analysis": "",
        "final_answer": "",
        "agent_id": "data_quality_tables",
        "session_id": uuid.uuid4().hex,
        "last_error": "",
        "available_tables": [],
        "available_columns": [],
        "date_columns": [],
    }


def _normalize_data_quality_state(payload: Optional[dict]) -> dict[str, Any]:
    state = _default_data_quality_state()
    if not isinstance(payload, dict):
        return state

    schema_info = payload.get("schema_info") if isinstance(payload.get("schema_info"), list) else payload.get("schemaInfo")
    state["stage"] = str(payload.get("stage") or state["stage"]).strip() or state["stage"]
    state["table"] = str(payload.get("table")).strip() if payload.get("table") else None
    state["columns"] = [
        str(item).strip()
        for item in (payload.get("columns") or [])
        if isinstance(item, str) and str(item).strip()
    ]
    sample_size = payload.get("sample_size") if "sample_size" in payload else payload.get("sampleSize")
    if isinstance(sample_size, (int, float)):
        state["sample_size"] = max(0, min(DATA_QUALITY_MAX_SAMPLE_ROWS, int(sample_size)))
    state["row_filter"] = str(payload.get("row_filter") or payload.get("rowFilter") or "").strip()
    time_column = payload.get("time_column") if "time_column" in payload else payload.get("timeColumn")
    state["time_column"] = str(time_column).strip() if isinstance(time_column, str) and time_column.strip() else None
    db_type = str(payload.get("db_type") or payload.get("dbType") or "clickhouse").lower()
    state["db_type"] = "oracle" if db_type == "oracle" else "clickhouse"
    state["schema_info"] = [
        {
            "name": str(column.get("name") or "").strip(),
            "type": str(column.get("type") or "").strip(),
            "category": classify_clickhouse_column_type(str(column.get("type") or "")),
        }
        for column in (schema_info or [])
        if isinstance(column, dict) and str(column.get("name") or "").strip()
    ]
    state["column_stats"] = payload.get("column_stats") if isinstance(payload.get("column_stats"), dict) else payload.get("columnStats") if isinstance(payload.get("columnStats"), dict) else {}
    state["volumetric_stats"] = payload.get("volumetric_stats") if isinstance(payload.get("volumetric_stats"), dict) else payload.get("volumetricStats") if isinstance(payload.get("volumetricStats"), dict) else None
    llm_analysis = payload.get("llm_analysis") if "llm_analysis" in payload else payload.get("llmAnalysis")
    if isinstance(llm_analysis, str):
        state["llm_analysis"] = llm_analysis
    elif isinstance(llm_analysis, dict):
        state["llm_analysis"] = json.dumps(llm_analysis, ensure_ascii=False)
    state["final_answer"] = str(payload.get("final_answer") or payload.get("finalAnswer") or "").strip()
    state["agent_id"] = str(payload.get("agent_id") or payload.get("agentId") or state["agent_id"]).strip() or state["agent_id"]
    state["session_id"] = str(payload.get("session_id") or payload.get("sessionId") or state["session_id"]).strip() or state["session_id"]
    state["last_error"] = str(payload.get("last_error") or payload.get("lastError") or "").strip()
    state["available_tables"] = [
        str(item).strip()
        for item in (payload.get("available_tables") or payload.get("availableTables") or [])
        if isinstance(item, str) and str(item).strip()
    ]
    state["available_columns"] = [
        str(item).strip()
        for item in (payload.get("available_columns") or payload.get("availableColumns") or [])
        if isinstance(item, str) and str(item).strip()
    ]
    state["date_columns"] = [
        str(item).strip()
        for item in (payload.get("date_columns") or payload.get("dateColumns") or [])
        if isinstance(item, str) and str(item).strip()
    ]
    return state


DATA_QUALITY_FOLLOWUP_STAGES = {
    "awaiting_table",
    "awaiting_columns_mode",
    "awaiting_custom_columns",
    "awaiting_sample_size",
    "awaiting_custom_sample_size",
    "awaiting_row_filter_mode",
    "awaiting_row_filter",
    "awaiting_time_column",
    "awaiting_review",
}


def _data_quality_state_needs_followup(state: dict[str, Any]) -> bool:
    return str(state.get("stage") or "").strip() in DATA_QUALITY_FOLLOWUP_STAGES


def _safe_float(value: Any) -> Optional[float]:
    if value is None:
        return None
    try:
        return float(value)
    except Exception:
        return None


def _round_metric(value: Any, digits: int = 4) -> Any:
    numeric = _safe_float(value)
    if numeric is None:
        return value
    if math.isfinite(numeric):
        return round(numeric, digits)
    return value


def _first_row(payload: dict[str, Any]) -> dict[str, Any]:
    data = payload.get("data") or []
    return data[0] if data else {}


def _data_quality_scan_limit(sample_size: int) -> int:
    if sample_size <= 0:
        return DATA_QUALITY_MAX_SAMPLE_ROWS
    return max(1, min(sample_size, DATA_QUALITY_MAX_SAMPLE_ROWS))


def _validate_data_quality_row_filter(row_filter: str) -> Optional[str]:
    text = (row_filter or "").strip()
    if not text:
        return None
    lowered = text.lower()
    if ";" in text or "--" in text or "/*" in text or "*/" in text:
        return "The row filter must be a single safe boolean expression without comments or semicolons."
    forbidden = ["drop", "delete", "insert", "update", "create", "alter", "exec", "union", "sleep"]
    if any(re.search(rf"\b{keyword}\b", lowered) for keyword in forbidden):
        return "The row filter contains blocked keywords and was rejected for safety reasons."
    return None


def _normalize_auto_ml_row_filter(row_filter: str) -> str:
    text = str(row_filter or "").strip()
    if not text:
        return ""
    text = text.strip().strip("`").strip()
    text = re.sub(r"^\s*where\s+", "", text, flags=re.IGNORECASE)
    return text.strip()


def _extract_auto_ml_row_filter(user_request: str) -> str:
    for line in str(user_request or "").splitlines():
        candidate = line.strip()
        if not candidate:
            continue
        lowered = candidate.lower()
        if lowered.startswith("apply this row filter:") or lowered.startswith("row filter:") or lowered.startswith("use this row filter:"):
            _, _, remainder = candidate.partition(":")
            return _normalize_auto_ml_row_filter(remainder)
    return ""


def _extract_guided_row_filter(user_request: str) -> str:
    return _extract_auto_ml_row_filter(user_request)


def _extract_auto_ml_sample_row_limit(user_request: str) -> Optional[int]:
    patterns = [
        r"\buse up to\s+(\d{2,5})\s+rows?\b",
        r"\brow limit\s*:\s*(\d{2,5})\b",
        r"\blimit(?: the benchmark)? to\s+(\d{2,5})\s+rows?\b",
        r"\butilise(?:r)?\s+(?:jusqu['’]a|jusqu'a)\s+(\d{2,5})\s+lignes?\b",
        r"\blimite(?:r)?\s+(?:a|à)\s+(\d{2,5})\s+lignes?\b",
    ]
    text = str(user_request or "")
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if not match:
            continue
        try:
            return max(100, min(10_000, int(match.group(1))))
        except Exception:
            return None
    return None


def _match_data_quality_columns(column_names: list[str], schema_info: list[dict[str, Any]]) -> list[str]:
    available = {
        str(column.get("name") or "").lower(): str(column.get("name") or "")
        for column in schema_info
        if column.get("name")
    }
    matched: list[str] = []
    for column_name in column_names:
        found = available.get(str(column_name or "").strip().lower())
        if found and found not in matched:
            matched.append(found)
    return matched


def _parse_custom_column_input(text: str) -> list[str]:
    chunks = re.split(r"[\n,]+", text or "")
    return [chunk.strip().strip("`") for chunk in chunks if chunk.strip()]


def _build_data_quality_source_sql(
    table_name: str,
    columns: list[str],
    row_filter: str,
    sample_size: int,
) -> str:
    selected_columns = ", ".join(quote_clickhouse_identifier(column) for column in columns) if columns else "*"
    sql = f"SELECT {selected_columns} FROM {quote_clickhouse_identifier(table_name)}"
    if row_filter:
        sql += f" WHERE ({row_filter})"
    sql += f"\nLIMIT {_data_quality_scan_limit(sample_size)}"
    return sql


def _data_quality_percentile(values: list[float], q: float) -> Optional[float]:
    if not values:
        return None
    ordered = sorted(values)
    if len(ordered) == 1:
        return ordered[0]
    position = (len(ordered) - 1) * q
    lower = math.floor(position)
    upper = math.ceil(position)
    if lower == upper:
        return ordered[lower]
    weight = position - lower
    return ordered[lower] * (1 - weight) + ordered[upper] * weight


def _data_quality_numeric_severity(stats: dict[str, Any]) -> tuple[str, list[str]]:
    reasons: list[str] = []
    null_pct = _safe_float(stats.get("null_pct")) or 0.0
    iqr_pct = _safe_float(stats.get("iqr_outlier_pct")) or 0.0
    zscore_pct = _safe_float(stats.get("zscore_outlier_pct")) or 0.0
    if null_pct > 20:
        reasons.append(f"Null rate is {round(null_pct, 2)}%, which is critical.")
    elif null_pct > 5:
        reasons.append(f"Null rate is {round(null_pct, 2)}%, which is a warning.")
    if iqr_pct > 10 or zscore_pct > 10:
        reasons.append("Outlier volume is above the critical threshold.")
    elif iqr_pct > 2 or zscore_pct > 2:
        reasons.append("Outlier volume is above the warning threshold.")
    severity = "ok"
    if any("critical" in reason.lower() for reason in reasons):
        severity = "critical"
    elif reasons:
        severity = "warning"
    return severity, reasons


def _data_quality_generic_severity(stats: dict[str, Any]) -> tuple[str, list[str]]:
    reasons: list[str] = []
    null_pct = _safe_float(stats.get("null_pct")) or 0.0
    sentinel_pct = _safe_float(stats.get("sentinel_pct")) or 0.0
    distinct_pct = _safe_float(stats.get("distinct_pct")) or 0.0
    if null_pct > 20:
        reasons.append(f"Null rate is {round(null_pct, 2)}%, which is critical.")
    elif null_pct > 5:
        reasons.append(f"Null rate is {round(null_pct, 2)}%, which is a warning.")
    if sentinel_pct > 5:
        reasons.append(f"Sentinel values represent {round(sentinel_pct, 2)}%, which is critical.")
    elif sentinel_pct > 0:
        reasons.append(f"Sentinel values are present at {round(sentinel_pct, 2)}%.")
    if distinct_pct > 95 and (stats.get("category") == "string"):
        reasons.append("Cardinality is very high for a text field and may indicate identifier-like data.")
    severity = "ok"
    if any("critical" in reason.lower() for reason in reasons):
        severity = "critical"
    elif reasons:
        severity = "warning"
    return severity, reasons


def _finalize_data_quality_stats(stats: dict[str, Any]) -> dict[str, Any]:
    row_count = int(stats.get("row_count") or 0)
    null_count = int(stats.get("null_count") or 0)
    nonnull_count = max(0, row_count - null_count)
    stats["nonnull_count"] = nonnull_count
    stats["null_pct"] = round((null_count / row_count) * 100, 2) if row_count else 0.0
    distinct_count = int(stats.get("distinct_count") or 0)
    stats["distinct_pct"] = round((distinct_count / nonnull_count) * 100, 2) if nonnull_count else 0.0

    if stats.get("category") == "numeric":
        severity, reasons = _data_quality_numeric_severity(stats)
    else:
        severity, reasons = _data_quality_generic_severity(stats)

    stats["severity_hint"] = severity
    stats["severity_icon"] = "🔴" if severity == "critical" else "🟡" if severity == "warning" else "🟢"
    stats["issues"] = reasons
    return stats


async def _data_quality_numeric_stats(
    config: "ClickHouseConfig",
    source_sql: str,
    column_name: str,
) -> dict[str, Any]:
    identifier = quote_clickhouse_identifier(column_name)
    summary = await execute_clickhouse_sql(
        config,
        f"""
        SELECT
          count() AS row_count,
          countIf(isNull(value)) AS null_count,
          uniqExact(value) AS distinct_count,
          min(value) AS min_value,
          max(value) AS max_value,
          avg(value) AS avg_value,
          stddevPop(value) AS stddev_value,
          quantilesExactInclusive(0.25, 0.5, 0.75)(value) AS quartiles,
          countIf(value = 0) AS zero_count,
          countIf(value < 0) AS negative_count
        FROM (
          SELECT toFloat64OrNull({identifier}) AS value
          FROM ({source_sql}) AS src
        ) AS profile
        """,
    )
    row = _first_row(summary)
    quartiles = row.get("quartiles") or [None, None, None]
    q1 = _safe_float(quartiles[0] if len(quartiles) > 0 else None)
    median = _safe_float(quartiles[1] if len(quartiles) > 1 else None)
    q3 = _safe_float(quartiles[2] if len(quartiles) > 2 else None)
    iqr = (q3 - q1) if q1 is not None and q3 is not None else None
    lower_fence = (q1 - 1.5 * iqr) if iqr is not None else None
    upper_fence = (q3 + 1.5 * iqr) if iqr is not None else None
    avg_value = _safe_float(row.get("avg_value"))
    stddev_value = _safe_float(row.get("stddev_value"))

    iqr_outlier_count = 0
    zscore_outlier_count = 0
    if lower_fence is not None and upper_fence is not None:
        outlier_query = await execute_clickhouse_sql(
            config,
            f"""
            SELECT
              countIf(value < {lower_fence} OR value > {upper_fence}) AS iqr_outlier_count,
              countIf(abs((value - {avg_value or 0}) / {stddev_value or 1}) > 3) AS zscore_outlier_count
            FROM (
              SELECT toFloat64OrNull({identifier}) AS value
              FROM ({source_sql}) AS src
            ) AS profile
            WHERE NOT isNull(value)
            """,
        )
        outlier_row = _first_row(outlier_query)
        iqr_outlier_count = int(outlier_row.get("iqr_outlier_count") or 0)
        zscore_outlier_count = int(outlier_row.get("zscore_outlier_count") or 0) if stddev_value and stddev_value > 0 else 0

    stats = {
        "category": "numeric",
        "row_count": int(row.get("row_count") or 0),
        "null_count": int(row.get("null_count") or 0),
        "distinct_count": int(row.get("distinct_count") or 0),
        "min": _round_metric(row.get("min_value")),
        "max": _round_metric(row.get("max_value")),
        "avg": _round_metric(avg_value),
        "stddev": _round_metric(stddev_value),
        "p25": _round_metric(q1),
        "p50": _round_metric(median),
        "p75": _round_metric(q3),
        "iqr": _round_metric(iqr),
        "lower_fence": _round_metric(lower_fence),
        "upper_fence": _round_metric(upper_fence),
        "zero_count": int(row.get("zero_count") or 0),
        "negative_count": int(row.get("negative_count") or 0),
        "iqr_outlier_count": iqr_outlier_count,
        "zscore_outlier_count": zscore_outlier_count,
    }
    nonnull_count = max(1, stats["row_count"] - stats["null_count"])
    stats["iqr_outlier_pct"] = round((iqr_outlier_count / nonnull_count) * 100, 2)
    stats["zscore_outlier_pct"] = round((zscore_outlier_count / nonnull_count) * 100, 2)
    stats["coeff_variation"] = round(abs((stddev_value or 0.0) / avg_value), 4) if avg_value not in (None, 0) else None
    stats["skewness_approx"] = round((3 * ((avg_value or 0.0) - (median or 0.0)) / stddev_value), 4) if stddev_value not in (None, 0) and median is not None else None
    return _finalize_data_quality_stats(stats)


async def _data_quality_string_stats(
    config: "ClickHouseConfig",
    source_sql: str,
    column_name: str,
) -> dict[str, Any]:
    identifier = quote_clickhouse_identifier(column_name)
    sentinel_values = ", ".join(quote_clickhouse_literal(value) for value in DATA_QUALITY_STRING_SENTINELS)
    query = await execute_clickhouse_sql(
        config,
        f"""
        SELECT
          (SELECT count() FROM ({source_sql}) AS src) AS row_count,
          (SELECT countIf(isNull(toNullable({identifier}))) FROM ({source_sql}) AS src) AS null_count,
          count() AS nonnull_count,
          countIf(trimmed = '') AS empty_count,
          uniqExact(trimmed) AS distinct_count,
          min(lengthUTF8(trimmed)) AS min_length,
          max(lengthUTF8(trimmed)) AS max_length,
          avg(lengthUTF8(trimmed)) AS avg_length,
          countIf(lengthUTF8(trimmed) > 1000) AS very_long_count,
          countIf(match(raw_text, '^\\s') OR match(raw_text, '\\s$')) AS edge_space_count,
          countIf(match(trimmed, '[A-Za-z]') AND trimmed = upperUTF8(trimmed)) AS uppercase_count,
          countIf(match(trimmed, '^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\\\\.[A-Za-z]{{2,}}$')) AS email_like_count,
          countIf(match(trimmed, '^[-+]?[0-9]+(\\\\.[0-9]+)?$')) AS numeric_like_count,
          countIf(lowerUTF8(trimmed) IN ({sentinel_values})) AS sentinel_count
        FROM (
          SELECT
            toString({identifier}) AS raw_text,
            replaceRegexpAll(toString({identifier}), '^\\s+|\\s+$', '') AS trimmed
          FROM ({source_sql}) AS src
          WHERE NOT isNull(toNullable({identifier}))
        ) AS profile
        """,
    )
    row = _first_row(query)
    stats = {
        "category": "string",
        "row_count": int(row.get("row_count") or 0),
        "null_count": int(row.get("null_count") or 0),
        "distinct_count": int(row.get("distinct_count") or 0),
        "empty_count": int(row.get("empty_count") or 0),
        "min_length": int(row.get("min_length") or 0) if row.get("min_length") is not None else None,
        "max_length": int(row.get("max_length") or 0) if row.get("max_length") is not None else None,
        "avg_length": _round_metric(row.get("avg_length")),
        "very_long_count": int(row.get("very_long_count") or 0),
        "edge_space_count": int(row.get("edge_space_count") or 0),
        "uppercase_count": int(row.get("uppercase_count") or 0),
        "email_like_count": int(row.get("email_like_count") or 0),
        "numeric_like_count": int(row.get("numeric_like_count") or 0),
        "sentinel_count": int(row.get("sentinel_count") or 0),
    }
    nonnull_count = max(1, stats["row_count"] - stats["null_count"])
    stats["sentinel_pct"] = round((stats["sentinel_count"] / nonnull_count) * 100, 2)
    return _finalize_data_quality_stats(stats)


async def _data_quality_date_stats(
    config: "ClickHouseConfig",
    source_sql: str,
    column_name: str,
) -> dict[str, Any]:
    identifier = quote_clickhouse_identifier(column_name)
    query = await execute_clickhouse_sql(
        config,
        f"""
        SELECT
          (SELECT count() FROM ({source_sql}) AS src) AS row_count,
          (SELECT countIf(isNull(toNullable({identifier}))) FROM ({source_sql}) AS src) AS null_count,
          count() AS nonnull_count,
          uniqExact(parsed_dt) AS distinct_count,
          min(parsed_dt) AS min_value,
          max(parsed_dt) AS max_value,
          countIf(parsed_dt > now()) AS future_count,
          countIf(parsed_dt < toDateTime('1970-01-02 00:00:00')) AS epoch_like_count,
          (
            SELECT count()
            FROM (
              SELECT toString({identifier}) AS raw_text
              FROM ({source_sql}) AS src
              WHERE NOT isNull(toNullable({identifier}))
            ) AS raw_profile
            WHERE match(raw_text, '^(?:0\\d{{3}}|1[0-8]\\d{{2}})-')
          ) AS pre_1900_count,
          countIf(toDayOfWeek(parsed_dt) >= 6) AS weekend_count
        FROM (
          SELECT parseDateTimeBestEffortOrNull(toString({identifier})) AS parsed_dt
          FROM ({source_sql}) AS src
          WHERE NOT isNull(toNullable({identifier}))
        ) AS profile
        WHERE NOT isNull(parsed_dt)
        """,
    )
    row = _first_row(query)
    stats = {
        "category": "date",
        "row_count": int(row.get("row_count") or 0),
        "null_count": int(row.get("null_count") or 0),
        "distinct_count": int(row.get("distinct_count") or 0),
        "min": row.get("min_value"),
        "max": row.get("max_value"),
        "future_count": int(row.get("future_count") or 0),
        "epoch_like_count": int(row.get("epoch_like_count") or 0),
        "pre_1900_count": int(row.get("pre_1900_count") or 0),
        "weekend_count": int(row.get("weekend_count") or 0),
        "sentinel_count": 0,
        "sentinel_pct": 0.0,
    }
    return _finalize_data_quality_stats(stats)


async def data_quality_schema_node(
    config: "ClickHouseConfig",
    state: dict[str, Any],
) -> dict[str, Any]:
    state["available_tables"] = await list_clickhouse_tables(config)
    if state.get("table"):
        raw_schema = await describe_clickhouse_table(config, state["table"])
        state["schema_info"] = [
            {
                "name": column["name"],
                "type": column["type"],
                "category": classify_clickhouse_column_type(column["type"]),
            }
            for column in raw_schema
        ]
        state["available_columns"] = [column["name"] for column in state["schema_info"]]
        state["date_columns"] = [column["name"] for column in state["schema_info"] if column["category"] == "date"]
    return state


async def data_quality_stats_node(
    config: "ClickHouseConfig",
    state: dict[str, Any],
) -> dict[str, Any]:
    if not state.get("table") or not state.get("columns"):
        raise ValueError("Table and columns are required before profiling can start.")
    source_sql = _build_data_quality_source_sql(
        state["table"],
        state["columns"],
        state.get("row_filter") or "",
        int(state.get("sample_size") or DATA_QUALITY_DEFAULT_SAMPLE_SIZE),
    )
    state["column_stats"] = {}
    schema_lookup = {
        column["name"]: column
        for column in state.get("schema_info", [])
        if column.get("name")
    }
    for column_name in state["columns"]:
        schema_entry = schema_lookup.get(column_name, {})
        category = schema_entry.get("category") or classify_clickhouse_column_type(schema_entry.get("type", ""))
        if category == "numeric":
            state["column_stats"][column_name] = await _data_quality_numeric_stats(config, source_sql, column_name)
        elif category == "string":
            state["column_stats"][column_name] = await _data_quality_string_stats(config, source_sql, column_name)
        elif category == "date":
            state["column_stats"][column_name] = await _data_quality_date_stats(config, source_sql, column_name)
        else:
            state["column_stats"][column_name] = {
                "category": "other",
                "severity_hint": "warning",
                "severity_icon": "🟡",
                "issues": ["This column type is not fully profiled yet."],
            }
    return state


async def data_quality_volumetric_node(
    config: "ClickHouseConfig",
    state: dict[str, Any],
) -> dict[str, Any]:
    time_column = state.get("time_column")
    if not time_column:
        state["volumetric_stats"] = None
        return state

    source_sql = _build_data_quality_source_sql(
        state["table"],
        [time_column],
        state.get("row_filter") or "",
        int(state.get("sample_size") or DATA_QUALITY_DEFAULT_SAMPLE_SIZE),
    )
    identifier = quote_clickhouse_identifier(time_column)
    range_query = await execute_clickhouse_sql(
        config,
        f"""
        SELECT
          min(parsed_time) AS min_time,
          max(parsed_time) AS max_time,
          count() AS profiled_rows
        FROM (
          SELECT parseDateTimeBestEffortOrNull(toString({identifier})) AS parsed_time
          FROM ({source_sql}) AS src
        ) AS profile
        WHERE NOT isNull(parsed_time)
        """,
    )
    range_row = _first_row(range_query)
    min_time = range_row.get("min_time")
    max_time = range_row.get("max_time")
    if not min_time or not max_time:
        state["volumetric_stats"] = None
        return state

    min_dt = _parse_iso_datetime(str(min_time)) if "T" in str(min_time) else None
    max_dt = _parse_iso_datetime(str(max_time)) if "T" in str(max_time) else None
    if min_dt and max_dt:
        span_days = max(0, int((max_dt - min_dt).total_seconds() // 86400))
    else:
        span_days = 8
    granularity = "hour" if span_days <= 7 else "day"
    bucket_expression = "toStartOfHour(parsed_time)" if granularity == "hour" else "toStartOfDay(parsed_time)"

    volume_query = await execute_clickhouse_sql(
        config,
        f"""
        SELECT
          {bucket_expression} AS bucket,
          count() AS volume
        FROM (
          SELECT parseDateTimeBestEffortOrNull(toString({identifier})) AS parsed_time
          FROM ({source_sql}) AS src
        ) AS profile
        WHERE NOT isNull(parsed_time)
        GROUP BY bucket
        ORDER BY bucket
        """,
    )
    buckets = [
        {
            "bucket": row.get("bucket"),
            "volume": int(row.get("volume") or 0),
        }
        for row in (volume_query.get("data") or [])
    ]
    volumes = [bucket["volume"] for bucket in buckets]
    if not volumes:
        state["volumetric_stats"] = None
        return state

    q1 = _data_quality_percentile([float(value) for value in volumes], 0.25)
    q3 = _data_quality_percentile([float(value) for value in volumes], 0.75)
    iqr = (q3 - q1) if q1 is not None and q3 is not None else None
    lower_fence = max(0.0, (q1 - 1.5 * iqr)) if iqr is not None else None
    avg_volume = sum(volumes) / len(volumes)
    stddev_volume = math.sqrt(sum((value - avg_volume) ** 2 for value in volumes) / len(volumes))
    anomalously_low = [
        bucket for bucket in buckets
        if lower_fence is not None and bucket["volume"] < lower_fence
    ]
    state["volumetric_stats"] = {
        "granularity": granularity,
        "min_time": min_time,
        "max_time": max_time,
        "bucket_count": len(buckets),
        "avg_volume": round(avg_volume, 2),
        "stddev_volume": round(stddev_volume, 2),
        "q1": _round_metric(q1),
        "q3": _round_metric(q3),
        "iqr": _round_metric(iqr),
        "lower_fence": _round_metric(lower_fence),
        "anomalously_low_periods": anomalously_low[:12],
        "series_preview": buckets[:24],
    }
    return state


def _compact_data_quality_payload(state: dict[str, Any]) -> dict[str, Any]:
    compact_columns = {}
    for column_name, stats in (state.get("column_stats") or {}).items():
        compact_columns[column_name] = {
            "category": stats.get("category"),
            "severity_hint": stats.get("severity_hint"),
            "null_pct": stats.get("null_pct"),
            "distinct_pct": stats.get("distinct_pct"),
            "issues": stats.get("issues") or [],
            "key_metrics": {
                key: value
                for key, value in stats.items()
                if key in {
                    "min", "max", "avg", "stddev", "p25", "p50", "p75", "iqr",
                    "zero_count", "negative_count", "iqr_outlier_pct", "zscore_outlier_pct",
                    "empty_count", "sentinel_pct", "avg_length", "very_long_count",
                    "future_count", "epoch_like_count", "pre_1900_count", "weekend_count",
                }
            },
        }
    return {
        "table": state.get("table"),
        "columns": state.get("columns"),
        "sample_size": state.get("sample_size"),
        "row_filter": state.get("row_filter") or "",
        "time_column": state.get("time_column"),
        "db_type": state.get("db_type"),
        "column_stats": compact_columns,
        "volumetric_stats": state.get("volumetric_stats"),
    }


def _data_quality_python_fallback_analysis(state: dict[str, Any]) -> dict[str, Any]:
    column_entries = []
    critical_count = 0
    warning_count = 0
    ok_count = 0
    score = 100

    for column_name, stats in (state.get("column_stats") or {}).items():
        severity = stats.get("severity_hint") or "ok"
        if severity == "critical":
            critical_count += 1
            score -= 15
        elif severity == "warning":
            warning_count += 1
            score -= 5
        else:
            ok_count += 1
        column_entries.append(
            {
                "name": column_name,
                "severity": severity,
                "headline": (
                    stats.get("issues")[0]
                    if stats.get("issues")
                    else "No major issue detected in the sampled profile."
                ),
                "business_risk": (
                    "Data quality issues may reduce reporting trust and downstream decision accuracy."
                    if severity != "ok"
                    else "No immediate business risk was detected in the sampled data."
                ),
            }
        )

    return {
        "global_score": max(0, score),
        "critical_count": critical_count,
        "warning_count": warning_count,
        "ok_count": ok_count,
        "columns": column_entries,
        "top_recommendations": [
            "Prioritize columns marked as critical before using this table for reporting or automation.",
            "Add validation rules for nulls, sentinel values, and format anomalies at ingestion time.",
            "Review row filters and sample size choices if you need a narrower business scope.",
        ],
        "volumetric_summary": (
            "No volumetric anomaly was computed."
            if not state.get("volumetric_stats")
            else "Volumetric analysis is included below and highlights low-volume periods."
        ),
    }


async def data_quality_llm_analysis_node(
    state: dict[str, Any],
    conversation_memory: str,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str],
) -> dict[str, Any]:
    payload = _compact_data_quality_payload(state)
    prompt = f"""
You are a senior data-quality analyst.
Analyze the profiling payload below and return JSON only with this exact shape:
{{
  "global_score": 0,
  "critical_count": 0,
  "warning_count": 0,
  "ok_count": 0,
  "columns": [
    {{
      "name": "column_name",
      "severity": "critical" | "warning" | "ok",
      "headline": "short finding summary",
      "business_risk": "short business risk"
    }}
  ],
  "top_recommendations": ["recommendation 1", "recommendation 2"],
  "volumetric_summary": "short summary if volumetric analysis exists"
}}

Scoring rules:
- Critical: null_pct > 20, outliers > 10, sentinel_pct > 5 => -15 points
- Warning: null_pct 5-20, outliers 2-10, suspicious cardinality => -5 points
- OK: no meaningful issue => 0 points

Keep everything in English.
Use the payload only. Do not invent metrics that are not provided.

Recent conversation memory:
{conversation_memory}

Payload:
{json.dumps(payload, ensure_ascii=False, indent=2)}
""".strip()

    try:
        raw = await llm_chat(
            [{"role": "user", "content": prompt}],
            llm_base_url,
            llm_model,
            llm_provider,
            llm_api_key,
            response_format="json",
        )
        parsed = extract_json_object(raw)
        if not parsed:
            raise ValueError("Empty data-quality analysis payload from the LLM.")
        state["llm_analysis"] = json.dumps(parsed, ensure_ascii=False)
        return parsed
    except Exception:
        fallback = _data_quality_python_fallback_analysis(state)
        state["llm_analysis"] = json.dumps(fallback, ensure_ascii=False)
        return fallback


def data_quality_synthesizer_node(
    state: dict[str, Any],
    llm_analysis: dict[str, Any],
) -> str:
    lines = [
        "## Executive Summary",
        f"Global score: {int(llm_analysis.get('global_score') or 0)}/100",
        (
            f"Critical issues: {int(llm_analysis.get('critical_count') or 0)} | "
            f"Warnings: {int(llm_analysis.get('warning_count') or 0)} | "
            f"OK: {int(llm_analysis.get('ok_count') or 0)}"
        ),
        "",
        "## Analysis by Column",
    ]

    analysis_by_name = {
        str(item.get("name")): item
        for item in (llm_analysis.get("columns") or [])
        if isinstance(item, dict) and item.get("name")
    }
    for column_name in state.get("columns") or []:
        stats = (state.get("column_stats") or {}).get(column_name, {})
        entry = analysis_by_name.get(column_name, {})
        severity = str(entry.get("severity") or stats.get("severity_hint") or "ok")
        icon = "🔴" if severity == "critical" else "🟡" if severity == "warning" else "🟢"
        label = "Critical" if severity == "critical" else "Warning" if severity == "warning" else "OK"
        lines.append(f"### `{column_name}` — {icon} [{label}]")
        key_bits = [
            f"null_pct={stats.get('null_pct', 0)}%",
            f"distinct_pct={stats.get('distinct_pct', 0)}%",
        ]
        if stats.get("category") == "numeric":
            key_bits.extend([
                f"min={stats.get('min')}",
                f"max={stats.get('max')}",
                f"avg={stats.get('avg')}",
                f"iqr_outliers={stats.get('iqr_outlier_pct', 0)}%",
            ])
        elif stats.get("category") == "string":
            key_bits.extend([
                f"avg_length={stats.get('avg_length')}",
                f"sentinel_pct={stats.get('sentinel_pct', 0)}%",
                f"empty_count={stats.get('empty_count', 0)}",
            ])
        elif stats.get("category") == "date":
            key_bits.extend([
                f"min={stats.get('min')}",
                f"max={stats.get('max')}",
                f"future_count={stats.get('future_count', 0)}",
            ])
        lines.append(f"Key metrics: {', '.join(key_bits)}")
        lines.append(str(entry.get("headline") or (stats.get("issues") or ["No major issue detected."])[0]))
        lines.append(str(entry.get("business_risk") or "No business risk summary was generated."))
        lines.append("")

    lines.extend(["## Top Recommendations"])
    recommendations = llm_analysis.get("top_recommendations") or []
    if recommendations:
        lines.extend(f"- {recommendation}" for recommendation in recommendations[:6])
    else:
        lines.append("- No recommendation was generated.")

    if state.get("volumetric_stats"):
        volumetric = state["volumetric_stats"]
        lines.extend([
            "",
            "## Volumetric Analysis",
            str(llm_analysis.get("volumetric_summary") or "Volumetric analysis was requested for the selected time column."),
            (
                f"Granularity: {volumetric.get('granularity')} | "
                f"Buckets: {volumetric.get('bucket_count')} | "
                f"Average volume: {volumetric.get('avg_volume')} | "
                f"Stddev: {volumetric.get('stddev_volume')}"
            ),
        ])
        low_periods = volumetric.get("anomalously_low_periods") or []
        if low_periods:
            lines.append("Low-volume periods:")
            lines.extend(
                f"- {item.get('bucket')}: {item.get('volume')}"
                for item in low_periods[:8]
            )

    return "\n".join(lines).strip()


DATA_QUALITY_ALL_COLUMNS_OPTION = "All columns"
DATA_QUALITY_NUMERIC_COLUMNS_OPTION = "Numeric columns only"
DATA_QUALITY_TEXT_COLUMNS_OPTION = "Text columns only"
DATA_QUALITY_DATE_COLUMNS_OPTION = "Date columns only"
DATA_QUALITY_CUSTOM_COLUMNS_OPTION = "Custom column list"
DATA_QUALITY_SKIP_TIME_OPTION = "Skip volumetric analysis"
DATA_QUALITY_SKIP_ROW_FILTER_OPTION = "Skip row filter"
DATA_QUALITY_ENTER_ROW_FILTER_OPTION = "Enter row filter manually"
DATA_QUALITY_SAMPLE_OPTIONS = {
    "50,000 rows": 50_000,
    "100,000 rows": 100_000,
    "500,000 rows": 500_000,
    f"Full scan (capped at {DATA_QUALITY_MAX_SAMPLE_ROWS:,} rows)": 0,
}
DATA_QUALITY_CUSTOM_SAMPLE_OPTION = "Custom sample size"
DATA_QUALITY_REVIEW_OPTIONS = [
    "Launch analysis",
    "Edit table",
    "Edit columns",
    "Edit sample size",
    "Edit row filter",
    "Edit time column",
    "Start over",
]


def _data_quality_agent_steps(
    stage_id: str,
    title: str,
    status: str,
    details: str,
    extra_steps: Optional[list[dict[str, Any]]] = None,
) -> list[dict[str, Any]]:
    steps = [
        {
            "id": stage_id,
            "title": title,
            "status": status,
            "details": details,
        }
    ]
    if extra_steps:
        steps.extend(extra_steps)
    return steps


def _data_quality_table_options(state: dict[str, Any]) -> list[str]:
    return (state.get("available_tables") or [])[:DATA_QUALITY_MAX_GUIDED_TABLES]


def _data_quality_column_mode_options(state: dict[str, Any]) -> list[str]:
    schema_info = state.get("schema_info") or []
    options = [DATA_QUALITY_ALL_COLUMNS_OPTION]
    if any(column.get("category") == "numeric" for column in schema_info):
        options.append(DATA_QUALITY_NUMERIC_COLUMNS_OPTION)
    if any(column.get("category") == "string" for column in schema_info):
        options.append(DATA_QUALITY_TEXT_COLUMNS_OPTION)
    if any(column.get("category") == "date" for column in schema_info):
        options.append(DATA_QUALITY_DATE_COLUMNS_OPTION)
    options.append(DATA_QUALITY_CUSTOM_COLUMNS_OPTION)
    return options


def _data_quality_columns_for_mode(mode: str, schema_info: list[dict[str, Any]]) -> list[str]:
    if mode == DATA_QUALITY_ALL_COLUMNS_OPTION:
        return [column["name"] for column in schema_info if column.get("name")]
    if mode == DATA_QUALITY_NUMERIC_COLUMNS_OPTION:
        return [column["name"] for column in schema_info if column.get("category") == "numeric"]
    if mode == DATA_QUALITY_TEXT_COLUMNS_OPTION:
        return [column["name"] for column in schema_info if column.get("category") == "string"]
    if mode == DATA_QUALITY_DATE_COLUMNS_OPTION:
        return [column["name"] for column in schema_info if column.get("category") == "date"]
    return []


def _data_quality_review_payload(state: dict[str, Any]) -> dict[str, Any]:
    raw_sample_size = state.get("sample_size")
    payload = {
        "__dq__": True,
        "table": state.get("table"),
        "columns": state.get("columns") or [],
        "sample_size": 0 if raw_sample_size == 0 else int(raw_sample_size or DATA_QUALITY_DEFAULT_SAMPLE_SIZE),
    }
    if state.get("row_filter"):
        payload["row_filter"] = state["row_filter"]
    if state.get("time_column"):
        payload["time_column"] = state["time_column"]
    return payload


def _data_quality_review_markdown(state: dict[str, Any]) -> str:
    row_filter = state.get("row_filter") or "No row filter"
    time_column = state.get("time_column") or "No volumetric analysis"
    selected_columns = state.get("columns") or []
    lines = [
        "## Data Quality Review",
        f"- Table: `{state.get('table') or 'Not selected yet'}`",
        f"- Columns: {', '.join(f'`{column}`' for column in selected_columns) if selected_columns else 'Not selected yet'}",
        (
            "- Sample size: "
            + (
                f"Full scan capped at {DATA_QUALITY_MAX_SAMPLE_ROWS:,} rows"
                if int(state.get("sample_size") or 0) == 0
                else f"{int(state.get('sample_size') or DATA_QUALITY_DEFAULT_SAMPLE_SIZE):,} rows"
            )
        ),
        f"- Row filter: {row_filter}",
        f"- Time column: {time_column}",
        "",
        "## Structured Payload",
        "```json",
        json.dumps(_data_quality_review_payload(state), ensure_ascii=False, indent=2),
        "```",
        "",
        "## Next Step",
        "Launch the analysis or edit one of the parameters below.",
    ]
    return "\n".join(lines)


def _data_quality_intro_markdown(database_name: str, table_options: list[str], total_tables: int) -> str:
    lines = [
        "## Data quality - Tables",
        (
            "I can guide you through a table quality assessment in English, then run statistical profiling "
            "and local-LLM scoring."
        ),
        "",
        "## What I Need",
        "- A target table",
        "- Which columns to profile",
        "- A sample size (or a capped full scan)",
        "- An optional row filter",
        "- An optional time column for volumetric analysis",
        "",
        f"I found {total_tables} table(s) in `{database_name}`.",
    ]
    if total_tables > len(table_options):
        lines.append(
            f"Choose from the suggestions below or type another exact table name from the remaining {total_tables - len(table_options)} table(s)."
        )
    else:
        lines.append("Choose the table you want to profile.")
    return "\n".join(lines)


def _data_quality_guess_table_from_message(user_message: str, available_tables: list[str]) -> Optional[str]:
    direct = resolve_user_choice(user_message, available_tables)
    if direct:
        return direct
    normalized = normalize_choice(user_message).lower()
    if not normalized:
        return None
    matches = [
        table for table in available_tables
        if re.search(rf"(?<![a-z0-9_]){re.escape(table.lower())}(?![a-z0-9_])", normalized)
    ]
    if len(matches) == 1:
        return matches[0]
    return None


def _data_quality_guess_columns_from_message(
    user_message: str,
    schema_info: list[dict[str, Any]],
) -> list[str]:
    normalized = normalize_choice(user_message).lower()
    if not normalized:
        return []
    exact = []
    for column in schema_info:
        name = str(column.get("name") or "").strip()
        if not name:
            continue
        if re.search(rf"(?<![a-z0-9_]){re.escape(name.lower())}(?![a-z0-9_])", normalized):
            exact.append(name)
    return exact


def _try_extract_data_quality_payload(user_message: str) -> Optional[dict[str, Any]]:
    parsed = extract_json_object(user_message)
    if isinstance(parsed, dict) and parsed.get("__dq__") is True:
        return parsed
    return None


CHART_CREATE_OPTION = "Create a chart"
CHART_SKIP_OPTION = "Keep text answer only"
CHART_TYPE_LABELS = {
    "bar": "Bar chart",
    "line": "Line chart",
    "area": "Area chart",
    "scatter": "Scatter plot",
}
CHART_TYPE_BY_LABEL = {label.lower(): key for key, label in CHART_TYPE_LABELS.items()}


def dump_clickhouse_agent_state(state: ClickHouseAgentState) -> dict[str, Any]:
    return state.model_dump(by_alias=True)


def enforce_clickhouse_preview_limit(sql: str, row_limit: int) -> str:
    cleaned = clean_sql_text(sql)
    safe_limit = max(1, min(int(row_limit or 1000), 10_000))
    lowered = cleaned.lower()
    if lowered.startswith(("show ", "describe ", "desc ", "exists ", "explain ")):
        return cleaned
    if re.search(r"\blimit\s+\d+\b", lowered):
        return cleaned
    return f"{cleaned}\nLIMIT {safe_limit}"


def reset_clickhouse_chart_state(state: ClickHouseAgentState) -> None:
    state.chart_requested = False
    state.chart_suggested = False
    state.chart_offer_options = []
    state.chart_x_options = []
    state.chart_y_options = []
    state.chart_type_options = []
    state.selected_chart_x = None
    state.selected_chart_y = None
    state.selected_chart_type = None


def reset_clickhouse_query_resolution(state: ClickHouseAgentState) -> None:
    state.candidate_fields = []
    state.date_fields = []
    state.selected_field = None
    state.selected_date_field = None
    reset_clickhouse_clarification(state)
    reset_clickhouse_chart_state(state)


def detect_chart_request(text: str) -> bool:
    normalized = (text or "").lower()
    return any(
        keyword in normalized
        for keyword in [
            "chart", "graph", "plot", "visual", "visualize", "visualise",
            "dashboard", "trend chart", "bar chart", "line chart", "scatter",
            "graphique", "graphe", "courbe", "histogram",
        ]
    )


def is_chart_followup_request(text: str) -> bool:
    normalized = (text or "").lower().strip()
    return any(
        phrase in normalized
        for phrase in [
            "create a chart",
            "generate a chart",
            "show a chart",
            "make a chart",
            "plot this",
            "graph this",
            "visualize this",
            "visualise this",
            "show me a graph",
            "create graph",
            "make graph",
        ]
    )


def is_affirmative_response(text: str) -> bool:
    normalized = normalize_choice(text).lower()
    return normalized in {
        "yes",
        "y",
        "ok",
        "okay",
        "sure",
        "please do",
        "do it",
        "go ahead",
        "why not",
        "yes please",
    }


def is_negative_response(text: str) -> bool:
    normalized = normalize_choice(text).lower()
    return normalized in {
        "no",
        "n",
        "nope",
        "not now",
        "skip",
        "cancel",
        "keep text",
        "text only",
        "no thanks",
    }


def detect_requested_chart_type(text: str) -> Optional[str]:
    normalized = (text or "").lower()
    if "scatter" in normalized:
        return "scatter"
    if "area" in normalized:
        return "area"
    if "line" in normalized or "curve" in normalized:
        return "line"
    if "bar" in normalized or "histogram" in normalized:
        return "bar"
    return None


def is_numeric_clickhouse_type(type_name: str) -> bool:
    lowered = (type_name or "").lower()
    return any(
        token in lowered
        for token in [
            "int", "float", "decimal", "numeric", "double", "real",
        ]
    ) and "interval" not in lowered


def is_temporal_clickhouse_type(type_name: str) -> bool:
    lowered = (type_name or "").lower()
    return "date" in lowered or "time" in lowered


def normalize_chart_value(value: Any) -> Optional[float]:
    if isinstance(value, bool):
        return float(value)
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        try:
            return float(value.replace(",", ""))
        except ValueError:
            return None
    return None


def infer_chart_options(meta: list[dict], rows: list[dict]) -> dict[str, Any]:
    if len(rows) < 2:
        return {
            "can_chart": False,
            "recommended": False,
            "x_options": [],
            "y_options": [],
            "type_options": [],
        }

    numeric_columns = [col["name"] for col in meta if is_numeric_clickhouse_type(col.get("type", ""))]
    temporal_columns = [col["name"] for col in meta if is_temporal_clickhouse_type(col.get("type", ""))]
    text_columns = [
        col["name"] for col in meta
        if col.get("name") not in numeric_columns and col.get("name") not in temporal_columns
    ]

    x_options = temporal_columns + text_columns
    if not x_options and len(numeric_columns) >= 2:
        x_options = numeric_columns[:-1]

    y_options = numeric_columns

    if not x_options or not y_options:
        return {
            "can_chart": False,
            "recommended": False,
            "x_options": [],
            "y_options": [],
            "type_options": [],
        }

    unique_counts = {
        column_name: len({str(row.get(column_name, "")) for row in rows if row.get(column_name) is not None})
        for column_name in x_options
    }

    filtered_x_options = [
        column_name for column_name in x_options
        if unique_counts.get(column_name, 0) <= min(40, len(rows))
    ] or x_options

    uses_temporal_x = any(column_name in temporal_columns for column_name in filtered_x_options)
    uses_numeric_x = all(column_name in numeric_columns for column_name in filtered_x_options)

    type_options = ["Bar chart", "Line chart", "Area chart"]
    if uses_numeric_x and len(numeric_columns) >= 2:
        type_options = ["Scatter plot", "Line chart", "Bar chart"]
    elif uses_temporal_x:
        type_options = ["Line chart", "Area chart", "Bar chart"]

    recommended = len(filtered_x_options) > 0 and len(y_options) > 0 and len(rows) >= 3
    return {
        "can_chart": True,
        "recommended": recommended,
        "x_options": filtered_x_options,
        "y_options": y_options,
        "type_options": type_options,
    }


def build_chart(
    rows: list[dict],
    x_field: str,
    y_field: str,
    chart_type: str,
) -> Optional[dict[str, Any]]:
    points: list[dict[str, Any]] = []
    for row in rows:
        x_raw = row.get(x_field)
        y_raw = normalize_chart_value(row.get(y_field))
        if x_raw is None or y_raw is None:
            continue
        points.append({"x": str(x_raw), "y": y_raw})

    if len(points) < 2:
        return None

    points = points[:30]
    return {
        "type": chart_type,
        "title": f"{y_field} by {x_field}",
        "xField": x_field,
        "yField": y_field,
        "points": points,
    }


def initialize_chart_selection(
    state: ClickHouseAgentState,
    x_options: list[str],
    y_options: list[str],
    type_options: list[str],
    requested_chart_type: Optional[str] = None,
) -> None:
    state.chart_requested = True
    state.chart_x_options = x_options
    state.chart_y_options = y_options
    state.chart_type_options = type_options

    if not state.selected_chart_x and len(x_options) == 1:
        state.selected_chart_x = x_options[0]

    filtered_y_options = [
        option for option in y_options
        if option != state.selected_chart_x
    ] or y_options

    if not state.selected_chart_y and len(filtered_y_options) == 1:
        state.selected_chart_y = filtered_y_options[0]

    if (
        not state.selected_chart_type
        and requested_chart_type
        and CHART_TYPE_LABELS.get(requested_chart_type) in type_options
    ):
        state.selected_chart_type = requested_chart_type

    if not state.selected_chart_type and len(type_options) == 1:
        state.selected_chart_type = CHART_TYPE_BY_LABEL.get(type_options[0].lower())


def next_chart_prompt(state: ClickHouseAgentState) -> Optional[dict[str, Any]]:
    x_options = state.chart_x_options
    y_options = [
        option for option in state.chart_y_options
        if option != state.selected_chart_x
    ] or state.chart_y_options

    if not state.selected_chart_x:
        state.stage = "awaiting_chart_x"
        return {
            "title": "Chart X Axis",
            "prompt": "Choose the field to use on the X axis.",
            "options": x_options,
            "step_id": "ch-chart-x",
            "step_title": "Waiting for X axis selection",
            "step_details": "The user must choose which field should drive the horizontal axis.",
        }

    if not state.selected_chart_y:
        state.stage = "awaiting_chart_y"
        return {
            "title": "Chart Y Axis",
            "prompt": "Choose the metric to use on the Y axis.",
            "options": y_options,
            "step_id": "ch-chart-y",
            "step_title": "Waiting for Y axis selection",
            "step_details": "The user must choose the metric to visualize.",
        }

    if not state.selected_chart_type:
        state.stage = "awaiting_chart_type"
        return {
            "title": "Chart Type",
            "prompt": "Choose the chart type.",
            "options": state.chart_type_options,
            "step_id": "ch-chart-type",
            "step_title": "Waiting for chart type",
            "step_details": "The user must choose how to visualize the selected axes.",
        }

    return None

# ── Text utilities ────────────────────────────────────────────────────────────

def keyword_score(query: str, text: str) -> float:
    terms = [t for t in re.split(r"\W+", query.lower()) if len(t) > 2]
    if not terms:
        return 0.0
    text_lower = text.lower()
    return sum(1 for t in terms if t in text_lower) / len(terms)


def chunk_text(text: str, max_words: int = 200, overlap_sentences: int = 2) -> list[str]:
    sentences = re.findall(r"[^.!?]+[.!?]+", text) or [text]
    result: list[str] = []
    current: list[str] = []
    current_words = 0
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        word_count = len(sentence.split())
        if current_words + word_count > max_words and current:
            result.append(" ".join(current))
            overlap = current[-overlap_sentences:]
            current = overlap + [sentence]
            current_words = sum(len(s.split()) for s in current)
        else:
            current.append(sentence)
            current_words += word_count
    if current:
        result.append(" ".join(current))
    return result or [text]


# ── Embedding helper ──────────────────────────────────────────────────────────

async def get_embedding(
    text: str,
    base_url: str,
    model: str,
    api_key: str = None,
    verify_ssl: bool = True,
) -> list[float]:
    """Get a vector embedding via an OpenAI-compatible /embeddings endpoint.

    base_url may be a base path (e.g. ``http://host/v1``) — ``/embeddings`` is
    appended automatically — or the full endpoint URL already ending with
    ``/embeddings`` or ``:embeddings`` (e.g. ``http://host/v1/openai/embeddings``
    or ``https://.../v2:embeddings``), used as-is.
    """
    stripped = _normalize_local_service_url(base_url)
    lowered = stripped.lower()
    is_direct_endpoint = lowered.endswith("/embeddings") or lowered.endswith(":embeddings")
    url = stripped if is_direct_endpoint else stripped + "/embeddings"
    headers = {"Content-Type": "application/json"}
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"
    payload: dict[str, Any] = {"input": text}
    if str(model or "").strip():
        payload["model"] = model
    try:
        async with httpx.AsyncClient(
            **_httpx_async_client_kwargs(url, timeout=60.0, verify=verify_ssl)
        ) as client:
            resp = await client.post(url, json=payload, headers=headers)
            resp.raise_for_status()
            data = _parse_http_json_response(
                resp,
                service_label="Embedding endpoint",
                endpoint=url,
            )
            try:
                return data["data"][0]["embedding"]
            except (KeyError, IndexError, TypeError) as exc:
                raise UpstreamServiceError(
                    f"Embedding endpoint returned an unexpected JSON payload at `{url}`."
                ) from exc
    except httpx.HTTPError as exc:
        raise UpstreamServiceError(
            f"Embedding endpoint error at `{url}`: {exc}. "
            "Check that the embedding service is running and reachable from `server.py`."
        ) from exc


def _is_direct_embedding_endpoint(url: str) -> bool:
    lowered = str(url or "").strip().rstrip("/").lower()
    return lowered.endswith("/embeddings") or lowered.endswith(":embeddings")


def _derive_embedding_models_endpoint(base_url: str) -> tuple[str, str]:
    """Resolve the best `/models` URL for an embedding configuration.

    The embedding setting may be either:
    - a base URL such as `https://host/v1`
    - a direct embedding endpoint such as `https://host/v1/embeddings`
    - a provider-specific endpoint such as `https://host/v2:embeddings`

    This helper returns the discovery endpoint plus a short explanation of how
    the URL was interpreted so the UI can give clear feedback to the user.
    """
    normalized = _normalize_local_service_url(base_url).rstrip("/")
    lowered = normalized.lower()
    if lowered.endswith("/embeddings"):
        return normalized[: -len("/embeddings")] + "/models", "derived from the direct `/embeddings` endpoint"
    if lowered.endswith(":embeddings"):
        return normalized[: -len(":embeddings")] + ":models", "derived from the direct `:embeddings` endpoint"
    return normalized + "/models", "derived from the base embedding URL"


# ── LLM helper ────────────────────────────────────────────────────────────────

def _truncate_body_preview(value: str, limit: int = 260) -> str:
    compact = re.sub(r"\s+", " ", (value or "")).strip()
    if not compact:
        return "<empty body>"
    if len(compact) <= limit:
        return compact
    return compact[: limit - 3] + "..."


def _parse_http_json_response(
    response: httpx.Response,
    *,
    service_label: str,
    endpoint: str,
) -> Any:
    try:
        return response.json()
    except json.JSONDecodeError as exc:
        content_type = response.headers.get("content-type", "unknown")
        preview = _truncate_body_preview(response.text)
        raise UpstreamServiceError(
            (
                f"{service_label} returned a non-JSON response at `{endpoint}` "
                f"(HTTP {response.status_code}, content-type `{content_type}`). "
                f"Body preview: {preview}. Check the configured Base URL and make sure "
                "the local service exposes an OpenAI/Ollama-compatible JSON API."
            )
        ) from exc


def _extract_ollama_message_content(payload: Any, endpoint: str) -> str:
    if not isinstance(payload, dict):
        raise UpstreamServiceError(
            f"Ollama endpoint returned an unexpected JSON payload at `{endpoint}`."
        )
    content = payload.get("message", {}).get("content", "")
    if isinstance(content, str):
        return content
    raise UpstreamServiceError(
        f"Ollama endpoint returned a JSON payload without `message.content` at `{endpoint}`."
    )


def _extract_openai_message_content(payload: Any, endpoint: str) -> str:
    if not isinstance(payload, dict):
        raise UpstreamServiceError(
            f"OpenAI-compatible LLM endpoint returned an unexpected JSON payload at `{endpoint}`."
        )
    try:
        content = payload["choices"][0]["message"]["content"]
    except (KeyError, IndexError, TypeError) as exc:
        raise UpstreamServiceError(
            (
                f"OpenAI-compatible LLM endpoint returned a JSON payload without "
                f"`choices[0].message.content` at `{endpoint}`."
            )
        ) from exc
    if isinstance(content, str):
        return content
    raise UpstreamServiceError(
        (
            f"OpenAI-compatible LLM endpoint returned a non-text "
            f"`choices[0].message.content` at `{endpoint}`."
        )
    )

async def llm_chat(
    messages: list[dict],
    base_url: str,
    model: str,
    provider: str = "ollama",
    api_key: str = None,
    response_format: str = None,
    disable_ssl_verification: Optional[bool] = None,
) -> str:
    headers = {"Content-Type": "application/json"}
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"
    normalized_base_url = _normalize_local_service_url(base_url)
    if disable_ssl_verification is None:
        try:
            state = await read_db_state()
            disable_ssl_verification = _ssl_verification_disabled(state.get("config") or {})
        except Exception:
            disable_ssl_verification = False
    effective_verify = False if disable_ssl_verification else None

    if provider == "ollama":
        payload: dict = {"model": model, "messages": messages, "stream": False}
        if response_format == "json":
            payload["format"] = "json"
        endpoint = normalized_base_url.rstrip("/") + "/api/chat"
        try:
            async with httpx.AsyncClient(
                **_httpx_async_client_kwargs(endpoint, timeout=120.0, verify=effective_verify)
            ) as client:
                resp = await client.post(endpoint, json=payload, headers=headers)
                resp.raise_for_status()
                data = _parse_http_json_response(
                    resp,
                    service_label="Ollama endpoint",
                    endpoint=endpoint,
                )
                return _extract_ollama_message_content(data, endpoint)
        except httpx.HTTPError as exc:
            raise UpstreamServiceError(
                f"Ollama endpoint error at `{endpoint}`: {exc}. "
                "Check the Base URL, the selected model, and that the local LLM server is running."
            ) from exc
    else:
        payload = {"model": model, "messages": messages, "stream": False}
        if response_format == "json":
            payload["response_format"] = {"type": "json_object"}
        endpoint = normalized_base_url.rstrip("/") + "/chat/completions"
        try:
            async with httpx.AsyncClient(
                **_httpx_async_client_kwargs(endpoint, timeout=120.0, verify=effective_verify)
            ) as client:
                resp = await client.post(endpoint, json=payload, headers=headers)
                resp.raise_for_status()
                data = _parse_http_json_response(
                    resp,
                    service_label="OpenAI-compatible LLM endpoint",
                    endpoint=endpoint,
                )
                return _extract_openai_message_content(data, endpoint)
        except httpx.HTTPError as exc:
            raise UpstreamServiceError(
                f"OpenAI-compatible LLM endpoint error at `{endpoint}`: {exc}. "
                "Check the Base URL, API key, and that the model server is reachable from `server.py`."
            ) from exc


# ── File Management agent helpers ────────────────────────────────────────────

FILE_PREVIEW_CHAR_LIMIT = 3000
FILE_TABULAR_PREVIEW_ROWS = 50
FILE_SEARCH_RESULTS_LIMIT = 20
FILE_MANAGER_MAX_ITERATIONS = 15
FILE_MANAGER_WRITE_TOOL_LIMIT = 1000
FILE_MANAGER_TEXT_EXTENSIONS = {
    ".txt", ".md", ".py", ".sql", ".json", ".yaml", ".yml", ".html", ".css",
    ".js", ".ts", ".tsx", ".jsx", ".csv", ".tsv", ".log", ".ini", ".toml",
    ".xml", ".sh", ".env", ".rst",
}
FILE_MANAGER_SPREADSHEET_EXTENSIONS = {".csv", ".tsv", ".xlsx", ".xls"}
FILE_MANAGER_CONFIRMATION_TOOLS = {
    "write_file",
    "delete_file",
    "delete_directory",
    "move_file",
    "write_excel_sheet",
    "edit_excel_cells",
    "delete_excel_sheet",
}


def _truncate_text_preview(text: str, limit: int = FILE_PREVIEW_CHAR_LIMIT) -> str:
    value = str(text or "")
    return value if len(value) <= limit else value[: limit - 1] + "…"


def _file_tool_result(
    summary: str,
    *,
    preview: str = "",
    data: Any = None,
    visited_path: str = "",
    requires_confirmation: bool = False,
    pending_action: Optional[dict[str, Any]] = None,
) -> dict[str, Any]:
    return {
        "summary": summary.strip(),
        "preview": _truncate_text_preview(preview),
        "data": data,
        "visited_path": visited_path,
        "requires_confirmation": requires_confirmation,
        "pending_action": pending_action,
    }


def _import_openpyxl():
    try:
        from openpyxl import Workbook, load_workbook
        return Workbook, load_workbook
    except ImportError as exc:
        raise ValueError("The optional dependency `openpyxl` is required for Excel tools.") from exc


def _import_xlrd():
    try:
        import xlrd
        return xlrd
    except ImportError as exc:
        raise ValueError("The optional dependency `xlrd` is required to read `.xls` files.") from exc


def _import_docx_document():
    try:
        from docx import Document
        return Document
    except ImportError as exc:
        raise ValueError("The optional dependency `python-docx` is required to read `.docx` files.") from exc


def _import_pyarrow_parquet():
    try:
        import pyarrow.parquet as pq
        return pq
    except ImportError as exc:
        raise ValueError("The optional dependency `pyarrow` is required to read `.parquet` files.") from exc


def _resolve_agent_path(path_value: str, base_path: str = "") -> Path:
    raw_path = str(path_value or ".").strip() or "."
    sandbox_root = Path(base_path).expanduser().resolve() if base_path else None

    if os.path.isabs(raw_path):
        resolved = Path(raw_path).expanduser().resolve()
    else:
        anchor = sandbox_root or Path.cwd().resolve()
        resolved = (anchor / raw_path).resolve()

    if sandbox_root:
        try:
            resolved.relative_to(sandbox_root)
        except ValueError as exc:
            raise ValueError("Path escapes the configured access root.") from exc

    return resolved


def _ensure_parent_directory(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def _read_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8", errors="replace")


def _format_table_rows(rows: list[dict[str, Any]]) -> str:
    if not rows:
        return "No rows found."
    headers = list(rows[0].keys())
    header_line = "| " + " | ".join(headers) + " |"
    divider = "| " + " | ".join(["---"] * len(headers)) + " |"
    body_lines = []
    for row in rows[:FILE_TABULAR_PREVIEW_ROWS]:
        body_lines.append("| " + " | ".join(str(row.get(header, "")) for header in headers) + " |")
    return "\n".join([header_line, divider, *body_lines])


def _normalize_excel_rows(headers: list[Any], rows: list[list[Any]]) -> list[dict[str, Any]]:
    safe_headers = []
    for index, header in enumerate(headers, start=1):
        value = str(header).strip() if header is not None else ""
        safe_headers.append(value or f"column_{index}")
    return [
        {safe_headers[idx]: row[idx] if idx < len(row) else "" for idx in range(len(safe_headers))}
        for row in rows
    ]


def _load_excel_sheet_preview(path: Path, sheet_name: Optional[str] = None) -> tuple[str, list[dict[str, Any]], str]:
    extension = path.suffix.lower()
    if extension == ".xlsx":
        _, load_workbook = _import_openpyxl()
        workbook = load_workbook(path, read_only=True, data_only=True)
        target_sheet = sheet_name or workbook.sheetnames[0]
        if target_sheet not in workbook.sheetnames:
            raise ValueError(f"Sheet `{target_sheet}` was not found.")
        worksheet = workbook[target_sheet]
        rows = list(worksheet.iter_rows(values_only=True))
        headers = list(rows[0]) if rows else []
        data_rows = [list(row) for row in rows[1:1 + FILE_TABULAR_PREVIEW_ROWS]] if len(rows) > 1 else []
        preview_rows = _normalize_excel_rows(headers, data_rows) if headers else []
        return target_sheet, preview_rows, ", ".join(workbook.sheetnames)

    if extension == ".xls":
        xlrd = _import_xlrd()
        workbook = xlrd.open_workbook(path)
        target_sheet = sheet_name or workbook.sheet_names()[0]
        worksheet = workbook.sheet_by_name(target_sheet)
        headers = worksheet.row_values(0) if worksheet.nrows > 0 else []
        data_rows = [worksheet.row_values(index) for index in range(1, min(worksheet.nrows, FILE_TABULAR_PREVIEW_ROWS + 1))]
        preview_rows = _normalize_excel_rows(headers, data_rows) if headers else []
        return target_sheet, preview_rows, ", ".join(workbook.sheet_names())

    raise ValueError("Excel tools support `.xlsx` and `.xls` files only.")


def _read_csv_rows(path: Path, delimiter: Optional[str] = None) -> tuple[list[str], list[dict[str, Any]], int]:
    detected_delimiter = delimiter or ("\t" if path.suffix.lower() == ".tsv" else ",")
    rows_preview: list[dict[str, Any]] = []
    total_rows = 0
    with path.open("r", encoding="utf-8", errors="replace", newline="") as handle:
        reader = csv.DictReader(handle, delimiter=detected_delimiter)
        headers = reader.fieldnames or []
        for row in reader:
            total_rows += 1
            if len(rows_preview) < FILE_TABULAR_PREVIEW_ROWS:
                rows_preview.append(dict(row))
    return headers, rows_preview, total_rows


def list_directory_tool(path: str = ".", recursive: bool = False, base_path: str = "") -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    if not target.exists():
        raise ValueError(f"Directory `{target}` does not exist.")
    if not target.is_dir():
        raise ValueError(f"`{target}` is not a directory.")

    iterator = target.rglob("*") if recursive else target.iterdir()
    entries = []
    for item in iterator:
        entry_type = "dir" if item.is_dir() else "file"
        size = item.stat().st_size if item.is_file() else 0
        entries.append(
            {
                "name": item.name,
                "path": str(item),
                "type": entry_type,
                "size": size,
            }
        )
        if len(entries) >= FILE_TABULAR_PREVIEW_ROWS:
            break

    entries.sort(key=lambda entry: (entry["type"] != "dir", entry["name"].lower(), entry["path"].lower()))

    preview_lines = [
        f"- `{Path(entry['path']).name}{'/' if entry['type'] == 'dir' else ''}` · {entry['type']} · {entry['size']} bytes"
        for entry in entries
    ] or ["No entries found."]
    return _file_tool_result(
        f"Listed `{target}`.",
        preview="\n".join(preview_lines),
        data=entries,
        visited_path=str(target),
    )


def get_file_info_tool(path: str, base_path: str = "") -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    if not target.exists():
        raise ValueError(f"Path `{target}` does not exist.")

    stats = target.stat()
    info = {
        "path": str(target),
        "name": target.name,
        "isDirectory": target.is_dir(),
        "size": stats.st_size,
        "modifiedAt": datetime.fromtimestamp(stats.st_mtime, tz=timezone.utc).isoformat(),
        "createdAt": datetime.fromtimestamp(stats.st_ctime, tz=timezone.utc).isoformat(),
        "suffix": target.suffix.lower(),
    }
    preview = "\n".join(f"- **{key}**: {value}" for key, value in info.items())
    return _file_tool_result(
        f"Loaded metadata for `{target.name}`.",
        preview=preview,
        data=info,
        visited_path=str(target if target.is_dir() else target.parent),
    )


def search_files_tool(path: str = ".", query: str = "", recursive: bool = True, base_path: str = "") -> dict[str, Any]:
    if not query.strip():
        raise ValueError("A search query is required.")

    root = _resolve_agent_path(path, base_path)
    if not root.exists() or not root.is_dir():
        raise ValueError(f"Directory `{root}` does not exist.")

    matches = []
    lowered_query = query.lower().strip()
    iterator = root.rglob("*") if recursive else root.iterdir()
    for item in iterator:
        if len(matches) >= FILE_SEARCH_RESULTS_LIMIT:
            break

        relative_path = str(item.relative_to(root))
        match_kind = None
        if lowered_query in relative_path.lower():
            match_kind = "name"
        elif item.is_file() and item.suffix.lower() in FILE_MANAGER_TEXT_EXTENSIONS and item.stat().st_size <= 1_000_000:
            content = _read_text_file(item)
            if lowered_query in content.lower():
                match_kind = "content"

        if match_kind:
            matches.append(
                {
                    "path": str(item),
                    "relativePath": relative_path,
                    "type": "dir" if item.is_dir() else "file",
                    "matchKind": match_kind,
                }
            )

    matches.sort(key=lambda item: (item["type"] != "dir", item["relativePath"].lower(), item["matchKind"]))

    preview = "\n".join(
        f"- `{item['relativePath']}` · {item['type']} · matched by {item['matchKind']}"
        for item in matches
    ) or "No match found."
    return _file_tool_result(
        f"Found {len(matches)} matching item(s) in `{root}`.",
        preview=preview,
        data=matches,
        visited_path=str(root),
    )


def read_file_tool(path: str, base_path: str = "") -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    if not target.exists() or not target.is_file():
        raise ValueError(f"File `{target}` does not exist.")

    suffix = target.suffix.lower()
    if suffix in FILE_MANAGER_TEXT_EXTENSIONS:
        content = _read_text_file(target)
        return _file_tool_result(
            f"Read text file `{target.name}`.",
            preview=f"```text\n{_truncate_text_preview(content)}\n```",
            data={"path": str(target), "content": _truncate_text_preview(content)},
            visited_path=str(target.parent),
        )

    if suffix in {".csv", ".tsv"}:
        headers, rows, total_rows = _read_csv_rows(target)
        preview = (
            f"Rows shown: {len(rows)} / {total_rows}\n\n"
            f"{_format_table_rows(rows)}"
        )
        return _file_tool_result(
            f"Read tabular file `{target.name}`.",
            preview=preview,
            data={"headers": headers, "rows": rows, "totalRows": total_rows},
            visited_path=str(target.parent),
        )

    if suffix in {".xlsx", ".xls"}:
        sheet_name, rows, sheet_list = _load_excel_sheet_preview(target)
        preview = (
            f"Sheet: `{sheet_name}`\n"
            f"Available sheets: {sheet_list}\n\n"
            f"{_format_table_rows(rows)}"
        )
        return _file_tool_result(
            f"Read Excel file `{target.name}`.",
            preview=preview,
            data={"sheet": sheet_name, "rows": rows, "sheetNames": sheet_list.split(", ") if sheet_list else []},
            visited_path=str(target.parent),
        )

    if suffix == ".docx":
        Document = _import_docx_document()
        document = Document(target)
        content = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        return _file_tool_result(
            f"Read Word document `{target.name}`.",
            preview=f"```text\n{_truncate_text_preview(content)}\n```",
            data={"path": str(target), "content": _truncate_text_preview(content)},
            visited_path=str(target.parent),
        )

    if suffix == ".parquet":
        pq = _import_pyarrow_parquet()
        parquet_file = pq.ParquetFile(target)
        batch = next(parquet_file.iter_batches(batch_size=FILE_TABULAR_PREVIEW_ROWS), None)
        rows = batch.to_pylist() if batch is not None else []
        preview = (
            f"Rows shown: {len(rows)} / {parquet_file.metadata.num_rows}\n\n"
            f"{_format_table_rows(rows)}"
        )
        return _file_tool_result(
            f"Read Parquet file `{target.name}`.",
            preview=preview,
            data={"rows": rows, "totalRows": parquet_file.metadata.num_rows},
            visited_path=str(target.parent),
        )

    raise ValueError(f"Unsupported file format `{suffix or 'unknown'}` for read_file.")


def read_csv_summary_tool(path: str, base_path: str = "") -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    if not target.exists() or not target.is_file():
        raise ValueError(f"File `{target}` does not exist.")
    if target.suffix.lower() not in {".csv", ".tsv"}:
        raise ValueError("read_csv_summary only supports `.csv` and `.tsv` files.")

    headers, rows, total_rows = _read_csv_rows(target)
    preview = (
        f"Columns: {', '.join(headers) if headers else 'No header'}\n"
        f"Rows shown: {len(rows)} / {total_rows}\n\n"
        f"{_format_table_rows(rows)}"
    )
    return _file_tool_result(
        f"Loaded a CSV summary for `{target.name}`.",
        preview=preview,
        data={"headers": headers, "rows": rows, "totalRows": total_rows},
        visited_path=str(target.parent),
    )


def create_directory_tool(path: str, base_path: str = "") -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    target.mkdir(parents=True, exist_ok=True)
    return _file_tool_result(
        f"Created directory `{target}`.",
        preview=f"Directory ready at `{target}`.",
        visited_path=str(target),
    )


def create_file_tool(path: str, content: str = "", base_path: str = "") -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    if target.exists():
        raise ValueError(f"File `{target}` already exists.")
    if target.suffix.lower() not in FILE_MANAGER_TEXT_EXTENSIONS:
        raise ValueError("create_file only supports text-based file formats.")
    _ensure_parent_directory(target)
    target.write_text(str(content), encoding="utf-8")
    return _file_tool_result(
        f"Created file `{target.name}`.",
        preview=f"```text\n{_truncate_text_preview(content)}\n```",
        visited_path=str(target.parent),
    )


def create_excel_file_tool(
    path: str,
    sheet_name: str = "Sheet1",
    headers: Optional[list[Any]] = None,
    rows: Optional[list[list[Any]]] = None,
    base_path: str = "",
) -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    if target.exists():
        raise ValueError(f"File `{target}` already exists.")
    if target.suffix.lower() != ".xlsx":
        raise ValueError("create_excel_file currently supports `.xlsx` files only.")
    Workbook, _ = _import_openpyxl()
    _ensure_parent_directory(target)
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = sheet_name or "Sheet1"
    if headers:
        worksheet.append(list(headers))
    for row in rows or []:
        worksheet.append(list(row))
    workbook.save(target)
    return _file_tool_result(
        f"Created Excel file `{target.name}` with sheet `{worksheet.title}`.",
        preview=f"Workbook saved to `{target}`.",
        visited_path=str(target.parent),
    )


def write_file_tool(path: str, content: str, confirmed: bool = False, base_path: str = "") -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    if target.suffix.lower() not in FILE_MANAGER_TEXT_EXTENSIONS:
        raise ValueError("write_file only supports text-based file formats.")
    existing_content = _read_text_file(target) if target.exists() else ""
    preview = (
        f"Target: `{target}`\n"
        f"Existing size: {len(existing_content)} characters\n"
        f"New size: {len(str(content))} characters\n\n"
        f"## New content preview\n```text\n{_truncate_text_preview(str(content))}\n```"
    )
    if not confirmed:
        return _file_tool_result(
            f"Writing `{target.name}` will overwrite the current content.",
            preview=preview,
            visited_path=str(target.parent),
            requires_confirmation=True,
            pending_action={
                "tool_name": "write_file",
                "tool_input": {"path": path, "content": str(content)},
            },
        )
    _ensure_parent_directory(target)
    target.write_text(str(content), encoding="utf-8")
    return _file_tool_result(
        f"Wrote `{target.name}` successfully.",
        preview=f"```text\n{_truncate_text_preview(str(content))}\n```",
        visited_path=str(target.parent),
    )


def _load_excel_for_write(path: Path):
    _, load_workbook = _import_openpyxl()
    if not path.exists():
        raise ValueError(f"Excel file `{path}` does not exist.")
    if path.suffix.lower() != ".xlsx":
        raise ValueError("Excel write operations currently support `.xlsx` files only.")
    return load_workbook(path)


def list_excel_sheets_tool(path: str, base_path: str = "") -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    if not target.exists() or not target.is_file():
        raise ValueError(f"Excel file `{target}` does not exist.")
    _, _, sheet_list = _load_excel_sheet_preview(target)
    preview = "\n".join(f"- `{sheet}`" for sheet in sheet_list.split(", ") if sheet) or "No sheet found."
    return _file_tool_result(
        f"Listed sheets for `{target.name}`.",
        preview=preview,
        data={"sheets": [sheet for sheet in sheet_list.split(", ") if sheet]},
        visited_path=str(target.parent),
    )


def read_excel_sheet_tool(path: str, sheet_name: Optional[str] = None, base_path: str = "") -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    if not target.exists() or not target.is_file():
        raise ValueError(f"Excel file `{target}` does not exist.")
    resolved_sheet_name, rows, sheet_list = _load_excel_sheet_preview(target, sheet_name)
    preview = (
        f"Sheet: `{resolved_sheet_name}`\n"
        f"Available sheets: {sheet_list}\n\n"
        f"{_format_table_rows(rows)}"
    )
    return _file_tool_result(
        f"Read sheet `{resolved_sheet_name}` from `{target.name}`.",
        preview=preview,
        data={"sheet": resolved_sheet_name, "rows": rows},
        visited_path=str(target.parent),
    )


def add_excel_sheet_tool(path: str, sheet_name: str, base_path: str = "") -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    workbook = _load_excel_for_write(target)
    if sheet_name in workbook.sheetnames:
        raise ValueError(f"Sheet `{sheet_name}` already exists.")
    workbook.create_sheet(title=sheet_name)
    workbook.save(target)
    return _file_tool_result(
        f"Added sheet `{sheet_name}` to `{target.name}`.",
        preview="\n".join(f"- `{sheet}`" for sheet in workbook.sheetnames),
        visited_path=str(target.parent),
    )


def rename_excel_sheet_tool(path: str, old_name: str, new_name: str, base_path: str = "") -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    workbook = _load_excel_for_write(target)
    if old_name not in workbook.sheetnames:
        raise ValueError(f"Sheet `{old_name}` does not exist.")
    if new_name in workbook.sheetnames:
        raise ValueError(f"Sheet `{new_name}` already exists.")
    workbook[old_name].title = new_name
    workbook.save(target)
    return _file_tool_result(
        f"Renamed sheet `{old_name}` to `{new_name}` in `{target.name}`.",
        preview="\n".join(f"- `{sheet}`" for sheet in workbook.sheetnames),
        visited_path=str(target.parent),
    )


def write_excel_sheet_tool(
    path: str,
    sheet_name: str,
    headers: Optional[list[Any]] = None,
    rows: Optional[list[list[Any]]] = None,
    confirmed: bool = False,
    base_path: str = "",
) -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    preview_rows = _normalize_excel_rows(list(headers or []), [list(row) for row in (rows or [])[:10]]) if headers else []
    preview = (
        f"Target workbook: `{target.name}`\n"
        f"Sheet: `{sheet_name}`\n"
        f"Rows to write: {len(rows or [])}\n\n"
        f"{_format_table_rows(preview_rows) if preview_rows else 'The sheet will be cleared and rewritten.'}"
    )
    if not confirmed:
        return _file_tool_result(
            f"Writing sheet `{sheet_name}` will replace its current content.",
            preview=preview,
            visited_path=str(target.parent),
            requires_confirmation=True,
            pending_action={
                "tool_name": "write_excel_sheet",
                "tool_input": {
                    "path": path,
                    "sheet_name": sheet_name,
                    "headers": headers or [],
                    "rows": rows or [],
                },
            },
        )

    workbook = _load_excel_for_write(target)
    worksheet = workbook[sheet_name] if sheet_name in workbook.sheetnames else workbook.create_sheet(sheet_name)
    worksheet.delete_rows(1, worksheet.max_row or 1)
    if headers:
        worksheet.append(list(headers))
    for row in rows or []:
        worksheet.append(list(row))
    workbook.save(target)
    return _file_tool_result(
        f"Wrote sheet `{sheet_name}` in `{target.name}`.",
        preview=preview,
        visited_path=str(target.parent),
    )


def edit_excel_cells_tool(
    path: str,
    sheet_name: str,
    updates: list[dict[str, Any]],
    confirmed: bool = False,
    base_path: str = "",
) -> dict[str, Any]:
    if not updates:
        raise ValueError("At least one cell update is required.")
    target = _resolve_agent_path(path, base_path)
    preview = "\n".join(
        f"- `{item.get('cell', '?')}` → `{item.get('value', '')}`"
        for item in updates[:20]
    )
    if not confirmed:
        return _file_tool_result(
            f"Editing cells in `{sheet_name}` will modify `{target.name}`.",
            preview=preview,
            visited_path=str(target.parent),
            requires_confirmation=True,
            pending_action={
                "tool_name": "edit_excel_cells",
                "tool_input": {
                    "path": path,
                    "sheet_name": sheet_name,
                    "updates": updates,
                },
            },
        )

    workbook = _load_excel_for_write(target)
    if sheet_name not in workbook.sheetnames:
        raise ValueError(f"Sheet `{sheet_name}` does not exist.")
    worksheet = workbook[sheet_name]
    for item in updates:
        cell = str(item.get("cell") or "").strip().upper()
        if not re.fullmatch(r"[A-Z]+[1-9][0-9]*", cell):
            raise ValueError(f"Invalid cell reference `{cell}`.")
        worksheet[cell] = item.get("value")
    workbook.save(target)
    return _file_tool_result(
        f"Updated {len(updates)} cell(s) in `{sheet_name}`.",
        preview=preview,
        visited_path=str(target.parent),
    )


def append_excel_rows_tool(path: str, sheet_name: str, rows: list[list[Any]], base_path: str = "") -> dict[str, Any]:
    if not rows:
        raise ValueError("At least one row is required.")
    target = _resolve_agent_path(path, base_path)
    workbook = _load_excel_for_write(target)
    if sheet_name not in workbook.sheetnames:
        raise ValueError(f"Sheet `{sheet_name}` does not exist.")
    worksheet = workbook[sheet_name]
    for row in rows:
        worksheet.append(list(row))
    workbook.save(target)
    return _file_tool_result(
        f"Appended {len(rows)} row(s) to `{sheet_name}`.",
        preview=f"Sheet `{sheet_name}` now has {worksheet.max_row} row(s).",
        visited_path=str(target.parent),
    )


def delete_excel_sheet_tool(path: str, sheet_name: str, confirmed: bool = False, base_path: str = "") -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    if not confirmed:
        return _file_tool_result(
            f"Deleting sheet `{sheet_name}` will modify `{target.name}`.",
            preview=f"Workbook: `{target}`\nSheet to delete: `{sheet_name}`",
            visited_path=str(target.parent),
            requires_confirmation=True,
            pending_action={
                "tool_name": "delete_excel_sheet",
                "tool_input": {"path": path, "sheet_name": sheet_name},
            },
        )

    workbook = _load_excel_for_write(target)
    if sheet_name not in workbook.sheetnames:
        raise ValueError(f"Sheet `{sheet_name}` does not exist.")
    if len(workbook.sheetnames) == 1:
        raise ValueError("An Excel workbook must keep at least one sheet.")
    worksheet = workbook[sheet_name]
    workbook.remove(worksheet)
    workbook.save(target)
    return _file_tool_result(
        f"Deleted sheet `{sheet_name}` from `{target.name}`.",
        preview="\n".join(f"- `{sheet}`" for sheet in workbook.sheetnames),
        visited_path=str(target.parent),
    )


def delete_file_tool(path: str, confirmed: bool = False, base_path: str = "") -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    if not target.exists() or not target.is_file():
        raise ValueError(f"File `{target}` does not exist.")
    preview = f"File: `{target}`\nSize: {target.stat().st_size} bytes"
    if not confirmed:
        return _file_tool_result(
            f"Deleting `{target.name}` is destructive.",
            preview=preview,
            visited_path=str(target.parent),
            requires_confirmation=True,
            pending_action={"tool_name": "delete_file", "tool_input": {"path": path}},
        )
    target.unlink()
    return _file_tool_result(
        f"Deleted file `{target.name}`.",
        preview=preview,
        visited_path=str(target.parent),
    )


def delete_directory_tool(path: str, recursive: bool = False, confirmed: bool = False, base_path: str = "") -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    sandbox_root = Path(base_path).expanduser().resolve() if base_path else None
    if sandbox_root and target == sandbox_root:
        raise ValueError("Deleting the configured access root is blocked for safety.")
    if not target.exists() or not target.is_dir():
        raise ValueError(f"Directory `{target}` does not exist.")
    child_count = sum(1 for _ in target.iterdir())
    preview = f"Directory: `{target}`\nEntries inside: {child_count}\nRecursive: {recursive}"
    if not confirmed:
        return _file_tool_result(
            f"Deleting directory `{target.name}` is destructive.",
            preview=preview,
            visited_path=str(target.parent),
            requires_confirmation=True,
            pending_action={
                "tool_name": "delete_directory",
                "tool_input": {"path": path, "recursive": recursive},
            },
        )
    if recursive:
        shutil.rmtree(target)
    else:
        target.rmdir()
    return _file_tool_result(
        f"Deleted directory `{target.name}`.",
        preview=preview,
        visited_path=str(target.parent),
    )


def move_file_tool(source_path: str, destination_path: str, confirmed: bool = False, base_path: str = "") -> dict[str, Any]:
    source = _resolve_agent_path(source_path, base_path)
    destination = _resolve_agent_path(destination_path, base_path)
    if not source.exists() or not source.is_file():
        raise ValueError(f"Source file `{source}` does not exist.")
    preview = (
        f"Source: `{source}`\n"
        f"Destination: `{destination}`\n"
        f"Destination exists: {destination.exists()}"
    )
    if not confirmed:
        return _file_tool_result(
            f"Moving `{source.name}` will change the filesystem state.",
            preview=preview,
            visited_path=str(source.parent),
            requires_confirmation=True,
            pending_action={
                "tool_name": "move_file",
                "tool_input": {"source_path": source_path, "destination_path": destination_path},
            },
        )
    _ensure_parent_directory(destination)
    if destination.exists():
        if destination.is_dir():
            raise ValueError("Destination already exists and is a directory.")
        destination.unlink()
    shutil.move(str(source), str(destination))
    return _file_tool_result(
        f"Moved `{source.name}` to `{destination}`.",
        preview=preview,
        visited_path=str(destination.parent),
    )


FILE_MANAGER_TOOLS: dict[str, dict[str, Any]] = {
    "list_directory": {
        "description": "List files and folders in a directory.",
        "parameters": {"path": "string, default '.'", "recursive": "boolean, default false"},
        "handler": list_directory_tool,
    },
    "get_file_info": {
        "description": "Get metadata for a file or directory.",
        "parameters": {"path": "string"},
        "handler": get_file_info_tool,
    },
    "search_files": {
        "description": "Search files by name and, for text files, by content.",
        "parameters": {"path": "string, default '.'", "query": "string", "recursive": "boolean, default true"},
        "handler": search_files_tool,
    },
    "read_file": {
        "description": "Read text, Word, Parquet, CSV, TSV, XLSX, or XLS content.",
        "parameters": {"path": "string"},
        "handler": read_file_tool,
    },
    "read_csv_summary": {
        "description": "Read a CSV or TSV file and summarize up to 50 rows.",
        "parameters": {"path": "string"},
        "handler": read_csv_summary_tool,
    },
    "create_directory": {
        "description": "Create a directory.",
        "parameters": {"path": "string"},
        "handler": create_directory_tool,
    },
    "create_file": {
        "description": "Create a new text file.",
        "parameters": {"path": "string", "content": "string, optional"},
        "handler": create_file_tool,
    },
    "create_excel_file": {
        "description": "Create a new XLSX workbook.",
        "parameters": {"path": "string ending with .xlsx", "sheet_name": "string, optional", "headers": "array, optional", "rows": "2D array, optional"},
        "handler": create_excel_file_tool,
    },
    "write_file": {
        "description": "Overwrite a text file. Requires confirmation first.",
        "parameters": {"path": "string", "content": "string", "confirmed": "boolean, default false"},
        "handler": write_file_tool,
    },
    "write_excel_sheet": {
        "description": "Overwrite an XLSX sheet with new rows. Requires confirmation first.",
        "parameters": {"path": "string", "sheet_name": "string", "headers": "array, optional", "rows": "2D array, optional", "confirmed": "boolean, default false"},
        "handler": write_excel_sheet_tool,
    },
    "edit_excel_cells": {
        "description": "Edit specific XLSX cells. Requires confirmation first.",
        "parameters": {"path": "string", "sheet_name": "string", "updates": "array of {cell, value}", "confirmed": "boolean, default false"},
        "handler": edit_excel_cells_tool,
    },
    "append_excel_rows": {
        "description": "Append rows to an existing XLSX sheet.",
        "parameters": {"path": "string", "sheet_name": "string", "rows": "2D array"},
        "handler": append_excel_rows_tool,
    },
    "delete_file": {
        "description": "Delete a file. Requires confirmation first.",
        "parameters": {"path": "string", "confirmed": "boolean, default false"},
        "handler": delete_file_tool,
    },
    "delete_directory": {
        "description": "Delete a directory. Requires confirmation first.",
        "parameters": {"path": "string", "recursive": "boolean, default false", "confirmed": "boolean, default false"},
        "handler": delete_directory_tool,
    },
    "delete_excel_sheet": {
        "description": "Delete an XLSX sheet. Requires confirmation first.",
        "parameters": {"path": "string", "sheet_name": "string", "confirmed": "boolean, default false"},
        "handler": delete_excel_sheet_tool,
    },
    "list_excel_sheets": {
        "description": "List sheets inside an Excel file.",
        "parameters": {"path": "string"},
        "handler": list_excel_sheets_tool,
    },
    "read_excel_sheet": {
        "description": "Read up to 50 rows from an Excel sheet.",
        "parameters": {"path": "string", "sheet_name": "string, optional"},
        "handler": read_excel_sheet_tool,
    },
    "add_excel_sheet": {
        "description": "Add a new sheet to an XLSX workbook.",
        "parameters": {"path": "string", "sheet_name": "string"},
        "handler": add_excel_sheet_tool,
    },
    "rename_excel_sheet": {
        "description": "Rename a sheet inside an XLSX workbook.",
        "parameters": {"path": "string", "old_name": "string", "new_name": "string"},
        "handler": rename_excel_sheet_tool,
    },
    "move_file": {
        "description": "Move a file to another location. Requires confirmation first.",
        "parameters": {"source_path": "string", "destination_path": "string", "confirmed": "boolean, default false"},
        "handler": move_file_tool,
    },
}


def _file_manager_tool_manifest() -> str:
    lines = []
    for tool_name, tool_spec in FILE_MANAGER_TOOLS.items():
        lines.append(f"- {tool_name}: {tool_spec['description']}")
        lines.append("  Parameters:")
        for key, value in tool_spec["parameters"].items():
            lines.append(f"    - {key}: {value}")
    return "\n".join(lines)


def execute_file_manager_tool(tool_name: str, tool_input: dict[str, Any], base_path: str = "") -> dict[str, Any]:
    if tool_name not in FILE_MANAGER_TOOLS:
        raise ValueError(f"Unknown file-management tool `{tool_name}`.")
    handler = FILE_MANAGER_TOOLS[tool_name]["handler"]
    safe_input = dict(tool_input or {})
    safe_input["base_path"] = base_path
    return handler(**safe_input)


def _try_extract_file_export_payload(user_message: str) -> Optional[dict[str, Any]]:
    text = str(user_message or "").strip()
    if not text:
        return None
    try:
        parsed = extract_json_object(text)
    except Exception:
        return None
    if not isinstance(parsed, dict) or not parsed.get("__file_export__"):
        return None

    export_format = str(parsed.get("format") or "").strip().lower()
    if export_format == "xls":
        export_format = "xlsx"
    if export_format not in {"csv", "tsv", "xlsx"}:
        raise ValueError("Unsupported file export format. Use `csv`, `tsv`, or `xlsx`.")

    path = str(parsed.get("path") or "").strip()
    if export_format == "xlsx" and path.lower().endswith(".xls"):
        path = path[:-4] + ".xlsx"
    headers = [str(item) for item in (parsed.get("headers") or []) if str(item).strip()]
    raw_rows = parsed.get("rows") or []
    if not isinstance(raw_rows, list):
        raise ValueError("The file export payload must provide `rows` as a list.")
    rows = []
    for row in raw_rows:
        if isinstance(row, list):
            rows.append(list(row))
        elif isinstance(row, dict) and headers:
            rows.append([row.get(header) for header in headers])
    return {
        "format": export_format,
        "path": path,
        "sheet_name": str(parsed.get("sheet_name") or parsed.get("sheetName") or "Results").strip() or "Results",
        "headers": headers,
        "rows": rows,
        "source_sql": str(parsed.get("source_sql") or parsed.get("sourceSql") or "").strip(),
        "source_request": str(parsed.get("source_request") or parsed.get("sourceRequest") or "").strip(),
    }


def _serialize_delimited_rows(headers: list[str], rows: list[list[Any]], delimiter: str) -> str:
    buffer = io.StringIO()
    writer = csv.writer(buffer, delimiter=delimiter)
    if headers:
        writer.writerow(headers)
    for row in rows:
        writer.writerow([
            json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else value
            for value in row
        ])
    return buffer.getvalue()


def _file_export_answer(result: dict[str, Any], payload: dict[str, Any]) -> str:
    preview = result.get("preview") or ""
    answer = (
        "## File Export\n"
        f"{result['summary']}\n\n"
        f"- **Format:** `{payload.get('format')}`\n"
        f"- **Target:** `{payload.get('path')}`"
    )
    if payload.get("source_sql"):
        answer += "\n- **Source SQL available:** yes"
    if preview:
        answer += f"\n\n## Preview\n{preview}"
    return answer


PDF_CREATOR_SOURCE_CHOICES = [
    "Use the latest analysis in this chat",
    "Paste the content in your next message",
]
PDF_CREATOR_MAX_TEXT_CHARS = 120_000
PDF_CREATOR_ACCENT_RGB = (0.306, 0.459, 0.949)
PDF_CREATOR_SLATE_RGB = (0.082, 0.117, 0.211)
PDF_CREATOR_TEXT_RGB = (0.188, 0.223, 0.305)
PDF_CREATOR_MUTED_RGB = (0.451, 0.494, 0.576)
PDF_CREATOR_CODE_BG_RGB = (0.943, 0.957, 0.984)
PDF_CREATOR_RULE_RGB = (0.835, 0.862, 0.921)
PDF_CREATOR_PAGE_BG_RGB = (0.972, 0.978, 0.992)
PDF_CREATOR_CARD_BG_RGB = (1.0, 1.0, 1.0)
PDF_CREATOR_CARD_BORDER_RGB = (0.879, 0.905, 0.952)
PDF_CREATOR_HERO_LEFT_RGB = (0.412, 0.231, 0.925)
PDF_CREATOR_HERO_CENTER_RGB = (0.267, 0.351, 0.933)
PDF_CREATOR_HERO_RIGHT_RGB = (0.173, 0.592, 0.898)
PDF_CREATOR_HERO_SOFT_RGB = (0.809, 0.863, 0.988)
PDF_CREATOR_PILL_BG_RGB = (0.516, 0.430, 0.969)
PDF_CREATOR_PILL_TEXT_RGB = (0.945, 0.953, 0.996)
PDF_CREATOR_META_LABEL_RGB = (0.761, 0.815, 0.969)
PDF_CREATOR_META_VALUE_RGB = (1.0, 1.0, 1.0)
PDF_CREATOR_SECTION_BG_RGB = (0.961, 0.969, 0.988)
PDF_CREATOR_SECTION_TEXT_RGB = (0.396, 0.454, 0.588)


def _pdf_escape_text(value: str) -> str:
    safe = unicodedata.normalize("NFKD", str(value or "")).encode("ascii", "ignore").decode("ascii")
    return (
        safe
        .replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
    )


def _slugify_filename(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", str(value or "")).encode("ascii", "ignore").decode("ascii")
    cleaned = re.sub(r"[^a-zA-Z0-9]+", "-", normalized.lower()).strip("-")
    return cleaned or "report"


def _markdown_inline_to_plain_text(value: str) -> str:
    text = str(value or "")
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    text = re.sub(r"<[^>]+>", "", text)
    return re.sub(r"\s+", " ", text).strip()


def _wrap_text_for_pdf(text: str, max_chars: int) -> list[str]:
    cleaned = str(text or "").strip()
    if not cleaned:
        return [""]
    words = cleaned.split()
    lines: list[str] = []
    current = ""
    for word in words:
        if not current:
            current = word
            continue
        candidate = f"{current} {word}"
        if len(candidate) <= max_chars:
            current = candidate
            continue
        lines.append(current)
        current = word
    if current:
        lines.append(current)
    return lines or [cleaned[:max_chars]]


def _parse_markdown_table_cells(line: str) -> list[str]:
    stripped = str(line or "").strip()
    if not stripped.startswith("|"):
        return []
    parts = [part.strip() for part in stripped.strip("|").split("|")]
    return [_markdown_inline_to_plain_text(part) for part in parts]


def _parse_markdown_for_pdf(markdown: str) -> list[dict[str, Any]]:
    text = re.sub(r"<!--[\s\S]*?-->", "", str(markdown or "")).replace("\r", "")
    lines = text.split("\n")
    blocks: list[dict[str, Any]] = []
    paragraph_lines: list[str] = []
    code_lines: list[str] = []
    in_code = False
    index = 0

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        paragraph = " ".join(line.strip() for line in paragraph_lines if line.strip())
        if paragraph:
            blocks.append({"kind": "paragraph", "text": _markdown_inline_to_plain_text(paragraph)})
        paragraph_lines = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            blocks.append({"kind": "code", "lines": [line.rstrip() for line in code_lines]})
        code_lines = []

    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.rstrip()
        if stripped.strip().startswith("```"):
            flush_paragraph()
            if in_code:
                flush_code()
            in_code = not in_code
            index += 1
            continue

        if in_code:
            code_lines.append(stripped)
            index += 1
            continue

        if not stripped.strip():
            flush_paragraph()
            blocks.append({"kind": "spacer", "height": 8})
            index += 1
            continue

        if (
            stripped.strip().startswith("|")
            and index + 1 < len(lines)
            and re.match(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*$", lines[index + 1].strip())
        ):
            flush_paragraph()
            headers = _parse_markdown_table_cells(stripped)
            table_rows: list[list[str]] = []
            cursor = index + 2
            while cursor < len(lines):
                table_line = lines[cursor].rstrip()
                if not table_line.strip().startswith("|"):
                    break
                cells = _parse_markdown_table_cells(table_line)
                if cells:
                    table_rows.append(cells)
                cursor += 1

            if headers and table_rows:
                if len(headers) == 2:
                    for row in table_rows:
                        label = row[0] if len(row) > 0 else ""
                        value = row[1] if len(row) > 1 else ""
                        blocks.append({"kind": "metric", "label": label, "value": value})
                else:
                    for row in table_rows:
                        row_pairs = []
                        for cell_index, header in enumerate(headers):
                            cell_value = row[cell_index] if cell_index < len(row) else ""
                            if cell_value:
                                row_pairs.append(f"{header}: {cell_value}")
                        if row_pairs:
                            blocks.append({"kind": "bullet", "text": " | ".join(row_pairs)})
            index = cursor
            continue

        heading_match = re.match(r"^\s*(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            blocks.append(
                {
                    "kind": f"heading_{level}",
                    "text": _markdown_inline_to_plain_text(heading_match.group(2)),
                }
            )
            index += 1
            continue

        bullet_match = re.match(r"^\s*(?:[-*]|\d+\.)\s+(?:\[[ xX]\]\s+)?(.+)$", stripped)
        if bullet_match:
            flush_paragraph()
            blocks.append({"kind": "bullet", "text": _markdown_inline_to_plain_text(bullet_match.group(1))})
            index += 1
            continue

        quote_match = re.match(r"^\s*>\s?(.*)$", stripped)
        if quote_match:
            paragraph_lines.append(quote_match.group(1))
            index += 1
            continue

        paragraph_lines.append(stripped)
        index += 1

    flush_paragraph()
    flush_code()
    return [block for block in blocks if block.get("text") or block.get("lines") or block.get("kind") == "spacer"]


def _pdf_content_stream_for_document(
    title: str,
    subtitle: str,
    body_markdown: str,
    generated_label: str,
) -> tuple[list[str], int]:
    page_width = 595.28
    page_height = 841.89
    page_margin = 36.0
    content_card_x = page_margin
    content_card_width = page_width - page_margin * 2
    content_left = content_card_x + 24.0
    content_width = content_card_width - 48.0
    content_right = content_left + content_width
    bottom_margin = 78.0
    first_page_top = 536.0
    other_page_top = 714.0
    pages: list[list[str]] = []
    page_index = -1
    cursor_y = 0.0

    def new_page() -> None:
        nonlocal page_index, cursor_y
        page_index += 1
        cursor_y = first_page_top if page_index == 0 else other_page_top
        ops: list[str] = []
        ops.extend(
            [
                f"{PDF_CREATOR_PAGE_BG_RGB[0]:.3f} {PDF_CREATOR_PAGE_BG_RGB[1]:.3f} {PDF_CREATOR_PAGE_BG_RGB[2]:.3f} rg",
                f"0 0 {page_width:.2f} {page_height:.2f} re f",
            ]
        )
        if page_index == 0:
            hero_x = page_margin
            hero_y = 610.0
            hero_height = 176.0
            hero_width = page_width - page_margin * 2
            ops.extend(
                [
                    f"{PDF_CREATOR_HERO_LEFT_RGB[0]:.3f} {PDF_CREATOR_HERO_LEFT_RGB[1]:.3f} {PDF_CREATOR_HERO_LEFT_RGB[2]:.3f} rg",
                    f"{hero_x:.2f} {hero_y:.2f} {hero_width * 0.36:.2f} {hero_height:.2f} re f",
                    f"{PDF_CREATOR_HERO_CENTER_RGB[0]:.3f} {PDF_CREATOR_HERO_CENTER_RGB[1]:.3f} {PDF_CREATOR_HERO_CENTER_RGB[2]:.3f} rg",
                    f"{hero_x + hero_width * 0.36:.2f} {hero_y:.2f} {hero_width * 0.32:.2f} {hero_height:.2f} re f",
                    f"{PDF_CREATOR_HERO_RIGHT_RGB[0]:.3f} {PDF_CREATOR_HERO_RIGHT_RGB[1]:.3f} {PDF_CREATOR_HERO_RIGHT_RGB[2]:.3f} rg",
                    f"{hero_x + hero_width * 0.68:.2f} {hero_y:.2f} {hero_width * 0.32:.2f} {hero_height:.2f} re f",
                    f"{PDF_CREATOR_PILL_BG_RGB[0]:.3f} {PDF_CREATOR_PILL_BG_RGB[1]:.3f} {PDF_CREATOR_PILL_BG_RGB[2]:.3f} rg",
                    f"{hero_x + 26:.2f} {hero_y + hero_height - 40:.2f} 196 18 re f",
                    f"{PDF_CREATOR_HERO_SOFT_RGB[0]:.3f} {PDF_CREATOR_HERO_SOFT_RGB[1]:.3f} {PDF_CREATOR_HERO_SOFT_RGB[2]:.3f} rg",
                    f"{hero_x + hero_width - 76:.2f} {hero_y + hero_height - 70:.2f} 40 40 re f",
                    "BT",
                    "/F2 7.5 Tf",
                    f"{PDF_CREATOR_PILL_TEXT_RGB[0]:.3f} {PDF_CREATOR_PILL_TEXT_RGB[1]:.3f} {PDF_CREATOR_PILL_TEXT_RGB[2]:.3f} rg",
                    f"1 0 0 1 {hero_x + 34:.2f} {hero_y + hero_height - 28:.2f} Tm",
                    "([EXECUTIVE REPORT  |  RAGNAROK]) Tj",
                    "ET",
                    "BT",
                    "/F2 24 Tf",
                    "1 1 1 rg",
                    f"1 0 0 1 {hero_x + 26:.2f} {hero_y + hero_height - 66:.2f} Tm",
                    f"({_pdf_escape_text(title)}) Tj",
                    "ET",
                    "BT",
                    "/F1 11 Tf",
                    "0.89 0.92 0.97 rg",
                    f"1 0 0 1 {hero_x + 26:.2f} {hero_y + hero_height - 90:.2f} Tm",
                    f"({_pdf_escape_text(subtitle)}) Tj",
                    "ET",
                    f"{PDF_CREATOR_HERO_SOFT_RGB[0]:.3f} {PDF_CREATOR_HERO_SOFT_RGB[1]:.3f} {PDF_CREATOR_HERO_SOFT_RGB[2]:.3f} rg",
                    f"{hero_x + 26:.2f} {hero_y + 72:.2f} {hero_width - 52:.2f} 1.4 re f",
                ]
            )
            metadata = [
                ("GENERATED", generated_label.replace("Generated on ", "")),
                ("TOOL", "RAGnarok PDF Creator"),
                ("CLASSIFICATION", "Internal executive report"),
            ]
            column_width = (hero_width - 52.0) / 3.0
            for metadata_index, (label, value) in enumerate(metadata):
                column_x = hero_x + 26.0 + metadata_index * column_width
                ops.extend(
                    [
                        "BT",
                        "/F2 8 Tf",
                        f"{PDF_CREATOR_META_LABEL_RGB[0]:.3f} {PDF_CREATOR_META_LABEL_RGB[1]:.3f} {PDF_CREATOR_META_LABEL_RGB[2]:.3f} rg",
                        f"1 0 0 1 {column_x:.2f} {hero_y + 52:.2f} Tm",
                        f"({_pdf_escape_text(label)}) Tj",
                        "ET",
                        "BT",
                        "/F2 10.5 Tf",
                        f"{PDF_CREATOR_META_VALUE_RGB[0]:.3f} {PDF_CREATOR_META_VALUE_RGB[1]:.3f} {PDF_CREATOR_META_VALUE_RGB[2]:.3f} rg",
                        f"1 0 0 1 {column_x:.2f} {hero_y + 34:.2f} Tm",
                        f"({_pdf_escape_text(value)}) Tj",
                        "ET",
                    ]
                )
            ops.extend(
                [
                    f"{PDF_CREATOR_SLATE_RGB[0]:.3f} {PDF_CREATOR_SLATE_RGB[1]:.3f} {PDF_CREATOR_SLATE_RGB[2]:.3f} rg",
                    f"{hero_x:.2f} {hero_y:.2f} {hero_width:.2f} 16 re f",
                    "BT",
                    "/F1 7.8 Tf",
                    f"{PDF_CREATOR_META_LABEL_RGB[0]:.3f} {PDF_CREATOR_META_LABEL_RGB[1]:.3f} {PDF_CREATOR_META_LABEL_RGB[2]:.3f} rg",
                    f"1 0 0 1 {hero_x + 26:.2f} {hero_y + 4.5:.2f} Tm",
                    "(Document generated automatically by RAGnarok  |  Reserved for internal reporting use) Tj",
                    "ET",
                ]
            )
            content_card_y = 56.0
            content_card_height = 520.0
        else:
            header_height = 46.0
            header_y = page_height - header_height - page_margin + 8.0
            ops.extend(
                [
                    f"{PDF_CREATOR_SLATE_RGB[0]:.3f} {PDF_CREATOR_SLATE_RGB[1]:.3f} {PDF_CREATOR_SLATE_RGB[2]:.3f} rg",
                    f"{page_margin:.2f} {header_y:.2f} {page_width - page_margin * 2:.2f} {header_height:.2f} re f",
                    "BT",
                    "/F2 11 Tf",
                    "1 1 1 rg",
                    f"1 0 0 1 {page_margin + 20:.2f} {header_y + 17:.2f} Tm",
                    f"({_pdf_escape_text(title)}) Tj",
                    "ET",
                ]
            )
            content_card_y = 56.0
            content_card_height = 676.0
        ops.extend(
            [
                f"{PDF_CREATOR_CARD_BG_RGB[0]:.3f} {PDF_CREATOR_CARD_BG_RGB[1]:.3f} {PDF_CREATOR_CARD_BG_RGB[2]:.3f} rg",
                f"{content_card_x:.2f} {content_card_y:.2f} {content_card_width:.2f} {content_card_height:.2f} re f",
                f"{PDF_CREATOR_CARD_BORDER_RGB[0]:.3f} {PDF_CREATOR_CARD_BORDER_RGB[1]:.3f} {PDF_CREATOR_CARD_BORDER_RGB[2]:.3f} RG",
                f"{content_card_x:.2f} {content_card_y:.2f} {content_card_width:.2f} {content_card_height:.2f} re S",
                f"{PDF_CREATOR_SECTION_BG_RGB[0]:.3f} {PDF_CREATOR_SECTION_BG_RGB[1]:.3f} {PDF_CREATOR_SECTION_BG_RGB[2]:.3f} rg",
                f"{content_card_x + 18:.2f} {content_card_y + content_card_height - 30:.2f} 164 18 re f",
                "BT",
                "/F2 7.8 Tf",
                f"{PDF_CREATOR_SECTION_TEXT_RGB[0]:.3f} {PDF_CREATOR_SECTION_TEXT_RGB[1]:.3f} {PDF_CREATOR_SECTION_TEXT_RGB[2]:.3f} rg",
                f"1 0 0 1 {content_card_x + 28:.2f} {content_card_y + content_card_height - 18:.2f} Tm",
                f"({_pdf_escape_text('ANALYSIS COMPLETE' if page_index == 0 else 'ANALYSIS CONTINUED')}) Tj",
                "ET",
            ]
        )
        ops.extend(
            [
                "BT",
                "/F1 9 Tf",
                f"{PDF_CREATOR_MUTED_RGB[0]:.3f} {PDF_CREATOR_MUTED_RGB[1]:.3f} {PDF_CREATOR_MUTED_RGB[2]:.3f} rg",
                f"1 0 0 1 {page_margin:.2f} 24.00 Tm",
                f"({_pdf_escape_text(f'{generated_label}  |  Page {page_index + 1}')}) Tj",
                "ET",
            ]
        )
        pages.append(ops)

    def ensure_space(height: float) -> None:
        nonlocal cursor_y
        if page_index < 0:
            new_page()
        if cursor_y - height < bottom_margin:
            new_page()

    def add_text_lines(lines: list[str], x: float, font: str, size: float, leading: float, color: tuple[float, float, float]) -> None:
        nonlocal cursor_y
        for line in lines:
            pages[page_index].extend(
                [
                    "BT",
                    f"/{font} {size:.2f} Tf",
                    f"{color[0]:.3f} {color[1]:.3f} {color[2]:.3f} rg",
                    f"1 0 0 1 {x:.2f} {cursor_y:.2f} Tm",
                    f"({_pdf_escape_text(line)}) Tj",
                    "ET",
                ]
            )
            cursor_y -= leading

    for block in _parse_markdown_for_pdf(body_markdown):
        kind = block.get("kind")
        if kind == "spacer":
            ensure_space(float(block.get("height") or 8))
            cursor_y -= float(block.get("height") or 8)
            continue
        if kind == "heading_1":
            lines = _wrap_text_for_pdf(str(block.get("text") or ""), 54)
            required = 14 + len(lines) * 24
            ensure_space(required)
            cursor_y -= 4
            add_text_lines(lines, content_left, "F2", 18, 22, PDF_CREATOR_SLATE_RGB)
            pages[page_index].append(
                f"{PDF_CREATOR_ACCENT_RGB[0]:.3f} {PDF_CREATOR_ACCENT_RGB[1]:.3f} {PDF_CREATOR_ACCENT_RGB[2]:.3f} rg\n"
                f"{content_left:.2f} {cursor_y + 8:.2f} 34 2 re f"
            )
            cursor_y -= 10
            continue
        if kind == "heading_2" or kind == "heading_3":
            lines = _wrap_text_for_pdf(str(block.get("text") or ""), 68)
            required = 10 + len(lines) * 19
            ensure_space(required)
            cursor_y -= 2
            add_text_lines(lines, content_left, "F2", 13 if kind == "heading_2" else 12, 17, PDF_CREATOR_SLATE_RGB)
            cursor_y -= 6
            continue
        if kind == "metric":
            label = _markdown_inline_to_plain_text(str(block.get("label") or ""))
            value = _markdown_inline_to_plain_text(str(block.get("value") or ""))
            label_lines = _wrap_text_for_pdf(label, 28)
            value_lines = _wrap_text_for_pdf(value, 42)
            line_count = max(len(label_lines), len(value_lines), 1)
            row_height = 14 + line_count * 13
            ensure_space(row_height + 6)
            row_bottom = cursor_y - row_height + 3
            pages[page_index].append(
                f"{PDF_CREATOR_SECTION_BG_RGB[0]:.3f} {PDF_CREATOR_SECTION_BG_RGB[1]:.3f} {PDF_CREATOR_SECTION_BG_RGB[2]:.3f} rg\n"
                f"{content_left:.2f} {row_bottom:.2f} {content_width:.2f} {row_height:.2f} re f"
            )
            pages[page_index].append(
                f"{PDF_CREATOR_CARD_BORDER_RGB[0]:.3f} {PDF_CREATOR_CARD_BORDER_RGB[1]:.3f} {PDF_CREATOR_CARD_BORDER_RGB[2]:.3f} RG\n"
                f"{content_left:.2f} {row_bottom:.2f} {content_width:.2f} {row_height:.2f} re S"
            )
            pages[page_index].append(
                f"{PDF_CREATOR_ACCENT_RGB[0]:.3f} {PDF_CREATOR_ACCENT_RGB[1]:.3f} {PDF_CREATOR_ACCENT_RGB[2]:.3f} rg\n"
                f"{content_left + 12:.2f} {row_bottom + 8:.2f} 3.5 {row_height - 16:.2f} re f"
            )
            row_start_y = cursor_y - 14
            for line_index in range(line_count):
                if line_index < len(label_lines):
                    pages[page_index].extend(
                        [
                            "BT",
                            "/F2 10.5 Tf",
                            f"{PDF_CREATOR_SLATE_RGB[0]:.3f} {PDF_CREATOR_SLATE_RGB[1]:.3f} {PDF_CREATOR_SLATE_RGB[2]:.3f} rg",
                            f"1 0 0 1 {content_left + 24:.2f} {row_start_y - line_index * 13:.2f} Tm",
                            f"({_pdf_escape_text(label_lines[line_index])}) Tj",
                            "ET",
                        ]
                    )
                if line_index < len(value_lines):
                    pages[page_index].extend(
                        [
                            "BT",
                            "/F1 10.5 Tf",
                            f"{PDF_CREATOR_TEXT_RGB[0]:.3f} {PDF_CREATOR_TEXT_RGB[1]:.3f} {PDF_CREATOR_TEXT_RGB[2]:.3f} rg",
                            f"1 0 0 1 {content_left + content_width * 0.43:.2f} {row_start_y - line_index * 13:.2f} Tm",
                            f"({_pdf_escape_text(value_lines[line_index])}) Tj",
                            "ET",
                        ]
                    )
            cursor_y -= row_height + 8
            continue
        if kind == "bullet":
            bullet_lines = _wrap_text_for_pdf(str(block.get("text") or ""), 74)
            required = len(bullet_lines) * 15 + 4
            ensure_space(required)
            if bullet_lines:
                pages[page_index].append(
                    f"{PDF_CREATOR_ACCENT_RGB[0]:.3f} {PDF_CREATOR_ACCENT_RGB[1]:.3f} {PDF_CREATOR_ACCENT_RGB[2]:.3f} rg\n"
                    f"{content_left:.2f} {cursor_y - 5:.2f} 5 5 re f"
                )
                add_text_lines([bullet_lines[0]], content_left + 14, "F1", 11, 15, PDF_CREATOR_TEXT_RGB)
                if len(bullet_lines) > 1:
                    add_text_lines(bullet_lines[1:], content_left + 14, "F1", 11, 15, PDF_CREATOR_TEXT_RGB)
            cursor_y -= 2
            continue
        if kind == "code":
            code_lines = [str(line)[:110] for line in (block.get("lines") or [])] or [""]
            leading = 12.0
            padding = 10.0
            required = padding * 2 + len(code_lines) * leading + 8
            ensure_space(required)
            rect_height = padding * 2 + len(code_lines) * leading
            rect_bottom = cursor_y - rect_height + 4
            pages[page_index].append(
                f"{PDF_CREATOR_CODE_BG_RGB[0]:.3f} {PDF_CREATOR_CODE_BG_RGB[1]:.3f} {PDF_CREATOR_CODE_BG_RGB[2]:.3f} rg\n"
                f"{content_left:.2f} {rect_bottom:.2f} {content_width:.2f} {rect_height:.2f} re f"
            )
            pages[page_index].append(
                f"{PDF_CREATOR_RULE_RGB[0]:.3f} {PDF_CREATOR_RULE_RGB[1]:.3f} {PDF_CREATOR_RULE_RGB[2]:.3f} RG\n"
                f"{content_left:.2f} {rect_bottom:.2f} {content_width:.2f} {rect_height:.2f} re S"
            )
            cursor_y -= padding
            add_text_lines(code_lines, content_left + 10, "F3", 9, leading, PDF_CREATOR_TEXT_RGB)
            cursor_y -= padding + 4
            continue

        paragraph_lines = _wrap_text_for_pdf(str(block.get("text") or ""), 84)
        required = len(paragraph_lines) * 15 + 6
        ensure_space(required)
        add_text_lines(paragraph_lines, content_left, "F1", 11, 15, PDF_CREATOR_TEXT_RGB)
        cursor_y -= 6

    if not pages:
        new_page()

    return ["\n".join(page_ops) for page_ops in pages], len(pages)


def build_professional_pdf_bytes(title: str, subtitle: str, body_markdown: str) -> tuple[bytes, int]:
    clean_title = _markdown_inline_to_plain_text(title) or "RAGnarok Report"
    clean_subtitle = _markdown_inline_to_plain_text(subtitle) or "Professional export generated from RAGnarok"
    generated_label = datetime.now().strftime("Generated on %Y-%m-%d %H:%M UTC")
    page_streams, page_count = _pdf_content_stream_for_document(clean_title, clean_subtitle, body_markdown, generated_label)
    page_width = 595.28
    page_height = 841.89
    objects: list[str] = []

    def add_object(content: str) -> int:
        objects.append(content)
        return len(objects)

    font_regular_id = add_object("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    font_bold_id = add_object("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>")
    font_code_id = add_object("<< /Type /Font /Subtype /Type1 /BaseFont /Courier >>")
    pages_id = add_object("<< /Type /Pages /Count 0 /Kids [] >>")
    page_ids: list[int] = []

    for stream in page_streams:
        content_id = add_object(
            f"<< /Length {len(stream.encode('latin-1', errors='ignore'))} >>\nstream\n{stream}\nendstream"
        )
        page_id = add_object(
            "<< /Type /Page "
            f"/Parent {pages_id} 0 R "
            f"/MediaBox [0 0 {page_width} {page_height}] "
            f"/Resources << /Font << /F1 {font_regular_id} 0 R /F2 {font_bold_id} 0 R /F3 {font_code_id} 0 R >> >> "
            f"/Contents {content_id} 0 R >>"
        )
        page_ids.append(page_id)

    objects[pages_id - 1] = f"<< /Type /Pages /Count {len(page_ids)} /Kids [{' '.join(f'{page_id} 0 R' for page_id in page_ids)}] >>"
    catalog_id = add_object(f"<< /Type /Catalog /Pages {pages_id} 0 R >>")

    pdf = "%PDF-1.4\n"
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf += f"{index} 0 obj\n{obj}\nendobj\n"

    xref_offset = len(pdf)
    pdf += f"xref\n0 {len(objects) + 1}\n"
    pdf += "0000000000 65535 f \n"
    for offset in offsets[1:]:
        pdf += f"{str(offset).rjust(10, '0')} 00000 n \n"
    pdf += f"trailer\n<< /Size {len(objects) + 1} /Root {catalog_id} 0 R >>\nstartxref\n{xref_offset}\n%%EOF"
    return pdf.encode("latin-1", errors="ignore"), page_count


def _extract_pdf_path(user_message: str) -> str:
    lowered = str(user_message or "").strip()
    quoted_match = re.search(r'["\']([^"\']+\.pdf)["\']', lowered, flags=re.IGNORECASE)
    if quoted_match:
        return quoted_match.group(1).strip()
    patterns = [
        r'(?:named|called|as|to|into|in|vers|dans|nomme|nommé|appele|appelé)\s+([A-Za-z0-9_./\\\\-]+\.pdf)',
    ]
    for pattern in patterns:
        match = re.search(pattern, lowered, flags=re.IGNORECASE)
        if match:
            return match.group(1).strip()
    matches = re.findall(r'([A-Za-z0-9_./\\\\-]+\.pdf)', lowered, flags=re.IGNORECASE)
    return matches[-1].strip() if matches else ""


def _extract_pdf_title(user_message: str, target_path: str = "") -> str:
    text = str(user_message or "").strip()
    title_match = re.search(r'(?:title|titled|intitule|intitulé)\s*[:=]?\s*["\']([^"\']+)["\']', text, flags=re.IGNORECASE)
    if title_match:
        return _markdown_inline_to_plain_text(title_match.group(1))[:120]
    about_match = re.search(r'(?:pdf|report|summary)\s+(?:for|about|of)\s+(.+)$', text, flags=re.IGNORECASE)
    if about_match:
        candidate = _markdown_inline_to_plain_text(about_match.group(1))
        candidate = re.sub(r'\b(?:as|named|called)\b.+$', "", candidate, flags=re.IGNORECASE).strip(" .:-")
        if candidate:
            return candidate[:120].title()
    if target_path:
        stem = Path(target_path).stem.replace("-", " ").replace("_", " ").strip()
        if stem:
            return stem.title()
    normalized = normalize_intent_text(text)
    if "data quality" in normalized:
        return "Data Quality Report"
    if any(token in normalized for token in ["clickhouse", "sql", "query", "analysis"]):
        return "Analysis Report"
    return "RAGnarok Report"


def _default_pdf_target_path(title: str) -> str:
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    return f"exports/{_slugify_filename(title)}-{timestamp}.pdf"


def _latest_exportable_assistant_message(history: list[dict[str, Any]]) -> str:
    for item in reversed(history):
        if str(item.get("role") or "") != "assistant":
            continue
        content = str(item.get("content") or "").strip()
        if not content:
            continue
        if "<!-- agent-intro:" in content:
            continue
        if content.startswith("# Welcome to RAGnarok"):
            continue
        return content
    return ""


def _try_extract_pdf_export_payload(user_message: str) -> Optional[dict[str, Any]]:
    text = str(user_message or "").strip()
    if not text:
        return None
    try:
        parsed = extract_json_object(text)
    except Exception:
        return None
    if not isinstance(parsed, dict) or not parsed.get("__pdf_export__"):
        return None

    title = _markdown_inline_to_plain_text(parsed.get("title") or "RAGnarok Report")[:120] or "RAGnarok Report"
    subtitle = _markdown_inline_to_plain_text(parsed.get("subtitle") or "Professional export generated from RAGnarok")[:180]
    target_path = str(parsed.get("path") or "").strip()
    source_markdown = str(parsed.get("source_markdown") or parsed.get("sourceMarkdown") or parsed.get("body_markdown") or parsed.get("bodyMarkdown") or "").strip()
    return {
        "title": title,
        "subtitle": subtitle,
        "path": target_path,
        "source_markdown": source_markdown[:PDF_CREATOR_MAX_TEXT_CHARS],
        "source_request": str(parsed.get("source_request") or parsed.get("sourceRequest") or "").strip(),
    }


def _build_pdf_creator_body_markdown(source_markdown: str, title: str, source_request: str = "") -> str:
    cleaned = str(source_markdown or "").strip()
    if not cleaned:
        return (
            "## Executive Summary\n"
            "No source analysis was available, so this PDF contains only the export metadata."
        )

    sections = [
        "## Executive Summary",
        f"This PDF captures the latest analysis prepared in RAGnarok for **{_markdown_inline_to_plain_text(title)}**.",
    ]
    if source_request:
        sections.append(f"Original request: `{_markdown_inline_to_plain_text(source_request)}`")
    sections.extend(
        [
            "",
            "## Detailed Result",
            cleaned[:PDF_CREATOR_MAX_TEXT_CHARS],
        ]
    )
    return "\n".join(sections)


def create_pdf_report_tool(
    path: str,
    title: str,
    subtitle: str,
    body_markdown: str,
    confirmed: bool = False,
    base_path: str = "",
) -> dict[str, Any]:
    target = _resolve_agent_path(path, base_path)
    preview = (
        f"- **Title:** {_markdown_inline_to_plain_text(title)}\n"
        f"- **Target:** `{target}`\n"
        f"- **Mode:** {'Overwrite existing PDF' if target.exists() else 'Create new PDF'}"
    )
    if target.exists() and not confirmed:
        return _file_tool_result(
            f"The PDF `{target.name}` already exists and needs confirmation before overwrite.",
            preview=preview,
            visited_path=str(target.parent),
            requires_confirmation=True,
            pending_action={
                "tool_name": "create_pdf_report",
                "tool_input": {
                    "path": path,
                    "title": title,
                    "subtitle": subtitle,
                    "body_markdown": body_markdown,
                },
            },
        )

    _ensure_parent_directory(target)
    pdf_bytes, page_count = build_professional_pdf_bytes(title, subtitle, body_markdown)
    target.write_bytes(pdf_bytes)
    file_size = target.stat().st_size if target.exists() else len(pdf_bytes)
    return _file_tool_result(
        f"Created PDF `{target.name}` with {page_count} page(s).",
        preview=(
            f"- **Title:** {_markdown_inline_to_plain_text(title)}\n"
            f"- **Saved to:** `{target}`\n"
            f"- **Pages:** {page_count}\n"
            f"- **Size:** {file_size} bytes"
        ),
        data={"path": str(target), "pageCount": page_count, "size": file_size},
        visited_path=str(target.parent),
    )


def _pdf_creator_confirmation_answer(state: dict[str, Any]) -> tuple[str, list[dict[str, Any]]]:
    pending = state.get("pending_confirmation") or {}
    preview = str(pending.get("preview") or "").strip()
    summary = str(pending.get("summary") or "This PDF export requires confirmation.").strip()
    return (
        "## Confirmation Needed\n"
        f"{summary}\n\n"
        f"{preview}\n\n"
        "Please confirm if you want me to overwrite the existing PDF."
    ), [
        {
            "id": "confirm-file-action",
            "label": "Confirm",
            "actionType": "confirm_file_action",
            "variant": "primary",
        },
        {
            "id": "cancel-file-action",
            "label": "Cancel",
            "actionType": "cancel_file_action",
            "variant": "secondary",
        },
    ]


def _pdf_creator_success_answer(result: dict[str, Any], title: str, target_path: str) -> str:
    preview = str(result.get("preview") or "").strip()
    answer = (
        "## PDF Created\n"
        f"{result.get('summary') or 'The PDF export is complete.'}\n\n"
        f"- **Title:** {_markdown_inline_to_plain_text(title)}\n"
        f"- **Saved to:** `{target_path}`\n"
        "- **Style:** Clean slate header, compact sections, and professional PDF layout."
    )
    if preview:
        answer += f"\n\n## Export Details\n{preview}"
    return answer

async def plan_file_manager_step(
    user_message: str,
    history: list[dict[str, Any]],
    scratchpad: list[dict[str, Any]],
    base_path: str,
    system_prompt: str,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str] = None,
) -> dict[str, Any]:
    trimmed_history = _normalized_history_messages(
        history,
        current_message=user_message,
        max_steps=CHAT_MEMORY_MAX_STEPS,
    )
    prompt = f"""
You are a ReAct-style file management agent.
Reply in English.
You must either:
1. return a tool action, or
2. return a final answer.

Return JSON only with this exact shape:
{{
  "reasoning": "short internal reasoning",
  "action": "tool" | "final",
  "tool_name": "list_directory",
  "tool_input": {{"path": "."}},
  "final_answer": "final answer when action is final"
}}

Rules:
- Never invent filesystem state. Use tools whenever the answer depends on the actual files.
- Use only one tool at a time.
- For overwrite, delete, and move actions, call the tool with confirmed=false first.
- Prefer short answers.
- If the task is already complete, return action="final".
- If a successful tool call already completed the requested filesystem change, stop and return action="final" instead of exploring again.
- Do not perform extra verification or follow-up tool calls unless the user explicitly asked for them or the task still has unfinished steps.
- The configured access root is: {base_path or "not set"}.
- If you need to inspect files before acting, do that first.

System prompt:
{_with_table_output_guidance(system_prompt)}

Available tools:
{_file_manager_tool_manifest()}

Recent conversation:
{json.dumps(trimmed_history, ensure_ascii=False, indent=2)}

Scratchpad from this request:
{json.dumps(scratchpad, ensure_ascii=False, indent=2)}

Current user request:
{user_message}
""".strip()

    raw = await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
        response_format="json",
    )
    parsed = extract_json_object(raw)
    action = str(parsed.get("action", "") or "").strip().lower()
    return {
        "reasoning": str(parsed.get("reasoning", "") or "").strip(),
        "action": action,
        "tool_name": str(parsed.get("tool_name", "") or "").strip(),
        "tool_input": parsed.get("tool_input") if isinstance(parsed.get("tool_input"), dict) else {},
        "final_answer": str(parsed.get("final_answer", "") or "").strip(),
    }


FILE_MANAGER_COMPLETION_FALLBACK_TOOLS = {
    "create_file",
    "create_excel_file",
    "write_file",
    "write_excel_sheet",
    "edit_excel_cells",
    "append_excel_rows",
    "delete_file",
    "delete_directory",
    "delete_excel_sheet",
    "move_file",
    "add_excel_sheet",
    "rename_excel_sheet",
}


def _file_manager_completion_answer_from_result(result: dict[str, Any]) -> str:
    summary = str(result.get("summary") or "The file-management task is complete.").strip()
    preview = str(result.get("preview") or "").strip()
    answer = f"## Answer\n{summary}"
    if preview:
        answer += f"\n\n## Preview\n{preview}"
    return answer


async def _evaluate_file_manager_completion(
    user_message: str,
    history: list[dict[str, Any]],
    scratchpad: list[dict[str, Any]],
    tool_name: str,
    tool_input: dict[str, Any],
    result: dict[str, Any],
    system_prompt: str,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str] = None,
    disable_ssl_verification: bool = False,
) -> dict[str, Any]:
    trimmed_history = _normalized_history_messages(
        history,
        current_message=user_message,
        max_steps=CHAT_MEMORY_MAX_STEPS,
    )
    recent_scratchpad = scratchpad[-6:]
    prompt = f"""
You are a completion checker for a file-management agent.
Reply in English.

Return JSON only with this exact shape:
{{
  "is_complete": true,
  "reasoning": "why the task is or is not complete",
  "final_answer": "final user-facing answer when is_complete=true"
}}

Rules:
- Set `is_complete=true` only when the user's request has already been fully satisfied.
- Set `is_complete=false` if another filesystem action is still required.
- After a successful create/write/move/delete/edit/export action, prefer `is_complete=true` unless the user explicitly asked for verification or for additional pending steps.
- Do not ask for extra verification unless the user explicitly requested it.
- Base the decision on actual tool results only.

System prompt:
{_with_table_output_guidance(system_prompt)}

Recent conversation:
{json.dumps(trimmed_history, ensure_ascii=False, indent=2)}

Scratchpad for this request:
{json.dumps(recent_scratchpad, ensure_ascii=False, indent=2)}

Latest successful tool:
{json.dumps({"tool_name": tool_name, "tool_input": tool_input, "result": result}, ensure_ascii=False, indent=2)}

Current user request:
{user_message}
""".strip()

    try:
        raw = await llm_chat(
            [{"role": "user", "content": prompt}],
            llm_base_url,
            llm_model,
            llm_provider,
            llm_api_key,
            response_format="json",
            disable_ssl_verification=disable_ssl_verification,
        )
        parsed = extract_json_object(raw)
        is_complete = bool(parsed.get("is_complete"))
        final_answer = str(parsed.get("final_answer") or "").strip()
        if is_complete and not final_answer:
            final_answer = _file_manager_completion_answer_from_result(result)
        return {
            "is_complete": is_complete,
            "reasoning": str(parsed.get("reasoning") or "").strip(),
            "final_answer": final_answer,
        }
    except Exception as exc:
        if tool_name in FILE_MANAGER_COMPLETION_FALLBACK_TOOLS:
            return {
                "is_complete": True,
                "reasoning": f"Fallback completion after `{tool_name}` because the completion checker failed: {exc}",
                "final_answer": _file_manager_completion_answer_from_result(result),
            }
        return {
            "is_complete": False,
            "reasoning": f"Completion checker failed after `{tool_name}`: {exc}",
            "final_answer": "",
        }


def _default_file_manager_state() -> dict[str, Any]:
    return {
        "pending_confirmation": None,
        "last_tool_result": "",
        "last_visited_path": "",
    }


def _normalize_file_manager_state(payload: Optional[dict]) -> dict[str, Any]:
    state = _default_file_manager_state()
    if not isinstance(payload, dict):
        return state

    pending = payload.get("pending_confirmation") or payload.get("pendingConfirmation")
    if isinstance(pending, dict):
        state["pending_confirmation"] = {
            "tool_name": str(pending.get("tool_name") or pending.get("toolName") or "").strip(),
            "tool_input": pending.get("tool_input") if isinstance(pending.get("tool_input"), dict) else pending.get("toolInput") if isinstance(pending.get("toolInput"), dict) else {},
            "preview": str(pending.get("preview") or "").strip(),
            "summary": str(pending.get("summary") or "").strip(),
            "requested_at": str(pending.get("requested_at") or pending.get("requestedAt") or "").strip(),
        }

    state["last_tool_result"] = str(payload.get("last_tool_result") or payload.get("lastToolResult") or "").strip()
    state["last_visited_path"] = str(payload.get("last_visited_path") or payload.get("lastVisitedPath") or "").strip()
    return state


def _normalize_file_manager_config(payload: Optional[dict]) -> dict[str, Any]:
    defaults = DEFAULT_APP_CONFIG["fileManagerConfig"]
    if not isinstance(payload, dict):
        return dict(defaults)
    max_iterations = payload.get("maxIterations") if "maxIterations" in payload else payload.get("max_iterations")
    return {
        "basePath": str(payload.get("basePath") or payload.get("base_path") or defaults["basePath"]).strip(),
        "maxIterations": max(1, min(FILE_MANAGER_MAX_ITERATIONS, int(max_iterations or defaults["maxIterations"]))),
        "systemPrompt": str(payload.get("systemPrompt") or payload.get("system_prompt") or defaults["systemPrompt"]).strip() or defaults["systemPrompt"],
    }


def _file_manager_confirmation_answer(state: dict[str, Any]) -> tuple[str, list[dict[str, Any]]]:
    pending = state.get("pending_confirmation") or {}
    preview = str(pending.get("preview") or "").strip()
    summary = str(pending.get("summary") or "This action requires confirmation.").strip()
    answer = (
        "## Confirmation Needed\n"
        f"{summary}\n\n"
        f"{preview}\n\n"
        "Please confirm if you want me to continue."
    )
    actions = [
        {
            "id": "confirm-file-action",
            "label": "Confirm",
            "actionType": "confirm_file_action",
            "variant": "primary",
        },
        {
            "id": "cancel-file-action",
            "label": "Cancel",
            "actionType": "cancel_file_action",
            "variant": "secondary",
        },
    ]
    return answer, actions


def _default_pdf_creator_state() -> dict[str, Any]:
    return {
        "stage": "idle",
        "pending_document": None,
        "pending_confirmation": None,
        "last_output_path": "",
        "last_title": "",
    }


def _normalize_pdf_creator_state(payload: Optional[dict]) -> dict[str, Any]:
    state = _default_pdf_creator_state()
    if not isinstance(payload, dict):
        return state

    stage = str(payload.get("stage") or "").strip()
    if stage in {"awaiting_source_choice", "awaiting_content"}:
        state["stage"] = stage

    pending_document = payload.get("pending_document") or payload.get("pendingDocument")
    if isinstance(pending_document, dict):
        state["pending_document"] = dict(pending_document)

    pending_confirmation = payload.get("pending_confirmation") or payload.get("pendingConfirmation")
    if isinstance(pending_confirmation, dict):
        state["pending_confirmation"] = {
            "preview": str(pending_confirmation.get("preview") or "").strip(),
            "summary": str(pending_confirmation.get("summary") or "").strip(),
            "requested_at": str(pending_confirmation.get("requested_at") or pending_confirmation.get("requestedAt") or "").strip(),
            "pending_action": dict(pending_confirmation.get("pending_action") or pending_confirmation.get("pendingAction") or {}),
        }

    state["last_output_path"] = str(payload.get("last_output_path") or payload.get("lastOutputPath") or "").strip()
    state["last_title"] = str(payload.get("last_title") or payload.get("lastTitle") or "").strip()
    return state


MANAGER_SPECIALIST_LABELS = {
    "clickhouse_query": "Clickhouse SQL",
    "data_analyst": "Data analyst",
    "auto_ml": "Auto-ML",
    "data_cleaner": "Data Cleaner",
    "anonymizer": "Anonymizer",
    "custom_agent": "Custom Agent",
    "file_management": "File management",
    "pdf_creator": "PDF creator",
    "oracle_analyst": "Oracle SQL",
}
MANAGER_CLICKHOUSE_FOLLOWUP_STAGES = {
    "awaiting_table",
    "awaiting_field",
    "awaiting_date",
    "awaiting_chart_offer",
    "awaiting_chart_x",
    "awaiting_chart_y",
    "awaiting_chart_type",
}


def _default_manager_agent_state() -> dict[str, Any]:
    return {
        "active_delegate": None,
        "last_routing_reason": "",
        "last_delegate_label": "",
        "pending_pipeline": None,
    }


def _normalize_manager_delegate_role(value: Any, allow_manager: bool = False) -> Optional[str]:
    if not isinstance(value, str):
        return None
    lowered = normalize_choice(value).lower()
    aliases = {
        "manager": "manager",
        "agent manager": "manager",
        "direct": "manager",
        "direct answer": "manager",
        "none": "manager",
        "clickhouse query": "clickhouse_query",
        "clickhouse sql": "clickhouse_query",
        "clickhouse_query": "clickhouse_query",
        "clickhouse": "clickhouse_query",
        "sql": "clickhouse_query",
        "data analyst": "data_analyst",
        "data_analyst": "data_analyst",
        "analyst": "data_analyst",
        "analysis agent": "data_analyst",
        "clickhouse analyst": "data_analyst",
        "analytics agent": "data_analyst",
        "auto ml": "auto_ml",
        "automl": "auto_ml",
        "auto_ml": "auto_ml",
        "ml": "auto_ml",
        "machine learning": "auto_ml",
        "model benchmark": "auto_ml",
        "data cleaner": "data_cleaner",
        "cleaner": "data_cleaner",
        "data_cleaner": "data_cleaner",
        "nettoyeur": "data_cleaner",
        "clean data": "data_cleaner",
        "anonymizer": "anonymizer",
        "anonymiser": "anonymizer",
        "anonymiseur": "anonymizer",
        "anonymizer agent": "anonymizer",
        "gdpr": "anonymizer",
        "rgpd": "anonymizer",
        "custom agent": "custom_agent",
        "custom_agent": "custom_agent",
        "file management": "file_management",
        "file manager": "file_management",
        "file_management": "file_management",
        "filesystem": "file_management",
        "files": "file_management",
        "pdf creator": "pdf_creator",
        "pdf_creator": "pdf_creator",
        "pdf": "pdf_creator",
        "pdf export": "pdf_creator",
        "report pdf": "pdf_creator",
        "oracle analyst": "oracle_analyst",
        "oracle sql": "oracle_analyst",
        "oracle_analyst": "oracle_analyst",
        "oracle": "oracle_analyst",
        "pl/sql": "oracle_analyst",
    }
    resolved = aliases.get(lowered)
    if resolved == "manager" and not allow_manager:
        return None
    return resolved


def _normalize_manager_pending_pipeline(payload: Any) -> Optional[dict[str, Any]]:
    if not isinstance(payload, dict):
        return None
    kind = str(payload.get("kind") or "").strip()
    stage = str(payload.get("stage") or "").strip()
    next_delegate = _normalize_manager_delegate_role(
        payload.get("next_delegate") or payload.get("nextDelegate"),
        allow_manager=False,
    )
    export_format = str(payload.get("export_format") or payload.get("exportFormat") or "").strip().lower() or None
    target_path = str(payload.get("target_path") or payload.get("targetPath") or "").strip()
    source_request = str(payload.get("source_request") or payload.get("sourceRequest") or "").strip()
    reason = str(payload.get("reason") or "").strip()
    if kind == "clickhouse_to_file" and next_delegate == "file_management":
        if stage not in {"awaiting_clickhouse", "awaiting_export_details"}:
            stage = "awaiting_clickhouse"
        if export_format not in {"csv", "tsv", "xlsx", None}:
            export_format = None
        return {
            "kind": "clickhouse_to_file",
            "stage": stage,
            "next_delegate": "file_management",
            "export_format": export_format,
            "target_path": target_path,
            "source_request": source_request,
            "reason": reason,
        }
    if kind == "clickhouse_to_pdf" and next_delegate == "pdf_creator":
        return {
            "kind": "clickhouse_to_pdf",
            "stage": "awaiting_clickhouse",
            "next_delegate": "pdf_creator",
            "target_path": target_path,
            "source_request": source_request,
            "reason": reason,
            "title": str(payload.get("title") or "").strip(),
        }
    return None


def _normalize_manager_agent_state(payload: Optional[dict]) -> dict[str, Any]:
    state = _default_manager_agent_state()
    if not isinstance(payload, dict):
        return state
    state["active_delegate"] = _normalize_manager_delegate_role(
        payload.get("active_delegate") or payload.get("activeDelegate"),
        allow_manager=False,
    )
    state["last_routing_reason"] = str(
        payload.get("last_routing_reason") or payload.get("lastRoutingReason") or ""
    ).strip()
    state["last_delegate_label"] = str(
        payload.get("last_delegate_label") or payload.get("lastDelegateLabel") or ""
    ).strip()
    state["pending_pipeline"] = _normalize_manager_pending_pipeline(
        payload.get("pending_pipeline") or payload.get("pendingPipeline")
    )
    return state


def _clickhouse_state_needs_followup(state: dict[str, Any]) -> bool:
    return str(state.get("stage") or "").strip() in MANAGER_CLICKHOUSE_FOLLOWUP_STAGES


def _file_manager_state_needs_followup(state: dict[str, Any]) -> bool:
    return isinstance(state.get("pending_confirmation"), dict)


def _pdf_creator_state_needs_followup(state: dict[str, Any]) -> bool:
    return (
        str(state.get("stage") or "").strip() in {"awaiting_source_choice", "awaiting_content"}
        or isinstance(state.get("pending_confirmation"), dict)
    )


def _manager_specialist_label(role: Optional[str]) -> str:
    if not role:
        return "Manager"
    return MANAGER_SPECIALIST_LABELS.get(role, role.replace("_", " ").title())


def _manager_trimmed_history(history: list[dict[str, Any]], limit: int = 10) -> list[dict[str, str]]:
    return _normalized_history_messages(history, max_steps=limit)


def _manager_specialist_state_summary(
    clickhouse_state: dict[str, Any],
    data_analyst_state: dict[str, Any],
    auto_ml_state: dict[str, Any],
    data_cleaner_state: dict[str, Any],
    anonymizer_state: dict[str, Any],
    custom_agents: list[dict[str, Any]],
    custom_agent_state: dict[str, Any],
    file_manager_state: dict[str, Any],
    pdf_creator_state: dict[str, Any],
    oracle_state: dict[str, Any],
    manager_state: Optional[dict[str, Any]] = None,
) -> str:
    clickhouse_summary = {
        "stage": clickhouse_state.get("stage") or "idle",
        "selected_table": clickhouse_state.get("selected_table"),
        "has_last_sql": bool(clickhouse_state.get("last_sql")),
        "has_last_rows": bool(clickhouse_state.get("last_result_rows")),
    }
    data_analyst_summary = {
        "stage": data_analyst_state.get("stage") or "idle",
        "selected_table": data_analyst_state.get("selected_table"),
        "last_sql_count": len(data_analyst_state.get("last_sqls") or []),
        "has_last_rows": bool(data_analyst_state.get("last_result_rows")),
        "last_export_path": data_analyst_state.get("last_export_path") or "",
    }
    auto_ml_summary = {
        "stage": auto_ml_state.get("stage") or "idle",
        "selected_table": auto_ml_state.get("selected_table"),
        "target_column": auto_ml_state.get("target_column"),
        "row_filter": auto_ml_state.get("row_filter") or "",
        "sample_row_limit": auto_ml_state.get("sample_row_limit") or "",
        "problem_type": auto_ml_state.get("problem_type") or "",
        "comparison_rows": len(auto_ml_state.get("comparison_rows") or []),
        "recommended_model": auto_ml_state.get("recommended_model") or "",
    }
    data_cleaner_summary = {
        "stage": data_cleaner_state.get("stage") or "idle",
        "selected_table": data_cleaner_state.get("selected_table"),
        "finding_count": len(data_cleaner_state.get("findings") or []),
        "script_count": len(data_cleaner_state.get("correction_scripts") or []),
    }
    anonymizer_summary = {
        "stage": anonymizer_state.get("stage") or "idle",
        "selected_table": anonymizer_state.get("selected_table"),
        "pii_findings": len(anonymizer_state.get("pii_findings") or []),
        "script_count": len(anonymizer_state.get("masking_scripts") or []),
    }
    custom_agent_summary = {
        "selected_agent_id": custom_agent_state.get("selected_agent_id"),
        "enabled_agents": [
            {
                "id": agent.get("id"),
                "title": agent.get("title"),
                "routing_hint": agent.get("managerRoutingHint") or "",
            }
            for agent in custom_agents[:12]
        ],
    }
    file_summary = {
        "pending_confirmation": bool(file_manager_state.get("pending_confirmation")),
        "last_visited_path": file_manager_state.get("last_visited_path") or "",
        "last_tool_result": _truncate_text_preview(
            str(file_manager_state.get("last_tool_result") or ""), 240
        ),
    }
    pdf_summary = {
        "stage": pdf_creator_state.get("stage") or "idle",
        "pending_confirmation": bool(pdf_creator_state.get("pending_confirmation")),
        "last_output_path": pdf_creator_state.get("last_output_path") or "",
        "last_title": pdf_creator_state.get("last_title") or "",
    }
    oracle_summary = {
        "stage": oracle_state.get("stage") or "idle",
        "selected_table": oracle_state.get("selected_table"),
        "has_last_sql": bool(oracle_state.get("last_sql")),
        "has_last_rows": bool(oracle_state.get("last_result_rows")),
        "awaiting_options": len(oracle_state.get("clarification_options") or []),
    }
    return json.dumps(
        {
            "manager": {
                "pending_pipeline": (manager_state or {}).get("pending_pipeline"),
            },
            "clickhouse": clickhouse_summary,
            "data_analyst": data_analyst_summary,
            "auto_ml": auto_ml_summary,
            "data_cleaner": data_cleaner_summary,
            "anonymizer": anonymizer_summary,
            "custom_agent": custom_agent_summary,
            "file_management": file_summary,
            "pdf_creator": pdf_summary,
            "oracle_analyst": oracle_summary,
        },
        ensure_ascii=False,
        indent=2,
    )


MANAGER_EXPORT_FORMAT_OPTIONS = ["CSV (.csv)", "Excel (.xlsx)"]
MANAGER_EXPORT_KEYWORDS = [
    "export",
    "save",
    "write",
    "create file",
    "create csv",
    "create excel",
    "generate csv",
    "generate excel",
    "download csv",
    "download excel",
    "sauvegarder",
    "exporter",
    "creer un csv",
    "creer un excel",
    "genere un csv",
    "genere un excel",
]
MANAGER_PDF_KEYWORDS = [
    "pdf",
    "report",
    "one pager",
    "one-pager",
    "brief pdf",
    "document pdf",
    "export pdf",
    "save as pdf",
    "create pdf",
    "generate pdf",
    "make a pdf",
    "exporter en pdf",
    "creer un pdf",
    "créer un pdf",
    "generer un pdf",
    "générer un pdf",
]


def _extract_manager_export_format(user_message: str) -> Optional[str]:
    normalized = normalize_intent_text(user_message)
    if ".csv" in normalized or re.search(r"\bcsv\b", normalized):
        return "csv"
    if ".tsv" in normalized or re.search(r"\btsv\b", normalized):
        return "tsv"
    if any(token in normalized for token in [".xlsx", ".xls"]) or re.search(r"\bexcel\b", normalized):
        return "xlsx"
    return None


def _extract_manager_export_path(user_message: str, export_format: Optional[str] = None) -> str:
    lowered = user_message.strip()
    quoted_match = re.search(r'["\']([^"\']+\.(?:csv|tsv|xlsx|xls))["\']', lowered, flags=re.IGNORECASE)
    if quoted_match:
        candidate = quoted_match.group(1).strip()
        if candidate.lower().endswith(".xls"):
            return candidate[:-4] + ".xlsx"
        return candidate

    patterns = [
        r'(?:named|called|as|to|into|in|vers|dans|nomme|nommé|appele|appelé)\s+([A-Za-z0-9_./\\\\-]+\.(?:csv|tsv|xlsx|xls))',
    ]
    extension = f".{export_format}" if export_format in {"csv", "tsv", "xlsx"} else ""
    for pattern in patterns:
        match = re.search(pattern, lowered, flags=re.IGNORECASE)
        if not match:
            continue
        candidate = match.group(1).strip()
        if not candidate:
            continue
        if extension and "." not in Path(candidate).name:
            candidate = f"{candidate}{extension}"
        if re.search(r"\.(csv|tsv|xlsx|xls)$", candidate, flags=re.IGNORECASE):
            if candidate.lower().endswith(".xls"):
                return candidate[:-4] + ".xlsx"
            return candidate

    matches = re.findall(r'([A-Za-z0-9_./\\\\-]+\.(?:csv|tsv|xlsx|xls))', lowered, flags=re.IGNORECASE)
    if matches:
        candidate = matches[-1].strip().strip('"').strip("'")
        if candidate.lower().endswith(".xls"):
            return candidate[:-4] + ".xlsx"
        return candidate
    return ""


def _extract_clickhouse_file_export_pipeline(user_message: str) -> Optional[dict[str, Any]]:
    normalized = normalize_intent_text(user_message)
    if any(token in normalized for token in MANAGER_PDF_KEYWORDS):
        return None
    has_export_signal = any(token in normalized for token in MANAGER_EXPORT_KEYWORDS) or any(
        token in normalized for token in ["csv", "tsv", "excel", "xlsx", "xls"]
    )
    if not has_export_signal:
        return None

    export_format = _extract_manager_export_format(user_message)
    target_path = _extract_manager_export_path(user_message, export_format)
    return {
        "kind": "clickhouse_to_file",
        "stage": "awaiting_clickhouse",
        "next_delegate": "file_management",
        "export_format": export_format,
        "target_path": target_path,
        "source_request": user_message.strip(),
        "reason": "The user wants the ClickHouse result to be exported as a file after the query runs.",
    }


def _extract_clickhouse_pdf_export_pipeline(user_message: str) -> Optional[dict[str, Any]]:
    normalized = normalize_intent_text(user_message)
    has_pdf_signal = any(token in normalized for token in MANAGER_PDF_KEYWORDS)
    if not has_pdf_signal:
        return None

    target_path = _extract_pdf_path(user_message)
    title = _extract_pdf_title(user_message, target_path)
    return {
        "kind": "clickhouse_to_pdf",
        "stage": "awaiting_clickhouse",
        "next_delegate": "pdf_creator",
        "target_path": target_path or _default_pdf_target_path(title),
        "source_request": user_message.strip(),
        "reason": "The user wants the ClickHouse result to be exported as a professional PDF after the query runs.",
        "title": title,
    }


def _manager_pending_pipeline_requires_details(pipeline: Optional[dict[str, Any]]) -> bool:
    if not isinstance(pipeline, dict):
        return False
    if pipeline.get("kind") != "clickhouse_to_file":
        return False
    return not pipeline.get("export_format") or not pipeline.get("target_path")


def _manager_export_details_prompt(pipeline: dict[str, Any]) -> str:
    missing_format = not pipeline.get("export_format")
    missing_path = not pipeline.get("target_path")
    if missing_format and missing_path:
        return append_choice_markdown(
            (
                "## File Export Details Needed\n"
                "I already have the ClickHouse result. To create the export file, I still need the output format and the target file name or path."
            ),
            "Format",
            "Choose the file format first.",
            MANAGER_EXPORT_FORMAT_OPTIONS,
        ) + "\n\nThen tell me the target file name or full path, for example `exports/result.csv`."
    if missing_format:
        return append_choice_markdown(
            (
                "## File Export Details Needed\n"
                f"I already have the ClickHouse result and the target path `{pipeline.get('target_path')}`."
            ),
            "Format",
            "Choose the file format for the export.",
            MANAGER_EXPORT_FORMAT_OPTIONS,
        )
    return (
        "## File Export Details Needed\n"
        f"I already have the ClickHouse result and the requested format **{pipeline.get('export_format')}**.\n\n"
        "Please tell me the target file name or full path, for example `exports/result."
        f"{pipeline.get('export_format')}`."
    )


def _manager_update_export_pipeline_from_reply(
    pipeline: dict[str, Any],
    user_message: str,
) -> dict[str, Any]:
    updated = dict(pipeline)
    format_choice = resolve_user_choice(user_message, MANAGER_EXPORT_FORMAT_OPTIONS)
    if format_choice == "CSV (.csv)":
        updated["export_format"] = "csv"
    elif format_choice == "Excel (.xlsx)":
        updated["export_format"] = "xlsx"
    elif not updated.get("export_format"):
        updated["export_format"] = _extract_manager_export_format(user_message)

    detected_path = _extract_manager_export_path(user_message, updated.get("export_format"))
    if detected_path:
        updated["target_path"] = detected_path

    if updated.get("export_format") and updated.get("target_path"):
        updated["stage"] = "awaiting_clickhouse"
    else:
        updated["stage"] = "awaiting_export_details"
    return updated


def _manager_export_headers_and_rows(clickhouse_state: dict[str, Any]) -> tuple[list[str], list[list[Any]]]:
    headers = [
        str(item.get("name") or "").strip()
        for item in (clickhouse_state.get("last_result_meta") or [])
        if str(item.get("name") or "").strip()
    ]
    rows_data = clickhouse_state.get("last_result_rows") or []
    if not headers and rows_data:
        headers = [str(key) for key in rows_data[0].keys()]

    def _cell(value: Any) -> Any:
        if value is None:
            return ""
        if isinstance(value, (dict, list)):
            return json.dumps(value, ensure_ascii=False)
        return value

    rows = [
        [_cell(row.get(header)) for header in headers]
        for row in rows_data
    ]
    return headers, rows


def _build_file_export_payload_from_clickhouse(
    pipeline: dict[str, Any],
    clickhouse_state: dict[str, Any],
) -> dict[str, Any]:
    headers, rows = _manager_export_headers_and_rows(clickhouse_state)
    export_format = pipeline.get("export_format") or "csv"
    target_path = str(pipeline.get("target_path") or "").strip()
    if export_format == "xlsx" and target_path.lower().endswith(".xls"):
        target_path = target_path[:-4] + ".xlsx"
    return {
        "__file_export__": True,
        "format": export_format,
        "path": target_path,
        "sheet_name": "Results",
        "headers": headers,
        "rows": rows,
        "source_sql": str(clickhouse_state.get("last_sql") or "").strip(),
        "source_request": str(pipeline.get("source_request") or "").strip(),
    }


def _build_pdf_export_payload_from_clickhouse(
    pipeline: dict[str, Any],
    clickhouse_answer: str,
) -> dict[str, Any]:
    title = str(pipeline.get("title") or "").strip() or _extract_pdf_title(
        str(pipeline.get("source_request") or ""),
        str(pipeline.get("target_path") or ""),
    )
    target_path = str(pipeline.get("target_path") or "").strip() or _default_pdf_target_path(title)
    return {
        "__pdf_export__": True,
        "title": title,
        "subtitle": "Analysis export generated by RAGnarok",
        "path": target_path,
        "source_markdown": str(clickhouse_answer or "").strip(),
        "source_request": str(pipeline.get("source_request") or "").strip(),
    }


def _manager_compose_chained_answer(
    primary_answer: str,
    secondary_answer: str,
    secondary_label: str,
) -> str:
    first = str(primary_answer or "").strip()
    second = str(secondary_answer or "").strip()
    if not first:
        return second
    if not second:
        return first
    return (
        f"{first}\n\n---\n\n"
        f"## {secondary_label}\n"
        "I then continued with the next specialist to complete the rest of the same request.\n\n"
        f"{second}"
    )


def _heuristic_manager_delegate(
    user_message: str,
    manager_state: dict[str, Any],
    clickhouse_state: dict[str, Any],
    data_analyst_state: dict[str, Any],
    auto_ml_state: dict[str, Any],
    data_cleaner_state: dict[str, Any],
    anonymizer_state: dict[str, Any],
    custom_agents: list[dict[str, Any]],
    custom_agent_state: dict[str, Any],
    file_manager_state: dict[str, Any],
    pdf_creator_state: dict[str, Any],
    oracle_state: dict[str, Any],
) -> Optional[tuple[str, str]]:
    normalized = normalize_intent_text(user_message)
    active_delegate = manager_state.get("active_delegate")
    pending_pipeline = manager_state.get("pending_pipeline")

    if active_delegate == "file_management" and _file_manager_state_needs_followup(file_manager_state):
        return "file_management", "Continuing the active file-management confirmation flow."
    if active_delegate == "clickhouse_query" and _clickhouse_state_needs_followup(clickhouse_state):
        return "clickhouse_query", "Continuing the active ClickHouse clarification flow."
    if active_delegate == "data_analyst" and _data_analyst_state_needs_followup(data_analyst_state):
        return "data_analyst", "Continuing the active Data analyst table-selection flow."
    if active_delegate == "auto_ml" and _auto_ml_state_needs_followup(auto_ml_state):
        return "auto_ml", "Continuing the active Auto-ML clarification flow."
    if active_delegate == "data_cleaner" and _data_cleaner_state_needs_followup(data_cleaner_state):
        return "data_cleaner", "Continuing the active Data Cleaner table-selection flow."
    if active_delegate == "anonymizer" and _anonymizer_state_needs_followup(anonymizer_state):
        return "anonymizer", "Continuing the active Anonymizer table-selection flow."
    if active_delegate == "pdf_creator" and _pdf_creator_state_needs_followup(pdf_creator_state):
        return "pdf_creator", "Continuing the active PDF-creation flow."
    if active_delegate == "oracle_analyst" and _oracle_state_needs_followup(oracle_state):
        return "oracle_analyst", "Continuing the active Oracle table-selection flow."
    if _file_manager_state_needs_followup(file_manager_state) and (
        is_affirmative_response(user_message)
        or is_negative_response(user_message)
        or normalized in {"confirm", "confirm file action", "cancel", "cancel file action"}
    ):
        return "file_management", "The user is answering a pending file-management confirmation."

    if _clickhouse_state_needs_followup(clickhouse_state):
        return "clickhouse_query", "The user is continuing a ClickHouse clarification step."
    if _data_analyst_state_needs_followup(data_analyst_state):
        return "data_analyst", "The user is continuing a Data analyst clarification step."
    if _auto_ml_state_needs_followup(auto_ml_state):
        return "auto_ml", "The user is continuing an Auto-ML clarification step."
    if _data_cleaner_state_needs_followup(data_cleaner_state):
        return "data_cleaner", "The user is continuing a Data Cleaner clarification step."
    if _anonymizer_state_needs_followup(anonymizer_state):
        return "anonymizer", "The user is continuing an Anonymizer clarification step."
    if _pdf_creator_state_needs_followup(pdf_creator_state):
        return "pdf_creator", "The user is continuing a PDF export clarification or confirmation."
    if _oracle_state_needs_followup(oracle_state):
        return "oracle_analyst", "The user is continuing an Oracle table-selection step."
    if (
        isinstance(pending_pipeline, dict)
        and pending_pipeline.get("kind") == "clickhouse_to_file"
        and pending_pipeline.get("stage") == "awaiting_export_details"
    ):
        return "manager", "The manager is waiting for the remaining file-export details before handing off to File management."

    if clickhouse_state.get("last_result_rows") and is_chart_followup_request(user_message):
        return "clickhouse_query", "The user is continuing the latest ClickHouse result with a chart follow-up."
    if data_analyst_state.get("last_result_rows") and _data_analyst_export_requested(user_message):
        return "data_analyst", "The user is asking the Data analyst to export the latest analytical dataset to CSV."

    file_tokens = [
        "file",
        "fichier",
        "folder",
        "dossier",
        "directory",
        "repertoire",
        "path",
        "chemin",
        "csv",
        "tsv",
        "txt",
        "md",
        "markdown",
        "json",
        "yaml",
        "yml",
        "sql",
        "py",
        "html",
        "xlsx",
        "xls",
        "excel",
        "docx",
        "parquet",
        "read",
        "open",
        "inspect",
        "save",
        "write",
        "create",
        "edit",
        "update",
        "append",
        "delete",
        "remove",
        "move",
        "rename",
        "list files",
        "search files",
        "lire",
        "ouvrir",
        "inspecter",
        "sauvegarder",
        "ecrire",
        "creer",
        "modifier",
        "mettre a jour",
        "ajouter",
        "supprimer",
        "effacer",
        "deplacer",
        "renommer",
        "lister les fichiers",
        "chercher des fichiers",
    ]
    clickhouse_tokens = [
        "clickhouse",
        "clickhouse sql",
        "clickhouse query",
        "chart",
        "graph",
        "graphique",
        "graphe",
        "courbe",
    ]
    generic_sql_tokens = [
        "sql",
        "table",
        "column",
        "database",
        "query",
        "rows",
        "count",
        "metrics",
        "measure",
        "aggregation",
        "trend",
        "schema",
        "requete",
        "base de donnees",
        "colonne",
        "lignes",
        "compte",
        "mesures",
        "agregation",
        "tendance",
        "schema",
    ]
    business_metric_tokens = [
        "revenue",
        "sales",
        "orders",
        "customers",
        "customer",
        "transactions",
        "transaction",
        "invoice",
        "invoices",
        "profit",
        "margin",
        "gmv",
        "users",
        "sessions",
        "events",
        "jobs",
        "retention",
        "funnel",
        "conversion",
        "kpi",
        "kpis",
        "ventes",
        "commandes",
        "clients",
        "client",
        "factures",
        "facture",
        "revenu",
        "revenus",
        "marge",
        "utilisateurs",
        "evenements",
        "événements",
    ]
    data_request_verbs = [
        "show",
        "list",
        "find",
        "count",
        "summarize",
        "summary",
        "compare",
        "display",
        "break down",
        "trend",
        "what",
        "which",
        "how many",
        "give me",
        "query",
        "montre",
        "affiche",
        "liste",
        "trouve",
        "compte",
        "resume",
        "résume",
        "compare",
        "donne moi",
    ]
    oracle_tokens = [
        "oracle",
        "oracle db",
        "oracle database",
        "oracle sql",
        "pl/sql",
        "ora-",
        "tns",
        "service name",
        "sid",
        "dual",
        "fetch first",
        "row_number() over",
    ]
    data_analyst_tokens = [
        "analyze",
        "analysis",
        "analyse",
        "analyses",
        "deep dive",
        "investigate",
        "investigation",
        "root cause",
        "why did",
        "why is",
        "diagnose",
        "diagnostic",
        "break down",
        "breakdown",
        "segment",
        "cohort",
        "retention",
        "funnel",
        "compare",
        "comparison",
        "driver",
        "drivers",
        "insight",
        "insights",
        "drill down",
        "trend analysis",
        "what changed",
        "pourquoi",
        "analyse approfondie",
        "analyse complexe",
        "investigue",
        "investiguer",
        "diagnostique",
        "cause racine",
        "segmente",
        "segmenter",
        "cohorte",
        "entonnoir",
        "comparer",
        "derive",
        "derivee",
        "insight",
        "insights",
    ]
    auto_ml_tokens = [
        "auto ml",
        "automl",
        "machine learning",
        "model benchmark",
        "benchmark models",
        "compare models",
        "classifier",
        "regression model",
        "random forest",
        "xgboost",
        "precision",
        "recall",
        "f1 score",
        "train model",
        "predictive model",
    ]
    data_cleaner_tokens = [
        "data cleaner", "clean data", "cleanup data", "duplicates", "duplicate rows", "missing values", "null values",
        "empty strings", "inconsistent format", "date format", "nettoyeur", "deduplicate", "dedup", "clean-up",
    ]
    anonymizer_tokens = [
        "anonymizer", "anonymiser", "anonymize", "anonymise", "pii", "personally identifiable", "personal data",
        "gdpr", "rgpd", "mask data", "hash data", "masking", "tokenize", "pseudonymize", "pseudonymise",
    ]
    file_action_tokens = [
        "create",
        "make",
        "generate",
        "save",
        "write",
        "edit",
        "update",
        "append",
        "move",
        "rename",
        "delete",
        "remove",
        "copy",
        "open",
        "read",
        "creer",
        "fais",
        "faire",
        "generer",
        "sauvegarder",
        "ecrire",
        "modifier",
        "mettre a jour",
        "ajouter",
        "deplacer",
        "renommer",
        "supprimer",
        "effacer",
        "copier",
        "ouvrir",
        "lire",
    ]
    file_target_tokens = [
        "file",
        "fichier",
        "folder",
        "dossier",
        "directory",
        "repertoire",
        "document",
        "spreadsheet",
        "workbook",
        "sheet",
        "excel",
        "csv",
        "parquet",
        "docx",
        "markdown",
        "json",
        "yaml",
        "yml",
        "sql",
        "python file",
        "texte",
    ]
    file_extension_hit = bool(
        re.search(r"\.(txt|md|csv|tsv|xlsx|xls|docx|parquet|json|ya?ml|html|sql|py|log)\b", normalized)
    )
    file_action_hit = any(token in normalized for token in file_action_tokens)
    file_target_hit = file_extension_hit or any(token in normalized for token in file_target_tokens)
    file_creation_or_edit_hit = file_action_hit and file_target_hit

    file_hit = file_creation_or_edit_hit or any(token in normalized for token in file_tokens)
    oracle_hit = any(token in normalized for token in oracle_tokens)
    clickhouse_hit = any(token in normalized for token in clickhouse_tokens)
    generic_sql_hit = any(token in normalized for token in generic_sql_tokens)
    business_query_hit = (
        any(token in normalized for token in business_metric_tokens)
        and any(token in normalized for token in data_request_verbs)
    )
    sql_execution_hit = oracle_hit or clickhouse_hit or generic_sql_hit or business_query_hit
    data_analyst_hit = any(token in normalized for token in data_analyst_tokens)
    auto_ml_hit = any(token in normalized for token in auto_ml_tokens)
    data_cleaner_hit = any(token in normalized for token in data_cleaner_tokens)
    anonymizer_hit = any(token in normalized for token in anonymizer_tokens)
    custom_agent_match = next((
        agent for agent in custom_agents
        if any(
            token and token in normalized
            for token in {
                normalize_intent_text(str(agent.get("title") or "")),
                normalize_intent_text(str(agent.get("managerRoutingHint") or "")),
            }
        )
    ), None)
    pdf_hit = any(token in normalized for token in MANAGER_PDF_KEYWORDS)
    export_pipeline = _extract_clickhouse_file_export_pipeline(user_message)
    pdf_pipeline = _extract_clickhouse_pdf_export_pipeline(user_message)
    clickhouse_default_hit = clickhouse_hit or (sql_execution_hit and not oracle_hit)

    if oracle_hit:
        return "oracle_analyst", "The request requires Oracle SQL execution or Oracle schema exploration, so Oracle SQL should handle it."
    if auto_ml_hit and not oracle_hit:
        return "auto_ml", "The request is explicitly about benchmarking machine-learning models on ClickHouse data."
    if anonymizer_hit and not oracle_hit:
        return "anonymizer", "The request is explicitly about PII detection, masking, or GDPR-style anonymization on ClickHouse data."
    if data_cleaner_hit and not oracle_hit:
        return "data_cleaner", "The request is explicitly about duplicates, missing values, or inconsistent formats on ClickHouse data."
    if custom_agent_match:
        return "custom_agent", f"The request explicitly matches the custom specialist `{custom_agent_match.get('title')}`."
    if export_pipeline and clickhouse_default_hit:
        return "clickhouse_query", "The request needs Clickhouse SQL first and then a file export from the query result."
    if pdf_pipeline and clickhouse_default_hit:
        return "clickhouse_query", "The request needs Clickhouse SQL first and then a PDF export from the query result."
    if clickhouse_default_hit and data_analyst_hit:
        return "data_analyst", "The request is a complex ClickHouse investigation, so Data analyst should handle it end to end."
    if pdf_hit and not sql_execution_hit:
        return "pdf_creator", "The request is explicitly about turning content or the latest analysis into a PDF."
    if file_creation_or_edit_hit:
        return "file_management", "The request explicitly asks to create or modify a file, so File management should handle it."
    if file_hit and not sql_execution_hit:
        return "file_management", "The request is explicitly about filesystem or spreadsheet actions."
    if clickhouse_default_hit:
        return "clickhouse_query", "The request requires Clickhouse SQL execution, schema inspection, or charting."
    return None


async def analyze_manager_routing(
    user_message: str,
    history: list[dict[str, Any]],
    manager_state: dict[str, Any],
    clickhouse_state: dict[str, Any],
    data_analyst_state: dict[str, Any],
    auto_ml_state: dict[str, Any],
    data_cleaner_state: dict[str, Any],
    anonymizer_state: dict[str, Any],
    custom_agents: list[dict[str, Any]],
    custom_agent_state: dict[str, Any],
    file_manager_state: dict[str, Any],
    pdf_creator_state: dict[str, Any],
    oracle_state: dict[str, Any],
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str] = None,
    functional_context: str = "",
) -> dict[str, str]:
    heuristic = _heuristic_manager_delegate(
        user_message,
        manager_state,
        clickhouse_state,
        data_analyst_state,
        auto_ml_state,
        data_cleaner_state,
        anonymizer_state,
        custom_agents,
        custom_agent_state,
        file_manager_state,
        pdf_creator_state,
        oracle_state,
    )
    if heuristic:
        result = {
            "delegate": heuristic[0],
            "reasoning": heuristic[1],
            "handoff_message": user_message,
        }
        if heuristic[0] == "custom_agent":
            normalized = normalize_intent_text(user_message)
            matched_custom = next((
                agent for agent in custom_agents
                if any(
                    token and token in normalized
                    for token in {
                        normalize_intent_text(str(agent.get("title") or "")),
                        normalize_intent_text(str(agent.get("managerRoutingHint") or "")),
                    }
                )
            ), None)
            if matched_custom:
                result["custom_agent_id"] = str(matched_custom.get("id") or "")
        return result

    trimmed_history = _manager_trimmed_history(history)
    prompt = f"""
You are the routing brain of the RAGnarok Agent Manager.
Choose which specialist should handle the next turn.

Available delegates:
- manager: use this when no specialist tool is needed and the manager can answer directly.
- clickhouse_query: use this for Clickhouse SQL execution, schema inspection, table selection, metrics, charts, and the first SQL step before a downstream export.
- data_analyst: use this only for complex multi-step investigations on ClickHouse data when a deeper end-to-end analysis is needed beyond a simple SQL answer.
- auto_ml: use this to benchmark several machine-learning models on ClickHouse data and compare their performance.
- data_cleaner: use this to audit ClickHouse data quality, detect duplicates and missing values, flag inconsistent formats, and propose correction scripts.
- anonymizer: use this to scan ClickHouse datasets for PII exposure and recommend masking, hashing, or tokenization patterns.
- custom_agent: use this when one of the enabled custom agents is the best fit for the request.
- file_management: use this for filesystem actions, directories, files, CSV/Excel/Word/Parquet handling, create/edit/move/delete operations.
- pdf_creator: use this to create a clean, professional PDF export from an analysis, a report, or the latest relevant result in the chat.
- oracle_analyst: use this for Oracle SQL execution, Oracle database discovery, schema inspection, query validation, and narrative business answers from Oracle data.

If more than one specialist could be relevant, choose the one that should act first.
Return JSON only with this exact shape:
{{
  "delegate": "manager" | "clickhouse_query" | "data_analyst" | "auto_ml" | "data_cleaner" | "anonymizer" | "custom_agent" | "file_management" | "pdf_creator" | "oracle_analyst",
  "reasoning": "short English explanation",
  "handoff_message": "short English specialist instruction preserving the user's intent",
  "custom_agent_id": "required when delegate is custom_agent"
}}

Rules:
- Keep the answer in English.
- Prefer a specialist when the request depends on real filesystem state, ClickHouse data, or Oracle data.
- If the request needs SQL execution, delegate to `clickhouse_query` for ClickHouse or `oracle_analyst` for Oracle.
- If the database is not explicitly Oracle and the task still needs SQL execution, default to `clickhouse_query`.
- Use `data_analyst` only when the task is a complex ClickHouse investigation that should run several analytical steps autonomously.
- Use `auto_ml` when the user wants to benchmark, compare, or score machine-learning models on ClickHouse data.
- Use `data_cleaner` when the request is about duplicates, missing values, inconsistent formats, or cleanup scripts.
- Use `anonymizer` when the request is about PII, masking, hashing, anonymization, or GDPR-style privacy safeguards.
- Use `custom_agent` only when an enabled custom agent is clearly a better fit than the built-in specialists.
- Never delegate Oracle work to `data_analyst`.
- Never delegate Oracle work to `auto_ml`.
- Never delegate Oracle work to `data_cleaner`.
- Never delegate Oracle work to `anonymizer`.
- If the user asks to create, write, save, edit, move, rename, or delete a file or folder, delegate to `file_management`.
- If the user asks for a PDF export, a report PDF, or a professional PDF document, delegate to `pdf_creator` unless a ClickHouse query must happen first.
- Keep `handoff_message` concise and actionable.
- If the manager can answer directly, set `delegate` to `manager`.

Current manager state:
{json.dumps(manager_state, ensure_ascii=False, indent=2)}

Current specialist state summary:
{_manager_specialist_state_summary(clickhouse_state, data_analyst_state, auto_ml_state, data_cleaner_state, anonymizer_state, custom_agents, custom_agent_state, file_manager_state, pdf_creator_state, oracle_state, manager_state)}

Functional knowledge context from RAG:
{functional_context or "No additional functional context loaded."}

Enabled custom agents:
{json.dumps([
    {
        "id": agent.get("id"),
        "title": agent.get("title"),
        "description": agent.get("description"),
        "manager_routing_hint": agent.get("managerRoutingHint"),
    }
    for agent in custom_agents[:12]
], ensure_ascii=False, indent=2)}

Recent conversation:
{json.dumps(trimmed_history, ensure_ascii=False, indent=2)}

User message:
{user_message}
""".strip()

    raw = await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
        response_format="json",
    )
    parsed = extract_json_object(raw)
    delegate = _normalize_manager_delegate_role(parsed.get("delegate"), allow_manager=True) or "manager"
    handoff_message = str(parsed.get("handoff_message") or user_message).strip() or user_message
    reasoning = str(parsed.get("reasoning") or "").strip() or "The manager selected the best available execution path."
    result = {
        "delegate": delegate,
        "reasoning": reasoning,
        "handoff_message": handoff_message,
    }
    if delegate == "custom_agent":
        custom_agent_id = str(parsed.get("custom_agent_id") or "").strip()
        if custom_agent_id and any(agent.get("id") == custom_agent_id for agent in custom_agents):
            result["custom_agent_id"] = custom_agent_id
    return result


async def _run_manager_direct_response(
    history: list[dict[str, Any]],
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str],
    system_prompt: str,
    functional_context: str = "",
) -> str:
    manager_system_prompt = (
        f"{_with_table_output_guidance(system_prompt)}\n\n"
        "You are the RAGnarok Agent Manager. Reply in English. Keep answers concise, "
        "use specialist agents only when needed, and explain clearly when you can answer directly."
    ).strip()
    messages = [{"role": "system", "content": manager_system_prompt}]
    if functional_context.strip():
        messages.append(
            {
                "role": "system",
                "content": "Functional knowledge context from the RAG knowledge base:\n" + functional_context.strip(),
            }
        )
    messages.extend(_manager_trimmed_history(history, limit=12))
    return await llm_chat(
        messages,
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
    )


def _prefix_agent_steps(
    steps: list[dict[str, Any]],
    specialist_role: str,
) -> list[dict[str, Any]]:
    label = _manager_specialist_label(specialist_role)
    prefixed: list[dict[str, Any]] = []
    for index, step in enumerate(steps):
        prefixed.append(
            {
                **step,
                "id": str(step.get("id") or f"{specialist_role}-{index}"),
                "title": f"{label} · {step.get('title') or 'Step'}",
            }
        )
    return prefixed


# ── ClickHouse agent LLM helpers ──────────────────────────────────────────────

async def analyze_clickhouse_schema(
    user_request: str,
    table_name: str,
    schema: list[dict[str, str]],
    conversation_memory: str,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str] = None,
) -> dict[str, Any]:
    schema_lines = "\n".join(
        f"- {column['name']}: {column['type']}"
        for column in schema[:120]
    )
    prompt = f"""
You are a ClickHouse query planner.
Analyze the user's request against a single ClickHouse table schema.
Return JSON only with this exact shape:
{{
  "field_candidates": ["col_a", "col_b"],
  "field_choice_required": true,
  "field_choice_prompt": "Which field should I use for ...?",
  "date_candidates": ["created_at", "updated_at"],
  "date_choice_required": true,
  "date_choice_prompt": "Which date column should I use?",
  "needs_date_column": true,
  "reasoning": "short explanation"
}}

Rules:
- Keep prompts in English.
- Only reference columns that exist in the schema.
- If the best field is obvious, return one field candidate and set field_choice_required to false.
- If the request does not need a date column, set needs_date_column to false and leave date_candidates empty.
- If there are multiple plausible date columns and a date context is needed, set date_choice_required to true.

Table: {table_name}
Schema:
{schema_lines}

Recent conversation memory:
{conversation_memory}

User request:
{user_request}
""".strip()

    raw = await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
        response_format="json",
    )
    parsed = extract_json_object(raw)
    return {
        "field_candidates": [
            item for item in parsed.get("field_candidates", [])
            if isinstance(item, str)
        ],
        "field_choice_required": bool(parsed.get("field_choice_required", False)),
        "field_choice_prompt": str(parsed.get("field_choice_prompt", "") or ""),
        "date_candidates": [
            item for item in parsed.get("date_candidates", [])
            if isinstance(item, str)
        ],
        "date_choice_required": bool(parsed.get("date_choice_required", False)),
        "date_choice_prompt": str(parsed.get("date_choice_prompt", "") or ""),
        "needs_date_column": bool(parsed.get("needs_date_column", False)),
        "reasoning": str(parsed.get("reasoning", "") or ""),
    }


async def analyze_clickhouse_tables(
    user_request: str,
    available_tables: list[str],
    conversation_memory: str,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str] = None,
) -> dict[str, Any]:
    tables_text = "\n".join(f"- {table_name}" for table_name in available_tables[:200])
    prompt = f"""
You are routing a user analytics request to the best ClickHouse table.
Return JSON only with this exact shape:
{{
  "selected_table": "orders",
  "table_candidates": ["orders", "order_items"],
  "table_choice_required": false,
  "table_choice_prompt": "Which table should I use for this request?",
  "reasoning": "short explanation"
}}

Rules:
- Use only table names from the provided list.
- If one table is clearly the best match, set selected_table and table_choice_required to false.
- If several tables are plausible, set table_choice_required to true and return the best candidates.
- Keep prompts in English.
- Prefer not asking the user unless there is a real ambiguity.

Available tables:
{tables_text}

Recent conversation memory:
{conversation_memory}

User request:
{user_request}
""".strip()

    raw = await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
        response_format="json",
    )
    parsed = extract_json_object(raw)
    return {
        "selected_table": str(parsed.get("selected_table", "") or ""),
        "table_candidates": [
            item for item in parsed.get("table_candidates", [])
            if isinstance(item, str)
        ],
        "table_choice_required": bool(parsed.get("table_choice_required", False)),
        "table_choice_prompt": str(parsed.get("table_choice_prompt", "") or ""),
        "reasoning": str(parsed.get("reasoning", "") or ""),
    }


async def generate_clickhouse_sql(
    user_request: str,
    table_name: str,
    schema: list[dict[str, str]],
    selected_field: Optional[str],
    selected_date_field: Optional[str],
    conversation_memory: str,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str],
    query_limit: int,
    error_feedback: str = "",
) -> dict[str, str]:
    schema_lines = "\n".join(
        f"- {column['name']}: {column['type']}"
        for column in schema[:120]
    )
    prompt = f"""
You are a senior ClickHouse analytics engineer.
Generate exactly one safe read-only SQL query for the user's request.
Return JSON only with:
{{
  "sql": "SELECT ...",
  "reasoning": "short explanation"
}}

Constraints:
- Use only table `{table_name}`.
- Use only columns from the schema below.
- Output English-friendly aliases when useful.
- Never use INSERT, UPDATE, DELETE, ALTER, DROP, CREATE, TRUNCATE, SYSTEM or multiple statements.
- Prefer explicit column names.
- Add a LIMIT when returning raw rows.
- Keep the SQL compatible with ClickHouse.
- If a date column is selected, prefer it when time filtering, grouping or ordering matters.

Selected field: {selected_field or "None"}
Selected date field: {selected_date_field or "None"}
Maximum row limit: {query_limit}
Schema:
{schema_lines}

Recent conversation memory:
{conversation_memory}

User request:
{user_request}

{("Previous execution error to fix: " + str(error_feedback)) if error_feedback else ""}
""".strip()

    raw = await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
        response_format="json",
    )
    parsed = extract_json_object(raw)
    sql = clean_sql_text(str(parsed.get("sql", "") or ""))
    reasoning = str(parsed.get("reasoning", "") or "")
    return {"sql": enforce_query_limit(sql, query_limit), "reasoning": reasoning}


async def summarize_clickhouse_result(
    user_request: str,
    executed_sql: str,
    reasoning: str,
    result_rows: list[dict[str, Any]],
    conversation_memory: str,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str],
) -> str:
    preview = json.dumps(result_rows[:10], ensure_ascii=False, indent=2)
    prompt = f"""
You are summarizing a ClickHouse query result for an end user.
Write the full answer in English and keep it concise.
The tone must be business-facing and functional, not technical.

Formatting rules:
- Return Markdown only.
- Write only the business result section body, with no title.
- Prefer either one short paragraph or 2 to 4 flat bullet points, depending on what is clearer.
- If the user explicitly asked for a table or tabular view, include a compact Markdown table in the answer body whenever the result is naturally tabular.
- Highlight the most important values, entities, dates, or thresholds with **bold**.
- Focus on what the result means for the user.
- Do not explain SQL generation, schema inspection, or tool usage.
- Do not include headings, SQL blocks, reasoning sections, or code fences.
- If the result is empty, say so clearly in one short sentence and give the most likely interpretation.

Context:
- User request: {user_request}
- Recent conversation memory:
{conversation_memory}
- Planner reasoning: {reasoning or "No extra reasoning provided."}
- Result rows preview:
{preview}
""".strip()

    return await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
    )


def _clean_clickhouse_summary_markdown(text: str) -> str:
    cleaned = (text or "").strip()
    if not cleaned:
        return "No clear business summary could be generated, but the query completed successfully."

    cleaned = re.sub(r"^\s*#{1,6}\s*(answer|result|summary|executive summary)\s*\n+", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"(?is)\n*#{1,6}\s*sql\b[\s\S]*$", "", cleaned)
    cleaned = re.sub(r"(?is)\n*#{1,6}\s*reasoning\b[\s\S]*$", "", cleaned)
    cleaned = re.sub(r"(?is)```sql[\s\S]*?```", "", cleaned)
    cleaned = re.sub(r"(?is)```[\s\S]*?```", "", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()

    return cleaned or "No clear business summary could be generated, but the query completed successfully."


def _build_clickhouse_sql_section(executed_sqls: list[str]) -> str:
    sql_statements = []
    seen = set()
    for statement in executed_sqls:
        normalized = str(statement or "").strip()
        if not normalized:
            continue
        if normalized in seen:
            continue
        seen.add(normalized)
        sql_statements.append(normalized)

    if not sql_statements:
        return ""

    if len(sql_statements) == 1:
        return f"## Executed SQL\n```sql\n{sql_statements[0]}\n```"

    blocks = []
    for index, statement in enumerate(sql_statements, start=1):
        blocks.append(f"### Query {index}\n```sql\n{statement}\n```")
    return "## Executed SQL\n\n" + "\n\n".join(blocks)


def build_clickhouse_response_markdown(
    result_markdown: str,
    executed_sqls: list[str],
    middle_sections: Optional[list[str]] = None,
) -> str:
    sections = [
        "## Result",
        _clean_clickhouse_summary_markdown(result_markdown),
    ]

    for section in middle_sections or []:
        normalized = str(section or "").strip()
        if normalized:
            sections.append(normalized)

    sql_section = _build_clickhouse_sql_section(executed_sqls)
    if sql_section:
        sections.append(sql_section)

    return "\n\n".join(sections)


def _default_data_analyst_state() -> dict[str, Any]:
    return {
        "stage": "idle",
        "pending_request": "",
        "available_tables": [],
        "selected_table": None,
        "table_schema": [],
        "clarification_prompt": "",
        "clarification_options": [],
        "analysis_goal": "",
        "analysis_plan": [],
        "success_criteria": [],
        "latest_review": "",
        "last_sqls": [],
        "last_result_meta": [],
        "last_result_rows": [],
        "final_answer": "",
        "last_error": "",
        "last_export_path": "",
        "knowledge_hits": [],
    }


def _normalize_data_analyst_state(payload: Optional[dict]) -> dict[str, Any]:
    state = _default_data_analyst_state()
    if not isinstance(payload, dict):
        return state

    stage = str(payload.get("stage") or "").strip()
    if stage in {"awaiting_table", "awaiting_row_intent", "ready"}:
        state["stage"] = stage

    state["pending_request"] = str(payload.get("pending_request") or payload.get("pendingRequest") or "").strip()
    available_tables = payload.get("available_tables") or payload.get("availableTables")
    if isinstance(available_tables, list):
        state["available_tables"] = [str(item).strip() for item in available_tables if str(item).strip()]
    selected_table = payload.get("selected_table") or payload.get("selectedTable")
    state["selected_table"] = (
        str(selected_table).strip() or None
        if selected_table is not None
        else None
    )
    table_schema = payload.get("table_schema") or payload.get("tableSchema") or payload.get("schema")
    if isinstance(table_schema, list):
        state["table_schema"] = [
            {
                "name": str(column.get("name") or "").strip(),
                "type": str(column.get("type") or "").strip(),
                "default_kind": str(column.get("default_kind") or column.get("defaultKind") or "").strip(),
                "default_expression": str(column.get("default_expression") or column.get("defaultExpression") or "").strip(),
            }
            for column in table_schema
            if isinstance(column, dict) and str(column.get("name") or "").strip()
        ]
    state["clarification_prompt"] = str(
        payload.get("clarification_prompt") or payload.get("clarificationPrompt") or ""
    ).strip()
    clarification_options = payload.get("clarification_options") or payload.get("clarificationOptions")
    if isinstance(clarification_options, list):
        state["clarification_options"] = [str(item).strip() for item in clarification_options if str(item).strip()]
    state["analysis_goal"] = str(payload.get("analysis_goal") or payload.get("analysisGoal") or "").strip()
    analysis_plan = payload.get("analysis_plan") or payload.get("analysisPlan")
    if isinstance(analysis_plan, list):
        state["analysis_plan"] = [str(item).strip() for item in analysis_plan if str(item).strip()]
    success_criteria = payload.get("success_criteria") or payload.get("successCriteria")
    if isinstance(success_criteria, list):
        state["success_criteria"] = [str(item).strip() for item in success_criteria if str(item).strip()]
    state["latest_review"] = str(payload.get("latest_review") or payload.get("latestReview") or "").strip()
    last_sqls = payload.get("last_sqls") or payload.get("lastSqls")
    if isinstance(last_sqls, list):
        state["last_sqls"] = [clean_sql_text(str(item)) for item in last_sqls if clean_sql_text(str(item))]
    last_result_meta = payload.get("last_result_meta") or payload.get("lastResultMeta")
    if isinstance(last_result_meta, list):
        state["last_result_meta"] = [
            {
                "name": str(item.get("name") or "").strip(),
                "type": str(item.get("type") or "").strip(),
            }
            for item in last_result_meta
            if isinstance(item, dict) and str(item.get("name") or "").strip()
        ]
    last_result_rows = payload.get("last_result_rows") or payload.get("lastResultRows")
    if isinstance(last_result_rows, list):
        state["last_result_rows"] = [row for row in last_result_rows if isinstance(row, dict)]
    state["final_answer"] = str(payload.get("final_answer") or payload.get("finalAnswer") or "").strip()
    state["last_error"] = str(payload.get("last_error") or payload.get("lastError") or "").strip()
    state["last_export_path"] = str(payload.get("last_export_path") or payload.get("lastExportPath") or "").strip()
    knowledge_hits = payload.get("knowledge_hits") or payload.get("knowledgeHits")
    if isinstance(knowledge_hits, list):
        state["knowledge_hits"] = [
            {
                "doc_name": str(item.get("doc_name") or item.get("docName") or "").strip(),
                "text": str(item.get("text") or "").strip(),
                "score": float(item.get("score") or 0.0),
            }
            for item in knowledge_hits
            if isinstance(item, dict)
        ]
    return state


def _data_analyst_state_needs_followup(state: dict[str, Any]) -> bool:
    return str(state.get("stage") or "").strip() in {"awaiting_table", "awaiting_row_intent"}


def _default_auto_ml_state() -> dict[str, Any]:
    return {
        "stage": "idle",
        "pending_request": "",
        "available_tables": [],
        "selected_table": None,
        "schema_info": [],
        "target_column": None,
        "row_filter": "",
        "sample_row_limit": 1000,
        "feature_columns": [],
        "clarification_prompt": "",
        "clarification_options": [],
        "comparison_rows": [],
        "problem_type": "",
        "recommended_model": "",
        "final_answer": "",
        "last_error": "",
    }


def _normalize_auto_ml_state(payload: Optional[dict]) -> dict[str, Any]:
    state = _default_auto_ml_state()
    if not isinstance(payload, dict):
        return state
    stage = str(payload.get("stage") or "").strip()
    if stage in {"awaiting_table", "awaiting_target", "ready"}:
        state["stage"] = stage
    state["pending_request"] = str(payload.get("pending_request") or payload.get("pendingRequest") or "").strip()
    available_tables = payload.get("available_tables") or payload.get("availableTables")
    if isinstance(available_tables, list):
        state["available_tables"] = [str(item).strip() for item in available_tables if str(item).strip()]
    selected_table = payload.get("selected_table") or payload.get("selectedTable")
    state["selected_table"] = str(selected_table).strip() or None if selected_table is not None else None
    schema_info = payload.get("schema_info") or payload.get("schemaInfo")
    if isinstance(schema_info, list):
        state["schema_info"] = [
            {
                "name": str(item.get("name") or "").strip(),
                "type": str(item.get("type") or "").strip(),
            }
            for item in schema_info
            if isinstance(item, dict) and str(item.get("name") or "").strip()
        ]
    target_column = payload.get("target_column") or payload.get("targetColumn")
    state["target_column"] = str(target_column).strip() or None if target_column is not None else None
    state["row_filter"] = str(payload.get("row_filter") or payload.get("rowFilter") or "").strip()
    raw_sample_row_limit = payload.get("sample_row_limit") if "sample_row_limit" in payload else payload.get("sampleRowLimit")
    if isinstance(raw_sample_row_limit, (int, float)):
        state["sample_row_limit"] = max(100, min(10_000, int(raw_sample_row_limit)))
    feature_columns = payload.get("feature_columns") or payload.get("featureColumns")
    if isinstance(feature_columns, list):
        state["feature_columns"] = [str(item).strip() for item in feature_columns if str(item).strip()]
    state["clarification_prompt"] = str(payload.get("clarification_prompt") or payload.get("clarificationPrompt") or "").strip()
    clarification_options = payload.get("clarification_options") or payload.get("clarificationOptions")
    if isinstance(clarification_options, list):
        state["clarification_options"] = [str(item).strip() for item in clarification_options if str(item).strip()]
    comparison_rows = payload.get("comparison_rows") or payload.get("comparisonRows")
    if isinstance(comparison_rows, list):
        state["comparison_rows"] = [row for row in comparison_rows if isinstance(row, dict)]
    raw_problem_type = str(payload.get("problem_type") or payload.get("problemType") or "").strip().lower()
    state["problem_type"] = raw_problem_type if raw_problem_type in {"classification", "regression"} else ""
    state["recommended_model"] = str(payload.get("recommended_model") or payload.get("recommendedModel") or "").strip()
    state["final_answer"] = str(payload.get("final_answer") or payload.get("finalAnswer") or "").strip()
    state["last_error"] = str(payload.get("last_error") or payload.get("lastError") or "").strip()
    return state


def _auto_ml_state_needs_followup(state: dict[str, Any]) -> bool:
    return str(state.get("stage") or "").strip() in {"awaiting_table", "awaiting_target"}


def _default_data_cleaner_state() -> dict[str, Any]:
    return {
        "stage": "idle",
        "pending_request": "",
        "available_tables": [],
        "selected_table": None,
        "schema_info": [],
        "row_filter": "",
        "clarification_prompt": "",
        "clarification_options": [],
        "findings": [],
        "correction_scripts": [],
        "final_answer": "",
        "last_error": "",
    }


def _normalize_data_cleaner_state(payload: Optional[dict]) -> dict[str, Any]:
    state = _default_data_cleaner_state()
    if not isinstance(payload, dict):
        return state
    stage = str(payload.get("stage") or "").strip()
    if stage in {"awaiting_table", "ready"}:
        state["stage"] = stage
    state["pending_request"] = str(payload.get("pending_request") or payload.get("pendingRequest") or "").strip()
    available_tables = payload.get("available_tables") or payload.get("availableTables")
    if isinstance(available_tables, list):
        state["available_tables"] = [str(item).strip() for item in available_tables if str(item).strip()]
    selected_table = payload.get("selected_table") or payload.get("selectedTable")
    state["selected_table"] = str(selected_table).strip() or None if selected_table is not None else None
    schema_info = payload.get("schema_info") or payload.get("schemaInfo")
    if isinstance(schema_info, list):
        state["schema_info"] = [
            {"name": str(item.get("name") or "").strip(), "type": str(item.get("type") or "").strip()}
            for item in schema_info
            if isinstance(item, dict) and str(item.get("name") or "").strip()
        ]
    state["row_filter"] = str(payload.get("row_filter") or payload.get("rowFilter") or "").strip()
    state["clarification_prompt"] = str(payload.get("clarification_prompt") or payload.get("clarificationPrompt") or "").strip()
    clarification_options = payload.get("clarification_options") or payload.get("clarificationOptions")
    if isinstance(clarification_options, list):
        state["clarification_options"] = [str(item).strip() for item in clarification_options if str(item).strip()]
    findings = payload.get("findings")
    if isinstance(findings, list):
        state["findings"] = [item for item in findings if isinstance(item, dict)]
    correction_scripts = payload.get("correction_scripts") or payload.get("correctionScripts")
    if isinstance(correction_scripts, list):
        state["correction_scripts"] = [item for item in correction_scripts if isinstance(item, dict)]
    state["final_answer"] = str(payload.get("final_answer") or payload.get("finalAnswer") or "").strip()
    state["last_error"] = str(payload.get("last_error") or payload.get("lastError") or "").strip()
    return state


def _data_cleaner_state_needs_followup(state: dict[str, Any]) -> bool:
    return str(state.get("stage") or "").strip() in {"awaiting_table"}


def _default_anonymizer_state() -> dict[str, Any]:
    return {
        "stage": "idle",
        "pending_request": "",
        "available_tables": [],
        "selected_table": None,
        "schema_info": [],
        "row_filter": "",
        "clarification_prompt": "",
        "clarification_options": [],
        "pii_findings": [],
        "masking_scripts": [],
        "final_answer": "",
        "last_error": "",
    }


def _normalize_anonymizer_state(payload: Optional[dict]) -> dict[str, Any]:
    state = _default_anonymizer_state()
    if not isinstance(payload, dict):
        return state
    stage = str(payload.get("stage") or "").strip()
    if stage in {"awaiting_table", "ready"}:
        state["stage"] = stage
    state["pending_request"] = str(payload.get("pending_request") or payload.get("pendingRequest") or "").strip()
    available_tables = payload.get("available_tables") or payload.get("availableTables")
    if isinstance(available_tables, list):
        state["available_tables"] = [str(item).strip() for item in available_tables if str(item).strip()]
    selected_table = payload.get("selected_table") or payload.get("selectedTable")
    state["selected_table"] = str(selected_table).strip() or None if selected_table is not None else None
    schema_info = payload.get("schema_info") or payload.get("schemaInfo")
    if isinstance(schema_info, list):
        state["schema_info"] = [
            {"name": str(item.get("name") or "").strip(), "type": str(item.get("type") or "").strip()}
            for item in schema_info
            if isinstance(item, dict) and str(item.get("name") or "").strip()
        ]
    state["row_filter"] = str(payload.get("row_filter") or payload.get("rowFilter") or "").strip()
    state["clarification_prompt"] = str(payload.get("clarification_prompt") or payload.get("clarificationPrompt") or "").strip()
    clarification_options = payload.get("clarification_options") or payload.get("clarificationOptions")
    if isinstance(clarification_options, list):
        state["clarification_options"] = [str(item).strip() for item in clarification_options if str(item).strip()]
    pii_findings = payload.get("pii_findings") or payload.get("piiFindings")
    if isinstance(pii_findings, list):
        state["pii_findings"] = [item for item in pii_findings if isinstance(item, dict)]
    masking_scripts = payload.get("masking_scripts") or payload.get("maskingScripts")
    if isinstance(masking_scripts, list):
        state["masking_scripts"] = [item for item in masking_scripts if isinstance(item, dict)]
    state["final_answer"] = str(payload.get("final_answer") or payload.get("finalAnswer") or "").strip()
    state["last_error"] = str(payload.get("last_error") or payload.get("lastError") or "").strip()
    return state


def _anonymizer_state_needs_followup(state: dict[str, Any]) -> bool:
    return str(state.get("stage") or "").strip() in {"awaiting_table"}


def _default_custom_agent_state() -> dict[str, Any]:
    return {
        "selected_agent_id": None,
        "final_answer": "",
        "last_error": "",
    }


def _normalize_custom_agent_state(payload: Optional[dict]) -> dict[str, Any]:
    state = _default_custom_agent_state()
    if not isinstance(payload, dict):
        return state
    selected_agent_id = payload.get("selected_agent_id") or payload.get("selectedAgentId")
    state["selected_agent_id"] = str(selected_agent_id).strip() or None if selected_agent_id is not None else None
    state["final_answer"] = str(payload.get("final_answer") or payload.get("finalAnswer") or "").strip()
    state["last_error"] = str(payload.get("last_error") or payload.get("lastError") or "").strip()
    return state


def _normalize_custom_agents_config(custom_agents: Any) -> list[dict[str, Any]]:
    if not isinstance(custom_agents, list):
        return []
    normalized: list[dict[str, Any]] = []
    for index, item in enumerate(custom_agents):
        if not isinstance(item, dict):
            continue
        normalized.append(
            {
                "id": str(item.get("id") or f"custom_agent_{index + 1}").strip() or f"custom_agent_{index + 1}",
                "title": str(item.get("title") or f"Custom Agent {index + 1}").strip() or f"Custom Agent {index + 1}",
                "description": str(item.get("description") or "").strip(),
                "pythonCode": str(item.get("pythonCode") or item.get("python_code") or "").strip(),
                "systemPrompt": str(item.get("systemPrompt") or item.get("system_prompt") or "").strip(),
                "managerRoutingHint": str(item.get("managerRoutingHint") or item.get("manager_routing_hint") or "").strip(),
                "status": str(item.get("status") or "draft").strip() or "draft",
                "statusMessage": str(item.get("statusMessage") or item.get("status_message") or "").strip(),
                "enabled": bool(item.get("enabled")),
                "badgeColor": str(item.get("badgeColor") or item.get("badge_color") or "zinc").strip() or "zinc",
            }
        )
    return normalized


def _enabled_custom_agents(custom_agents: Any) -> list[dict[str, Any]]:
    return [
        agent
        for agent in _normalize_custom_agents_config(custom_agents)
        if agent.get("enabled") and agent.get("status") == "ready"
    ]


def _find_custom_agent(custom_agents: Any, agent_id: str) -> Optional[dict[str, Any]]:
    for agent in _normalize_custom_agents_config(custom_agents):
        if agent.get("id") == agent_id:
            return agent
    return None


def _custom_agent_color_token(candidate: str) -> str:
    allowed = {"zinc", "slate", "gray", "blue", "sky", "cyan", "teal", "emerald", "green", "lime", "amber", "orange", "rose", "pink", "fuchsia", "violet", "indigo"}
    return candidate if candidate in allowed else "zinc"


async def _analyze_custom_agent_code(
    python_code: str,
    title: str,
    description: str,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str],
) -> dict[str, Any]:
    prompt = f"""
You are reviewing a Python agent implementation draft for RAGnarok.
Return JSON only with this exact shape:
{{
  "title": "short agent title",
  "description": "2-4 sentence user-facing description in English",
  "system_prompt": "complete system prompt in English for the agent",
  "manager_routing_hint": "short English hint so the Agent Manager knows when to use it",
  "badge_color": "zinc|slate|gray|blue|sky|cyan|teal|emerald|green|lime|amber|orange|rose|pink|fuchsia|violet|indigo",
  "status_message": "short implementation status in English"
}}

Rules:
- The agent must stay compatible with a local-LLM-only Python backend.
- Infer the functional scope from the code and preserve its specific behavior as much as possible.
- The description should be concise, clear, and ready to show in the chat intro card.
- The system prompt should tell the agent how to behave, what it should focus on, and how to answer.
- The manager routing hint must be short and operational.
- If the code is incomplete, still produce the best coherent agent profile and say so in status_message.

Provided title:
{title or "(empty)"}

Provided description:
{description or "(empty)"}

Python code:
```python
{python_code}
```
""".strip()
    raw = await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
        response_format="json",
    )
    parsed = extract_json_object(raw)
    if not parsed:
        raise ValueError("The local LLM did not return a valid custom-agent profile.")
    return {
        "title": str(parsed.get("title") or title or "Custom Agent").strip() or "Custom Agent",
        "description": str(parsed.get("description") or description or "").strip(),
        "system_prompt": str(parsed.get("system_prompt") or "").strip(),
        "managerRoutingHint": str(parsed.get("manager_routing_hint") or "").strip(),
        "badgeColor": _custom_agent_color_token(str(parsed.get("badge_color") or "zinc").strip().lower()),
        "statusMessage": str(parsed.get("status_message") or "The custom agent profile is ready.").strip() or "The custom agent profile is ready.",
    }


def _fallback_custom_agent_profile(python_code: str, title: str, description: str) -> dict[str, Any]:
    inferred_title = title.strip() or "Custom Agent"
    inferred_description = (
        description.strip()
        or "This custom agent was generated from Python code and keeps its behavior anchored in the uploaded implementation draft."
    )
    return {
        "title": inferred_title,
        "description": inferred_description,
        "system_prompt": (
            f"You are {inferred_title}. Reply in English. Follow the uploaded Python implementation draft as your specification. "
            "Stay factual, keep the workflow coherent with the code, and explain any limitation explicitly."
        ),
        "managerRoutingHint": f"Use this agent when the user explicitly asks for {inferred_title.lower()} capabilities.",
        "badgeColor": "zinc",
        "statusMessage": "The profile was generated with a fallback because the local LLM response was incomplete.",
    }


def _data_analyst_date_columns(schema: list[dict[str, Any]]) -> list[str]:
    return [
        str(column.get("name") or "").strip()
        for column in schema
        if str(column.get("name") or "").strip()
        and classify_clickhouse_column_type(str(column.get("type") or "")) == "date"
    ]


def _data_analyst_numeric_columns(schema: list[dict[str, Any]]) -> list[str]:
    return [
        str(column.get("name") or "").strip()
        for column in schema
        if str(column.get("name") or "").strip()
        and classify_clickhouse_column_type(str(column.get("type") or "")) == "numeric"
    ]


def _data_analyst_dimension_columns(schema: list[dict[str, Any]]) -> list[str]:
    return [
        str(column.get("name") or "").strip()
        for column in schema
        if str(column.get("name") or "").strip()
        and classify_clickhouse_column_type(str(column.get("type") or "")) in {"string", "other"}
    ]


def _data_analyst_schema_brief(schema: list[dict[str, Any]]) -> str:
    date_columns = _data_analyst_date_columns(schema)
    numeric_columns = _data_analyst_numeric_columns(schema)
    dimension_columns = _data_analyst_dimension_columns(schema)
    lines = [
        f"- Date columns: {', '.join(date_columns[:10]) if date_columns else 'none detected'}",
        f"- Numeric columns: {', '.join(numeric_columns[:12]) if numeric_columns else 'none detected'}",
        f"- Dimension-like columns: {', '.join(dimension_columns[:12]) if dimension_columns else 'none detected'}",
    ]
    return "\n".join(lines)


def _data_analyst_target_step_count(user_request: str, max_steps: int) -> int:
    normalized = normalize_intent_text(user_request)
    if not normalized:
        return min(max_steps, 4)

    score = 3
    if any(
        token in normalized
        for token in [
            "why",
            "root cause",
            "investigate",
            "analysis",
            "analyze",
            "deep dive",
            "deeper",
            "diagnose",
            "explain",
            "understand",
        ]
    ):
        score += 2
    if any(
        token in normalized
        for token in [
            "trend",
            "evolution",
            "over time",
            "compare",
            "versus",
            " vs ",
            "breakdown",
            "segment",
            "channel",
            "country",
            "product",
            "customer",
            "cohort",
            "week",
            "month",
            "quarter",
            "day",
            "hour",
            "top",
            "driver",
            "drivers",
            "drop",
            "decline",
            "growth",
            "anomaly",
        ]
    ):
        score += 2
    if any(token in normalized for token in [" and ", " plus ", " puis ", " then ", "after that"]):
        score += 1
    return max(3, min(max_steps, score))


def _data_analyst_requested_limit(user_request: str) -> Optional[int]:
    normalized = normalize_intent_text(user_request)
    if not normalized:
        return None

    patterns = [
        r"\b(?:top|first|limit|latest)\s+(\d{1,4})\b",
        r"\b(\d{1,4})\s*(rows?|lines?|lignes?)\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, normalized)
        if match:
            try:
                return max(1, min(int(match.group(1)), DATA_ANALYST_MAX_RESULT_ROWS))
            except Exception:
                return None
    return None


def _data_analyst_dynamic_query_limit(user_request: str, sql: str, base_limit: int) -> int:
    requested = _data_analyst_requested_limit(user_request)
    if requested is not None:
        return max(1, min(requested, base_limit, DATA_ANALYST_MAX_RESULT_ROWS))

    normalized_request = normalize_intent_text(user_request)
    lowered_sql = clean_sql_text(sql).lower()
    safe_cap = max(1, min(int(base_limit or 200), DATA_ANALYST_MAX_RESULT_ROWS))

    if _data_analyst_export_requested(user_request):
        return min(safe_cap, 500)

    if any(token in normalized_request for token in ["sample rows", "show rows", "preview rows", "raw rows", "detail rows"]):
        return min(safe_cap, 100)

    if " group by " in f" {lowered_sql} " or any(token in normalized_request for token in ["breakdown", "compare", "segment", "ranking", "top ", "top-", "leaderboard"]):
        return min(safe_cap, 100)

    if any(token in normalized_request for token in ["trend", "over time", "daily", "weekly", "monthly", "quarterly", "cohort", "retention"]):
        return min(safe_cap, 200)

    if re.search(r"\b(count|sum|avg|min|max|uniq|quantile)", lowered_sql) and " group by " not in f" {lowered_sql} ":
        return min(safe_cap, 10)

    return min(safe_cap, 50)


def _data_analyst_has_forbidden_date_functions(sql: str) -> bool:
    lowered = clean_sql_text(sql).lower()
    return any(token in lowered for token in DATA_ANALYST_FORBIDDEN_DATE_FUNCTIONS)


def _data_analyst_uses_date_literals_without_between(sql: str, schema: list[dict[str, Any]]) -> bool:
    cleaned = clean_sql_text(sql)
    lowered = cleaned.lower()
    if " between " in f" {lowered} ":
        return False
    if not re.search(r"'\d{4}-\d{2}-\d{2}", cleaned):
        return False
    for column_name in _data_analyst_date_columns(schema):
        pattern = rf"\b{re.escape(column_name.lower())}\b"
        if re.search(pattern, lowered):
            return True
    return False


def _data_analyst_sql_is_valid(sql: str, schema: Optional[list[dict[str, Any]]] = None) -> bool:
    cleaned = clean_sql_text(sql)
    if not is_safe_read_only_sql(cleaned):
        return False
    if re.search(r"\bselect\s+\*", cleaned, flags=re.IGNORECASE):
        return False
    if _data_analyst_has_forbidden_date_functions(cleaned):
        return False
    if schema and _data_analyst_uses_date_literals_without_between(cleaned, schema):
        return False
    return True


async def _data_analyst_define_goal(
    user_request: str,
    selected_table: str,
    schema: list[dict[str, Any]],
    conversation_memory: str,
    max_steps: int,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str],
) -> dict[str, Any]:
    schema_lines = "\n".join(
        f"- {column.get('name')}: {column.get('type')}"
        for column in schema[:120]
    )
    prompt = f"""
You are preparing a senior ClickHouse data analyst mission.
Define a compact investigation goal before the ReAct loop starts.

Return JSON only:
{{
  "analysis_goal": "one-sentence mission",
  "analysis_plan": ["step 1", "step 2", "step 3"],
  "success_criteria": ["criterion 1", "criterion 2", "criterion 3"]
}}

Rules:
- Keep everything in English.
- Be concrete and business-oriented.
- The plan should describe distinct analytical angles to test.
- The success criteria should explain when the analyst has enough evidence to finish.
- Assume at most {max_steps} credited analytical actions.

Selected table: {selected_table}
Schema:
{schema_lines}

Conversation memory:
{conversation_memory}

User request:
{user_request}
""".strip()

    raw = await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
        response_format="json",
    )
    parsed = extract_json_object(raw)
    return {
        "analysis_goal": str(parsed.get("analysis_goal") or parsed.get("goal") or "").strip(),
        "analysis_plan": [
            str(item).strip()
            for item in (parsed.get("analysis_plan") or parsed.get("plan") or [])
            if str(item).strip()
        ][:5],
        "success_criteria": [
            str(item).strip()
            for item in (parsed.get("success_criteria") or parsed.get("criteria") or [])
            if str(item).strip()
        ][:5],
    }


async def _data_analyst_review_progress(
    user_request: str,
    selected_table: str,
    analysis_goal: str,
    analysis_plan: list[str],
    success_criteria: list[str],
    step_log: list[dict[str, Any]],
    last_result_rows: list[dict[str, Any]],
    knowledge_hits: list[dict[str, Any]],
    max_steps: int,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str],
) -> dict[str, Any]:
    prompt = f"""
You are reviewing a ClickHouse Data Analyst investigation in progress.
Your job is to decide whether the current evidence is already sufficient and what gap remains.

Return JSON only:
{{
  "enough_evidence": true,
  "review_summary": "short executive review",
  "next_focus": "what the next step should test",
  "recommended_action": "query|search_knowledge|finish"
}}

Rules:
- Keep everything in English.
- Be strict: only say `enough_evidence=true` if the core business question is already answered well enough for leadership.
- Prefer `finish` only when the main goal and most success criteria are covered.
- If there is still a meaningful gap, explain the next focus precisely.

User request:
{user_request}

Selected table:
{selected_table}

Analysis goal:
{analysis_goal}

Investigation plan:
{json.dumps(analysis_plan, ensure_ascii=False, indent=2)}

Success criteria:
{json.dumps(success_criteria, ensure_ascii=False, indent=2)}

Completed steps:
{_data_analyst_steps_context(step_log)}

Latest query preview:
{_data_analyst_result_preview_text(last_result_rows)}

Knowledge hits:
{json.dumps([
    {"doc_name": item.get("doc_name"), "score": item.get("score")}
    for item in knowledge_hits[:DATA_ANALYST_MAX_KNOWLEDGE_RESULTS]
], ensure_ascii=False, indent=2)}

Remaining credited actions:
{max(0, max_steps - len(step_log))}
""".strip()

    raw = await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
        response_format="json",
    )
    parsed = extract_json_object(raw)
    return {
        "enough_evidence": bool(parsed.get("enough_evidence")),
        "review_summary": str(parsed.get("review_summary") or "").strip(),
        "next_focus": str(parsed.get("next_focus") or "").strip(),
        "recommended_action": str(parsed.get("recommended_action") or "").strip().lower(),
    }


def _data_analyst_step_label(action_type: str) -> str:
    labels = {
        "query": "Query",
        "search_knowledge": "Knowledge search",
        "export_csv": "CSV export",
        "finish": "Final answer",
    }
    return labels.get(action_type, action_type.replace("_", " ").title())


def _data_analyst_format_step(
    step_number: int,
    action_type: str,
    reasoning: str,
    result_summary: str,
    row_count: int,
    ok: bool,
    sql: str = "",
    retried: bool = False,
    suggested_path: str = "",
) -> dict[str, Any]:
    details_lines = []
    if reasoning:
        details_lines.append(f"Reasoning: {reasoning}")
    if result_summary:
        details_lines.append(f"Result: {result_summary}")
    details_lines.append(f"Rows: {row_count}")
    if suggested_path:
        details_lines.append(f"Suggested path: {suggested_path}")
    if retried:
        details_lines.append("Retry: The first SQL failed, so the agent repaired it automatically.")
    if sql:
        details_lines.append(f"SQL:\n{clean_sql_text(sql)}")
    return {
        "id": f"data-analyst-step-{step_number}",
        "title": f"Step {step_number} · {_data_analyst_step_label(action_type)}",
        "status": "success" if ok else "error",
        "details": "\n\n".join(line for line in details_lines if line),
        "step": step_number,
        "type": action_type,
        "reasoning": reasoning,
        "sql": clean_sql_text(sql) if sql else "",
        "result_summary": result_summary,
        "row_count": row_count,
        "ok": ok,
        "retried": retried,
        "suggested_path": suggested_path,
    }


def _data_analyst_steps_context(steps: list[dict[str, Any]]) -> str:
    if not steps:
        return "No prior analytical steps."
    compact = []
    for item in steps[-DATA_ANALYST_MAX_STEPS:]:
        compact.append(
            {
                "step": item.get("step"),
                "type": item.get("type"),
                "reasoning": _truncate_text_preview(str(item.get("reasoning") or ""), 220),
                "result_summary": _truncate_text_preview(str(item.get("result_summary") or ""), 420),
                "row_count": int(item.get("row_count") or 0),
                "ok": bool(item.get("ok")),
                "retried": bool(item.get("retried")),
                "suggested_path": str(item.get("suggested_path") or "").strip(),
            }
        )
    return json.dumps(compact, ensure_ascii=False, indent=2)


def _data_analyst_step_highlights(steps: list[dict[str, Any]], limit: int = 6) -> str:
    if not steps:
        return ""
    lines = []
    for item in steps[:limit]:
        lines.append(
            f"- **Step {item.get('step')} · {_data_analyst_step_label(str(item.get('type') or 'step'))}:** "
            f"{_truncate_text_preview(str(item.get('result_summary') or ''), 180) or 'No summary available.'}"
        )
    return "\n".join(lines)


def _data_analyst_needs_more_evidence(
    user_request: str,
    step_log: list[dict[str, Any]],
    max_steps: int,
    latest_review: str = "",
) -> bool:
    target_steps = _data_analyst_target_step_count(user_request, max_steps)
    successful_queries = sum(1 for step in step_log if step.get("ok") and step.get("type") == "query")
    successful_steps = sum(1 for step in step_log if step.get("ok"))
    lowered_review = str(latest_review or "").lower()
    if "enough evidence" in lowered_review or "ready to conclude" in lowered_review:
        return False
    if successful_queries == 0:
        return True
    if target_steps >= 5 and successful_queries < 2:
        return True
    if target_steps >= 7 and successful_queries < 3 and successful_steps < target_steps:
        return True
    return False


def _data_analyst_result_preview_text(rows: list[dict[str, Any]], max_rows: int = DATA_ANALYST_RESULT_PREVIEW_ROWS) -> str:
    preview_rows = rows[:max_rows]
    if not preview_rows:
        return "[]"
    return json.dumps(preview_rows, ensure_ascii=False, indent=2)


def _data_analyst_tabular_preview(
    meta: list[dict[str, Any]],
    rows: list[dict[str, Any]],
    max_rows: int = 8,
) -> str:
    preview_rows = rows[:max_rows]
    if not preview_rows:
        return ""
    headers = [
        str(item.get("name") or "").strip()
        for item in meta
        if str(item.get("name") or "").strip()
    ]
    if not headers:
        headers = [str(key) for key in preview_rows[0].keys()]
    if not headers:
        return ""

    def _cell(value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, (dict, list)):
            return json.dumps(value, ensure_ascii=False)
        return str(value)

    header_line = "| " + " | ".join(headers) + " |"
    divider_line = "| " + " | ".join(["---"] * len(headers)) + " |"
    body = [
        "| " + " | ".join(_cell(row.get(header)) for header in headers) + " |"
        for row in preview_rows
    ]
    table = "\n".join([header_line, divider_line, *body])
    if len(rows) > max_rows:
        table += f"\n\n_Showing the first {max_rows} rows out of {len(rows)}._"
    return table


def _data_analyst_export_requested(user_message: str) -> bool:
    normalized = normalize_intent_text(user_message)
    if not normalized:
        return False
    export_tokens = [
        "export csv",
        "csv export",
        "save csv",
        "write csv",
        "create csv",
        "download csv",
        "exporter csv",
        "exporter en csv",
        "creer un csv",
        "créer un csv",
        "sauvegarder en csv",
    ]
    return any(token in normalized for token in export_tokens) or bool(re.search(r"\bcsv\b", normalized))


def _data_analyst_suggest_export_path(user_message: str, table_name: str) -> str:
    detected = _extract_manager_export_path(user_message, "csv")
    if detected:
        return detected if detected.lower().endswith(".csv") else f"{detected}.csv"
    safe_table = re.sub(r"[^A-Za-z0-9._-]+", "-", table_name or "analysis").strip("-") or "analysis"
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
    return f"exports/{safe_table}-analysis-{timestamp}.csv"


def _data_analyst_unique_export_path(path: str) -> str:
    candidate_path = path if path.lower().endswith(".csv") else f"{path}.csv"
    target = _resolve_agent_path(candidate_path)
    if not target.exists():
        return candidate_path

    suffix = target.suffix or ".csv"
    stem = target.stem or "analysis"
    parent = target.parent
    counter = 2
    while True:
        next_target = parent / f"{stem}-{counter}{suffix}"
        if not next_target.exists():
            try:
                return str(next_target.relative_to(Path.cwd()))
            except Exception:
                return str(next_target)
        counter += 1


def _data_analyst_compact_query_summary(
    meta: list[dict[str, Any]],
    rows: list[dict[str, Any]],
) -> str:
    if not rows:
        return "The query returned no rows."
    headers = [
        str(item.get("name") or "").strip()
        for item in meta
        if str(item.get("name") or "").strip()
    ]
    if not headers:
        headers = [str(key) for key in rows[0].keys()]
    preview = rows[:2]
    preview_text = "; ".join(
        ", ".join(f"{key}={json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else value}" for key, value in row.items())
        for row in preview
    )
    return (
        f"Returned {len(rows)} row(s)"
        + (f" with columns {', '.join(headers[:8])}." if headers else ".")
        + (f" Preview: {preview_text}" if preview_text else "")
    )


def _data_analyst_confidence_score(
    step_log: list[dict[str, Any]],
    forced_finish: bool,
) -> tuple[int, str]:
    score = 72
    success_count = sum(1 for step in step_log if step.get("ok"))
    query_success_count = sum(1 for step in step_log if step.get("ok") and step.get("type") == "query")
    retry_count = sum(1 for step in step_log if step.get("retried"))
    error_count = sum(1 for step in step_log if not step.get("ok"))
    knowledge_count = sum(1 for step in step_log if step.get("type") == "search_knowledge" and step.get("ok"))
    export_count = sum(1 for step in step_log if step.get("type") == "export_csv" and step.get("ok"))

    score += min(12, query_success_count * 6)
    score += min(6, knowledge_count * 3)
    score += min(4, export_count * 2)
    score -= retry_count * 6
    score -= error_count * 8
    if forced_finish:
        score -= 10
    score = max(35, min(98, score))

    if score >= 88:
        reason = "Multiple successful evidence-gathering steps converged on a stable answer."
    elif score >= 74:
        reason = "The answer is supported by query results, with only limited uncertainty."
    elif score >= 60:
        reason = "The answer is useful, but there were retries or partial signals during the analysis."
    else:
        reason = "The answer is directional and should be double-checked with a narrower follow-up query."
    return score, reason


def _app_opensearch_config(app_config: dict) -> Optional[OSConfig]:
    url = str(app_config.get("elasticsearchUrl") or "").strip()
    index = str(app_config.get("elasticsearchIndex") or "").strip()
    if not url or not index:
        return None
    return OSConfig(
        url=url,
        index=index,
        username=str(app_config.get("elasticsearchUsername") or "").strip() or None,
        password=str(app_config.get("elasticsearchPassword") or "").strip() or None,
    )


def _app_embedding_config(app_config: dict) -> dict[str, Any]:
    disable_ssl_verification = _ssl_verification_disabled(app_config)
    return {
        "embedding_base_url": str(app_config.get("embeddingBaseUrl") or "http://localhost:11434/v1").strip() or "http://localhost:11434/v1",
        "embedding_api_key": str(app_config.get("embeddingApiKey") or "").strip() or None,
        "embedding_model": str(app_config.get("embeddingModel") or "nomic-embed-text").strip() or "nomic-embed-text",
        "embedding_verify_ssl": _effective_verify_ssl(bool(app_config.get("embeddingVerifySsl", True)), disable_ssl_verification),
        "knn_neighbors": max(1, min(int(app_config.get("knnNeighbors") or 50), 100)),
    }


async def _data_analyst_search_knowledge(
    query_text: str,
) -> dict[str, Any]:
    state = await read_db_state()
    app_config = state.get("config") or {}
    os_config = _app_opensearch_config(app_config)
    if not os_config:
        return {
            "ok": False,
            "summary": "Knowledge-base search is not configured in RAGnarok.",
            "results": [],
            "context": "",
        }

    embedding_config = _app_embedding_config(app_config)
    try:
        query_vector = await get_embedding(
            query_text,
            embedding_config["embedding_base_url"],
            embedding_config["embedding_model"],
            embedding_config["embedding_api_key"],
            verify_ssl=embedding_config["embedding_verify_ssl"],
        )
    except Exception as exc:
        return {
            "ok": False,
            "summary": f"Knowledge search could not compute an embedding: {exc}",
            "results": [],
            "context": "",
        }

    def _search() -> list[dict[str, Any]]:
        client = get_os_client(os_config.url, os_config.username, os_config.password)
        if not client.indices.exists(index=os_config.index):
            return []
        response = client.search(
            index=os_config.index,
            body={
                "size": embedding_config["knn_neighbors"],
                "query": {
                    "knn": {
                        "embedding": {
                            "vector": query_vector,
                            "k": embedding_config["knn_neighbors"],
                        }
                    }
                },
                "_source": ["doc_name", "text"],
            },
        )
        return [
            {
                "doc_name": hit.get("_source", {}).get("doc_name", "document"),
                "text": hit.get("_source", {}).get("text", ""),
                "score": float(hit.get("_score") or 0.0),
            }
            for hit in response.get("hits", {}).get("hits", [])
        ]

    try:
        results = await asyncio.to_thread(_search)
    except Exception as exc:
        return {
            "ok": False,
            "summary": f"Knowledge search failed: {exc}",
            "results": [],
            "context": "",
        }

    if not results:
        return {
            "ok": True,
            "summary": "No relevant knowledge-base documents were found.",
            "results": [],
            "context": "",
        }

    for item in results:
        item["score"] = item["score"] * 0.7 + keyword_score(query_text, item.get("text", "")) * 0.3
    ranked = sorted(results, key=lambda item: item.get("score", 0.0), reverse=True)[:DATA_ANALYST_MAX_KNOWLEDGE_RESULTS]
    summary = "Top knowledge signals: " + "; ".join(
        f"{item['doc_name']} (score {item['score']:.2f})"
        for item in ranked
    )
    context = "\n\n".join(
        f"[{item['doc_name']}]\n{_truncate_text_preview(str(item.get('text') or ''), 700)}"
        for item in ranked
    )
    return {
        "ok": True,
        "summary": summary,
        "results": ranked,
        "context": context,
    }


def _manager_rag_context_enabled(app_config: dict[str, Any]) -> bool:
    return bool((app_config or {}).get("managerUseRagFunctionalContext"))


async def _manager_functional_context_from_rag(
    query_text: str,
    app_config: dict[str, Any],
) -> dict[str, Any]:
    if not _manager_rag_context_enabled(app_config):
        return {
            "enabled": False,
            "used": False,
            "summary": "",
            "context": "",
            "results": [],
        }
    knowledge = await _data_analyst_search_knowledge(query_text)
    return {
        "enabled": True,
        "used": bool(knowledge.get("ok") and knowledge.get("results")),
        "summary": str(knowledge.get("summary") or ""),
        "context": str(knowledge.get("context") or ""),
        "results": knowledge.get("results") or [],
    }


async def plan_data_analyst_step(
    user_request: str,
    selected_table: str,
    schema: list[dict[str, Any]],
    conversation_memory: str,
    analysis_goal: str,
    analysis_plan: list[str],
    success_criteria: list[str],
    latest_review: str,
    step_log: list[dict[str, Any]],
    max_steps: int,
    used_steps: int,
    export_requested: bool,
    knowledge_enabled: bool,
    last_result_rows: list[dict[str, Any]],
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str],
    must_continue: bool = False,
) -> dict[str, Any]:
    # This node is the "thought" step of the ReAct loop: it does not execute
    # anything itself. It only decides the next evidence-gathering action based
    # on what is already known and what is still missing.
    target_steps = _data_analyst_target_step_count(user_request, max_steps)
    schema_lines = "\n".join(
        f"- {column.get('name')}: {column.get('type')}"
        for column in schema[:160]
    )
    prompt = f"""
You are an autonomous ClickHouse data analyst agent.
Your goal is to answer the user's question through a sequence of targeted analytical actions.
Operate as a true ReAct analyst:
- Observation: inspect the latest evidence and detect what is still unknown.
- Hypothesis: decide what business question or driver to test next.
- Action: choose exactly one next action that brings genuinely new evidence.
- Review: before finishing, check whether the goal and success criteria are truly covered.

You may use up to {max_steps} credited actions. Each action must add genuinely new evidence.
For this request, aim for about {target_steps} credited actions before finishing unless the evidence becomes decisive earlier.
If the request is investigative, comparative, diagnostic, segmented, or trend-oriented, do not finish after a single obvious query.

Return JSON only with this exact shape:
{{
  "action": "query|search_knowledge|export_csv|finish",
  "reasoning": "short English explanation",
  "sql": "SELECT ...",
  "knowledge_query": "optional short knowledge-base search query",
  "suggested_path": "optional csv path",
  "final_answer": "markdown answer when action=finish"
}}

Rules:
- Reply in English.
- Use only the ClickHouse table `{selected_table}` for SQL in this turn.
- NEVER use SELECT *.
- Keep SQL read-only and explicit.
- Add LIMIT when returning raw rows.
- You may use simple subqueries or derived tables when they materially improve the analysis.
- Subqueries are especially appropriate for top-N isolation, de-duplication, threshold filtering, ranking subsets, comparing segments, or building a clean aggregation before a final select.
- Avoid unnecessary complexity, but do not avoid subqueries when they are the clearest correct approach.
- Date handling is strict: for date filters, use `date_column BETWEEN 'YYYY-MM-DD' AND 'YYYY-MM-DD'`.
- Do not use complex date functions for filtering logic: no toStartOf..., year(), month(), day(), dateDiff(), addDays(), toDate(), toDateTime(), today(), or yesterday().
- `search_knowledge` is {"enabled" if knowledge_enabled else "disabled"}.
- `export_csv` is {"allowed because the user explicitly requested CSV export" if export_requested else "disabled unless the user explicitly requests CSV export"}.
- Use `export_csv` only after at least one successful query produced rows.
- When the evidence is sufficient, choose `finish`.
- Before finishing, prefer to cover at least two different analytical angles when the request asks for an explanation, root cause, trend, comparison, or driver analysis.
- If you have already used {used_steps} of {max_steps} actions, be economical.
- If {used_steps} >= {max_steps}, you MUST choose `finish`.
{"- You are not allowed to choose `finish` in this turn. Gather one more distinct piece of evidence." if must_continue else ""}

Database table: {selected_table}
Schema:
{schema_lines}

Schema profile:
{_data_analyst_schema_brief(schema)}

Analysis goal:
{analysis_goal or "Not defined yet."}

Investigation plan:
{json.dumps(analysis_plan, ensure_ascii=False, indent=2)}

Success criteria:
{json.dumps(success_criteria, ensure_ascii=False, indent=2)}

Recent conversation memory:
{conversation_memory}

Completed analytical steps:
{_data_analyst_steps_context(step_log)}

Latest review:
{latest_review or "No formal review yet."}

Latest query preview:
{_data_analyst_result_preview_text(last_result_rows)}

User request:
{user_request}
""".strip()

    raw = await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
        response_format="json",
    )
    parsed = extract_json_object(raw)
    action = str(parsed.get("action") or "").strip().lower()
    if must_continue and action == "finish":
        action = ""
    return {
        "action": action,
        "reasoning": str(parsed.get("reasoning") or "").strip(),
        "sql": clean_sql_text(str(parsed.get("sql") or "")),
        "knowledge_query": str(parsed.get("knowledge_query") or parsed.get("knowledgeQuery") or "").strip(),
        "suggested_path": str(parsed.get("suggested_path") or parsed.get("suggestedPath") or "").strip(),
        "final_answer": str(parsed.get("final_answer") or parsed.get("finalAnswer") or "").strip(),
    }


async def repair_data_analyst_sql(
    user_request: str,
    selected_table: str,
    schema: list[dict[str, Any]],
    failed_sql: str,
    error_feedback: str,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str],
    query_limit: int,
) -> dict[str, str]:
    schema_lines = "\n".join(
        f"- {column.get('name')}: {column.get('type')}"
        for column in schema[:160]
    )
    prompt = f"""
You are repairing a failed ClickHouse SQL query.
Generate a simpler safe replacement query and return JSON only:
{{
  "sql": "SELECT ...",
  "reasoning": "short English explanation"
}}

Rules:
- Use only table `{selected_table}`.
- NEVER use SELECT *.
- Keep the query read-only.
- Prefer COUNT, SUM, AVG, MIN, MAX, GROUP BY, ORDER BY, and BETWEEN.
- A simple subquery or derived table is allowed when it is needed to keep the logic correct.
- Avoid deep nesting and avoid advanced or brittle functions.
- Date handling is strict: if you filter on a date column, use `BETWEEN 'YYYY-MM-DD' AND 'YYYY-MM-DD'`.
- Do not use toStartOf..., year(), month(), day(), dateDiff(), addDays(), toDate(), toDateTime(), today(), or yesterday().
- Add LIMIT {max(1, min(int(query_limit or 200), 1000))} if raw rows are returned.

Schema:
{schema_lines}

User request:
{user_request}

Failed SQL:
{failed_sql}

Database error:
{error_feedback}
""".strip()

    raw = await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
        response_format="json",
    )
    parsed = extract_json_object(raw)
    sql = enforce_query_limit(clean_sql_text(str(parsed.get("sql") or "")), query_limit)
    return {
        "sql": sql,
        "reasoning": str(parsed.get("reasoning") or "").strip(),
    }


async def synthesize_data_analyst_answer(
    user_request: str,
    selected_table: str,
    analysis_goal: str,
    success_criteria: list[str],
    latest_review: str,
    conversation_memory: str,
    step_log: list[dict[str, Any]],
    last_result_meta: list[dict[str, Any]],
    last_result_rows: list[dict[str, Any]],
    knowledge_hits: list[dict[str, Any]],
    forced_finish: bool,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str],
) -> str:
    # First-pass executive synthesis. A second polishing pass can further
    # tighten tone and clarity, but this prompt already asks for a COMEX-ready
    # structure grounded in the collected evidence.
    prompt = f"""
You are finishing a complex ClickHouse analysis for an executive audience.
Write the final answer in English as a developed business-facing Markdown analysis that reads like a senior data analyst briefing the COMEX.

Required structure:
- `## Executive Summary`
- `## What The Data Says`
- `## Business Implications`
- `## Recommended Actions`
- `## Caveats`

Rules:
- Keep the tone functional, precise, concise, and executive-friendly.
- Prefer short paragraphs with a few targeted bullets only when they genuinely improve readability.
- If the user explicitly asked for a table or tabular comparison, include one compact Markdown table in `## What The Data Says` when the result is naturally tabular.
- Highlight the most important values with **bold**.
- Explicitly explain what the evidence suggests, not only what was queried.
- Reference concrete metrics, rankings, percentages, counts, deltas, segments, or date ranges whenever available.
- Translate the evidence into business consequences, opportunities, or risks.
- Distinguish clearly between firm findings, likely drivers, and open questions.
- Do not include SQL sections or code fences here.
- If the evidence is partial, say so honestly.
- Avoid generic filler. Every section should add analytical value.
- Sound like a senior analyst: measured, specific, and decision-oriented.

Selected table: {selected_table}
Forced finish: {"yes" if forced_finish else "no"}

Analysis goal:
{analysis_goal or "Not defined explicitly."}

Success criteria:
{json.dumps(success_criteria, ensure_ascii=False, indent=2)}

Latest review:
{latest_review or "No formal review was captured."}

Recent conversation memory:
{conversation_memory}

Analytical steps:
{_data_analyst_steps_context(step_log)}

Latest query preview:
{_data_analyst_result_preview_text(last_result_rows)}

Knowledge signals:
{json.dumps([
    {
        "doc_name": item.get("doc_name"),
        "score": item.get("score"),
        "text": _truncate_text_preview(str(item.get("text") or ""), 300),
    }
    for item in knowledge_hits[:DATA_ANALYST_MAX_KNOWLEDGE_RESULTS]
], ensure_ascii=False, indent=2)}

User request:
{user_request}
""".strip()

    return await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
    )


async def polish_data_analyst_answer(
    user_request: str,
    selected_table: str,
    analysis_goal: str,
    success_criteria: list[str],
    latest_review: str,
    draft_answer: str,
    conversation_memory: str,
    step_log: list[dict[str, Any]],
    last_result_rows: list[dict[str, Any]],
    knowledge_hits: list[dict[str, Any]],
    forced_finish: bool,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str],
) -> str:
    # Final editorial pass: keep the facts, improve the executive narrative.
    # This lets the agent stay evidence-driven during the loop, then switch to
    # a cleaner senior-analyst writing style right before returning to the UI.
    draft = str(draft_answer or "").strip()
    if not draft:
        return draft

    prompt = f"""
You are polishing the final answer of a ClickHouse Data Analyst agent.
Rewrite the draft into a stronger executive memo for a COMEX audience.

Required structure:
- `## Executive Summary`
- `## What The Data Says`
- `## Business Implications`
- `## Recommended Actions`
- `## Caveats`

Rules:
- Keep the answer in English.
- Make it more developed, more functional, and more decision-oriented than the draft.
- Use precise statements grounded in the executed evidence.
- If the user explicitly asked for a table or tabular comparison, keep one compact Markdown table in `## What The Data Says` when the evidence is naturally tabular.
- Mention concrete figures, segments, comparisons, rankings, or time ranges whenever available.
- Prefer short paragraphs. Use bullets only for tightly scoped recommendations or caveats.
- Do not add SQL, code fences, or implementation details.
- Do not invent facts that are not supported by the evidence.
- Distinguish facts, likely explanations, and residual uncertainty.
- The visible part should read like a senior data analyst's written debrief to leadership.

Selected table: {selected_table}
Forced finish: {"yes" if forced_finish else "no"}

Analysis goal:
{analysis_goal or "Not defined explicitly."}

Success criteria:
{json.dumps(success_criteria, ensure_ascii=False, indent=2)}

Latest review:
{latest_review or "No formal review was captured."}

Recent conversation memory:
{conversation_memory}

Analytical steps:
{_data_analyst_steps_context(step_log)}

Latest query preview:
{_data_analyst_result_preview_text(last_result_rows)}

Knowledge signals:
{json.dumps([
    {
        "doc_name": item.get("doc_name"),
        "score": item.get("score"),
        "text": _truncate_text_preview(str(item.get("text") or ""), 260),
    }
    for item in knowledge_hits[:DATA_ANALYST_MAX_KNOWLEDGE_RESULTS]
], ensure_ascii=False, indent=2)}

User request:
{user_request}

Draft answer to improve:
{draft}
""".strip()

    return await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
    )


def build_data_analyst_response_markdown(
    final_body: str,
    executed_sqls: list[str],
    last_result_meta: list[dict[str, Any]],
    last_result_rows: list[dict[str, Any]],
    knowledge_hits: list[dict[str, Any]],
    last_export_path: str,
    confidence_score: int,
    confidence_reason: str,
    step_log: Optional[list[dict[str, Any]]] = None,
    force_visible_table: bool = False,
) -> str:
    # Keep the user-facing memo first, then append technical evidence blocks.
    # The chat UI collapses most of these sections automatically, so analysts
    # still get traceability without overwhelming the primary narrative.
    sections = [str(final_body or "").strip() or "## Executive Summary\nThe analysis completed, but no final narrative could be generated."]
    if step_log:
        step_highlights = _data_analyst_step_highlights(step_log)
        if step_highlights:
            sections.append("## Evidence Trail\n" + step_highlights)
    preview_table = _data_analyst_tabular_preview(last_result_meta, last_result_rows)
    if preview_table and force_visible_table:
        sections.append("## Result Table\n" + preview_table)
    if preview_table:
        sections.append("## Data Preview\n" + preview_table)
    if knowledge_hits:
        knowledge_lines = "\n".join(
            f"- **{item.get('doc_name') or 'Document'}**: {_truncate_text_preview(str(item.get('text') or ''), 180)}"
            for item in knowledge_hits[:DATA_ANALYST_MAX_KNOWLEDGE_RESULTS]
        )
        sections.append("## Knowledge Signals\n" + knowledge_lines)
    if last_export_path:
        sections.append(
            "## CSV Export\n"
            f"The latest dataset was exported to `{last_export_path}`."
        )
    sections.append(
        "## Confidence\n"
        f"**Score:** {confidence_score}/100\n\n{confidence_reason}"
    )
    sql_section = _build_clickhouse_sql_section(executed_sqls)
    if sql_section:
        sections.append(sql_section)
    return "\n\n".join(section for section in sections if section.strip())


AUTO_ML_TABLE_OPTION_LIMIT = 8
AUTO_ML_TARGET_OPTION_LIMIT = 8
AUTO_ML_SAMPLE_ROWS = 1000
AUTO_ML_MAX_FEATURE_COLUMNS = 24

def _records_preview_markdown(rows: list[dict[str, Any]], limit: int = 5) -> str:
    if not rows:
        return "No sample rows available."
    preview_rows = rows[:limit]
    columns = list(preview_rows[0].keys())[:8]
    if not columns:
        return "No sample rows available."
    header = "| " + " | ".join(columns) + " |"
    divider = "| " + " | ".join(["---"] * len(columns)) + " |"
    body = []
    for row in preview_rows:
        body.append("| " + " | ".join(_markdown_table_cell(row.get(col)) for col in columns) + " |")
    return "\n".join([header, divider, *body])


async def _fetch_clickhouse_sample_rows(
    clickhouse: ClickHouseConfig,
    table_name: str,
    columns: list[str],
    row_limit: int,
    row_filter: str = "",
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], str]:
    selected_columns = columns[:AUTO_ML_MAX_FEATURE_COLUMNS]
    select_sql = ", ".join(quote_clickhouse_identifier(column) for column in selected_columns) if selected_columns else "*"
    normalized_row_filter = _normalize_auto_ml_row_filter(row_filter)
    sql = (
        f"SELECT {select_sql}\n"
        f"FROM {quote_clickhouse_identifier(table_name)}\n"
        + (f"WHERE ({normalized_row_filter})\n" if normalized_row_filter else "")
        + f"LIMIT {max(1, min(int(row_limit), 10_000))}"
    )
    result = await execute_clickhouse_sql(clickhouse, sql, max_result_rows_override=max(1, min(int(row_limit), 10_000)))
    meta = result.get("meta", []) if isinstance(result.get("meta"), list) else []
    rows = result.get("data", []) if isinstance(result.get("data"), list) else []
    normalized_meta = [
        {
            "name": str(item.get("name") or "").strip(),
            "type": str(item.get("type") or "").strip(),
        }
        for item in meta
        if isinstance(item, dict) and str(item.get("name") or "").strip()
    ]
    normalized_rows = [row for row in rows if isinstance(row, dict)]
    return normalized_meta, normalized_rows, sql


def _infer_target_column(user_request: str, schema: list[dict[str, Any]]) -> Optional[str]:
    normalized = normalize_intent_text(user_request)
    schema_names = [str(column.get("name") or "").strip() for column in schema if str(column.get("name") or "").strip()]
    direct = resolve_user_choice(normalized, schema_names)
    if direct:
        return direct
    target_match = re.search(r"(?:predict|forecast|estimate|model|target|label|classifier|regression|predire|prédire|cible)\s+([a-zA-Z0-9_]+)", normalized)
    if target_match:
        hinted = target_match.group(1)
        for column_name in schema_names:
            if normalize_choice(column_name).lower() == normalize_choice(hinted).lower():
                return column_name
    return None


def _auto_ml_target_candidates(schema: list[dict[str, Any]]) -> list[str]:
    preferred = []
    for column in schema:
        name = str(column.get("name") or "").strip()
        if not name:
            continue
        category = classify_clickhouse_column_type(str(column.get("type") or ""))
        if category in {"numeric", "string", "date"}:
            preferred.append(name)
    return preferred


def _normalize_ml_value(value: Any) -> Any:
    if value is None:
        return None
    if isinstance(value, bool):
        return int(value)
    if isinstance(value, (int, float)):
        return value
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return None
        try:
            if re.fullmatch(r"-?\d+", text):
                return int(text)
            if re.fullmatch(r"-?\d+(?:\.\d+)?", text):
                return float(text)
        except Exception:
            pass
        return text
    return str(value)


def _auto_ml_problem_type(target_values: list[Any], target_type: str) -> str:
    non_null = [value for value in target_values if value is not None]
    if not non_null:
        return "classification"
    unique_count = len({str(value) for value in non_null})
    if target_type == "numeric" and unique_count > 12:
        return "regression"
    if unique_count > 20 and all(isinstance(value, (int, float)) for value in non_null):
        return "regression"
    return "classification"


def _prepare_automl_dataset(
    rows: list[dict[str, Any]],
    schema: list[dict[str, Any]],
    target_column: str,
) -> tuple[list[dict[str, Any]], list[Any], list[str], str]:
    schema_map = {
        str(column.get("name") or "").strip(): classify_clickhouse_column_type(str(column.get("type") or ""))
        for column in schema
        if str(column.get("name") or "").strip()
    }
    feature_columns = [
        name
        for name in schema_map.keys()
        if name != target_column
    ][:AUTO_ML_MAX_FEATURE_COLUMNS]
    records: list[dict[str, Any]] = []
    targets: list[Any] = []
    for row in rows:
        if target_column not in row:
            continue
        normalized_target = _normalize_ml_value(row.get(target_column))
        if normalized_target is None:
            continue
        record: dict[str, Any] = {}
        for column in feature_columns:
            value = _normalize_ml_value(row.get(column))
            category = schema_map.get(column) or "other"
            if category == "numeric":
                record[column] = value if isinstance(value, (int, float)) else None
            else:
                record[column] = "__missing__" if value is None else str(value)
        records.append(record)
        targets.append(normalized_target)

    problem_type = _auto_ml_problem_type(targets, schema_map.get(target_column) or "other")
    if problem_type == "classification":
        targets = [str(value) for value in targets]
    else:
        targets = [float(value) for value in targets if isinstance(value, (int, float))]
        records = records[:len(targets)]
    return records, targets, feature_columns, problem_type


def _prepare_vectorized_features(records: list[dict[str, Any]]) -> Any:
    if not HAVE_SKLEARN:
        raise RuntimeError("scikit-learn is not installed on this server.")
    vectorizer = DictVectorizer(sparse=False)
    matrix = vectorizer.fit_transform(records)
    imputer = SimpleImputer(strategy="constant", fill_value=0)
    return imputer.fit_transform(matrix)


def _run_automl_benchmark(
    records: list[dict[str, Any]],
    targets: list[Any],
    problem_type: str,
) -> tuple[list[dict[str, Any]], str]:
    if not HAVE_SKLEARN:
        raise RuntimeError("scikit-learn is not installed on this server.")
    if len(records) < 60 or len(targets) < 60:
        raise RuntimeError("Auto-ML needs at least 60 usable rows after cleaning to benchmark models reliably.")

    features = _prepare_vectorized_features(records)
    stratify = targets if problem_type == "classification" and len(set(targets)) > 1 else None
    X_train, X_test, y_train, y_test = train_test_split(
        features,
        targets,
        test_size=0.25,
        random_state=42,
        stratify=stratify,
    )

    candidate_models: list[tuple[str, Any]] = []
    if problem_type == "classification":
        candidate_models = [
            ("Logistic Regression", LogisticRegression(max_iter=1000)),
            ("Random Forest", RandomForestClassifier(n_estimators=180, random_state=42)),
        ]
        if HAVE_XGBOOST:
            candidate_models.append(("XGBoost", XGBClassifier(n_estimators=160, max_depth=6, learning_rate=0.08, subsample=0.9, colsample_bytree=0.9, eval_metric="logloss", random_state=42)))
        else:
            candidate_models.append(("Gradient Boosting", GradientBoostingClassifier(random_state=42)))
    else:
        candidate_models = [
            ("Linear Regression", LinearRegression()),
            ("Random Forest", RandomForestRegressor(n_estimators=180, random_state=42)),
        ]
        if HAVE_XGBOOST:
            candidate_models.append(("XGBoost", XGBRegressor(n_estimators=160, max_depth=6, learning_rate=0.08, subsample=0.9, colsample_bytree=0.9, random_state=42)))
        else:
            candidate_models.append(("Gradient Boosting", GradientBoostingRegressor(random_state=42)))

    comparison_rows: list[dict[str, Any]] = []
    for model_name, model in candidate_models:
        try:
            model.fit(X_train, y_train)
            predictions = model.predict(X_test)
            if problem_type == "classification":
                row = {
                    "Model": model_name,
                    "Accuracy": round(float(accuracy_score(y_test, predictions)), 4),
                    "Precision": round(float(precision_score(y_test, predictions, average="weighted", zero_division=0)), 4),
                    "Recall": round(float(recall_score(y_test, predictions, average="weighted", zero_division=0)), 4),
                    "F1": round(float(f1_score(y_test, predictions, average="weighted", zero_division=0)), 4),
                }
            else:
                row = {
                    "Model": model_name,
                    "RMSE": round(float(math.sqrt(mean_squared_error(y_test, predictions))), 4),
                    "MAE": round(float(mean_absolute_error(y_test, predictions)), 4),
                    "R2": round(float(r2_score(y_test, predictions)), 4),
                }
            comparison_rows.append(row)
        except Exception as exc:
            comparison_rows.append({"Model": model_name, "Status": f"Failed: {exc}"})

    scored_rows = [row for row in comparison_rows if any(metric in row for metric in ["Accuracy", "RMSE"])]
    if not scored_rows:
        raise RuntimeError("All candidate models failed during training.")
    if problem_type == "classification":
        best = max(scored_rows, key=lambda row: (row.get("F1") or 0, row.get("Accuracy") or 0))
    else:
        best = max(scored_rows, key=lambda row: (row.get("R2") or -999, -float(row.get("RMSE") or 999999)))
    return comparison_rows, str(best.get("Model") or "")


def _markdown_table_from_rows(rows: list[dict[str, Any]]) -> str:
    if not rows:
        return ""
    columns: list[str] = []
    for row in rows:
        for key in row.keys():
            if key not in columns:
                columns.append(key)
    header = "| " + " | ".join(columns) + " |"
    divider = "| " + " | ".join(["---"] * len(columns)) + " |"
    body = []
    for row in rows:
        body.append("| " + " | ".join(_markdown_table_cell(row.get(column)) for column in columns) + " |")
    return "\n".join([header, divider, *body])


async def _auto_ml_narrative(
    user_request: str,
    selected_table: str,
    target_column: str,
    row_filter: str,
    sample_row_limit: int,
    problem_type: str,
    feature_columns: list[str],
    comparison_rows: list[dict[str, Any]],
    conversation_memory: str,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str],
) -> str:
    prompt = f"""
You are presenting an Auto-ML benchmark to a business stakeholder.
Write a concise but practical English Markdown summary.

Rules:
- Use the following sections:
  - `## Executive Summary`
  - `## Benchmark Scope`
  - `## Model Comparison`
  - `## Recommendation`
- Mention the selected target, the problem type, and the winning model.
- Mention the applied row scope and training sample limit in plain business language.
- Explain the decision in functional language.
- Do not include SQL or code.

User request:
{user_request}

Selected table:
{selected_table}

Target column:
{target_column}

Row filter:
{row_filter or "Full table"}

Sample row limit:
{sample_row_limit}

Problem type:
{problem_type}

Feature columns:
{json.dumps(feature_columns, ensure_ascii=False)}

Comparison rows:
{json.dumps(comparison_rows, ensure_ascii=False, indent=2)}

Conversation memory:
{conversation_memory}
""".strip()
    return await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
    )


DATA_CLEANER_SAMPLE_ROWS = 500
ANONYMIZER_SAMPLE_ROWS = 300


def _schema_column_names(schema: list[dict[str, Any]]) -> list[str]:
    return [str(column.get("name") or "").strip() for column in schema if str(column.get("name") or "").strip()]


def _pick_candidate_key_columns(schema: list[dict[str, Any]]) -> list[str]:
    candidates = _schema_column_names(schema)
    preferred_patterns = [
        r"(^|_)(id|uuid|key)$",
        r"(^|_)(user_id|customer_id|order_id|transaction_id|invoice_id|email)$",
    ]
    matched: list[str] = []
    for column in candidates:
        normalized = column.lower()
        if any(re.search(pattern, normalized) for pattern in preferred_patterns):
            matched.append(column)
    return matched[:3]


def _detect_mixed_date_formats(rows: list[dict[str, Any]], schema: list[dict[str, Any]]) -> list[dict[str, str]]:
    findings: list[dict[str, str]] = []
    string_columns = [
        str(column.get("name") or "").strip()
        for column in schema
        if classify_clickhouse_column_type(str(column.get("type") or "")) == "string"
    ]
    for column_name in string_columns[:12]:
        format_hits: set[str] = set()
        for row in rows:
            value = row.get(column_name)
            if value is None:
                continue
            text = str(value).strip()
            if not text:
                continue
            if re.fullmatch(r"\d{4}-\d{2}-\d{2}", text):
                format_hits.add("ISO")
            elif re.fullmatch(r"\d{2}/\d{2}/\d{4}", text):
                left, right = text.split("/")[:2]
                try:
                    if int(left) > 12:
                        format_hits.add("FR")
                    elif int(right) > 12:
                        format_hits.add("US")
                    else:
                        format_hits.add("Ambiguous slash date")
                except Exception:
                    pass
        if len(format_hits) > 1 or "Ambiguous slash date" in format_hits:
            findings.append(
                {
                    "level": "warning",
                    "title": f"Mixed date formats in `{column_name}`",
                    "detail": f"Sample values suggest multiple date formats: {', '.join(sorted(format_hits))}.",
                }
            )
    return findings


async def _data_cleaner_profile(
    clickhouse: ClickHouseConfig,
    table_name: str,
    schema: list[dict[str, Any]],
    row_filter: str = "",
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]], list[dict[str, str]], str]:
    column_names = _schema_column_names(schema)
    sampled_columns = column_names[: min(20, len(column_names))]
    normalized_row_filter = _normalize_auto_ml_row_filter(row_filter)
    _, sample_rows, sample_sql = await _fetch_clickhouse_sample_rows(
        clickhouse,
        table_name,
        sampled_columns,
        DATA_CLEANER_SAMPLE_ROWS,
        normalized_row_filter,
    )

    findings: list[dict[str, Any]] = []
    scripts: list[dict[str, Any]] = []
    steps: list[dict[str, Any]] = []

    total_sql = (
        f"SELECT count() AS row_count FROM {quote_clickhouse_identifier(table_name)}"
        + (f" WHERE ({normalized_row_filter})" if normalized_row_filter else "")
    )
    total_payload = await execute_clickhouse_sql(clickhouse, total_sql)
    total_rows = int(_first_row(total_payload).get("row_count") or 0)

    key_columns = _pick_candidate_key_columns(schema)
    if key_columns:
        tuple_sql = ", ".join(quote_clickhouse_identifier(column) for column in key_columns)
        duplicate_sql = (
            f"SELECT count() AS row_count, countDistinct(tuple({tuple_sql})) AS distinct_rows "
            f"FROM {quote_clickhouse_identifier(table_name)}"
            + (f" WHERE ({normalized_row_filter})" if normalized_row_filter else "")
        )
        duplicate_payload = await execute_clickhouse_sql(clickhouse, duplicate_sql)
        first = _first_row(duplicate_payload)
        duplicate_rows = max(0, int(first.get("row_count") or 0) - int(first.get("distinct_rows") or 0))
        if duplicate_rows > 0:
            findings.append({
                "level": "critical" if duplicate_rows > max(1, total_rows * 0.01) else "warning",
                "title": "Potential duplicates detected",
                "detail": f"`{duplicate_rows}` rows appear duplicated based on `{', '.join(key_columns)}`.",
            })
            partition_by = ", ".join(quote_clickhouse_identifier(column) for column in key_columns)
            scripts.append({
                "title": "Deduplicate on the candidate business key",
                "sql": (
                    f"SELECT *\nFROM (\n"
                    f"  SELECT *, row_number() OVER (PARTITION BY {partition_by} ORDER BY {partition_by}) AS _rn\n"
                    f"  FROM {quote_clickhouse_identifier(table_name)}\n"
                    f")\nWHERE _rn = 1"
                ),
            })

    profile_columns = sampled_columns[:12]
    if profile_columns:
        profile_exprs: list[str] = []
        for column_name in profile_columns:
            identifier = quote_clickhouse_identifier(column_name)
            profile_exprs.append(f"countIf(isNull(toNullable({identifier}))) AS {quote_clickhouse_identifier(f'{column_name}__nulls')}")
            if any(col.get("name") == column_name and classify_clickhouse_column_type(str(col.get("type") or "")) == "string" for col in schema):
                profile_exprs.append(f"countIf(trim(BOTH ' ' FROM toString({identifier})) = '') AS {quote_clickhouse_identifier(f'{column_name}__empties')}")
        profile_sql = (
            f"SELECT\n  " + ",\n  ".join(profile_exprs) + f"\nFROM {quote_clickhouse_identifier(table_name)}"
            + (f"\nWHERE ({normalized_row_filter})" if normalized_row_filter else "")
        )
        profile_payload = await execute_clickhouse_sql(clickhouse, profile_sql)
        profile_row = _first_row(profile_payload)
        for column_name in profile_columns:
            null_count = int(profile_row.get(f"{column_name}__nulls") or 0)
            empty_count = int(profile_row.get(f"{column_name}__empties") or 0)
            if null_count > 0 or empty_count > 0:
                findings.append({
                    "level": "critical" if null_count > max(5, total_rows * 0.2) else "warning",
                    "title": f"Missing values in `{column_name}`",
                    "detail": f"`{null_count}` nulls and `{empty_count}` empty strings were detected in `{column_name}`.",
                })
                identifier = quote_clickhouse_identifier(column_name)
                scripts.append({
                    "title": f"Standardize missing values in `{column_name}`",
                    "sql": f"SELECT nullIf(trim(BOTH ' ' FROM toString({identifier})), '') AS {identifier}\nFROM {quote_clickhouse_identifier(table_name)}",
                })

    findings.extend(_detect_mixed_date_formats(sample_rows, schema))
    for finding in findings:
        if "Mixed date formats" in str(finding.get("title") or ""):
            column_name = re.findall(r"`([^`]+)`", str(finding.get("title") or ""))
            if column_name:
                identifier = quote_clickhouse_identifier(column_name[0])
                scripts.append({
                    "title": f"Normalize `{column_name[0]}` into ISO dates",
                    "sql": (
                        f"SELECT\n"
                        f"  multiIf(\n"
                        f"    match(toString({identifier}), '^\\\\d{{2}}/\\\\d{{2}}/\\\\d{{4}}$'), parseDateTimeBestEffortOrNull(toString({identifier})),\n"
                        f"    match(toString({identifier}), '^\\\\d{{4}}-\\\\d{{2}}-\\\\d{{2}}$'), parseDateTimeBestEffortOrNull(toString({identifier})),\n"
                        f"    NULL\n"
                        f"  ) AS {identifier}\n"
                        f"FROM {quote_clickhouse_identifier(table_name)}"
                    ),
                })

    if not findings:
        findings.append({
            "level": "info",
            "title": "No critical data-quality issue detected in the sampled audit",
            "detail": "The table does not show obvious duplicates, missing-value spikes, or mixed date formats in the inspected scope.",
        })

    return findings, scripts[:8], sample_rows, [{"sql": sample_sql}, {"sql": total_sql}], _records_preview_markdown(sample_rows)


EMAIL_PATTERN = re.compile(r"^[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}$", re.IGNORECASE)
PHONE_PATTERN = re.compile(r"^\+?[0-9][0-9()\-\s]{6,}$")
IP_PATTERN = re.compile(r"^(?:\d{1,3}\.){3}\d{1,3}$")


def _anonymizer_candidate_by_name(column_name: str) -> Optional[tuple[str, str, str]]:
    lowered = column_name.lower()
    rules = [
        ("email", "email", "high", "Hash or tokenize email addresses before any external sharing."),
        ("phone", "phone number", "high", "Mask or hash phone numbers and keep the raw value in a restricted zone only."),
        ("mobile", "phone number", "high", "Mask or hash phone numbers and keep the raw value in a restricted zone only."),
        ("name", "person name", "medium", "Mask names or keep initials only when identity is not required."),
        ("first_name", "person name", "medium", "Mask names or keep initials only when identity is not required."),
        ("last_name", "person name", "medium", "Mask names or keep initials only when identity is not required."),
        ("address", "postal address", "high", "Remove house-level address data or replace it with a generalized location."),
        ("street", "postal address", "high", "Remove house-level address data or replace it with a generalized location."),
        ("ssn", "government identifier", "high", "Hash or tokenize the identifier and restrict raw access."),
        ("iban", "bank identifier", "high", "Mask the identifier except for the last visible characters."),
        ("passport", "identity document", "high", "Hash or tokenize the identifier and restrict raw access."),
        ("birth", "date of birth", "medium", "Generalize the date to month or year unless exact day precision is required."),
        ("dob", "date of birth", "medium", "Generalize the date to month or year unless exact day precision is required."),
        ("ip", "IP address", "medium", "Mask part of the IP address or hash it when user-level tracking is not required."),
    ]
    for token, pii_type, risk, recommendation in rules:
        if token in lowered:
            return pii_type, risk, recommendation
    return None


def _build_anonymizer_sql(table_name: str, column_name: str, pii_type: str) -> dict[str, str]:
    identifier = quote_clickhouse_identifier(column_name)
    source = quote_clickhouse_identifier(table_name)
    if pii_type == "email":
        return {
            "title": f"Hash `{column_name}`",
            "sql": f"SELECT lower(hex(SHA256(toString({identifier})))) AS {identifier}_hash\nFROM {source}",
        }
    if pii_type == "phone number":
        return {
            "title": f"Mask `{column_name}`",
            "sql": f"SELECT concat(substringUTF8(toString({identifier}), 1, 2), '******') AS {identifier}_masked\nFROM {source}",
        }
    if pii_type == "person name":
        return {
            "title": f"Mask `{column_name}`",
            "sql": f"SELECT concat(substringUTF8(toString({identifier}), 1, 1), '***') AS {identifier}_masked\nFROM {source}",
        }
    if pii_type == "postal address":
        return {
            "title": f"Generalize `{column_name}`",
            "sql": f"SELECT lower(hex(SHA256(toString({identifier})))) AS {identifier}_token\nFROM {source}",
        }
    if pii_type == "date of birth":
        return {
            "title": f"Generalize `{column_name}`",
            "sql": f"SELECT formatDateTime(toDate({identifier}), '%Y-%m') AS {identifier}_month\nFROM {source}",
        }
    if pii_type == "IP address":
        return {
            "title": f"Mask `{column_name}`",
            "sql": f"SELECT replaceRegexpOne(toString({identifier}), '(\\\\d+)$', '0') AS {identifier}_masked\nFROM {source}",
        }
    return {
        "title": f"Hash `{column_name}`",
        "sql": f"SELECT lower(hex(SHA256(toString({identifier})))) AS {identifier}_hash\nFROM {source}",
    }


async def _anonymizer_profile(
    clickhouse: ClickHouseConfig,
    table_name: str,
    schema: list[dict[str, Any]],
    row_filter: str = "",
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]], str]:
    sampled_columns = _schema_column_names(schema)[: min(20, len(schema))]
    normalized_row_filter = _normalize_auto_ml_row_filter(row_filter)
    _, sample_rows, _ = await _fetch_clickhouse_sample_rows(
        clickhouse,
        table_name,
        sampled_columns,
        ANONYMIZER_SAMPLE_ROWS,
        normalized_row_filter,
    )
    findings: list[dict[str, Any]] = []
    scripts: list[dict[str, Any]] = []

    for column_name in sampled_columns:
        detected = _anonymizer_candidate_by_name(column_name)
        evidence = ""
        if not detected:
            values = [str(row.get(column_name) or "").strip() for row in sample_rows[:40] if str(row.get(column_name) or "").strip()]
            if any(EMAIL_PATTERN.match(value) for value in values):
                detected = ("email", "high", "Hash or tokenize email addresses before any external sharing.")
                evidence = "Email-like values were found in the sample."
            elif any(PHONE_PATTERN.match(value) for value in values):
                detected = ("phone number", "high", "Mask or hash phone numbers and keep the raw value in a restricted zone only.")
                evidence = "Phone-like values were found in the sample."
            elif any(IP_PATTERN.match(value) for value in values):
                detected = ("IP address", "medium", "Mask part of the IP address or hash it when user-level tracking is not required.")
                evidence = "IP-like values were found in the sample."
        if not detected:
            continue
        pii_type, risk, recommendation = detected
        findings.append({
            "column": column_name,
            "pii_type": pii_type,
            "risk": risk,
            "recommendation": recommendation,
            "evidence": evidence or f"Column name `{column_name}` strongly suggests {pii_type}.",
        })
        scripts.append(_build_anonymizer_sql(table_name, column_name, pii_type))

    return findings[:12], scripts[:8], sample_rows, _records_preview_markdown(sample_rows)


async def _data_cleaner_narrative(
    user_request: str,
    table_name: str,
    row_filter: str,
    findings: list[dict[str, Any]],
    scripts: list[dict[str, Any]],
    preview_markdown: str,
    conversation_memory: str,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str],
) -> str:
    prompt = f"""
You are a senior data quality consultant reporting to a business stakeholder.
Write a polished English Markdown answer.

Rules:
- Use the following visible sections only:
  - `## Executive Summary`
  - `## Key Quality Risks`
  - `## Recommended Fix Path`
- Keep the tone practical and business-oriented.
- Mention only the most relevant risks.
- Do not output SQL in the main narrative.

User request:
{user_request}

Table:
{table_name}

Row filter:
{row_filter or "Full table"}

Findings:
{json.dumps(findings, ensure_ascii=False, indent=2)}

Available correction scripts:
{json.dumps(scripts, ensure_ascii=False, indent=2)}

Sample preview:
{preview_markdown}

Conversation memory:
{conversation_memory}
""".strip()
    return await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
    )


async def _anonymizer_narrative(
    user_request: str,
    table_name: str,
    row_filter: str,
    pii_findings: list[dict[str, Any]],
    masking_scripts: list[dict[str, Any]],
    preview_markdown: str,
    conversation_memory: str,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str],
) -> str:
    prompt = f"""
You are a privacy and governance specialist reporting to a stakeholder about PII exposure.
Write a polished English Markdown answer.

Rules:
- Use the following visible sections only:
  - `## Executive Summary`
  - `## PII Exposure Review`
  - `## Recommended Masking Strategy`
- Keep the tone pragmatic, concise, and compliant.
- Do not output SQL in the main narrative.

User request:
{user_request}

Table:
{table_name}

Row filter:
{row_filter or "Full table"}

PII findings:
{json.dumps(pii_findings, ensure_ascii=False, indent=2)}

Available masking scripts:
{json.dumps(masking_scripts, ensure_ascii=False, indent=2)}

Sample preview:
{preview_markdown}

Conversation memory:
{conversation_memory}
""".strip()
    return await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
    )


def _default_oracle_analyst_state() -> dict[str, Any]:
    return {
        "stage": "idle",
        "pending_request": "",
        "available_tables": [],
        "selected_table": None,
        "schema_info": [],
        "clarification_prompt": "",
        "clarification_options": [],
        "last_sql": "",
        "last_result_meta": [],
        "last_result_rows": [],
        "final_answer": "",
        "action_log": [],
        "last_error": "",
    }


def _normalize_oracle_analyst_state(payload: Optional[dict]) -> dict[str, Any]:
    state = _default_oracle_analyst_state()
    if not isinstance(payload, dict):
        return state

    stage = str(payload.get("stage") or "").strip()
    state["stage"] = stage if stage in {"idle", "awaiting_table", "ready"} else "idle"
    state["pending_request"] = str(payload.get("pending_request") or payload.get("pendingRequest") or "").strip()
    state["available_tables"] = [
        str(item).strip()
        for item in (payload.get("available_tables") or payload.get("availableTables") or [])
        if str(item).strip()
    ]
    state["selected_table"] = str(payload.get("selected_table") or payload.get("selectedTable") or "").strip() or None
    schema_info = payload.get("schema_info") if isinstance(payload.get("schema_info"), list) else payload.get("schemaInfo")
    if isinstance(schema_info, list):
        state["schema_info"] = [
            {
                "name": str(column.get("name") or "").strip(),
                "type": str(column.get("type") or "").strip(),
                "nullable": bool(column.get("nullable")),
            }
            for column in schema_info
            if isinstance(column, dict) and str(column.get("name") or "").strip()
        ]
    state["clarification_prompt"] = str(payload.get("clarification_prompt") or payload.get("clarificationPrompt") or "").strip()
    state["clarification_options"] = [
        str(item).strip()
        for item in (payload.get("clarification_options") or payload.get("clarificationOptions") or [])
        if str(item).strip()
    ]
    state["last_sql"] = str(payload.get("last_sql") or payload.get("lastSql") or "").strip()
    last_result_meta = payload.get("last_result_meta") if isinstance(payload.get("last_result_meta"), list) else payload.get("lastResultMeta")
    if isinstance(last_result_meta, list):
        state["last_result_meta"] = [
            {
                "name": str(column.get("name") or "").strip(),
                "type": str(column.get("type") or "").strip(),
            }
            for column in last_result_meta
            if isinstance(column, dict) and str(column.get("name") or "").strip()
        ]
    last_result_rows = payload.get("last_result_rows") if isinstance(payload.get("last_result_rows"), list) else payload.get("lastResultRows")
    if isinstance(last_result_rows, list):
        state["last_result_rows"] = [row for row in last_result_rows if isinstance(row, dict)]
    state["final_answer"] = str(payload.get("final_answer") or payload.get("finalAnswer") or "").strip()
    state["action_log"] = [
        str(item).strip()
        for item in (payload.get("action_log") or payload.get("actionLog") or [])
        if str(item).strip()
    ]
    state["last_error"] = str(payload.get("last_error") or payload.get("lastError") or "").strip()
    return state


def _normalize_oracle_connection_payload(payload: Any, index: int = 0) -> dict[str, Any]:
    default = DEFAULT_APP_CONFIG["oracleConnections"][0]
    if not isinstance(payload, dict):
        return {**default, "id": f"oracle_{index + 1}", "label": f"Oracle {index + 1}"}
    return {
        "id": str(payload.get("id") or f"oracle_{index + 1}").strip() or f"oracle_{index + 1}",
        "label": str(payload.get("label") or f"Oracle {index + 1}").strip() or f"Oracle {index + 1}",
        "host": str(payload.get("host") or default["host"]).strip(),
        "port": max(1, int(payload.get("port") or default["port"])),
        "serviceName": str(payload.get("serviceName") or payload.get("service_name") or "").strip(),
        "sid": str(payload.get("sid") or "").strip(),
        "dsn": str(payload.get("dsn") or "").strip(),
        "username": str(payload.get("username") or "").strip(),
        "password": str(payload.get("password") or ""),
    }


def _normalize_oracle_connections_payload(payload: Any) -> list[dict[str, Any]]:
    if not isinstance(payload, list):
        return json.loads(json.dumps(DEFAULT_APP_CONFIG["oracleConnections"]))
    normalized = [
        _normalize_oracle_connection_payload(item, index)
        for index, item in enumerate(payload)
        if isinstance(item, dict)
    ]
    return normalized or json.loads(json.dumps(DEFAULT_APP_CONFIG["oracleConnections"]))


def _normalize_oracle_analyst_config(payload: Optional[dict]) -> dict[str, Any]:
    defaults = DEFAULT_APP_CONFIG["oracleAnalystConfig"]
    if not isinstance(payload, dict):
        return dict(defaults)
    return {
        "connectionId": str(payload.get("connectionId") or payload.get("connection_id") or defaults["connectionId"]).strip() or defaults["connectionId"],
        "rowLimit": max(1, min(ORACLE_MAX_ROW_LIMIT, int(payload.get("rowLimit") or payload.get("row_limit") or defaults["rowLimit"]))),
        "maxRetries": max(1, min(10, int(payload.get("maxRetries") or payload.get("max_retries") or defaults["maxRetries"]))),
        "maxIterations": max(1, min(20, int(payload.get("maxIterations") or payload.get("max_iterations") or defaults["maxIterations"]))),
        "toolkitId": str(payload.get("toolkitId") or payload.get("toolkit_id") or defaults["toolkitId"]).strip(),
        "systemPrompt": str(payload.get("systemPrompt") or payload.get("system_prompt") or defaults["systemPrompt"]).strip() or defaults["systemPrompt"],
    }


def _resolve_oracle_connection(
    connections: list[OracleConnectionConfig],
    config: OracleAnalystConfigModel | dict[str, Any],
) -> OracleConnectionConfig:
    normalized_connections = [
        OracleConnectionConfig(
            id=str(connection.id).strip(),
            label=str(connection.label).strip(),
            host=str(connection.host).strip(),
            port=int(connection.port),
            service_name=str(connection.service_name).strip(),
            sid=str(connection.sid).strip(),
            dsn=str(connection.dsn).strip(),
            username=str(connection.username).strip(),
            password=str(connection.password),
        )
        if isinstance(connection, OracleConnectionConfig)
        else OracleConnectionConfig(
            id=str(connection.get("id") or "").strip(),
            label=str(connection.get("label") or "").strip(),
            host=str(connection.get("host") or "").strip(),
            port=int(connection.get("port") or 1521),
            service_name=str(connection.get("service_name") or connection.get("serviceName") or "").strip(),
            sid=str(connection.get("sid") or "").strip(),
            dsn=str(connection.get("dsn") or "").strip(),
            username=str(connection.get("username") or "").strip(),
            password=str(connection.get("password") or ""),
        )
        for connection in connections
    ]
    if not normalized_connections:
        raise HTTPException(status_code=400, detail="No Oracle connection is configured.")
    connection_id = (
        getattr(config, "connection_id", None)
        if isinstance(config, OracleAnalystConfigModel)
        else str((config or {}).get("connectionId") or (config or {}).get("connection_id") or "")
    )
    for connection in normalized_connections:
        if connection.id == connection_id:
            return connection
    return normalized_connections[0]


def _oracle_markdown_table(rows: list[dict[str, Any]], limit: int = ORACLE_RESULT_PREVIEW_ROWS) -> str:
    preview_rows = rows[:limit]
    if not preview_rows:
        return "No rows returned."
    headers = list(preview_rows[0].keys())
    header_line = "| " + " | ".join(headers) + " |"
    divider = "| " + " | ".join(["---"] * len(headers)) + " |"
    body = []
    for row in preview_rows:
        body.append("| " + " | ".join(str(row.get(header, "")) for header in headers) + " |")
    return "\n".join([header_line, divider, *body])


def _oracle_actions_markdown(actions: list[str]) -> str:
    if not actions:
        return "- No tool call was needed."
    return "\n".join(f"- {action}" for action in actions)


async def plan_oracle_react_step(
    user_message: str,
    history: list[dict[str, Any]],
    scratchpad: list[dict[str, Any]],
    state: dict[str, Any],
    oracle_config: dict[str, Any],
    connection_label: str,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str] = None,
) -> dict[str, Any]:
    prompt = f"""
You are the Oracle SQL agent in RAGnarok.
Your job is to answer the user's question by reasoning step by step and deciding whether to use one Oracle tool.

Return JSON only with this exact shape:
{{
  "reasoning": "short English explanation",
  "action": "tool" | "clarify_table" | "final",
  "tool_name": "list_tables" | "get_schema" | "check_query" | "execute_query",
  "tool_input": {{}},
  "clarification_prompt": "short English question",
  "clarification_options": ["SCHEMA.TABLE_A", "SCHEMA.TABLE_B"],
  "final_answer": "Markdown final answer if action=final"
}}

Oracle SQL rules:
- NEVER use SELECT *
- Prefer indexed columns in WHERE clauses when possible
- Use TRUNC(), TO_DATE(..., 'YYYY-MM-DD'), and SYSDATE for date logic
- Use ROW_NUMBER() OVER (PARTITION BY ... ORDER BY ...) for ranking when needed
- Use FETCH FIRST n ROWS ONLY for pagination
- Use SUBSTR(), NVL(), TO_CHAR(), and DECODE() when helpful
- Use only read-only SELECT or WITH queries
- If the user explicitly asks for a table, preserve a tabular result preview in the final answer.

Working rules:
- At most one tool call per iteration
- Start with list_tables when you do not know the right table
- Use get_schema after the table is known
- Use get_schema with columns_filter when you only need a focused type lookup
- Prefer check_query before execute_query when SQL was just created or repaired
- If a tool result contains an Oracle error, repair the SQL and continue
- Ask for table clarification only when multiple tables remain plausible
- Keep everything in English

Connection label: {connection_label}
Configured row limit: {oracle_config.get("rowLimit", ORACLE_DEFAULT_ROW_LIMIT)}
Toolkit id: {oracle_config.get("toolkitId") or "default"}
Custom system prompt:
{_with_table_output_guidance(str(oracle_config.get("systemPrompt") or ""))}

Current agent state:
{json.dumps(state, ensure_ascii=False, indent=2)}

Recent conversation memory:
{_conversation_memory_markdown(history, current_message=user_message)}

Scratchpad:
{json.dumps(scratchpad[-10:], ensure_ascii=False, indent=2)}

User request:
{user_message}
""".strip()

    raw = await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
        response_format="json",
    )
    parsed = extract_json_object(raw)
    action = str(parsed.get("action") or "").strip().lower()
    return {
        "reasoning": str(parsed.get("reasoning") or "").strip(),
        "action": action,
        "tool_name": str(parsed.get("tool_name") or "").strip(),
        "tool_input": parsed.get("tool_input") if isinstance(parsed.get("tool_input"), dict) else {},
        "clarification_prompt": str(parsed.get("clarification_prompt") or "").strip(),
        "clarification_options": [
            str(item).strip()
            for item in (parsed.get("clarification_options") or [])
            if str(item).strip()
        ],
        "final_answer": str(parsed.get("final_answer") or "").strip(),
    }


async def summarize_oracle_result(
    user_request: str,
    executed_sql: str,
    result_rows: list[dict[str, Any]],
    action_log: list[str],
    conversation_memory: str,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str] = None,
) -> dict[str, Any]:
    preview = json.dumps(result_rows[:ORACLE_RESULT_PREVIEW_ROWS], ensure_ascii=False, indent=2)
    prompt = f"""
You are summarizing an Oracle SQL analysis result for an end user.
Return JSON only with this exact shape:
{{
  "executive_summary": "2 to 5 sentence narrative in English",
  "key_metrics": [
    {{"label": "Total revenue", "value": "123456"}},
    {{"label": "Top region", "value": "EMEA"}}
  ],
  "insights": ["short insight", "short recommendation"],
  "confidence_score": 0,
  "confidence_reason": "short explanation"
}}

Rules:
- Keep the tone business-facing and precise
- Use the actual result values when possible
- Keep confidence between 0 and 100
- If there are no rows, explain that clearly

User request:
{user_request}

Executed SQL:
{executed_sql}

Recent conversation memory:
{conversation_memory}

Action log:
{json.dumps(action_log, ensure_ascii=False)}

Result preview:
{preview}
""".strip()

    raw = await llm_chat(
        [{"role": "user", "content": prompt}],
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
        response_format="json",
    )
    parsed = extract_json_object(raw)
    key_metrics = []
    for item in parsed.get("key_metrics", []) if isinstance(parsed.get("key_metrics"), list) else []:
        if not isinstance(item, dict):
            continue
        label = str(item.get("label") or "").strip()
        value = str(item.get("value") or "").strip()
        if label and value:
            key_metrics.append({"label": label, "value": value})

    insights = [
        str(item).strip()
        for item in (parsed.get("insights") or [])
        if str(item).strip()
    ] if isinstance(parsed.get("insights"), list) else []

    score = parsed.get("confidence_score")
    try:
        confidence_score = max(0, min(100, int(score)))
    except Exception:
        confidence_score = 78 if result_rows else 62

    return {
        "executive_summary": str(parsed.get("executive_summary") or "").strip() or "The Oracle query completed successfully.",
        "key_metrics": key_metrics,
        "insights": insights,
        "confidence_score": confidence_score,
        "confidence_reason": str(parsed.get("confidence_reason") or "").strip() or "Confidence is based on schema grounding and successful query execution.",
    }


def build_oracle_response_markdown(
    summary: dict[str, Any],
    sql: str,
    rows: list[dict[str, Any]],
    action_log: list[str],
) -> str:
    executive_summary = str(summary.get("executive_summary") or "The Oracle query completed successfully.").strip()
    key_metrics = summary.get("key_metrics") or []
    insights = summary.get("insights") or []
    confidence_score = int(summary.get("confidence_score") or 0)
    confidence_reason = str(summary.get("confidence_reason") or "").strip()

    sections = [
        "## Executive Summary",
        executive_summary,
        "",
        "## Key Metrics",
    ]

    if key_metrics:
        sections.extend(
            f"- **{metric['label']}**: {metric['value']}"
            for metric in key_metrics
            if metric.get("label") and metric.get("value")
        )
    else:
        sections.append(f"- **Rows returned**: {len(rows)}")

    sections.extend(
        [
            "",
            "## SQL Used",
            "```sql",
            clean_sql_text(sql),
            "```",
            "",
            "## Data Table",
            _oracle_markdown_table(rows),
            "",
            "## Insights & Recommendations",
        ]
    )

    if insights:
        sections.extend(f"- {insight}" for insight in insights)
    else:
        sections.append("- No additional recommendation was required beyond the query result.")

    sections.extend(
        [
            "",
            "## Actions Performed",
            _oracle_actions_markdown(action_log),
            "",
            "## Confidence Score",
            f"Score: **{confidence_score}/100** — {confidence_reason or 'Confidence is based on successful Oracle tool execution.'}",
        ]
    )
    return "\n".join(sections)

# ── LangGraph planning helpers ────────────────────────────────────────────────

PLANNING_PRODUCT_NAME = "LangGraph Planning"

PLANNING_ROLE_PROMPTS = {
    "manager": (
        "You are an operations manager agent. Produce a concise operational brief with "
        "clear priorities, risks, and next actions."
    ),
    "file_management": (
        "You are a file management agent. Use filesystem facts only, keep answers concise, "
        "and never imply that a destructive action ran without explicit confirmation."
    ),
}


def _safe_zoneinfo(timezone_name: str) -> ZoneInfo:
    try:
        return ZoneInfo(timezone_name or "UTC")
    except ZoneInfoNotFoundError:
        return ZoneInfo("UTC")


def _parse_iso_datetime(value: Optional[str]) -> Optional[datetime]:
    if not value or not isinstance(value, str):
        return None
    try:
        text = value.strip().replace("Z", "+00:00")
        parsed = datetime.fromisoformat(text)
    except Exception:
        return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=timezone.utc)
    return parsed.astimezone(timezone.utc)


def _to_iso_datetime(value: datetime) -> str:
    return value.astimezone(timezone.utc).isoformat()


def _parse_time_of_day(value: str) -> tuple[int, int]:
    match = re.fullmatch(r"(\d{1,2}):(\d{2})", (value or "").strip())
    if not match:
        return 9, 0
    hour = max(0, min(23, int(match.group(1))))
    minute = max(0, min(59, int(match.group(2))))
    return hour, minute


def _weekday_to_index(day: str) -> int:
    return PLANNER_WEEKDAYS.index(day) if day in PLANNER_WEEKDAYS else 0


def _build_localized_datetime(date_value: datetime, time_of_day: str, timezone_name: str) -> datetime:
    tz = _safe_zoneinfo(timezone_name)
    localized = date_value.astimezone(tz)
    hour, minute = _parse_time_of_day(time_of_day)
    return localized.replace(hour=hour, minute=minute, second=0, microsecond=0)


def _compute_plan_next_run_at(plan: dict, reference_dt: Optional[datetime] = None) -> Optional[str]:
    trigger = plan.get("trigger") or {}
    if plan.get("status") != "active":
        return None

    now_dt = (reference_dt or datetime.now(timezone.utc)).astimezone(timezone.utc)
    kind = trigger.get("kind")
    timezone_name = str(trigger.get("timezone") or "UTC")

    if kind == "once":
        one_time = _parse_iso_datetime(trigger.get("oneTimeAt"))
        return _to_iso_datetime(one_time) if one_time and one_time > now_dt else None

    if kind == "daily":
        candidate = _build_localized_datetime(now_dt, trigger.get("timeOfDay") or "09:00", timezone_name)
        if candidate.astimezone(timezone.utc) <= now_dt:
            candidate += timedelta(days=1)
        return _to_iso_datetime(candidate.astimezone(timezone.utc))

    if kind == "weekly":
        weekdays = [
            day for day in trigger.get("weekdays", [])
            if isinstance(day, str) and day in PLANNER_WEEKDAYS
        ] or ["mon"]
        current_local = now_dt.astimezone(_safe_zoneinfo(timezone_name))
        hour, minute = _parse_time_of_day(trigger.get("timeOfDay") or "09:00")
        for offset in range(0, 8):
            candidate_date = current_local + timedelta(days=offset)
            candidate_day = PLANNER_WEEKDAYS[candidate_date.weekday()]
            if candidate_day not in weekdays:
                continue
            candidate = candidate_date.replace(hour=hour, minute=minute, second=0, microsecond=0)
            if candidate.astimezone(timezone.utc) > now_dt:
                return _to_iso_datetime(candidate.astimezone(timezone.utc))
        fallback = current_local + timedelta(days=7)
        fallback = fallback.replace(hour=hour, minute=minute, second=0, microsecond=0)
        return _to_iso_datetime(fallback.astimezone(timezone.utc))

    if kind == "interval":
        interval_minutes = max(1, int(trigger.get("intervalMinutes") or 60))
        last_run = _parse_iso_datetime(plan.get("lastRunAt"))
        anchor = last_run or now_dt
        if not last_run and plan.get("nextRunAt"):
            return plan.get("nextRunAt")
        return _to_iso_datetime(anchor + timedelta(minutes=interval_minutes))

    return None


def _refresh_planning_plan(plan: dict, reference_dt: Optional[datetime] = None) -> dict:
    trigger = plan.get("trigger") or {}
    if trigger.get("kind") in {"clickhouse_watch", "file_watch"}:
        plan["nextRunAt"] = None
        return plan

    plan["nextRunAt"] = _compute_plan_next_run_at(plan, reference_dt)
    return plan


def _planning_state_from_db(state: dict) -> dict:
    planning = _normalize_planning_state(state.get("planning"))
    for plan in planning["plans"]:
        _refresh_planning_plan(plan)
    planning["runs"] = planning["runs"][:PLANNER_MAX_RUNS]
    return planning


def _default_planning_draft(timezone_name: str = "UTC") -> dict:
    return {
        "name": "",
        "prompt": "",
        "agents": [],
        "status": "active",
        "trigger": {
            "kind": "daily",
            "timezone": timezone_name or "UTC",
            "oneTimeAt": "",
            "timeOfDay": "09:00",
            "weekdays": ["mon"],
            "intervalMinutes": 60,
            "pollMinutes": 5,
            "watchSql": "",
            "watchMode": "result_changes",
            "directory": "",
            "pattern": "*",
            "recursive": False,
        },
    }


def _normalize_planner_agent_role(value: Any) -> Optional[str]:
    if not isinstance(value, str):
        return None
    lowered = normalize_choice(value).lower()
    aliases = {
        "manager": "manager",
        "agent manager": "manager",
        "clickhouse query": "clickhouse_query",
        "clickhouse sql": "clickhouse_query",
        "clickhouse_query": "clickhouse_query",
        "clickhouse": "clickhouse_query",
        "file management": "file_management",
        "file manager": "file_management",
        "file_management": "file_management",
    }
    return aliases.get(lowered)


def _merge_planning_draft(current_draft: Optional[dict], updates: Optional[dict], timezone_name: str = "UTC") -> dict:
    current = _default_planning_draft(timezone_name)
    if isinstance(current_draft, dict):
        current.update({k: v for k, v in current_draft.items() if k != "trigger"})
        current["trigger"].update((current_draft.get("trigger") or {}))

    if not isinstance(updates, dict):
        return current

    if updates.get("name"):
        current["name"] = str(updates.get("name")).strip()
    if updates.get("prompt"):
        current["prompt"] = str(updates.get("prompt")).strip()

    if isinstance(updates.get("agents"), list):
        existing_agents = [agent for agent in current.get("agents", []) if agent in PLANNER_AGENT_ROLES]
        for agent in updates["agents"]:
            normalized_agent = _normalize_planner_agent_role(agent)
            if normalized_agent and normalized_agent not in existing_agents:
                existing_agents.append(normalized_agent)
        current["agents"] = existing_agents

    if updates.get("status") in {"active", "paused"}:
        current["status"] = updates["status"]

    incoming_trigger = updates.get("trigger")
    if isinstance(incoming_trigger, dict):
        for key, value in incoming_trigger.items():
            if value in (None, "", []):
                continue
            current["trigger"][key] = value

    current["trigger"] = _normalize_planning_trigger(current["trigger"])
    if not current["name"] and current["prompt"]:
        prompt_words = current["prompt"].split()
        current["name"] = " ".join(prompt_words[:6])[:64]

    return current


def _validate_planning_draft(draft: dict) -> list[str]:
    missing: list[str] = []
    if not draft.get("prompt"):
        missing.append("prompt")
    if not draft.get("agents"):
        missing.append("agents")

    trigger = draft.get("trigger") or {}
    kind = trigger.get("kind")
    if kind not in PLANNER_TRIGGER_KINDS:
        missing.append("trigger_kind")
        return missing

    if kind == "once" and not trigger.get("oneTimeAt"):
        missing.append("one_time_at")
    elif kind == "daily" and not trigger.get("timeOfDay"):
        missing.append("time_of_day")
    elif kind == "weekly":
        if not trigger.get("weekdays"):
            missing.append("weekdays")
        if not trigger.get("timeOfDay"):
            missing.append("time_of_day")
    elif kind == "interval" and int(trigger.get("intervalMinutes") or 0) <= 0:
        missing.append("interval_minutes")
    elif kind == "clickhouse_watch":
        if not trigger.get("watchSql"):
            missing.append("watch_sql")
        if int(trigger.get("pollMinutes") or 0) <= 0:
            missing.append("poll_minutes")
    elif kind == "file_watch":
        if not trigger.get("directory"):
            missing.append("directory")
        if int(trigger.get("pollMinutes") or 0) <= 0:
            missing.append("poll_minutes")

    return missing


def _planning_missing_prompt(missing_fields: list[str]) -> str:
    labels = {
        "prompt": "what the automation should do",
        "agents": "which existing agent(s) should run",
        "trigger_kind": "what trigger type you want",
        "one_time_at": "the exact date and time",
        "time_of_day": "the time of day",
        "weekdays": "which weekday(s) should run",
        "interval_minutes": "the repeat interval in minutes",
        "watch_sql": "the ClickHouse watch SQL",
        "poll_minutes": "the polling frequency",
        "directory": "the directory to watch",
    }
    readable = [labels.get(field, field.replace("_", " ")) for field in missing_fields]
    if not readable:
        return ""
    if len(readable) == 1:
        return readable[0]
    return ", ".join(readable[:-1]) + f" and {readable[-1]}"


def _planning_summary_markdown(draft: dict, missing_fields: list[str]) -> str:
    trigger = draft.get("trigger") or {}
    agents = draft.get("agents") or []
    lines = [
        "## Draft Summary",
        f"- Name: {draft.get('name') or 'Untitled plan'}",
        f"- Agents: {', '.join(agents) if agents else 'Not selected yet'}",
        f"- Trigger: {trigger.get('kind') or 'Not selected yet'}",
    ]
    if trigger.get("kind") == "once":
        lines.append(f"- Run at: {trigger.get('oneTimeAt') or 'Missing'}")
    elif trigger.get("kind") == "daily":
        lines.append(f"- Time: {trigger.get('timeOfDay') or 'Missing'} ({trigger.get('timezone') or 'UTC'})")
    elif trigger.get("kind") == "weekly":
        weekdays = ", ".join(trigger.get("weekdays") or []) or "Missing"
        lines.append(f"- Weekdays: {weekdays}")
        lines.append(f"- Time: {trigger.get('timeOfDay') or 'Missing'} ({trigger.get('timezone') or 'UTC'})")
    elif trigger.get("kind") == "interval":
        lines.append(f"- Every: {trigger.get('intervalMinutes') or 'Missing'} minute(s)")
    elif trigger.get("kind") == "clickhouse_watch":
        lines.append(f"- Watch SQL: `{(trigger.get('watchSql') or '').strip()[:100] or 'Missing'}`")
        lines.append(f"- Polling: every {trigger.get('pollMinutes') or 'Missing'} minute(s)")
    elif trigger.get("kind") == "file_watch":
        lines.append(f"- Directory: `{trigger.get('directory') or 'Missing'}`")
        lines.append(f"- Pattern: `{trigger.get('pattern') or '*'}`")
        lines.append(f"- Polling: every {trigger.get('pollMinutes') or 'Missing'} minute(s)")

    lines.extend(
        [
            "",
            "## Objective",
            draft.get("prompt") or "Missing",
        ]
    )

    if missing_fields:
        lines.extend(
            [
                "",
                "## Next Step",
                f"I still need { _planning_missing_prompt(missing_fields) } before this automation is ready.",
            ]
        )
    else:
        lines.extend(
            [
                "",
                "## Next Step",
                "The draft is ready. Open the planner form to review and save it.",
            ]
        )
    return "\n".join(lines)


PLANNING_ANALYSIS_PROMPT = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            (
                "You are helping a user configure a LangGraph-based automation planner for existing agents. "
                "Return JSON only with the keys: draft, clarification_question, should_open_form, and reasoning.\n"
                "Inside draft, use these keys only: name, prompt, agents, status, and trigger.\n"
                "Inside trigger, use these keys only: kind, timezone, oneTimeAt, timeOfDay, weekdays, "
                "intervalMinutes, pollMinutes, watchSql, watchMode, directory, pattern, recursive.\n"
                "Allowed agents: manager, clickhouse_query, file_management.\n"
                "Allowed trigger kinds: once, daily, weekly, interval, clickhouse_watch, file_watch.\n"
                "Allowed watch modes: returns_rows, count_increases, result_changes.\n"
                "Allowed weekdays: mon, tue, wed, thu, fri, sat, sun.\n"
                "Keep the answer in English.\n"
                "If the user explicitly asks to open the form, set should_open_form to true.\n"
                "Do not invent unsupported trigger kinds or agent names.\n"
                "Only fill fields when the user clearly implies them or they already exist in the current draft."
            ),
        ),
        (
            "human",
            (
                "Current draft:\n{current_draft_json}\n\n"
                "Recent conversation memory:\n{history_markdown}\n\n"
                "User message:\n{user_message}"
            ),
        ),
    ]
)

PLANNING_RUN_SUMMARY_PROMPT = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            (
                "You are summarizing the result of a LangGraph-orchestrated automation. "
                "Write the answer in English with exactly these markdown sections:\n"
                "## Summary\nOne concise summary.\n\n"
                "## Trigger\nOne concise explanation of why it ran.\n\n"
                "## Agent Outputs\nOne-line summary per agent."
            ),
        ),
        (
            "human",
            (
                "Automation name: {plan_name}\n"
                "Objective:\n{plan_prompt}\n\n"
                "Trigger context:\n{trigger_markdown}\n\n"
                "Outputs:\n{outputs_json}"
            ),
        ),
    ]
)


class PlanningChatGraphState(TypedDict, total=False):
    user_message: str
    history: list[dict[str, Any]]
    current_draft: dict[str, Any]
    timezone_name: str
    planning_state: dict[str, Any]
    llm_config: dict[str, Any]
    mode: Literal["empty", "reset", "list", "open_form", "analyze"]
    analysis: dict[str, Any]
    response: dict[str, Any]


class PlanningExecutionGraphState(TypedDict, total=False):
    plan: dict[str, Any]
    trigger_context: dict[str, Any]
    app_config: dict[str, Any]
    pending_agents: list[str]
    outputs: list[dict[str, Any]]
    overall_status: str
    summary: str


_planning_chat_graph: Optional[Any] = None
_planning_execution_graph: Optional[Any] = None


def _prompt_content_as_text(content: Any) -> str:
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        fragments: list[str] = []
        for item in content:
            if isinstance(item, dict):
                fragments.append(str(item.get("text") or item.get("content") or ""))
            else:
                fragments.append(str(item))
        return "\n".join(fragment for fragment in fragments if fragment).strip()
    return str(content or "")


def _prompt_messages_from_template(
    template: ChatPromptTemplate,
    **values: Any,
) -> list[dict[str, str]]:
    prompt_value = template.invoke(values)
    rendered_messages: list[dict[str, str]] = []
    for message in prompt_value.to_messages():
        message_type = getattr(message, "type", "human")
        role = "user"
        if message_type == "system":
            role = "system"
        elif message_type == "ai":
            role = "assistant"
        rendered_messages.append(
            {
                "role": role,
                "content": _prompt_content_as_text(getattr(message, "content", "")),
            }
        )
    return rendered_messages


async def _llm_chat_from_prompt_template(
    template: ChatPromptTemplate,
    *,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str] = None,
    response_format: Optional[str] = None,
    disable_ssl_verification: Optional[bool] = None,
    **values: Any,
) -> str:
    return await llm_chat(
        _prompt_messages_from_template(template, **values),
        llm_base_url,
        llm_model,
        llm_provider,
        llm_api_key,
        response_format=response_format,
        disable_ssl_verification=disable_ssl_verification,
    )


def _planning_open_form_action() -> dict[str, str]:
    return {
        "id": "open-planning-form",
        "label": "Open planning form",
        "actionType": "open_planning_form",
        "variant": "primary",
    }


def _planning_agent_state_payload(
    draft: dict[str, Any],
    missing_fields: list[str],
    last_question: str,
) -> dict[str, Any]:
    return {
        "draft": draft,
        "missing_fields": missing_fields,
        "last_question": last_question,
        "ready_to_review": len(missing_fields) == 0,
    }


def _planning_response_payload(
    draft: dict[str, Any],
    missing_fields: list[str],
    answer: str,
    steps: list[dict[str, Any]],
    *,
    last_question: str = "",
) -> dict[str, Any]:
    return {
        "answer": answer,
        "agent_state": _planning_agent_state_payload(draft, missing_fields, last_question),
        "actions": [_planning_open_form_action()],
        "steps": steps,
    }


async def analyze_planning_request(
    user_message: str,
    history: list[dict[str, Any]],
    current_draft: dict,
    llm_base_url: str,
    llm_model: str,
    llm_provider: str,
    llm_api_key: Optional[str] = None,
    disable_ssl_verification: bool = False,
) -> dict[str, Any]:
    raw = await _llm_chat_from_prompt_template(
        PLANNING_ANALYSIS_PROMPT,
        llm_base_url=llm_base_url,
        llm_model=llm_model,
        llm_provider=llm_provider,
        llm_api_key=llm_api_key,
        response_format="json",
        disable_ssl_verification=disable_ssl_verification,
        current_draft_json=json.dumps(current_draft, ensure_ascii=False, indent=2),
        history_markdown=_conversation_memory_markdown(history, current_message=user_message),
        user_message=user_message,
    )
    parsed = extract_json_object(raw)
    return {
        "draft": parsed.get("draft") if isinstance(parsed.get("draft"), dict) else {},
        "clarification_question": str(parsed.get("clarification_question", "") or "").strip(),
        "should_open_form": bool(parsed.get("should_open_form", False)),
        "reasoning": str(parsed.get("reasoning", "") or "").strip(),
    }


def _planning_detect_mode(user_message: str) -> Literal["empty", "reset", "list", "open_form", "analyze"]:
    message = (user_message or "").strip()
    if not message:
        return "empty"

    lowered = message.lower()
    if any(token in lowered for token in ["reset", "start over", "clear draft"]):
        return "reset"
    if any(token in lowered for token in ["list plans", "show plans", "existing plans", "what plans"]):
        return "list"
    if any(token in lowered for token in ["open form", "open planner", "planner form", "show form"]):
        return "open_form"
    return "analyze"


def _planning_detect_mode_node(state: PlanningChatGraphState) -> PlanningChatGraphState:
    return {"mode": _planning_detect_mode(str(state.get("user_message") or ""))}


def _planning_chat_route(state: PlanningChatGraphState) -> str:
    return "analysis" if state.get("mode") == "analyze" else "direct"


def _planning_direct_response_node(state: PlanningChatGraphState) -> PlanningChatGraphState:
    mode = state.get("mode") or "empty"
    timezone_name = str(state.get("timezone_name") or "UTC")
    current_draft = _merge_planning_draft(state.get("current_draft"), None, timezone_name)
    planning_state = state.get("planning_state") or {"plans": [], "runs": []}
    missing_fields = _validate_planning_draft(current_draft)

    if mode == "reset":
        fresh_draft = _default_planning_draft(timezone_name)
        response = _planning_response_payload(
            fresh_draft,
            _validate_planning_draft(fresh_draft),
            (
                f"## {PLANNING_PRODUCT_NAME}\n"
                "The planning draft has been reset. Tell me what you want to automate, or open the form."
            ),
            [
                {
                    "id": "planning-reset",
                    "title": "Reset planning draft",
                    "status": "success",
                    "details": "The previous draft was cleared.",
                }
            ],
        )
        return {"response": response}

    if mode == "list":
        response = _planning_response_payload(
            current_draft,
            missing_fields,
            _planning_state_markdown(planning_state),
            [
                {
                    "id": "planning-list",
                    "title": "Loaded existing plans",
                    "status": "success",
                    "details": f"Found {len((planning_state or {}).get('plans', []))} saved planning job(s).",
                }
            ],
        )
        return {"response": response}

    if mode == "open_form":
        response = _planning_response_payload(
            current_draft,
            missing_fields,
            _planning_summary_markdown(current_draft, missing_fields),
            [
                {
                    "id": "planning-open-form",
                    "title": "Prepared planner form",
                    "status": "success",
                    "details": "The current draft is ready to be reviewed in the planning form.",
                }
            ],
        )
        return {"response": response}

    response = _planning_response_payload(
        current_draft,
        missing_fields,
        (
            f"## {PLANNING_PRODUCT_NAME}\n"
            "Describe the automation you want in natural language, or open the planning form "
            "to configure triggers, existing agents, and monitoring rules step by step."
        ),
        [
            {
                "id": "planning-ready",
                "title": "Planner ready",
                "status": "success",
                "details": "The planner can work from natural language or from the full-screen form.",
            }
        ],
    )
    return {"response": response}


async def _planning_analysis_node(state: PlanningChatGraphState) -> PlanningChatGraphState:
    llm_config = state.get("llm_config") or {}
    analysis = await analyze_planning_request(
        str(state.get("user_message") or ""),
        list(state.get("history") or []),
        dict(state.get("current_draft") or {}),
        str(llm_config.get("llm_base_url") or "http://localhost:11434"),
        str(llm_config.get("llm_model") or "llama3"),
        str(llm_config.get("llm_provider") or "ollama"),
        llm_config.get("llm_api_key") or None,
        bool(llm_config.get("disable_ssl_verification", False)),
    )
    return {"analysis": analysis}


def _planning_finalize_response_node(state: PlanningChatGraphState) -> PlanningChatGraphState:
    timezone_name = str(state.get("timezone_name") or "UTC")
    current_draft = dict(state.get("current_draft") or {})
    analysis = state.get("analysis") or {}
    merged_draft = _merge_planning_draft(current_draft, analysis.get("draft"), timezone_name)
    missing_fields = _validate_planning_draft(merged_draft)
    ready_to_review = len(missing_fields) == 0
    clarification_question = (
        str(analysis.get("clarification_question") or "").strip()
        or (
            f"Please tell me { _planning_missing_prompt(missing_fields) }."
            if missing_fields
            else "The draft is ready. Open the planning form to review and save it."
        )
    )
    answer = (
        f"{_planning_summary_markdown(merged_draft, missing_fields)}\n\n"
        "## Guidance\n"
        f"{clarification_question}"
    )
    response = _planning_response_payload(
        merged_draft,
        missing_fields,
        answer,
        [
            {
                "id": "planning-parse",
                "title": "Parsed planning request",
                "status": "success",
                "details": (
                    analysis.get("reasoning")
                    or "The local LLM extracted the planning intent into a structured draft through LangGraph."
                ),
            },
            {
                "id": "planning-review",
                "title": "Draft readiness",
                "status": "success" if ready_to_review else "running",
                "details": (
                    "The draft is complete and ready for form review."
                    if ready_to_review
                    else f"Still missing { _planning_missing_prompt(missing_fields) }."
                ),
            },
        ],
        last_question=clarification_question,
    )
    return {"response": response}


def _build_planning_chat_graph():
    graph = StateGraph(PlanningChatGraphState)
    graph.add_node("detect_mode", _planning_detect_mode_node)
    graph.add_node("direct_response", _planning_direct_response_node)
    graph.add_node("llm_analysis", _planning_analysis_node)
    graph.add_node("finalize_response", _planning_finalize_response_node)
    graph.add_edge(START, "detect_mode")
    graph.add_conditional_edges(
        "detect_mode",
        _planning_chat_route,
        {
            "direct": "direct_response",
            "analysis": "llm_analysis",
        },
    )
    graph.add_edge("direct_response", END)
    graph.add_edge("llm_analysis", "finalize_response")
    graph.add_edge("finalize_response", END)
    return graph.compile()


def _get_planning_chat_graph():
    global _planning_chat_graph
    if _planning_chat_graph is None:
        _planning_chat_graph = _build_planning_chat_graph()
    return _planning_chat_graph


def _planning_state_markdown(planning_state: dict) -> str:
    plans = planning_state.get("plans", [])
    if not plans:
        return (
            f"## {PLANNING_PRODUCT_NAME}\n"
            "No automation has been saved yet.\n\n"
            "Use natural language to describe a schedule or open the planning form."
        )

    lines = [
        "## Existing Plans",
    ]
    for plan in plans[:8]:
        lines.append(
            f"- **{plan.get('name') or 'Untitled plan'}** · {plan.get('status')} · "
            f"{(plan.get('trigger') or {}).get('kind', 'unknown')}"
        )
    if len(plans) > 8:
        lines.append(f"- ...and {len(plans) - 8} more plan(s)")
    return "\n".join(lines)


def _planning_trigger_label(trigger_context: dict) -> str:
    kind = trigger_context.get("kind") or "manual"
    if kind == "once":
        return "One-time schedule"
    if kind == "daily":
        return "Daily schedule"
    if kind == "weekly":
        return "Weekly schedule"
    if kind == "interval":
        return "Interval schedule"
    if kind == "clickhouse_watch":
        return "ClickHouse watch"
    if kind == "file_watch":
        return "File watcher"
    return "Manual run"


def _truncate_json(value: Any, max_chars: int = 2000) -> str:
    text = json.dumps(value, ensure_ascii=False, indent=2, default=str)
    return text if len(text) <= max_chars else text[: max_chars - 1] + "…"


def _build_trigger_context_markdown(trigger_context: dict) -> str:
    kind = trigger_context.get("kind") or "manual"
    lines = [
        f"Trigger kind: {kind}",
        f"Trigger label: {_planning_trigger_label(trigger_context)}",
    ]
    if kind == "clickhouse_watch":
        lines.append("Watch SQL:")
        lines.append(str(trigger_context.get("sql") or "").strip() or "N/A")
        preview = trigger_context.get("rows") or []
        if preview:
            lines.append("Watch result preview:")
            lines.append(_truncate_json(preview[:5]))
    elif kind == "file_watch":
        files = trigger_context.get("files") or []
        if files:
            lines.append("New files:")
            lines.extend(str(path) for path in files[:10])
    return "\n".join(lines)


def _app_llm_config(app_config: dict) -> dict:
    return {
        "llm_base_url": str(app_config.get("baseUrl") or "http://localhost:11434"),
        "llm_model": str(app_config.get("model") or "llama3"),
        "llm_provider": str(app_config.get("provider") or "ollama"),
        "llm_api_key": app_config.get("apiKey") or None,
        "disable_ssl_verification": _ssl_verification_disabled(app_config),
    }


def _app_clickhouse_config(app_config: dict) -> "ClickHouseConfig":
    disable_ssl_verification = _ssl_verification_disabled(app_config)
    return ClickHouseConfig(
        host=str(app_config.get("clickhouseHost") or "localhost"),
        port=int(app_config.get("clickhousePort") or 8123),
        database=str(app_config.get("clickhouseDatabase") or "default"),
        username=str(app_config.get("clickhouseUsername") or "default"),
        password=str(app_config.get("clickhousePassword") or ""),
        secure=bool(app_config.get("clickhouseSecure", False)),
        verify_ssl=_effective_verify_ssl(bool(app_config.get("clickhouseVerifySsl", True)), disable_ssl_verification),
        http_path=str(app_config.get("clickhouseHttpPath") or ""),
        query_limit=int(app_config.get("clickhouseQueryLimit") or 200),
    )


def _app_file_manager_config(app_config: dict) -> dict[str, Any]:
    return _normalize_file_manager_config(app_config.get("fileManagerConfig"))


async def _run_local_role_agent(
    role: str,
    plan: dict,
    trigger_context: dict,
    app_config: dict,
) -> dict[str, Any]:
    llm_config = _app_llm_config(app_config)
    system_prompt = PLANNING_ROLE_PROMPTS.get(role, PLANNING_ROLE_PROMPTS["manager"])
    user_prompt = f"""
You are executing a scheduled automation for RAGnarok.
Write the answer in English and keep it concise but useful.

Automation name: {plan.get("name") or "Untitled plan"}
Automation objective:
{plan.get("prompt") or ""}

Trigger context:
{_build_trigger_context_markdown(trigger_context)}
""".strip()

    content = await llm_chat(
        [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        llm_config["llm_base_url"],
        llm_config["llm_model"],
        llm_config["llm_provider"],
        llm_config["llm_api_key"],
    )
    return {"agent": role, "status": "success", "content": content}


async def _run_file_management_planning_agent(
    plan: dict,
    trigger_context: dict,
    app_config: dict,
) -> dict[str, Any]:
    llm_config = _app_llm_config(app_config)
    file_manager_config = _app_file_manager_config(app_config)

    request_text = (plan.get("prompt") or "").strip()
    if trigger_context.get("kind") == "file_watch":
        files = trigger_context.get("files") or []
        if files:
            request_text += (
                "\n\nAdditional context from the file watcher:\n"
                + "\n".join(str(path) for path in files[:10])
            )
    elif trigger_context.get("kind") == "clickhouse_watch":
        request_text += (
            "\n\nAdditional context from the ClickHouse watcher:\n"
            f"{_truncate_json(trigger_context.get('rows') or [])}"
        )

    scratchpad: list[dict[str, Any]] = []
    max_iterations = max(
        1,
        min(FILE_MANAGER_MAX_ITERATIONS, int(file_manager_config["maxIterations"])),
    )
    last_error = ""

    for _ in range(max_iterations):
        planned = await plan_file_manager_step(
            request_text,
            [],
            scratchpad,
            file_manager_config["basePath"],
            file_manager_config["systemPrompt"],
            llm_config["llm_base_url"],
            llm_config["llm_model"],
            llm_config["llm_provider"],
            llm_config["llm_api_key"],
        )

        action = planned.get("action")
        if action == "final":
            content = planned.get("final_answer") or "## Answer\nThe scheduled file-management task is complete."
            return {"agent": "file_management", "status": "success", "content": content}

        tool_name = planned.get("tool_name") or ""
        tool_input = dict(planned.get("tool_input") or {})
        if action != "tool" or tool_name not in FILE_MANAGER_TOOLS:
            last_error = "The file-management planner returned an invalid tool action."
            scratchpad.append({"type": "error", "error": last_error})
            continue

        if tool_name in FILE_MANAGER_CONFIRMATION_TOOLS and "confirmed" not in tool_input:
            tool_input["confirmed"] = False

        try:
            result = execute_file_manager_tool(tool_name, tool_input, file_manager_config["basePath"])
        except Exception as exc:
            last_error = str(exc)
            scratchpad.append(
                {
                    "type": "tool_error",
                    "tool": tool_name,
                    "input": tool_input,
                    "error": last_error,
                }
            )
            continue

        scratchpad.append(
            {
                "type": "tool_result",
                "tool": tool_name,
                "input": tool_input,
                "summary": result.get("summary") or "",
                "preview": result.get("preview") or "",
            }
        )

        if result.get("requires_confirmation"):
            preview = result.get("preview") or ""
            content = (
                "## Answer\n"
                "The scheduled file-management task stopped because it requires explicit user confirmation.\n\n"
                "## Reasoning\n"
                f"{result.get('summary') or 'A destructive or overwrite action was requested.'}"
            )
            if preview:
                content += f"\n\n## Preview\n{preview}"
            return {"agent": "file_management", "status": "error", "content": content}

    content = "## Answer\nThe scheduled file-management task reached its iteration limit."
    if last_error:
        content += f"\n\n```text\n{last_error}\n```"
    return {"agent": "file_management", "status": "error", "content": content}


async def _run_clickhouse_planning_agent(
    plan: dict,
    trigger_context: dict,
    app_config: dict,
) -> dict[str, Any]:
    clickhouse = _app_clickhouse_config(app_config)
    llm_config = _app_llm_config(app_config)
    request_text = (plan.get("prompt") or "").strip()
    if trigger_context.get("kind") == "clickhouse_watch":
        request_text += (
            "\n\nAdditional context from the trigger watch query:\n"
            f"{_truncate_json(trigger_context.get('rows') or [])}"
        )
    elif trigger_context.get("kind") == "file_watch":
        request_text += (
            "\n\nAdditional context from new files:\n"
            + "\n".join(str(path) for path in (trigger_context.get("files") or [])[:10])
        )

    tables = await list_clickhouse_tables(clickhouse)
    if not tables:
        raise ValueError("No ClickHouse tables are available for the scheduled ClickHouse agent.")

    table_analysis = await analyze_clickhouse_tables(
        request_text,
        tables,
        "No recent memory. This is a scheduled execution.",
        llm_config["llm_base_url"],
        llm_config["llm_model"],
        llm_config["llm_provider"],
        llm_config["llm_api_key"],
    )
    matched_selected = match_available_options([table_analysis.get("selected_table", "")], tables)
    matched_candidates = match_available_options(table_analysis.get("table_candidates", []), tables)
    selected_table = (
        (matched_selected[0] if matched_selected else None)
        or (matched_candidates[0] if matched_candidates else None)
        or tables[0]
    )

    schema = await describe_clickhouse_table(clickhouse, selected_table)
    if not schema:
        raise ValueError(f"Table '{selected_table}' has no readable schema.")

    analysis = await analyze_clickhouse_schema(
        request_text,
        selected_table,
        schema,
        "No recent memory. This is a scheduled execution.",
        llm_config["llm_base_url"],
        llm_config["llm_model"],
        llm_config["llm_provider"],
        llm_config["llm_api_key"],
    )
    candidate_fields = match_schema_columns(analysis["field_candidates"], schema)
    date_fields = match_schema_columns(analysis["date_candidates"], schema) or find_date_columns(schema)
    selected_field = candidate_fields[0] if candidate_fields else None
    selected_date_field = date_fields[0] if date_fields else None

    generated = await generate_clickhouse_sql(
        request_text,
        selected_table,
        schema,
        selected_field,
        selected_date_field,
        "No recent memory. This is a scheduled execution.",
        llm_config["llm_base_url"],
        llm_config["llm_model"],
        llm_config["llm_provider"],
        llm_config["llm_api_key"],
        clickhouse.query_limit,
    )

    sql = generated["sql"]
    if not is_safe_read_only_sql(sql):
        raise ValueError("Scheduled ClickHouse SQL was rejected because it is not read-only.")

    try:
        await execute_clickhouse_sql(clickhouse, f"EXPLAIN SYNTAX {sql}", readonly=False, json_format=False)
        result = await execute_clickhouse_sql(clickhouse, sql)
    except Exception as first_error:
        repaired = await generate_clickhouse_sql(
            request_text,
            selected_table,
            schema,
            selected_field,
            selected_date_field,
            "No recent memory. This is a scheduled execution.",
            llm_config["llm_base_url"],
            llm_config["llm_model"],
            llm_config["llm_provider"],
            llm_config["llm_api_key"],
            clickhouse.query_limit,
            error_feedback=str(first_error),
        )
        sql = repaired["sql"]
        if not is_safe_read_only_sql(sql):
            raise ValueError("Scheduled ClickHouse SQL repair was rejected because it is not read-only.")
        generated["reasoning"] = repaired["reasoning"] or generated["reasoning"]
        result = await execute_clickhouse_sql(clickhouse, sql)

    result_summary = await summarize_clickhouse_result(
        request_text,
        sql,
        generated["reasoning"],
        result.get("data", []),
        "No recent memory. This is a scheduled execution.",
        llm_config["llm_base_url"],
        llm_config["llm_model"],
        llm_config["llm_provider"],
        llm_config["llm_api_key"],
    )
    content = build_clickhouse_response_markdown(result_summary, [sql])
    return {"agent": "clickhouse_query", "status": "success", "content": content}


async def _summarize_planning_outputs(
    plan: dict,
    outputs: list[dict[str, Any]],
    trigger_context: dict,
    app_config: dict,
) -> str:
    llm_config = _app_llm_config(app_config)
    return await _llm_chat_from_prompt_template(
        PLANNING_RUN_SUMMARY_PROMPT,
        llm_base_url=llm_config["llm_base_url"],
        llm_model=llm_config["llm_model"],
        llm_provider=llm_config["llm_provider"],
        llm_api_key=llm_config["llm_api_key"],
        disable_ssl_verification=bool(llm_config.get("disable_ssl_verification", False)),
        plan_name=plan.get("name") or "Untitled plan",
        plan_prompt=plan.get("prompt") or "",
        trigger_markdown=_build_trigger_context_markdown(trigger_context),
        outputs_json=_truncate_json(outputs, max_chars=5000),
    )


def _clickhouse_watch_metric(result: dict[str, Any]) -> Optional[float]:
    rows = result.get("data") or []
    if not rows:
        return 0.0
    first_row = rows[0]
    if isinstance(first_row, dict):
        for value in first_row.values():
            numeric = normalize_chart_value(value)
            if numeric is not None:
                return numeric
    return float(len(rows))


async def _evaluate_clickhouse_watch(plan: dict, app_config: dict, now_dt: datetime) -> Optional[dict[str, Any]]:
    trigger = plan.get("trigger") or {}
    runtime = plan.get("runtime") or _default_planning_runtime()
    last_checked = _parse_iso_datetime(runtime.get("lastCheckedAt"))
    poll_minutes = max(1, int(trigger.get("pollMinutes") or 5))
    if last_checked and last_checked + timedelta(minutes=poll_minutes) > now_dt:
        return None

    sql = clean_sql_text(trigger.get("watchSql") or "")
    runtime["lastCheckedAt"] = _to_iso_datetime(now_dt)
    if not sql or not is_safe_read_only_sql(sql):
        return None

    clickhouse = _app_clickhouse_config(app_config)
    result = await execute_clickhouse_sql(clickhouse, sql)
    rows = result.get("data", [])[:10]
    fingerprint = hashlib.sha256(
        json.dumps(rows, ensure_ascii=False, sort_keys=True, default=str).encode("utf-8")
    ).hexdigest()
    metric = _clickhouse_watch_metric(result)

    previous_fingerprint = str(runtime.get("lastSeenFingerprint") or "")
    previous_metric = runtime.get("lastSeenMetric")
    runtime["lastSeenFingerprint"] = fingerprint
    runtime["lastSeenMetric"] = metric
    plan["runtime"] = runtime

    watch_mode = trigger.get("watchMode") or "result_changes"
    if previous_fingerprint == "" and previous_metric is None:
        return None
    if watch_mode == "returns_rows":
        if rows and previous_fingerprint != fingerprint:
            return {"kind": "clickhouse_watch", "sql": sql, "rows": rows}
        return None
    if watch_mode == "count_increases":
        if (
            metric is not None
            and isinstance(previous_metric, (int, float))
            and metric > previous_metric
        ):
            return {"kind": "clickhouse_watch", "sql": sql, "rows": rows, "metric": metric}
        return None
    if previous_fingerprint != fingerprint:
        return {"kind": "clickhouse_watch", "sql": sql, "rows": rows}
    return None


def _scan_directory_files(directory: str, pattern: str, recursive: bool) -> list[str]:
    root = Path(directory).expanduser()
    if not root.exists() or not root.is_dir():
        return []

    matched: list[str] = []
    iterator = root.rglob("*") if recursive else root.glob("*")
    for path in iterator:
        if not path.is_file():
            continue
        relative_name = str(path.relative_to(root))
        if fnmatch.fnmatch(relative_name, pattern) or fnmatch.fnmatch(path.name, pattern):
            matched.append(str(path.resolve()))
        if len(matched) >= PLANNER_MAX_KNOWN_FILES:
            break
    return sorted(matched)


def _evaluate_file_watch(plan: dict, now_dt: datetime) -> Optional[dict[str, Any]]:
    trigger = plan.get("trigger") or {}
    runtime = plan.get("runtime") or _default_planning_runtime()
    last_checked = _parse_iso_datetime(runtime.get("lastCheckedAt"))
    poll_minutes = max(1, int(trigger.get("pollMinutes") or 5))
    if last_checked and last_checked + timedelta(minutes=poll_minutes) > now_dt:
        return None

    runtime["lastCheckedAt"] = _to_iso_datetime(now_dt)
    current_files = _scan_directory_files(
        str(trigger.get("directory") or ""),
        str(trigger.get("pattern") or "*"),
        bool(trigger.get("recursive", False)),
    )
    previous_files = set(runtime.get("knownFiles") or [])
    runtime["knownFiles"] = current_files[:PLANNER_MAX_KNOWN_FILES]
    plan["runtime"] = runtime

    if not previous_files:
        return None

    new_files = [path for path in current_files if path not in previous_files]
    if new_files:
        return {"kind": "file_watch", "files": new_files[:20]}
    return None


async def _due_trigger_context(plan: dict, app_config: dict, now_dt: datetime) -> Optional[dict[str, Any]]:
    trigger = plan.get("trigger") or {}
    kind = trigger.get("kind")
    if plan.get("status") != "active":
        return None

    if kind in {"once", "daily", "weekly", "interval"}:
        next_run = _parse_iso_datetime(plan.get("nextRunAt"))
        if next_run and next_run <= now_dt:
            return {"kind": kind}
        return None
    if kind == "clickhouse_watch":
        return await _evaluate_clickhouse_watch(plan, app_config, now_dt)
    if kind == "file_watch":
        return _evaluate_file_watch(plan, now_dt)
    return None


def _planning_execution_route(state: PlanningExecutionGraphState) -> str:
    pending_agents = state.get("pending_agents") or []
    return "run_agent" if pending_agents else "summarize"


def _planning_execution_init_node(state: PlanningExecutionGraphState) -> PlanningExecutionGraphState:
    return {
        "pending_agents": list((state.get("plan") or {}).get("agents") or []),
        "outputs": [],
        "overall_status": "success",
    }


async def _planning_execution_run_agent_node(
    state: PlanningExecutionGraphState,
) -> PlanningExecutionGraphState:
    pending_agents = list(state.get("pending_agents") or [])
    outputs = list(state.get("outputs") or [])
    overall_status = str(state.get("overall_status") or "success")
    if not pending_agents:
        return {
            "pending_agents": pending_agents,
            "outputs": outputs,
            "overall_status": overall_status,
        }

    agent = pending_agents.pop(0)
    plan = dict(state.get("plan") or {})
    trigger_context = dict(state.get("trigger_context") or {"kind": "manual"})
    app_config = dict(state.get("app_config") or {})

    try:
        if agent == "clickhouse_query":
            output = await _run_clickhouse_planning_agent(plan, trigger_context, app_config)
        elif agent == "file_management":
            output = await _run_file_management_planning_agent(plan, trigger_context, app_config)
        else:
            output = await _run_local_role_agent(agent, plan, trigger_context, app_config)
    except Exception as agent_error:
        output = {
            "agent": agent,
            "status": "error",
            "content": f"Agent execution failed: {agent_error}",
        }
        overall_status = "error"

    outputs.append(output)
    if output.get("status") == "error":
        overall_status = "error"

    return {
        "pending_agents": pending_agents,
        "outputs": outputs,
        "overall_status": overall_status,
    }


async def _planning_execution_summarize_node(
    state: PlanningExecutionGraphState,
) -> PlanningExecutionGraphState:
    plan = dict(state.get("plan") or {})
    trigger_context = dict(state.get("trigger_context") or {"kind": "manual"})
    app_config = dict(state.get("app_config") or {})
    outputs = list(state.get("outputs") or [])
    overall_status = str(state.get("overall_status") or "success")

    try:
        summary = await _summarize_planning_outputs(plan, outputs, trigger_context, app_config)
    except Exception as summary_error:
        fallback_lines = [
            "## Summary",
            "The automation finished, but the final narrative summary could not be generated.",
            "",
            "## Trigger",
            _planning_trigger_label(trigger_context),
            "",
            "## Agent Outputs",
        ]
        if outputs:
            fallback_lines.extend(
                f"- **{output.get('agent') or 'agent'}** · {output.get('status') or 'unknown'}"
                for output in outputs
            )
        else:
            fallback_lines.append("- No agent output was captured.")
            overall_status = "error"
        fallback_lines.extend(
            [
                "",
                "<details><summary>Expand technical details</summary>",
                "",
                "```text",
                str(summary_error),
                "```",
                "",
                "</details>",
            ]
        )
        summary = "\n".join(fallback_lines)

    return {
        "summary": summary,
        "overall_status": overall_status,
    }


def _build_planning_execution_graph():
    graph = StateGraph(PlanningExecutionGraphState)
    graph.add_node("initialize", _planning_execution_init_node)
    graph.add_node("run_agent", _planning_execution_run_agent_node)
    graph.add_node("summarize", _planning_execution_summarize_node)
    graph.add_edge(START, "initialize")
    graph.add_conditional_edges(
        "initialize",
        _planning_execution_route,
        {
            "run_agent": "run_agent",
            "summarize": "summarize",
        },
    )
    graph.add_conditional_edges(
        "run_agent",
        _planning_execution_route,
        {
            "run_agent": "run_agent",
            "summarize": "summarize",
        },
    )
    graph.add_edge("summarize", END)
    return graph.compile()


def _get_planning_execution_graph():
    global _planning_execution_graph
    if _planning_execution_graph is None:
        _planning_execution_graph = _build_planning_execution_graph()
    return _planning_execution_graph


async def execute_planning_run(
    plan_id: str,
    trigger_context: Optional[dict[str, Any]] = None,
    manual: bool = False,
) -> dict[str, Any]:
    initial_state = await read_db_state()
    planning = _planning_state_from_db(initial_state)
    plan = next((item for item in planning["plans"] if item.get("id") == plan_id), None)
    if not plan:
        raise HTTPException(status_code=404, detail="Planning job not found.")
    if plan.get("status") != "active" and not manual:
        raise HTTPException(status_code=400, detail="The selected planning job is paused.")

    app_config = initial_state.get("config") or {}
    now_dt = datetime.now(timezone.utc)
    context = trigger_context or {"kind": "manual"}
    run_id = uuid.uuid4().hex
    run_record = {
        "id": run_id,
        "planId": plan.get("id"),
        "planName": plan.get("name") or "Untitled plan",
        "triggerKind": context.get("kind") or "manual",
        "triggerLabel": _planning_trigger_label(context),
        "startedAt": _to_iso_datetime(now_dt),
        "finishedAt": None,
        "status": "running",
        "summary": "",
        "outputs": [],
    }

    planning["runs"].insert(0, run_record)
    planning["runs"] = planning["runs"][:PLANNER_MAX_RUNS]
    plan["lastStatus"] = "running"
    initial_state["planning"] = planning
    await write_db_state(initial_state)

    outputs: list[dict[str, Any]] = []
    overall_status = "success"
    summary = ""
    try:
        graph_result = await _get_planning_execution_graph().ainvoke(
            {
                "plan": plan,
                "trigger_context": context,
                "app_config": app_config,
            }
        )
        outputs = list(graph_result.get("outputs") or [])
        overall_status = str(graph_result.get("overall_status") or "success")
        summary = str(graph_result.get("summary") or "").strip()
        if not summary:
            summary = await _summarize_planning_outputs(plan, outputs, context, app_config)
    except Exception as execution_error:
        overall_status = "error"
        summary = (
            "## Summary\n"
            "The automation failed.\n\n"
            f"## Trigger\n{_planning_trigger_label(context)}\n\n"
            "## Agent Outputs\n"
            f"{execution_error}"
        )

    latest_state = await read_db_state()
    latest_planning = _planning_state_from_db(latest_state)
    latest_plan = next((item for item in latest_planning["plans"] if item.get("id") == plan_id), None)
    latest_run = next((item for item in latest_planning["runs"] if item.get("id") == run_id), None)
    if not latest_plan or not latest_run:
        return {"status": overall_status, "summary": summary, "outputs": outputs}

    finished_at = _to_iso_datetime(datetime.now(timezone.utc))
    latest_run["finishedAt"] = finished_at
    latest_run["status"] = overall_status
    latest_run["summary"] = summary
    latest_run["outputs"] = outputs

    latest_plan["lastRunAt"] = finished_at
    latest_plan["lastStatus"] = overall_status
    latest_plan["lastSummary"] = summary
    if (latest_plan.get("trigger") or {}).get("kind") == "once":
        latest_plan["status"] = "paused"
        latest_plan["nextRunAt"] = None
    else:
        latest_plan["nextRunAt"] = _compute_plan_next_run_at(latest_plan, datetime.now(timezone.utc))

    latest_state["planning"] = latest_planning
    await write_db_state(latest_state)
    return {"status": overall_status, "summary": summary, "outputs": outputs}


async def process_planning_jobs() -> None:
    state = await read_db_state()
    original_snapshot = json.dumps(
        _normalize_planning_state(state.get("planning")),
        sort_keys=True,
        ensure_ascii=False,
        default=str,
    )
    planning = _planning_state_from_db(state)
    now_dt = datetime.now(timezone.utc)
    app_config = state.get("config") or {}
    due_runs: list[tuple[str, dict[str, Any]]] = []

    for plan in planning["plans"]:
        context = await _due_trigger_context(plan, app_config, now_dt)
        if context:
            due_runs.append((plan.get("id"), context))

    updated_snapshot = json.dumps(planning, sort_keys=True, ensure_ascii=False, default=str)
    if updated_snapshot != original_snapshot:
        state["planning"] = planning
        await write_db_state(state)

    for plan_id, trigger_context in due_runs:
        try:
            await execute_planning_run(plan_id, trigger_context=trigger_context, manual=False)
        except Exception:
            continue


async def planning_scheduler_loop(stop_event: asyncio.Event) -> None:
    while not stop_event.is_set():
        try:
            await process_planning_jobs()
        except Exception:
            pass
        try:
            await asyncio.wait_for(stop_event.wait(), timeout=PLANNER_LOOP_INTERVAL_SECONDS)
        except asyncio.TimeoutError:
            continue

# ── Pydantic models ───────────────────────────────────────────────────────────

class MCPTestRequest(BaseModel):
    url: str
    disable_ssl_verification: bool = False


class MCPChatRequest(BaseModel):
    message: str
    history: list[dict] = []
    mcp_url: str
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"
    system_prompt: str = ""
    disable_ssl_verification: bool = False


class LlmModelsRequest(BaseModel):
    provider: str = "ollama"
    base_url: str = "http://localhost:11434"
    api_key: Optional[str] = None
    disable_ssl_verification: bool = False


class EmbeddingModelsRequest(BaseModel):
    base_url: str = "http://localhost:11434/v1"
    api_key: Optional[str] = None
    disable_ssl_verification: bool = False


class OSConfig(BaseModel):
    url: str
    index: str = "rag_documents"
    username: Optional[str] = None
    password: Optional[str] = None


class ClickHouseConfig(BaseModel):
    host: str = "localhost"
    port: int = 8123
    database: str = "default"
    username: str = "default"
    password: str = ""
    secure: bool = False
    verify_ssl: bool = True
    http_path: str = ""
    query_limit: int = 200


class OracleConnectionConfig(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    id: str = "oracle_default"
    label: str = "Default Oracle"
    host: str = "localhost"
    port: int = 1521
    service_name: str = Field(
        default="",
        validation_alias=AliasChoices("service_name", "serviceName"),
        serialization_alias="serviceName",
    )
    sid: str = ""
    dsn: str = ""
    username: str = ""
    password: str = ""


class OracleAnalystConfigModel(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    connection_id: str = Field(
        default="oracle_default",
        validation_alias=AliasChoices("connection_id", "connectionId"),
        serialization_alias="connectionId",
    )
    row_limit: int = Field(
        default=1000,
        validation_alias=AliasChoices("row_limit", "rowLimit"),
        serialization_alias="rowLimit",
    )
    max_retries: int = Field(
        default=3,
        validation_alias=AliasChoices("max_retries", "maxRetries"),
        serialization_alias="maxRetries",
    )
    max_iterations: int = Field(
        default=8,
        validation_alias=AliasChoices("max_iterations", "maxIterations"),
        serialization_alias="maxIterations",
    )
    toolkit_id: str = Field(
        default="",
        validation_alias=AliasChoices("toolkit_id", "toolkitId"),
        serialization_alias="toolkitId",
    )
    system_prompt: str = Field(
        default=DEFAULT_APP_CONFIG["oracleAnalystConfig"]["systemPrompt"],
        validation_alias=AliasChoices("system_prompt", "systemPrompt"),
        serialization_alias="systemPrompt",
    )


class TestConnectionRequest(BaseModel):
    url: str
    username: Optional[str] = None
    password: Optional[str] = None


class SetupIndexRequest(BaseModel):
    opensearch: OSConfig
    embedding_dimension: int = 768


class IngestRequest(BaseModel):
    text: str
    doc_name: str
    opensearch: OSConfig
    embedding_base_url: str = "http://localhost:11434/v1"
    embedding_api_key: Optional[str] = None
    embedding_model: str = "nomic-embed-text"
    embedding_verify_ssl: bool = True
    chunk_size: int = 200
    chunk_overlap: int = 2
    embedding_dimension: int = 768


class ChatRequest(BaseModel):
    message: str
    history: list[dict] = []
    opensearch: OSConfig
    embedding_base_url: str = "http://localhost:11434/v1"
    embedding_api_key: Optional[str] = None
    embedding_model: str = "nomic-embed-text"
    embedding_verify_ssl: bool = True
    knn_neighbors: int = 10
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"
    disable_ssl_verification: bool = False


class ClickHouseTestRequest(BaseModel):
    clickhouse: ClickHouseConfig


class OracleTestRequest(BaseModel):
    connection: OracleConnectionConfig


class ClickHouseAgentState(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    stage: str = "idle"
    pending_request: str = ""
    available_tables: list[str] = Field(default_factory=list)
    selected_table: Optional[str] = None
    table_schema: list[dict] = Field(
        default_factory=list,
        validation_alias=AliasChoices("schema", "table_schema"),
        serialization_alias="schema",
    )
    candidate_fields: list[str] = Field(default_factory=list)
    date_fields: list[str] = Field(default_factory=list)
    selected_field: Optional[str] = None
    selected_date_field: Optional[str] = None
    clarification_prompt: str = ""
    clarification_options: list[str] = Field(default_factory=list)
    last_sql: str = ""
    last_result_meta: list[dict] = Field(default_factory=list)
    last_result_rows: list[dict] = Field(default_factory=list)
    chart_requested: bool = False
    chart_suggested: bool = False
    chart_offer_options: list[str] = Field(default_factory=list)
    chart_x_options: list[str] = Field(default_factory=list)
    chart_y_options: list[str] = Field(default_factory=list)
    chart_type_options: list[str] = Field(default_factory=list)
    selected_chart_x: Optional[str] = None
    selected_chart_y: Optional[str] = None
    selected_chart_type: Optional[str] = None


class ClickHouseAgentRequest(BaseModel):
    message: str
    history: list[dict] = []
    clickhouse: ClickHouseConfig
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"
    disable_ssl_verification: bool = False
    agent_state: ClickHouseAgentState = Field(default_factory=ClickHouseAgentState)


class ManagerAgentStateModel(BaseModel):
    active_delegate: Optional[str] = None
    last_routing_reason: str = ""
    last_delegate_label: str = ""
    pending_pipeline: Optional[dict] = None


class FileManagerAgentConfigModel(BaseModel):
    base_path: str = ""
    max_iterations: int = 10
    system_prompt: str = DEFAULT_APP_CONFIG["fileManagerConfig"]["systemPrompt"]


class FileManagerAgentStateModel(BaseModel):
    pending_confirmation: Optional[dict] = None
    last_tool_result: str = ""
    last_visited_path: str = ""


class FileManagerAgentRequest(BaseModel):
    message: str
    history: list[dict] = []
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"
    disable_ssl_verification: bool = False
    agent_state: FileManagerAgentStateModel = Field(default_factory=FileManagerAgentStateModel)
    file_manager_config: FileManagerAgentConfigModel = Field(default_factory=FileManagerAgentConfigModel)


class FileManagerPathTestRequest(BaseModel):
    file_manager_config: FileManagerAgentConfigModel = Field(default_factory=FileManagerAgentConfigModel)


class PdfCreatorAgentStateModel(BaseModel):
    stage: str = "idle"
    pending_document: Optional[dict] = None
    pending_confirmation: Optional[dict] = None
    last_output_path: str = ""
    last_title: str = ""


class PdfCreatorAgentRequest(BaseModel):
    message: str
    history: list[dict] = []
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"
    disable_ssl_verification: bool = False
    agent_state: PdfCreatorAgentStateModel = Field(default_factory=PdfCreatorAgentStateModel)
    file_manager_config: FileManagerAgentConfigModel = Field(default_factory=FileManagerAgentConfigModel)


class OracleAnalystAgentStateModel(BaseModel):
    stage: str = "idle"
    pending_request: str = ""
    available_tables: list[str] = Field(default_factory=list)
    selected_table: Optional[str] = None
    schema_info: list[dict] = Field(default_factory=list)
    clarification_prompt: str = ""
    clarification_options: list[str] = Field(default_factory=list)
    last_sql: str = ""
    last_result_meta: list[dict] = Field(default_factory=list)
    last_result_rows: list[dict] = Field(default_factory=list)
    final_answer: str = ""
    action_log: list[str] = Field(default_factory=list)
    last_error: str = ""


class OracleAnalystAgentRequest(BaseModel):
    message: str
    history: list[dict] = []
    oracle_connections: list[OracleConnectionConfig] = Field(default_factory=list)
    oracle_analyst_config: OracleAnalystConfigModel = Field(default_factory=OracleAnalystConfigModel)
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"
    disable_ssl_verification: bool = False
    agent_state: OracleAnalystAgentStateModel = Field(default_factory=OracleAnalystAgentStateModel)


class DataQualityAgentStateModel(BaseModel):
    stage: str = "idle"
    table: Optional[str] = None
    columns: list[str] = Field(default_factory=list)
    sample_size: int = DATA_QUALITY_DEFAULT_SAMPLE_SIZE
    row_filter: str = ""
    time_column: Optional[str] = None
    db_type: str = "clickhouse"
    schema_info: list[dict] = Field(default_factory=list)
    column_stats: dict = Field(default_factory=dict)
    volumetric_stats: Optional[dict] = None
    llm_analysis: str = ""
    final_answer: str = ""
    agent_id: str = "data_quality_tables"
    session_id: str = ""
    last_error: str = ""
    available_tables: list[str] = Field(default_factory=list)
    available_columns: list[str] = Field(default_factory=list)
    date_columns: list[str] = Field(default_factory=list)


class DataAnalystAgentStateModel(BaseModel):
    stage: str = "idle"
    pending_request: str = ""
    available_tables: list[str] = Field(default_factory=list)
    selected_table: Optional[str] = None
    table_schema: list[dict] = Field(default_factory=list)
    clarification_prompt: str = ""
    clarification_options: list[str] = Field(default_factory=list)
    analysis_goal: str = ""
    analysis_plan: list[str] = Field(default_factory=list)
    success_criteria: list[str] = Field(default_factory=list)
    latest_review: str = ""
    last_sqls: list[str] = Field(default_factory=list)
    last_result_meta: list[dict] = Field(default_factory=list)
    last_result_rows: list[dict] = Field(default_factory=list)
    final_answer: str = ""
    last_error: str = ""
    last_export_path: str = ""
    knowledge_hits: list[dict] = Field(default_factory=list)


class AutoMlAgentStateModel(BaseModel):
    stage: str = "idle"
    pending_request: str = ""
    available_tables: list[str] = Field(default_factory=list)
    selected_table: Optional[str] = None
    schema_info: list[dict] = Field(default_factory=list)
    target_column: Optional[str] = None
    row_filter: str = ""
    sample_row_limit: int = 1000
    feature_columns: list[str] = Field(default_factory=list)
    clarification_prompt: str = ""
    clarification_options: list[str] = Field(default_factory=list)
    comparison_rows: list[dict] = Field(default_factory=list)
    problem_type: str = ""
    recommended_model: str = ""
    final_answer: str = ""
    last_error: str = ""


class DataCleanerAgentStateModel(BaseModel):
    stage: str = "idle"
    pending_request: str = ""
    available_tables: list[str] = Field(default_factory=list)
    selected_table: Optional[str] = None
    schema_info: list[dict] = Field(default_factory=list)
    row_filter: str = ""
    clarification_prompt: str = ""
    clarification_options: list[str] = Field(default_factory=list)
    findings: list[dict] = Field(default_factory=list)
    correction_scripts: list[dict] = Field(default_factory=list)
    final_answer: str = ""
    last_error: str = ""


class AnonymizerAgentStateModel(BaseModel):
    stage: str = "idle"
    pending_request: str = ""
    available_tables: list[str] = Field(default_factory=list)
    selected_table: Optional[str] = None
    schema_info: list[dict] = Field(default_factory=list)
    row_filter: str = ""
    clarification_prompt: str = ""
    clarification_options: list[str] = Field(default_factory=list)
    pii_findings: list[dict] = Field(default_factory=list)
    masking_scripts: list[dict] = Field(default_factory=list)
    final_answer: str = ""
    last_error: str = ""


class DataQualityAgentRequest(BaseModel):
    message: str
    history: list[dict] = []
    clickhouse: ClickHouseConfig
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"
    disable_ssl_verification: bool = False
    agent_state: DataQualityAgentStateModel = Field(default_factory=DataQualityAgentStateModel)


class DataAnalystAgentRequest(BaseModel):
    message: str
    history: list[dict] = []
    clickhouse: ClickHouseConfig
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"
    disable_ssl_verification: bool = False
    max_steps: int = DATA_ANALYST_DEFAULT_MAX_STEPS
    agent_state: DataAnalystAgentStateModel = Field(default_factory=DataAnalystAgentStateModel)


class AutoMlAgentRequest(BaseModel):
    message: str
    history: list[dict] = []
    clickhouse: ClickHouseConfig
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"
    disable_ssl_verification: bool = False
    agent_state: AutoMlAgentStateModel = Field(default_factory=AutoMlAgentStateModel)


class DataCleanerAgentRequest(BaseModel):
    message: str
    history: list[dict] = []
    clickhouse: ClickHouseConfig
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"
    disable_ssl_verification: bool = False
    agent_state: DataCleanerAgentStateModel = Field(default_factory=DataCleanerAgentStateModel)


class AnonymizerAgentRequest(BaseModel):
    message: str
    history: list[dict] = []
    clickhouse: ClickHouseConfig
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"
    disable_ssl_verification: bool = False
    agent_state: AnonymizerAgentStateModel = Field(default_factory=AnonymizerAgentStateModel)


class CustomAgentConfigModel(BaseModel):
    id: str = ""
    title: str = "Custom Agent"
    description: str = ""
    python_code: str = ""
    system_prompt: str = ""
    manager_routing_hint: str = ""
    status: str = "draft"
    status_message: str = ""
    enabled: bool = False
    badge_color: str = "zinc"


class CustomAgentStateModel(BaseModel):
    selected_agent_id: Optional[str] = None
    final_answer: str = ""
    last_error: str = ""


class CustomAgentAnalysisRequest(BaseModel):
    title: str = ""
    description: str = ""
    python_code: str
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"
    disable_ssl_verification: bool = False


class CustomAgentChatRequest(BaseModel):
    message: str
    history: list[dict] = []
    custom_agent: CustomAgentConfigModel
    agent_state: CustomAgentStateModel = Field(default_factory=CustomAgentStateModel)
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"
    disable_ssl_verification: bool = False


class AutoMlFilterSuggestionRequest(BaseModel):
    table: str
    target_column: Optional[str] = None
    goal: str = ""
    notes: str = ""
    schema_info: list[dict] = Field(default_factory=list)
    clickhouse: ClickHouseConfig
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"


class ClickHouseGuideMetadataRequest(BaseModel):
    clickhouse: ClickHouseConfig
    table: Optional[str] = None


class ClickHouseScopeSuggestionRequest(BaseModel):
    table: str
    agent_kind: str = "generic"
    target_column: Optional[str] = None
    goal: str = ""
    notes: str = ""
    schema_info: list[dict] = Field(default_factory=list)
    clickhouse: ClickHouseConfig
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"


class SqlDraftPreviewRequest(BaseModel):
    sql: str
    engine: Literal["clickhouse", "oracle"] = "clickhouse"
    row_limit: int = 1000
    clickhouse: ClickHouseConfig = Field(default_factory=ClickHouseConfig)
    oracle_connections: list[OracleConnectionConfig] = Field(default_factory=list)
    oracle_analyst_config: OracleAnalystConfigModel = Field(default_factory=OracleAnalystConfigModel)


class ManagerAgentRequest(BaseModel):
    message: str
    history: list[dict] = []
    clickhouse: ClickHouseConfig = Field(default_factory=ClickHouseConfig)
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"
    disable_ssl_verification: bool = False
    system_prompt: str = DEFAULT_APP_CONFIG["systemPrompt"]
    manager_state: ManagerAgentStateModel = Field(default_factory=ManagerAgentStateModel)
    clickhouse_state: ClickHouseAgentState = Field(default_factory=ClickHouseAgentState)
    file_manager_state: FileManagerAgentStateModel = Field(default_factory=FileManagerAgentStateModel)
    pdf_creator_state: PdfCreatorAgentStateModel = Field(default_factory=PdfCreatorAgentStateModel)
    oracle_analyst_state: OracleAnalystAgentStateModel = Field(default_factory=OracleAnalystAgentStateModel)
    data_analyst_state: DataAnalystAgentStateModel = Field(default_factory=DataAnalystAgentStateModel)
    auto_ml_state: AutoMlAgentStateModel = Field(default_factory=AutoMlAgentStateModel)
    data_cleaner_state: DataCleanerAgentStateModel = Field(default_factory=DataCleanerAgentStateModel)
    anonymizer_state: AnonymizerAgentStateModel = Field(default_factory=AnonymizerAgentStateModel)
    custom_agent_state: CustomAgentStateModel = Field(default_factory=CustomAgentStateModel)
    custom_agents: list[CustomAgentConfigModel] = Field(default_factory=list)
    oracle_connections: list[OracleConnectionConfig] = Field(default_factory=list)
    oracle_analyst_config: OracleAnalystConfigModel = Field(default_factory=OracleAnalystConfigModel)
    file_manager_config: FileManagerAgentConfigModel = Field(default_factory=FileManagerAgentConfigModel)


class PlanningAgentStateModel(BaseModel):
    draft: dict = Field(default_factory=lambda: _default_planning_draft("UTC"))
    missing_fields: list[str] = Field(default_factory=list)
    last_question: str = ""
    ready_to_review: bool = False


class PlanningChatRequest(BaseModel):
    message: str
    history: list[dict] = []
    llm_base_url: str = "http://localhost:11434"
    llm_model: str = "llama3"
    llm_api_key: Optional[str] = None
    llm_provider: str = "ollama"
    disable_ssl_verification: bool = False
    agent_state: PlanningAgentStateModel = Field(default_factory=PlanningAgentStateModel)


class PlanningPlanRequest(BaseModel):
    plan: dict


class PlanningPlanStatusRequest(BaseModel):
    status: str


class EmbeddingTestRequest(BaseModel):
    embedding_base_url: str
    embedding_model: str
    embedding_api_key: Optional[str] = None
    embedding_verify_ssl: bool = True
    opensearch: Optional[OSConfig] = None


class PersistedPreferences(BaseModel):
    darkMode: bool = False
    currentConversationId: Optional[str] = None
    workflow: str = "LLM"
    agentRole: str = "manager"
    selectedMcpToolId: str = ""
    selectedCustomAgentId: str = ""
    page: str = "landing"


class PersistedStateRequest(BaseModel):
    config: dict = Field(default_factory=lambda: json.loads(json.dumps(DEFAULT_APP_CONFIG)))
    conversations: list[dict] = Field(default_factory=list)
    preferences: PersistedPreferences = Field(default_factory=PersistedPreferences)
    include_config: bool = False


# ── App state persistence endpoints ───────────────────────────────────────────

def _request_db_user_id(request: Request) -> str:
    return _normalize_db_user_id(request.headers.get("X-RAGnarok-User-Id"))


@app.get("/api/db/state")
async def get_app_state(request: Request):
    full_state = await read_db_state()
    user_id = _request_db_user_id(request)
    ensured_state, changed = _ensure_user_db_state(full_state, user_id)
    if changed:
        ensured_state = await write_db_state(ensured_state)
    return _db_state_for_user(ensured_state, user_id)


@app.put("/api/db/state")
async def save_app_state(req: PersistedStateRequest, request: Request):
    existing_state = await read_db_state()
    user_id = _request_db_user_id(request)
    payload, _ = _ensure_user_db_state(existing_state, user_id)
    payload["schemaVersion"] = payload.get("schemaVersion", 2)
    if req.include_config:
        payload["config"] = req.config
    payload.setdefault("users", {})
    payload["users"][user_id] = _normalize_user_db_state(
        {
            "conversations": req.conversations,
            "preferences": req.preferences.model_dump(),
            "updatedAt": _utc_now_iso(),
        }
    )
    payload["planning"] = existing_state.get("planning", _default_planning_state())
    saved = await write_db_state(payload)
    return _db_state_for_user(saved, user_id)


@app.get("/api/db/export")
async def export_app_state():
    state = await read_db_state()
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    return JSONResponse(
        content=state,
        headers={
            "Content-Disposition": f'attachment; filename="ragnarok-db-backup-{timestamp}.json"'
        },
    )


@app.post("/api/db/import")
async def import_app_state(payload: dict, request: Request):
    if not isinstance(payload, dict):
        raise HTTPException(status_code=400, detail="Invalid JSON payload for DB import.")
    saved = await write_db_state(payload)
    user_id = _request_db_user_id(request)
    ensured_state, changed = _ensure_user_db_state(saved, user_id)
    if changed:
        ensured_state = await write_db_state(ensured_state)
    return _db_state_for_user(ensured_state, user_id)


@app.get("/api/planning/state")
async def get_planning_state():
    state = await read_db_state()
    planning = _planning_state_from_db(state)
    if json.dumps(planning, sort_keys=True, ensure_ascii=False, default=str) != json.dumps(
        _normalize_planning_state(state.get("planning")),
        sort_keys=True,
        ensure_ascii=False,
        default=str,
    ):
        state["planning"] = planning
        await write_db_state(state)
    return planning


@app.post("/api/planning/plans")
async def upsert_planning_plan(req: PlanningPlanRequest):
    normalized_plan = _normalize_planning_plan(req.plan)
    if not normalized_plan.get("prompt"):
        raise HTTPException(status_code=400, detail="A planning job objective is required.")
    if not normalized_plan.get("agents"):
        raise HTTPException(status_code=400, detail="Select at least one existing agent.")

    trigger = normalized_plan.get("trigger") or {}
    kind = trigger.get("kind")
    if kind == "clickhouse_watch" and not is_safe_read_only_sql(trigger.get("watchSql") or ""):
        raise HTTPException(status_code=400, detail="The ClickHouse watch SQL must be a safe read-only query.")
    if kind == "file_watch" and not trigger.get("directory"):
        raise HTTPException(status_code=400, detail="A directory is required for file watch jobs.")

    state = await read_db_state()
    planning = _planning_state_from_db(state)
    now_iso = _utc_now_iso()
    existing_plan = next((plan for plan in planning["plans"] if plan.get("id") == normalized_plan["id"]), None)
    if existing_plan:
        normalized_plan["createdAt"] = existing_plan.get("createdAt") or now_iso
    else:
        normalized_plan["createdAt"] = now_iso
    normalized_plan["updatedAt"] = now_iso
    normalized_plan["nextRunAt"] = _compute_plan_next_run_at(normalized_plan, datetime.now(timezone.utc))
    normalized_plan["runtime"] = existing_plan.get("runtime") if existing_plan else _default_planning_runtime()
    if kind in {"clickhouse_watch", "file_watch"}:
        normalized_plan["nextRunAt"] = None

    if existing_plan:
        index = planning["plans"].index(existing_plan)
        planning["plans"][index] = normalized_plan
    else:
        planning["plans"].insert(0, normalized_plan)

    state["planning"] = planning
    await write_db_state(state)
    return planning


@app.post("/api/planning/plans/{plan_id}/status")
async def set_planning_plan_status(plan_id: str, req: PlanningPlanStatusRequest):
    if req.status not in {"active", "paused"}:
        raise HTTPException(status_code=400, detail="Invalid planning job status.")

    state = await read_db_state()
    planning = _planning_state_from_db(state)
    plan = next((item for item in planning["plans"] if item.get("id") == plan_id), None)
    if not plan:
        raise HTTPException(status_code=404, detail="Planning job not found.")

    plan["status"] = req.status
    plan["updatedAt"] = _utc_now_iso()
    plan["nextRunAt"] = _compute_plan_next_run_at(plan, datetime.now(timezone.utc))
    if (plan.get("trigger") or {}).get("kind") in {"clickhouse_watch", "file_watch"}:
        plan["nextRunAt"] = None

    state["planning"] = planning
    await write_db_state(state)
    return planning


@app.delete("/api/planning/plans/{plan_id}")
async def delete_planning_plan(plan_id: str):
    state = await read_db_state()
    planning = _planning_state_from_db(state)
    existing = next((item for item in planning["plans"] if item.get("id") == plan_id), None)
    if not existing:
        raise HTTPException(status_code=404, detail="Planning job not found.")

    planning["plans"] = [plan for plan in planning["plans"] if plan.get("id") != plan_id]
    state["planning"] = planning
    await write_db_state(state)
    return planning


@app.post("/api/planning/plans/{plan_id}/run")
async def run_planning_plan_now(plan_id: str):
    await execute_planning_run(plan_id, trigger_context={"kind": "manual"}, manual=True)
    state = await read_db_state()
    return _planning_state_from_db(state)


@app.post("/api/chat/crewai-planning")
async def chat_crewai_planning(req: PlanningChatRequest):
    user_message = (req.message or "").strip()
    timezone_name = str(
        ((req.agent_state.draft or {}).get("trigger") or {}).get("timezone") or "UTC"
    )
    current_draft = _merge_planning_draft(req.agent_state.draft, None, timezone_name)
    planning_state = _planning_state_from_db(await read_db_state())
    result = await _get_planning_chat_graph().ainvoke(
        {
            "user_message": user_message,
            "history": req.history,
            "current_draft": current_draft,
            "timezone_name": timezone_name,
            "planning_state": planning_state,
            "llm_config": {
                "llm_base_url": req.llm_base_url,
                "llm_model": req.llm_model,
                "llm_provider": req.llm_provider,
                "llm_api_key": req.llm_api_key,
                "disable_ssl_verification": req.disable_ssl_verification,
            },
        }
    )
    return result.get("response") or _planning_response_payload(
        current_draft,
        _validate_planning_draft(current_draft),
        f"## {PLANNING_PRODUCT_NAME}\nThe planner could not build a response.",
        [
            {
                "id": "planning-fallback",
                "title": "Fallback response",
                "status": "error",
                "details": "The planning graph returned an empty response.",
            }
        ],
    )


@app.post("/api/chat/file-manager-agent")
async def chat_file_manager_agent(req: FileManagerAgentRequest):
    _emit_log("info", "file_manager", "Received request", {"query": req.message})
    user_message = (req.message or "").strip()
    state = _normalize_file_manager_state(req.agent_state.model_dump())
    config = _normalize_file_manager_config(req.file_manager_config.model_dump())
    normalized_choice = normalize_choice(user_message).lower()
    export_payload = _try_extract_file_export_payload(user_message)


    # Empty turns are used as a lightweight bootstrap for the dedicated agent
    # view: return a compact capability reminder instead of entering the ReAct loop.
    if not user_message:
        return {
            "answer": (
                "## File Management Agent\n"
                "Ask me to inspect, search, create, edit, move, or delete files.\n\n"
                f"Current access root: `{config['basePath'] or 'full server-visible workspace'}`"
            ),
            "agent_state": state,
            "steps": [
                {
                    "id": "fm-ready",
                    "title": "Ready for file operations",
                    "status": "success",
                    "details": "The agent is ready to reason over filesystem tasks with tool calls.",
                }
                ],
            }

    if export_payload:
        if not export_payload.get("path"):
            return {
                "answer": (
                    "## File Export\n"
                    "I have the dataset ready, but I still need the target file name or path.\n\n"
                    f"Requested format: **{export_payload.get('format')}**"
                ),
                "agent_state": state,
                "steps": [
                    {
                        "id": "fm-export-path",
                        "title": "Waiting for export path",
                        "status": "running",
                        "details": "The export payload is ready, but the target file path is missing.",
                    }
                ],
            }

        try:
            if export_payload["format"] in {"csv", "tsv"}:
                content = _serialize_delimited_rows(
                    export_payload.get("headers") or [],
                    export_payload.get("rows") or [],
                    "\t" if export_payload["format"] == "tsv" else ",",
                )
                target = _resolve_agent_path(export_payload["path"], config["basePath"])
                result = (
                    write_file_tool(export_payload["path"], content, confirmed=False, base_path=config["basePath"])
                    if target.exists()
                    else create_file_tool(export_payload["path"], content, base_path=config["basePath"])
                )
            else:
                target = _resolve_agent_path(export_payload["path"], config["basePath"])
                result = (
                    write_excel_sheet_tool(
                        export_payload["path"],
                        export_payload.get("sheet_name") or "Results",
                        headers=export_payload.get("headers") or [],
                        rows=export_payload.get("rows") or [],
                        confirmed=False,
                        base_path=config["basePath"],
                    )
                    if target.exists()
                    else create_excel_file_tool(
                        export_payload["path"],
                        sheet_name=export_payload.get("sheet_name") or "Results",
                        headers=export_payload.get("headers") or [],
                        rows=export_payload.get("rows") or [],
                        base_path=config["basePath"],
                    )
                )
        except Exception as exc:
            return {
                "answer": f"## File Export\nI could not prepare the export.\n\n```text\n{exc}\n```",
                "agent_state": state,
                "steps": [
                    {
                        "id": "fm-export-error",
                        "title": "Export failed",
                        "status": "error",
                        "details": str(exc),
                    }
                ],
            }

        state["last_tool_result"] = result["summary"]
        state["last_visited_path"] = result.get("visited_path") or state.get("last_visited_path", "")
        if result.get("requires_confirmation"):
            pending_action = dict(result.get("pending_action") or {})
            state["pending_confirmation"] = {
                "tool_name": pending_action.get("tool_name"),
                "tool_input": pending_action.get("tool_input") or {},
                "preview": result.get("preview") or "",
                "summary": result.get("summary") or "",
                "requested_at": _utc_now_iso(),
            }
            answer, actions = _file_manager_confirmation_answer(state)
            return {
                "answer": answer,
                "actions": actions,
                "agent_state": state,
                "steps": [
                    {
                        "id": "fm-export-confirm",
                        "title": "Confirmation required",
                        "status": "running",
                        "details": result["summary"],
                    }
                ],
            }

        return {
            "answer": _file_export_answer(result, export_payload),
            "agent_state": state,
            "steps": [
                {
                    "id": "fm-export-dataset",
                    "title": "Exported ClickHouse result",
                    "status": "success",
                    "details": result["summary"],
                }
            ],
        }

    if state.get("pending_confirmation"):
        if is_negative_response(user_message) or normalized_choice in {"cancel", "cancel file action"}:
            state["pending_confirmation"] = None
            return {
                "answer": "## Answer\nThe pending file operation was cancelled.",
                "agent_state": state,
                "steps": [
                    {
                        "id": "fm-cancel",
                        "title": "Cancelled pending action",
                        "status": "success",
                        "details": "The destructive or overwrite operation was not executed.",
                    }
                ],
            }

        if is_affirmative_response(user_message) or normalized_choice in {"confirm", "confirm file action"}:
            pending = state["pending_confirmation"] or {}
            tool_name = str(pending.get("tool_name") or "").strip()
            tool_input = dict(pending.get("tool_input") or {})
            tool_input["confirmed"] = True
            try:
                result = execute_file_manager_tool(tool_name, tool_input, config["basePath"])
            except Exception as exc:
                state["pending_confirmation"] = None
                return {
                    "answer": f"## Answer\nI could not complete the confirmed action.\n\n```text\n{exc}\n```",
                    "agent_state": state,
                    "steps": [
                        {
                            "id": "fm-confirm-error",
                            "title": "Confirmed action failed",
                            "status": "error",
                            "details": str(exc),
                        }
                    ],
                }

            state["pending_confirmation"] = None
            state["last_tool_result"] = result["summary"]
            state["last_visited_path"] = result.get("visited_path") or state.get("last_visited_path", "")
            preview = result.get("preview") or ""
            answer = f"## Answer\n{result['summary']}"
            if preview:
                answer += f"\n\n## Preview\n{preview}"
            return {
                "answer": answer,
                "agent_state": state,
                "steps": [
                    {
                        "id": "fm-confirmed-action",
                        "title": f"Executed `{tool_name}`",
                        "status": "success",
                        "details": result["summary"],
                    }
                ],
            }

        answer, actions = _file_manager_confirmation_answer(state)
        return {
            "answer": answer,
            "actions": actions,
            "agent_state": state,
            "steps": [
                {
                    "id": "fm-await-confirmation",
                    "title": "Waiting for confirmation",
                    "status": "running",
                    "details": "The requested file operation needs explicit user confirmation.",
                }
            ],
        }

    scratchpad: list[dict[str, Any]] = []
    steps: list[dict[str, Any]] = []
    last_error = ""
    max_iterations = max(1, min(FILE_MANAGER_MAX_ITERATIONS, int(config["maxIterations"])))

    for iteration in range(max_iterations):
        planned = await plan_file_manager_step(
            user_message,
            req.history,
            scratchpad,
            config["basePath"],
            config["systemPrompt"],
            req.llm_base_url,
            req.llm_model,
            req.llm_provider,
            req.llm_api_key,
        )
        reasoning = planned.get("reasoning") or "The local LLM selected the next best action."
        action = planned.get("action")

        if action == "final":
            final_answer = planned.get("final_answer") or "## Answer\nThe file-management task is complete."
            steps.append(
                {
                    "id": f"fm-final-{iteration}",
                    "title": "Prepared final answer",
                    "status": "success",
                    "details": reasoning,
                }
            )
            state["last_tool_result"] = final_answer
            return {
                "answer": final_answer,
                "agent_state": state,
                "steps": steps,
            }

        tool_name = planned.get("tool_name") or ""
        tool_input = dict(planned.get("tool_input") or {})
        if action != "tool" or tool_name not in FILE_MANAGER_TOOLS:
            last_error = "The planner returned an invalid tool action."
            scratchpad.append({"type": "error", "error": last_error})
            steps.append(
                {
                    "id": f"fm-invalid-{iteration}",
                    "title": "Planner action invalid",
                    "status": "error",
                    "details": last_error,
                }
            )
            continue

        if tool_name in FILE_MANAGER_CONFIRMATION_TOOLS and "confirmed" not in tool_input:
            tool_input["confirmed"] = False

        try:
            start_exec = _time.time()
            result = execute_file_manager_tool(tool_name, tool_input, config["basePath"])
            elapsed = _time.time() - start_exec
            _emit_log("tool_call", "file_manager", f"Executed {tool_name} ({elapsed:.2f}s)", {"tool": tool_name, "input": tool_input, "summary": result.get("summary")})
        except Exception as exc:
            last_error = str(exc)
            _emit_log("error", "file_manager", f"Tool {tool_name} failed: {exc}", {"tool": tool_name, "error": last_error})
            scratchpad.append(
                {
                    "type": "tool_error",
                    "tool": tool_name,
                    "input": tool_input,
                    "error": last_error,
                }
            )
            steps.append(
                {
                    "id": f"fm-tool-error-{iteration}",
                    "title": f"Tool `{tool_name}` failed",
                    "status": "error",
                    "details": last_error,
                }
            )
            continue


        state["last_tool_result"] = result["summary"]
        state["last_visited_path"] = result.get("visited_path") or state.get("last_visited_path", "")
        scratchpad.append(
            {
                "type": "tool_result",
                "tool": tool_name,
                "input": tool_input,
                "summary": result["summary"],
                "preview": result.get("preview") or "",
            }
        )
        steps.append(
            {
                "id": f"fm-tool-{iteration}",
                "title": f"Used `{tool_name}`",
                "status": "success",
                "details": result["summary"],
            }
        )

        if result.get("requires_confirmation"):
            pending_action = dict(result.get("pending_action") or {})
            state["pending_confirmation"] = {
                "tool_name": pending_action.get("tool_name"),
                "tool_input": pending_action.get("tool_input") or {},
                "preview": result.get("preview") or "",
                "summary": result.get("summary") or "",
                "requested_at": _utc_now_iso(),
            }
            answer, actions = _file_manager_confirmation_answer(state)
            steps.append(
                {
                    "id": f"fm-confirm-{iteration}",
                    "title": "Confirmation required",
                    "status": "running",
                    "details": result["summary"],
                }
            )
            return {
                "answer": answer,
                "actions": actions,
                "agent_state": state,
                "steps": steps,
            }

        completion = await _evaluate_file_manager_completion(
            user_message,
            req.history,
            scratchpad,
            tool_name,
            tool_input,
            result,
            config["systemPrompt"],
            req.llm_base_url,
            req.llm_model,
            req.llm_provider,
            req.llm_api_key,
            req.disable_ssl_verification,
        )
        if completion.get("is_complete"):
            final_answer = str(completion.get("final_answer") or "").strip() or _file_manager_completion_answer_from_result(result)
            steps.append(
                {
                    "id": f"fm-complete-{iteration}",
                    "title": "Stopped after completed task",
                    "status": "success",
                    "details": str(completion.get("reasoning") or result["summary"]).strip(),
                }
            )
            state["last_tool_result"] = final_answer
            return {
                "answer": final_answer,
                "agent_state": state,
                "steps": steps,
            }

    answer = (
        "## Answer\n"
        "I reached the file-management iteration limit before I could finish the task."
    )
    if last_error:
        answer += f"\n\n```text\n{last_error}\n```"
    return {
        "answer": answer,
        "agent_state": state,
        "steps": steps or [
            {
                "id": "fm-timeout",
                "title": "Reached iteration limit",
                "status": "error",
                "details": "The agent stopped to avoid an infinite loop.",
            }
        ],
    }


@app.post("/api/chat/pdf-creator-agent")
async def chat_pdf_creator_agent(req: PdfCreatorAgentRequest):
    _emit_log("info", "pdf_creator", "Received request", {"query": req.message})
    user_message = (req.message or "").strip()
    state = _normalize_pdf_creator_state(req.agent_state.model_dump())
    config = _normalize_file_manager_config(req.file_manager_config.model_dump())
    normalized_choice = normalize_choice(user_message).lower()
    payload = _try_extract_pdf_export_payload(user_message)


    def _choice_prompt(current_state: dict[str, Any], reason: str) -> dict[str, Any]:
        current_state["stage"] = "awaiting_source_choice"
        answer = build_choice_markdown(
            "PDF Creator",
            (
                "I can create a polished PDF for you, but I still need the content source.\n\n"
                f"{reason}"
            ),
            PDF_CREATOR_SOURCE_CHOICES,
        )
        return {
            "answer": answer,
            "agent_state": current_state,
            "steps": [
                {
                    "id": "pdf-source-choice",
                    "title": "Waiting for content source",
                    "status": "running",
                    "details": "The PDF creator needs to know whether it should use the latest chat analysis or new content pasted by the user.",
                }
            ],
        }

    if not user_message:
        return {
            "answer": (
                "## PDF Creator Agent\n"
                "Ask me to turn the latest analysis or any pasted content into a clean, professional PDF.\n\n"
                f"- **Sandbox base path:** `{config['basePath'] or 'not restricted'}`\n"
                "- **Style:** Slate header, compact sections, and export-ready formatting aligned with the UI."
            ),
            "agent_state": state,
            "steps": [
                {
                    "id": "pdf-ready",
                    "title": "Ready for PDF export",
                    "status": "success",
                    "details": "The agent can build a PDF from the latest analysis or from text the user provides.",
                }
            ],
        }

    if state.get("pending_confirmation"):
        if is_negative_response(user_message) or normalized_choice in {"cancel", "cancel file action"}:
            state["pending_confirmation"] = None
            state["stage"] = "idle"
            return {
                "answer": "## PDF Creator\nThe pending PDF overwrite was cancelled.",
                "agent_state": state,
                "steps": [
                    {
                        "id": "pdf-cancel",
                        "title": "Cancelled overwrite",
                        "status": "success",
                        "details": "The existing PDF file was left unchanged.",
                    }
                ],
            }

        if is_affirmative_response(user_message) or normalized_choice in {"confirm", "confirm file action"}:
            pending = state.get("pending_confirmation") or {}
            pending_action = dict(pending.get("pending_action") or {})
            try:
                result = create_pdf_report_tool(
                    path=str(pending_action.get("path") or "").strip(),
                    title=str(pending_action.get("title") or "RAGnarok Report").strip(),
                    subtitle=str(pending_action.get("subtitle") or "Professional export generated from RAGnarok").strip(),
                    body_markdown=str(pending_action.get("body_markdown") or "").strip(),
                    confirmed=True,
                    base_path=config["basePath"],
                )
            except Exception as exc:
                state["pending_confirmation"] = None
                state["stage"] = "idle"
                return {
                    "answer": f"## PDF Creator\nI could not complete the confirmed PDF export.\n\n```text\n{exc}\n```",
                    "agent_state": state,
                    "steps": [
                        {
                            "id": "pdf-confirm-error",
                            "title": "Confirmed PDF export failed",
                            "status": "error",
                            "details": str(exc),
                        }
                    ],
                }

            state["pending_confirmation"] = None
            state["pending_document"] = None
            state["stage"] = "idle"
            state["last_output_path"] = str((result.get("data") or {}).get("path") or "")
            state["last_title"] = str(pending_action.get("title") or "").strip()
            return {
                "answer": _pdf_creator_success_answer(result, state["last_title"], state["last_output_path"]),
                "agent_state": state,
                "steps": [
                    {
                        "id": "pdf-overwrite-complete",
                        "title": "Overwrote PDF",
                        "status": "success",
                        "details": result.get("summary") or "The PDF export completed successfully.",
                    }
                ],
            }

        answer, actions = _pdf_creator_confirmation_answer(state)
        return {
            "answer": answer,
            "actions": actions,
            "agent_state": state,
            "steps": [
                {
                    "id": "pdf-await-confirmation",
                    "title": "Waiting for confirmation",
                    "status": "running",
                    "details": "The requested PDF target already exists, so overwrite confirmation is required.",
                }
            ],
        }

    if state.get("stage") == "awaiting_source_choice":
        pending_document = dict(state.get("pending_document") or {})
        selected_choice = resolve_user_choice(user_message, PDF_CREATOR_SOURCE_CHOICES)
        if selected_choice == PDF_CREATOR_SOURCE_CHOICES[0]:
            latest_analysis = _latest_exportable_assistant_message(req.history)
            if not latest_analysis:
                state["stage"] = "awaiting_content"
                return {
                    "answer": (
                        "## PDF Creator\n"
                        "I could not find a usable analysis in the recent chat history.\n\n"
                        "Please paste the content you want me to turn into a PDF in your next message."
                    ),
                    "agent_state": state,
                    "steps": [
                        {
                            "id": "pdf-await-content",
                            "title": "Waiting for pasted content",
                            "status": "running",
                            "details": "No recent exportable analysis was found, so the agent is asking the user to paste the content manually.",
                        }
                    ],
                }
            payload = {
                "title": str(pending_document.get("title") or "RAGnarok Report").strip(),
                "subtitle": str(pending_document.get("subtitle") or "Professional export generated from RAGnarok").strip(),
                "path": str(pending_document.get("path") or "").strip(),
                "source_markdown": latest_analysis,
                "source_request": str(pending_document.get("source_request") or "").strip(),
            }
            state["stage"] = "idle"
            state["pending_document"] = None
        elif selected_choice == PDF_CREATOR_SOURCE_CHOICES[1]:
            state["stage"] = "awaiting_content"
            return {
                "answer": (
                    "## PDF Creator\n"
                    "Please paste the content you want me to export.\n\n"
                    "I will turn your next message into a professional PDF."
                ),
                "agent_state": state,
                "steps": [
                    {
                        "id": "pdf-await-content",
                        "title": "Waiting for pasted content",
                        "status": "running",
                        "details": "The agent is waiting for the user to paste the content to export as PDF.",
                    }
                ],
            }
        else:
            return _choice_prompt(state, "Choose how I should gather the content for the PDF.")

    if state.get("stage") == "awaiting_content" and not payload:
        pending_document = dict(state.get("pending_document") or {})
        payload = {
            "title": str(pending_document.get("title") or "RAGnarok Report").strip(),
            "subtitle": str(pending_document.get("subtitle") or "Professional export generated from RAGnarok").strip(),
            "path": str(pending_document.get("path") or "").strip(),
            "source_markdown": user_message[:PDF_CREATOR_MAX_TEXT_CHARS],
            "source_request": str(pending_document.get("source_request") or "").strip(),
        }
        state["stage"] = "idle"
        state["pending_document"] = None

    if not payload:
        target_path = _extract_pdf_path(user_message)
        title = _extract_pdf_title(user_message, target_path)
        source_markdown = _latest_exportable_assistant_message(req.history)
        if not source_markdown:
            state["pending_document"] = {
                "title": title,
                "subtitle": "Professional export generated from RAGnarok",
                "path": target_path or _default_pdf_target_path(title),
                "source_request": user_message,
            }
            return _choice_prompt(
                state,
                "I did not find an explicit body in your request, and there is no recent assistant analysis I can safely export yet.",
            )

        payload = {
            "title": title,
            "subtitle": "Professional export generated from RAGnarok",
            "path": target_path or _default_pdf_target_path(title),
            "source_markdown": source_markdown,
            "source_request": user_message,
        }

    title = str(payload.get("title") or "RAGnarok Report").strip() or "RAGnarok Report"
    subtitle = str(payload.get("subtitle") or "Professional export generated from RAGnarok").strip()
    target_path = str(payload.get("path") or "").strip() or _default_pdf_target_path(title)
    source_markdown = str(payload.get("source_markdown") or "").strip()
    body_markdown = _build_pdf_creator_body_markdown(
        source_markdown,
        title,
        str(payload.get("source_request") or "").strip(),
    )

    try:
        start_exec = _time.time()
        result = create_pdf_report_tool(
            path=target_path,
            title=title,
            subtitle=subtitle,
            body_markdown=body_markdown,
            confirmed=False,
            base_path=config["basePath"],
        )
        elapsed = _time.time() - start_exec
        _emit_log("tool_call", "pdf_creator", f"Generated PDF ({elapsed:.2f}s)", {"path": target_path})
    except Exception as exc:
        _emit_log("error", "pdf_creator", f"PDF generation failed: {exc}", {"error": str(exc)})

        state["stage"] = "idle"
        state["pending_document"] = None
        return {
            "answer": f"## PDF Creator\nI could not prepare the PDF export.\n\n```text\n{exc}\n```",
            "agent_state": state,
            "steps": [
                {
                    "id": "pdf-error",
                    "title": "PDF export failed",
                    "status": "error",
                    "details": str(exc),
                }
            ],
        }

    if result.get("requires_confirmation"):
        pending_action = dict((result.get("pending_action") or {}))
        tool_input = dict(pending_action.get("tool_input") or {})
        state["pending_confirmation"] = {
            "preview": str(result.get("preview") or "").strip(),
            "summary": str(result.get("summary") or "").strip(),
            "requested_at": _utc_now_iso(),
            "pending_action": {
                "path": str(tool_input.get("path") or target_path).strip(),
                "title": str(tool_input.get("title") or title).strip(),
                "subtitle": str(tool_input.get("subtitle") or subtitle).strip(),
                "body_markdown": str(tool_input.get("body_markdown") or body_markdown),
            },
        }
        state["last_title"] = title
        state["last_output_path"] = target_path
        answer, actions = _pdf_creator_confirmation_answer(state)
        return {
            "answer": answer,
            "actions": actions,
            "agent_state": state,
            "steps": [
                {
                    "id": "pdf-confirm",
                    "title": "Confirmation required",
                    "status": "running",
                    "details": str(result.get("summary") or "Overwrite confirmation is required."),
                }
            ],
        }

    state["stage"] = "idle"
    state["pending_document"] = None
    state["pending_confirmation"] = None
    state["last_output_path"] = str((result.get("data") or {}).get("path") or target_path)
    state["last_title"] = title
    return {
        "answer": _pdf_creator_success_answer(result, title, state["last_output_path"]),
        "agent_state": state,
        "steps": [
            {
                "id": "pdf-created",
                "title": "Created PDF",
                "status": "success",
                "details": result.get("summary") or "The PDF export completed successfully.",
            }
        ],
    }


def _oracle_match_table_choice(user_message: str, options: list[str]) -> Optional[str]:
    direct = resolve_user_choice(user_message, options)
    if direct:
        return direct
    normalized = normalize_intent_text(user_message)
    if not normalized:
        return None
    matches = []
    for option in options:
        lowered = option.lower()
        short_name = lowered.split(".")[-1]
        if lowered in normalized or re.search(rf"(?<![a-z0-9_]){re.escape(short_name)}(?![a-z0-9_])", normalized):
            matches.append(option)
    return matches[0] if len(matches) == 1 else None


def _oracle_state_needs_followup(state: dict[str, Any]) -> bool:
    return str(state.get("stage") or "").strip() == "awaiting_table"


def _oracle_tool_preview(value: Any, max_items: int = 20) -> str:
    if isinstance(value, list):
        preview_items = value[:max_items]
        return json.dumps(preview_items, ensure_ascii=False, indent=2)
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False, indent=2)
    return str(value)


@app.post("/api/chat/oracle-analyst-agent")
async def chat_oracle_analyst_agent(req: OracleAnalystAgentRequest):
    _emit_log("info", "oracle", "Received request", {"query": req.message})
    user_message = (req.message or "").strip()
    state = _normalize_oracle_analyst_state(req.agent_state.model_dump())
    oracle_config = _normalize_oracle_analyst_config(req.oracle_analyst_config.model_dump())
    connections = req.oracle_connections or [OracleConnectionConfig(**DEFAULT_APP_CONFIG["oracleConnections"][0])]
    connection = _resolve_oracle_connection(connections, req.oracle_analyst_config)

    if not user_message:
        return {
            "answer": (
                "## Oracle SQL Agent\n"
                "Ask a business question in English and I will inspect the Oracle schema, generate optimized SQL, and return a narrative answer.\n\n"
                f"- **Connection:** {connection.label or connection.id}\n"
                f"- **Row limit:** {oracle_config['rowLimit']}\n"
                "- **Flow:** list tables → inspect schema → validate SQL → execute query → summarize result."
            ),
            "agent_state": state,
            "steps": [
                {
                    "id": "oracle-ready",
                    "title": "Ready for Oracle analysis",
                    "status": "success",
                    "details": "The Oracle SQL agent is ready to query Oracle safely with a ReAct loop.",
                }
            ],
        }

    if state.get("stage") == "awaiting_table":
        selected_table = _oracle_match_table_choice(user_message, state.get("clarification_options") or [])
        if not selected_table:
            return {
                "answer": build_choice_markdown(
                    "Oracle Table Selection",
                    state.get("clarification_prompt") or "Choose the Oracle table that best matches your request.",
                    state.get("clarification_options") or state.get("available_tables")[:ORACLE_TABLE_PREVIEW_LIMIT],
                ),
                "agent_state": state,
                "steps": [
                    {
                        "id": "oracle-await-table",
                        "title": "Waiting for table selection",
                        "status": "running",
                        "details": "The agent is waiting for the user to choose the Oracle table.",
                    }
                ],
            }
        state["selected_table"] = selected_table
        state["stage"] = "idle"
        state["clarification_prompt"] = ""
        state["clarification_options"] = []
        user_message = state.get("pending_request") or user_message

    scratchpad: list[dict[str, Any]] = []
    steps: list[dict[str, Any]] = []
    current_request = state.get("pending_request") or user_message

    for iteration in range(ORACLE_REACT_MAX_ITERATIONS):
        planned = await plan_oracle_react_step(
            current_request,
            req.history,
            scratchpad,
            state,
            oracle_config,
            connection.label or connection.id,
            req.llm_base_url,
            req.llm_model,
            req.llm_provider,
            req.llm_api_key,
        )
        reasoning = planned.get("reasoning") or "The local LLM selected the next Oracle step."
        action = planned.get("action")

        if action == "clarify_table":
            options = [
                option for option in planned.get("clarification_options", [])
                if option in (state.get("available_tables") or [])
            ][:6]
            if not options:
                if not state.get("available_tables"):
                    try:
                        state["available_tables"] = await list_oracle_tables(connection)
                    except Exception as exc:
                        raise HTTPException(status_code=400, detail=f"Oracle connection error: {exc}") from exc
                options = (state.get("available_tables") or [])[:6]
            state["stage"] = "awaiting_table"
            state["pending_request"] = current_request
            state["clarification_prompt"] = planned.get("clarification_prompt") or "Choose the Oracle table that best matches your request."
            state["clarification_options"] = options
            return {
                "answer": build_choice_markdown(
                    "Oracle Table Selection",
                    state["clarification_prompt"],
                    options,
                ),
                "agent_state": state,
                "steps": steps + [
                    {
                        "id": f"oracle-clarify-{iteration}",
                        "title": "Waiting for table selection",
                        "status": "running",
                        "details": reasoning,
                    }
                ],
            }

        if action == "final":
            state["stage"] = "ready"
            state["pending_request"] = ""
            state["final_answer"] = planned.get("final_answer") or state.get("final_answer") or "## Executive Summary\nThe Oracle analysis is complete."
            return {
                "answer": state["final_answer"],
                "agent_state": state,
                "steps": steps + [
                    {
                        "id": f"oracle-final-{iteration}",
                        "title": "Prepared final answer",
                        "status": "success",
                        "details": reasoning,
                    }
                ],
            }

        tool_name = str(planned.get("tool_name") or "").strip()
        tool_input = dict(planned.get("tool_input") or {})
        if action != "tool" or tool_name not in ORACLE_REACT_TOOL_NAMES:
            scratchpad.append({"type": "planner_error", "error": "The Oracle planner returned an invalid action."})
            steps.append(
                {
                    "id": f"oracle-invalid-{iteration}",
                    "title": "Planner action invalid",
                    "status": "error",
                    "details": "The Oracle planner returned an invalid action.",
                }
            )
            continue

        try:
            if tool_name == "list_tables":
                tables = await list_oracle_tables(connection)
                state["available_tables"] = tables
                result_payload = {
                    "tables": tables,
                    "preview": tables[:ORACLE_TABLE_PREVIEW_LIMIT],
                }
                action_label = "list_tables → Agent"
                step_detail = f"Loaded {len(tables)} accessible Oracle table(s)."
            elif tool_name == "get_schema":
                table_name = str(tool_input.get("table") or state.get("selected_table") or "").strip()
                if not table_name:
                    raise ValueError("get_schema requires a target table.")
                columns_filter = str(tool_input.get("columns_filter") or tool_input.get("columnsFilter") or "").strip()
                schema = await get_oracle_schema(connection, table_name, columns_filter)
                if not schema:
                    raise ValueError(f"No readable schema was found for `{table_name}`.")
                state["selected_table"] = table_name
                state["schema_info"] = schema
                result_payload = {
                    "table": table_name,
                    "columns_filter": columns_filter,
                    "schema": schema,
                }
                columns_filter_suffix = f', columns_filter="{columns_filter}"' if columns_filter else ""
                action_label = f"get_schema('{table_name}'{columns_filter_suffix}) → Agent"
                step_detail = f"Loaded {len(schema)} column(s) from `{table_name}`."
            elif tool_name == "check_query":
                sql = str(tool_input.get("sql") or "").strip()
                if not sql:
                    raise ValueError("check_query requires SQL.")
                checked = await check_oracle_query(connection, sql)
                result_payload = checked
                action_label = "check_query → Agent"
                step_detail = f"Validated the Oracle query using `{checked.get('mode', 'parse_only')}` mode."
            else:
                sql = str(tool_input.get("sql") or "").strip()
                if not sql:
                    raise ValueError("execute_query requires SQL.")
                start_exec = _time.time()
                result = await execute_oracle_query(connection, sql, oracle_config["rowLimit"])
                elapsed = _time.time() - start_exec
                _emit_log("sql", "oracle", f"Executed query successfully ({elapsed:.2f}s)", {"sql": sql, "rows": result["row_count"]})
                state["last_sql"] = result["sql"]

                state["last_result_meta"] = result["columns"]
                state["last_result_rows"] = result["rows"]
                state["action_log"].append("execute_query → Agent")
                summary = await summarize_oracle_result(
                    current_request,
                    result["sql"],
                    result["rows"],
                    state["action_log"],
                    _conversation_memory_markdown(req.history, current_message=current_request),
                    req.llm_base_url,
                    req.llm_model,
                    req.llm_provider,
                    req.llm_api_key,
                )
                state["final_answer"] = build_oracle_response_markdown(
                    summary,
                    result["sql"],
                    result["rows"],
                    state["action_log"],
                )
                state["stage"] = "ready"
                state["pending_request"] = ""
                steps.append(
                    {
                        "id": f"oracle-execute-{iteration}",
                        "title": "Executed Oracle query",
                        "status": "success",
                        "details": f"Returned {result['row_count']} row(s).",
                    }
                )
                return {
                    "answer": state["final_answer"],
                    "agent_state": state,
                    "steps": steps,
                }

            if action_label and action_label not in state["action_log"]:
                state["action_log"].append(action_label)
            scratchpad.append(
                {
                    "type": "tool_result",
                    "tool": tool_name,
                    "input": tool_input,
                    "result": result_payload,
                }
            )
            steps.append(
                {
                    "id": f"oracle-tool-{iteration}",
                    "title": f"Used `{tool_name}`",
                    "status": "success",
                    "details": step_detail,
                }
            )
        except Exception as exc:
            state["last_error"] = str(exc)
            scratchpad.append(
                {
                    "type": "tool_error",
                    "tool": tool_name,
                    "input": tool_input,
                    "error": str(exc),
                }
            )
            steps.append(
                {
                    "id": f"oracle-tool-error-{iteration}",
                    "title": f"Tool `{tool_name}` failed",
                    "status": "error",
                    "details": str(exc),
                }
            )

    fallback_answer = state.get("final_answer") or (
        "## Executive Summary\n"
        "I reached the Oracle reasoning limit before I could safely finish the analysis."
    )
    return {
        "answer": fallback_answer,
        "agent_state": state,
        "steps": steps or [
            {
                "id": "oracle-iteration-limit",
                "title": "Reached iteration limit",
                "status": "error",
                "details": "The Oracle ReAct loop stopped after the maximum number of iterations.",
            }
        ],
    }


@app.post("/api/data-quality/options")
async def get_data_quality_options(req: DataQualityMetadataRequest):
    raise HTTPException(status_code=410, detail="Data Quality - Tables has been removed from this application.")


@app.post("/api/clickhouse/guide-metadata")
async def get_clickhouse_guide_metadata(req: ClickHouseGuideMetadataRequest):
    try:
        available_tables = await list_clickhouse_tables(req.clickhouse)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"ClickHouse metadata loading failed: {exc}") from exc

    selected_table = str(req.table or "").strip()
    schema_info: list[dict[str, Any]] = []
    target_candidates: list[str] = []

    if selected_table:
        if selected_table not in available_tables:
            raise HTTPException(status_code=400, detail=f"Unknown ClickHouse table: {selected_table}")
        try:
            raw_schema = await describe_clickhouse_table(req.clickhouse, selected_table)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Failed to inspect schema for {selected_table}: {exc}") from exc

        schema_info = [
            {
                "name": str(column.get("name") or ""),
                "type": str(column.get("type") or ""),
                "category": _clickhouse_schema_category(str(column.get("type") or "")),
            }
            for column in raw_schema
            if str(column.get("name") or "").strip()
        ]
        target_candidates = _auto_ml_target_candidates(schema_info)

    return {
        "available_tables": available_tables,
        "schema_info": schema_info,
        "target_candidates": target_candidates,
    }


@app.post("/api/auto-ml/filter-suggestion")
async def suggest_auto_ml_filter(req: AutoMlFilterSuggestionRequest):
    table_name = str(req.table or "").strip()
    if not table_name:
        raise HTTPException(status_code=400, detail="A ClickHouse table is required before the AI can suggest a row filter.")

    try:
        available_tables = await list_clickhouse_tables(req.clickhouse)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"ClickHouse metadata loading failed: {exc}") from exc

    if table_name not in available_tables:
        raise HTTPException(status_code=400, detail=f"Unknown ClickHouse table: {table_name}")

    schema_info = [
        {
            "name": str(column.get("name") or "").strip(),
            "type": str(column.get("type") or "").strip(),
            "category": str(column.get("category") or "").strip(),
        }
        for column in req.schema_info
        if isinstance(column, dict) and str(column.get("name") or "").strip()
    ]
    if not schema_info:
        try:
            raw_schema = await describe_clickhouse_table(req.clickhouse, table_name)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Failed to inspect schema for {table_name}: {exc}") from exc
        schema_info = [
            {
                "name": str(column.get("name") or "").strip(),
                "type": str(column.get("type") or "").strip(),
                "category": _clickhouse_schema_category(str(column.get("type") or "")),
            }
            for column in raw_schema
            if str(column.get("name") or "").strip()
        ]

    prompt = f"""
You are helping a business user narrow an Auto-ML benchmark on ClickHouse.
Suggest one safe SQL boolean WHERE clause to focus the training dataset.

Rules:
- Return ONLY valid JSON.
- Use this exact shape:
  {{
    "where_clause": "country = 'FR' AND event_date BETWEEN '2025-01-01' AND '2025-03-31'",
    "rationale": "Short explanation"
  }}
- The clause must be a single boolean expression only.
- Do NOT include SELECT, FROM, JOIN, ORDER BY, LIMIT, comments, semicolons, or markdown.
- Do NOT prefix the expression with WHERE.
- Prefer a practical business scope such as geography, product segment, or a recent explicit date range if the schema supports it.
- If you cannot infer a safe useful scope, return an empty string for `where_clause`.

Table:
{table_name}

Target column:
{str(req.target_column or "").strip() or "Not provided"}

Goal:
{str(req.goal or "").strip() or "Not provided"}

Notes:
{str(req.notes or "").strip() or "Not provided"}

Schema:
{json.dumps(schema_info, ensure_ascii=False, indent=2)}
""".strip()

    try:
        raw = await llm_chat(
            [{"role": "user", "content": prompt}],
            req.llm_base_url,
            req.llm_model,
            req.llm_provider,
            req.llm_api_key,
        )
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Local LLM filter suggestion failed: {exc}") from exc

    parsed = extract_json_object(raw)
    where_clause = _normalize_auto_ml_row_filter(str(parsed.get("where_clause") or parsed.get("where") or "").strip())
    rationale = str(parsed.get("rationale") or parsed.get("reasoning") or "").strip()
    validation_error = _validate_data_quality_row_filter(where_clause)
    if validation_error:
        where_clause = ""
        rationale = rationale or validation_error

    return {
        "where_clause": where_clause,
        "rationale": rationale,
    }


@app.post("/api/clickhouse/filter-suggestion")
async def suggest_clickhouse_scope_filter(req: ClickHouseScopeSuggestionRequest):
    table_name = str(req.table or "").strip()
    if not table_name:
        raise HTTPException(status_code=400, detail="A ClickHouse table is required before the AI can suggest a row filter.")

    try:
        available_tables = await list_clickhouse_tables(req.clickhouse)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"ClickHouse metadata loading failed: {exc}") from exc

    if table_name not in available_tables:
        raise HTTPException(status_code=400, detail=f"Unknown ClickHouse table: {table_name}")

    schema_info = [
        {
            "name": str(column.get("name") or "").strip(),
            "type": str(column.get("type") or "").strip(),
            "category": str(column.get("category") or "").strip(),
        }
        for column in req.schema_info
        if isinstance(column, dict) and str(column.get("name") or "").strip()
    ]
    if not schema_info:
        try:
            raw_schema = await describe_clickhouse_table(req.clickhouse, table_name)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Failed to inspect schema for {table_name}: {exc}") from exc
        schema_info = [
            {
                "name": str(column.get("name") or "").strip(),
                "type": str(column.get("type") or "").strip(),
                "category": _clickhouse_schema_category(str(column.get("type") or "")),
            }
            for column in raw_schema
            if str(column.get("name") or "").strip()
        ]

    agent_kind = str(req.agent_kind or "generic").strip().lower()
    agent_brief = {
        "data_cleaner": "You are helping a Data Cleaner focus a quality audit on the most relevant business slice.",
        "anonymizer": "You are helping an Anonymizer focus a privacy scan on the most relevant business slice.",
        "auto_ml": "You are helping Auto-ML focus model training on the most relevant business slice.",
    }.get(agent_kind, "You are helping a ClickHouse specialist focus an analysis on the most relevant business slice.")

    prompt = f"""
{agent_brief}
Suggest one safe SQL boolean WHERE clause to narrow the dataset.

Rules:
- Return ONLY valid JSON.
- Use this exact shape:
  {{
    "where_clause": "country = 'FR' AND event_date BETWEEN '2025-01-01' AND '2025-03-31'",
    "rationale": "Short explanation"
  }}
- The clause must be a single boolean expression only.
- Do NOT include SELECT, FROM, JOIN, ORDER BY, LIMIT, comments, semicolons, or markdown.
- Do NOT prefix the expression with WHERE.
- Prefer a practical scope such as geography, product segment, status, or an explicit recent date range when the schema supports it.
- If you cannot infer a safe useful scope, return an empty string for `where_clause`.

Table:
{table_name}

Target column:
{str(req.target_column or "").strip() or "Not provided"}

Goal:
{str(req.goal or "").strip() or "Not provided"}

Notes:
{str(req.notes or "").strip() or "Not provided"}

Schema:
{json.dumps(schema_info, ensure_ascii=False, indent=2)}
""".strip()

    try:
        raw = await llm_chat(
            [{"role": "user", "content": prompt}],
            req.llm_base_url,
            req.llm_model,
            req.llm_provider,
            req.llm_api_key,
        )
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Local LLM filter suggestion failed: {exc}") from exc

    parsed = extract_json_object(raw)
    where_clause = _normalize_auto_ml_row_filter(str(parsed.get("where_clause") or parsed.get("where") or "").strip())
    rationale = str(parsed.get("rationale") or parsed.get("reasoning") or "").strip()
    validation_error = _validate_data_quality_row_filter(where_clause)
    if validation_error:
        where_clause = ""
        rationale = rationale or validation_error

    return {
        "where_clause": where_clause,
        "rationale": rationale,
    }


@app.post("/api/chat/data-quality-agent")
async def chat_data_quality_agent(req: DataQualityAgentRequest):
    raise HTTPException(status_code=410, detail="Data Quality - Tables has been removed from this application.")

    _emit_log("info", "data_quality", "Received request", {"query": req.message})
    user_message = (req.message or "").strip()
    normalized_choice = normalize_choice(user_message)
    normalized_lower = normalized_choice.lower()
    conversation_memory = _conversation_memory_markdown(req.history, current_message=user_message)
    state = _normalize_data_quality_state(req.agent_state.model_dump())
    state["agent_id"] = "data_quality_tables"
    state["db_type"] = "clickhouse"
    state["session_id"] = state.get("session_id") or uuid.uuid4().hex
    state["last_error"] = ""


    async def _reload_tables(current_state: dict[str, Any]) -> dict[str, Any]:
        try:
            return await data_quality_schema_node(req.clickhouse, current_state)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"ClickHouse schema discovery failed: {exc}") from exc

    async def _run_data_quality_analysis(current_state: dict[str, Any]) -> dict[str, Any]:
        execution_state = dict(current_state)
        steps = [
            {
                "id": "dq-schema-node",
                "title": "schema_node",
                "status": "success",
                "details": f"Loaded schema metadata for `{execution_state.get('table')}`.",
            }
        ]

        execution_state = await data_quality_stats_node(req.clickhouse, execution_state)
        steps.append(
            {
                "id": "dq-stats-node",
                "title": "stats_node",
                "status": "success",
                "details": f"Profiled {len(execution_state.get('columns') or [])} column(s) with statistical SQL.",
            }
        )

        if execution_state.get("time_column"):
            execution_state = await data_quality_volumetric_node(req.clickhouse, execution_state)
            steps.append(
                {
                    "id": "dq-volumetric-node",
                    "title": "volumetric_node",
                    "status": "success",
                    "details": (
                        "Computed volumetric patterns for the selected time column."
                        if execution_state.get("volumetric_stats")
                        else "Volumetric analysis was requested, but no usable time buckets were found."
                    ),
                }
            )

        start_llm = _time.time()
        llm_analysis = await data_quality_llm_analysis_node(
            execution_state,
            conversation_memory,
            req.llm_base_url,
            req.llm_model,
            req.llm_provider,
            req.llm_api_key,
        )
        llm_elapsed = _time.time() - start_llm
        _emit_log("llm", "data_quality", f"Generated LLM analysis ({llm_elapsed:.2f}s)", {})
        
        steps.append(
            {
                "id": "dq-llm-analysis-node",
                "title": "llm_analysis_node",
                "status": "success",
                "details": "The local LLM scored the dataset and generated recommendations.",
            }
        )


        execution_state["final_answer"] = data_quality_synthesizer_node(execution_state, llm_analysis)
        execution_state["stage"] = "ready"
        execution_state["last_error"] = ""
        steps.append(
            {
                "id": "dq-synthesizer-node",
                "title": "synthesizer_node",
                "status": "success",
                "details": "Built the final Markdown report in English.",
            }
        )
        table_slug = re.sub(r"[^a-z0-9]+", "-", str(execution_state.get("table") or "data-quality").strip().lower()).strip("-") or "data-quality"
        return {
            "answer": execution_state["final_answer"],
            "agent_state": execution_state,
            "steps": steps,
            "actions": [
                {
                    "id": "dq-export-pdf",
                    "label": "Export Summary PDF",
                    "actionType": "export_data_quality_pdf",
                    "variant": "secondary",
                    "payload": {
                        "fileName": f"data-quality-{table_slug}.pdf",
                        "title": f"Data Quality Summary - {execution_state.get('table') or 'Table'}",
                    },
                }
            ],
        }

    async def _run_data_quality_analysis_safe(current_state: dict[str, Any]) -> dict[str, Any]:
        try:
            return await _run_data_quality_analysis(current_state)
        except HTTPException:
            raise
        except Exception as exc:
            current_state["last_error"] = str(exc)
            raise HTTPException(status_code=400, detail=f"Data-quality analysis failed: {exc}") from exc

    start_over_requested = any(
        token in normalized_lower for token in ["start over", "reset", "new analysis", "clear analysis"]
    )
    if start_over_requested:
        state = _default_data_quality_state()
        state["agent_id"] = "data_quality_tables"
        state["db_type"] = "clickhouse"
        state["session_id"] = uuid.uuid4().hex

    state = await _reload_tables(state)
    if not state.get("available_tables"):
        raise HTTPException(status_code=400, detail="No tables were found in the configured ClickHouse database.")

    direct_payload = _try_extract_data_quality_payload(user_message)
    if direct_payload:
        table_name = str(direct_payload.get("table") or "").strip()
        if not table_name:
            raise HTTPException(status_code=400, detail="The structured data-quality payload must include a table.")
        matched_table = _data_quality_guess_table_from_message(table_name, state["available_tables"])
        if not matched_table:
            raise HTTPException(status_code=400, detail=f"Unknown table in the data-quality payload: {table_name}")

        state["table"] = matched_table
        payload_sample_size = direct_payload.get("sample_size")
        if isinstance(payload_sample_size, (int, float)):
            state["sample_size"] = max(0, min(DATA_QUALITY_MAX_SAMPLE_ROWS, int(payload_sample_size)))
        else:
            state["sample_size"] = DATA_QUALITY_DEFAULT_SAMPLE_SIZE
        state["row_filter"] = str(direct_payload.get("row_filter") or "").strip()
        validation_error = _validate_data_quality_row_filter(state["row_filter"])
        if validation_error:
            raise HTTPException(status_code=400, detail=validation_error)

        state = await _reload_tables(state)
        requested_columns = direct_payload.get("columns")
        if isinstance(requested_columns, list) and requested_columns:
            state["columns"] = _match_data_quality_columns(
                [str(item).strip() for item in requested_columns if isinstance(item, str)],
                state["schema_info"],
            )
            if not state["columns"]:
                raise HTTPException(status_code=400, detail="None of the requested columns were found in the selected table.")
        else:
            state["columns"] = [column["name"] for column in state.get("schema_info") or [] if column.get("name")]
            if not state["columns"]:
                raise HTTPException(status_code=400, detail="The selected table has no readable columns.")

        requested_time_column = str(direct_payload.get("time_column") or "").strip()
        if requested_time_column:
            matched_time = resolve_user_choice(requested_time_column, state.get("date_columns") or [])
            if not matched_time:
                raise HTTPException(status_code=400, detail="The structured time_column must match a date-like column in the selected table.")
            state["time_column"] = matched_time
        else:
            state["time_column"] = None

        return await _run_data_quality_analysis_safe(state)

    if state.get("stage") == "ready" and user_message:
        state = _default_data_quality_state()
        state["agent_id"] = "data_quality_tables"
        state["db_type"] = "clickhouse"
        state["session_id"] = uuid.uuid4().hex
        state = await _reload_tables(state)

    table_options = _data_quality_table_options(state)

    if not user_message or normalized_lower in {"start guided setup", "guide me", "guided setup", "start setup", "start"}:
        state["stage"] = "awaiting_table"
        answer = append_choice_markdown(
            _data_quality_intro_markdown(req.clickhouse.database, table_options, len(state["available_tables"])),
            "Table Selection",
            "Choose the table you want to profile first.",
            table_options,
        )
        return {
            "answer": answer,
            "agent_state": state,
            "steps": _data_quality_agent_steps(
                "dq-guided-start",
                "Started guided setup",
                "running",
                f"Loaded {len(state['available_tables'])} table(s) and waiting for the table selection.",
            ),
        }

    if not state.get("table"):
        guessed_table = _data_quality_guess_table_from_message(user_message, state["available_tables"])
        if guessed_table:
            state["table"] = guessed_table
            state["stage"] = "awaiting_columns_mode"
            state["columns"] = []
            state["column_stats"] = {}
            state["volumetric_stats"] = None
            state["llm_analysis"] = ""
            state["final_answer"] = ""
            state["time_column"] = None
            state = await _reload_tables(state)
        else:
            state["stage"] = "awaiting_table"
            answer = append_choice_markdown(
                _data_quality_intro_markdown(req.clickhouse.database, table_options, len(state["available_tables"])),
                "Table Selection",
                "Choose the table you want to profile first.",
                table_options,
            )
            return {
                "answer": answer,
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-await-table",
                    "Waiting for table",
                    "running",
                    "The agent needs the target table before it can configure profiling parameters.",
                ),
            }

    if state.get("table") and not state.get("schema_info"):
        state = await _reload_tables(state)
    if state.get("table") and not state.get("schema_info"):
        raise HTTPException(status_code=400, detail=f"Table '{state['table']}' has no readable schema.")

    if state.get("stage") == "awaiting_table":
        selected_table = _data_quality_guess_table_from_message(user_message, state["available_tables"])
        if not selected_table:
            return {
                "answer": append_choice_markdown(
                    _data_quality_intro_markdown(req.clickhouse.database, table_options, len(state["available_tables"])),
                    "Table Selection",
                    "Choose the table you want to profile first.",
                    table_options,
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-await-table",
                    "Waiting for table",
                    "running",
                    "The selected agent is waiting for the table choice.",
                ),
            }
        state["table"] = selected_table
        state["stage"] = "awaiting_columns_mode"
        state["columns"] = []
        state["column_stats"] = {}
        state["volumetric_stats"] = None
        state["llm_analysis"] = ""
        state["final_answer"] = ""
        state["time_column"] = None
        state = await _reload_tables(state)

    guessed_columns = []
    if state.get("stage") in {"idle", "awaiting_columns_mode"}:
        guessed_columns = _data_quality_guess_columns_from_message(user_message, state.get("schema_info") or [])
        if guessed_columns and not state.get("columns"):
            state["columns"] = guessed_columns
            state["stage"] = "awaiting_sample_size"

    if state.get("stage") == "awaiting_columns_mode" and not state.get("columns"):
        column_mode_options = _data_quality_column_mode_options(state)
        selected_mode = resolve_user_choice(user_message, column_mode_options)
        if selected_mode == DATA_QUALITY_CUSTOM_COLUMNS_OPTION:
            state["stage"] = "awaiting_custom_columns"
            preview_columns = (state.get("available_columns") or [])[:20]
            answer = (
                "## Custom Column Selection\n"
                "Type the exact column names separated by commas or new lines.\n\n"
                "Available columns preview:\n"
                + "\n".join(f"- `{column}`" for column in preview_columns)
            )
            return {
                "answer": answer,
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-custom-columns",
                    "Waiting for custom columns",
                    "running",
                    "The user chose to provide a custom column list.",
                ),
            }

        if not selected_mode and not guessed_columns:
            return {
                "answer": build_choice_markdown(
                    "Column Scope",
                    f"I loaded `{len(state.get('available_columns') or [])}` column(s) from `{state['table']}`. Choose the profiling scope.",
                    column_mode_options,
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-columns-mode",
                    "Waiting for column scope",
                    "running",
                    "The agent is waiting for the user to define which columns should be profiled.",
                ),
            }

        if selected_mode:
            selected_columns = _data_quality_columns_for_mode(selected_mode, state.get("schema_info") or [])
            if not selected_columns:
                return {
                    "answer": (
                        "## Column Scope\n"
                        "The selected scope did not resolve to any column in this table. Please choose another option."
                    ),
                    "agent_state": state,
                    "steps": _data_quality_agent_steps(
                        "dq-columns-empty",
                        "Column scope empty",
                        "error",
                        "The selected column category returned no usable column.",
                    ),
                }
            state["columns"] = selected_columns
            state["stage"] = "awaiting_sample_size"

    if state.get("stage") == "awaiting_custom_columns":
        matched_columns = _match_data_quality_columns(
            _parse_custom_column_input(user_message),
            state.get("schema_info") or [],
        )
        if not matched_columns:
            preview_columns = (state.get("available_columns") or [])[:20]
            answer = (
                "## Custom Column Selection\n"
                "I could not match those names to the selected table. Please type exact column names separated by commas or new lines.\n\n"
                "Available columns preview:\n"
                + "\n".join(f"- `{column}`" for column in preview_columns)
            )
            return {
                "answer": answer,
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-custom-columns",
                    "Waiting for custom columns",
                    "running",
                    "The provided custom columns did not match the table schema.",
                ),
            }
        state["columns"] = matched_columns
        state["stage"] = "awaiting_sample_size"

    if state.get("stage") == "awaiting_sample_size":
        sample_options = list(DATA_QUALITY_SAMPLE_OPTIONS.keys()) + [DATA_QUALITY_CUSTOM_SAMPLE_OPTION]
        selected_sample = resolve_user_choice(user_message, sample_options)
        if not selected_sample:
            number_match = re.search(r"\b(\d[\d\s_,]*)\b", user_message)
            if number_match:
                parsed_number = int(re.sub(r"[^\d]", "", number_match.group(1)))
                state["sample_size"] = max(0, min(DATA_QUALITY_MAX_SAMPLE_ROWS, parsed_number))
                state["stage"] = "awaiting_row_filter_mode"
                return {
                    "answer": build_choice_markdown(
                        "Row Filter",
                        "Choose whether to profile all rows or enter a manual row filter.",
                        [DATA_QUALITY_SKIP_ROW_FILTER_OPTION, DATA_QUALITY_ENTER_ROW_FILTER_OPTION],
                    ),
                    "agent_state": state,
                    "steps": _data_quality_agent_steps(
                        "dq-row-filter-mode",
                        "Waiting for row filter mode",
                        "running",
                        "The sampling strategy is set, and the agent is waiting for the optional row-filter choice.",
                    ),
                }
            else:
                return {
                    "answer": build_choice_markdown(
                        "Sample Size",
                        "Choose how many rows should be profiled. Full scans are safety-capped.",
                        sample_options,
                    ),
                    "agent_state": state,
                    "steps": _data_quality_agent_steps(
                        "dq-sample",
                        "Waiting for sample size",
                        "running",
                        "The agent is waiting for the sampling strategy.",
                    ),
                }
        elif selected_sample == DATA_QUALITY_CUSTOM_SAMPLE_OPTION:
            state["stage"] = "awaiting_custom_sample_size"
            return {
                "answer": (
                    "## Custom Sample Size\n"
                    f"Type the number of rows to profile. Use `0` for a capped full scan up to {DATA_QUALITY_MAX_SAMPLE_ROWS:,} rows."
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-custom-sample",
                    "Waiting for custom sample size",
                    "running",
                    "The user chose to enter a custom sample size.",
                ),
            }
        else:
            state["sample_size"] = DATA_QUALITY_SAMPLE_OPTIONS[selected_sample]
            state["stage"] = "awaiting_row_filter_mode"
            return {
                "answer": build_choice_markdown(
                    "Row Filter",
                    "Choose whether to profile all rows or enter a manual row filter.",
                    [DATA_QUALITY_SKIP_ROW_FILTER_OPTION, DATA_QUALITY_ENTER_ROW_FILTER_OPTION],
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-row-filter-mode",
                    "Waiting for row filter mode",
                    "running",
                    "The sampling strategy is set, and the agent is waiting for the optional row-filter choice.",
                ),
            }

    if state.get("stage") == "awaiting_custom_sample_size":
        number_match = re.search(r"\b(\d[\d\s_,]*)\b", user_message)
        if not number_match and normalized_lower not in {"0", "full scan"}:
            return {
                "answer": (
                    "## Custom Sample Size\n"
                    f"Please enter a numeric row count. Use `0` for a capped full scan up to {DATA_QUALITY_MAX_SAMPLE_ROWS:,} rows."
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-custom-sample",
                    "Waiting for custom sample size",
                    "running",
                    "The custom sample size must be numeric.",
                ),
            }
        parsed_number = 0 if normalized_lower in {"0", "full scan"} else int(re.sub(r"[^\d]", "", number_match.group(1)))
        state["sample_size"] = max(0, min(DATA_QUALITY_MAX_SAMPLE_ROWS, parsed_number))
        state["stage"] = "awaiting_row_filter_mode"
        return {
            "answer": build_choice_markdown(
                "Row Filter",
                "Choose whether to profile all rows or enter a manual row filter.",
                [DATA_QUALITY_SKIP_ROW_FILTER_OPTION, DATA_QUALITY_ENTER_ROW_FILTER_OPTION],
            ),
            "agent_state": state,
            "steps": _data_quality_agent_steps(
                "dq-row-filter-mode",
                "Waiting for row filter mode",
                "running",
                "The custom sampling strategy is set, and the agent is waiting for the optional row-filter choice.",
            ),
        }

    if state.get("stage") == "awaiting_row_filter_mode":
        row_filter_options = [DATA_QUALITY_SKIP_ROW_FILTER_OPTION, DATA_QUALITY_ENTER_ROW_FILTER_OPTION]
        selected_mode = resolve_user_choice(user_message, row_filter_options)

        if not selected_mode and normalized_lower in {"skip", "no filter", "without filter", "none", "all rows"}:
            selected_mode = DATA_QUALITY_SKIP_ROW_FILTER_OPTION

        if selected_mode == DATA_QUALITY_SKIP_ROW_FILTER_OPTION:
            state["row_filter"] = ""
            state["stage"] = "awaiting_time_column" if state.get("date_columns") else "awaiting_review"

            if state.get("stage") == "awaiting_time_column":
                time_options = [DATA_QUALITY_SKIP_TIME_OPTION] + list(state.get("date_columns") or [])
                return {
                    "answer": build_choice_markdown(
                        "Volumetric Analysis",
                        "Choose a time column if you also want a volume-over-time analysis.",
                        time_options,
                    ),
                    "agent_state": state,
                    "steps": _data_quality_agent_steps(
                        "dq-time-column",
                        "Waiting for time column",
                        "running",
                        "The row filter was skipped, and the agent is waiting for the optional volumetric-analysis choice.",
                    ),
                }

            return {
                "answer": append_choice_markdown(
                    _data_quality_review_markdown(state),
                    "Review Actions",
                    "Choose the next action for this data-quality run.",
                    DATA_QUALITY_REVIEW_OPTIONS,
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-review",
                    "Ready to launch",
                    "running",
                    "The row filter choice is complete, and the run is ready for review.",
                ),
            }

        if selected_mode == DATA_QUALITY_ENTER_ROW_FILTER_OPTION:
            state["stage"] = "awaiting_row_filter"
            return {
                "answer": append_choice_markdown(
                    (
                        "## Row Filter\n"
                        "Type a safe boolean expression such as `region = 'FR'`."
                    ),
                    "Quick Choice",
                    "Or skip the row filter for this run.",
                    [DATA_QUALITY_SKIP_ROW_FILTER_OPTION],
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-row-filter",
                    "Waiting for row filter",
                    "running",
                    "The agent is waiting for a manual row filter expression.",
                ),
            }

        if user_message:
            validation_error = _validate_data_quality_row_filter(user_message)
            if not validation_error:
                state["row_filter"] = user_message.strip()
                state["stage"] = "awaiting_time_column" if state.get("date_columns") else "awaiting_review"

                if state.get("stage") == "awaiting_time_column":
                    time_options = [DATA_QUALITY_SKIP_TIME_OPTION] + list(state.get("date_columns") or [])
                    return {
                        "answer": build_choice_markdown(
                            "Volumetric Analysis",
                            "Choose a time column if you also want a volume-over-time analysis.",
                            time_options,
                        ),
                        "agent_state": state,
                        "steps": _data_quality_agent_steps(
                            "dq-time-column",
                            "Waiting for time column",
                            "running",
                            "The row filter is set, and the agent is waiting for the optional volumetric-analysis choice.",
                        ),
                    }

                return {
                    "answer": append_choice_markdown(
                        _data_quality_review_markdown(state),
                        "Review Actions",
                        "Choose the next action for this data-quality run.",
                        DATA_QUALITY_REVIEW_OPTIONS,
                    ),
                    "agent_state": state,
                    "steps": _data_quality_agent_steps(
                        "dq-review",
                        "Ready to launch",
                        "running",
                        "The row filter is set, and the run is ready for review.",
                    ),
                }

        return {
            "answer": build_choice_markdown(
                "Row Filter",
                "Choose whether to profile all rows or enter a manual row filter.",
                row_filter_options,
            ),
            "agent_state": state,
            "steps": _data_quality_agent_steps(
                "dq-row-filter-mode",
                "Waiting for row filter mode",
                "running",
                "The agent is waiting for the optional row-filter choice.",
            ),
        }

    if state.get("stage") == "awaiting_row_filter":
        if (
            not user_message
            or normalized_lower in {"skip", "no filter", "without filter", "none", "all rows"}
            or resolve_user_choice(user_message, [DATA_QUALITY_SKIP_ROW_FILTER_OPTION]) == DATA_QUALITY_SKIP_ROW_FILTER_OPTION
        ):
            state["row_filter"] = ""
            state["stage"] = "awaiting_time_column" if state.get("date_columns") else "awaiting_review"
        else:
            validation_error = _validate_data_quality_row_filter(user_message)
            if validation_error:
                return {
                    "answer": append_choice_markdown(
                        (
                            "## Row Filter\n"
                            f"{validation_error}\n\n"
                            "Please type a safe boolean expression."
                        ),
                        "Quick Choice",
                        "Or skip the row filter for this run.",
                        [DATA_QUALITY_SKIP_ROW_FILTER_OPTION],
                    ),
                    "agent_state": state,
                    "steps": _data_quality_agent_steps(
                        "dq-row-filter",
                        "Waiting for row filter",
                        "running",
                        "The proposed row filter did not pass the safety validation.",
                    ),
                }
            state["row_filter"] = user_message.strip()
            state["stage"] = "awaiting_time_column" if state.get("date_columns") else "awaiting_review"

        if state.get("stage") == "awaiting_time_column":
            time_options = [DATA_QUALITY_SKIP_TIME_OPTION] + list(state.get("date_columns") or [])
            return {
                "answer": build_choice_markdown(
                    "Volumetric Analysis",
                    "Choose a time column if you also want a volume-over-time analysis.",
                    time_options,
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-time-column",
                    "Waiting for time column",
                    "running",
                    "The agent is waiting for the optional volumetric-analysis choice.",
                ),
            }

    if state.get("stage") == "awaiting_time_column":
        time_options = [DATA_QUALITY_SKIP_TIME_OPTION] + list(state.get("date_columns") or [])
        selected_time = resolve_user_choice(user_message, time_options)
        if not selected_time:
            return {
                "answer": build_choice_markdown(
                    "Volumetric Analysis",
                    "Choose a time column if you also want a volume-over-time analysis.",
                    time_options,
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-time-column",
                    "Waiting for time column",
                    "running",
                    "The agent is waiting for the optional volumetric-analysis choice.",
                ),
            }
        state["time_column"] = None if selected_time == DATA_QUALITY_SKIP_TIME_OPTION else selected_time
        state["stage"] = "awaiting_review"
        return {
            "answer": append_choice_markdown(
                _data_quality_review_markdown(state),
                "Review Actions",
                "Choose the next action for this data-quality run.",
                DATA_QUALITY_REVIEW_OPTIONS,
            ),
            "agent_state": state,
            "steps": _data_quality_agent_steps(
                "dq-review",
                "Ready to launch",
                "running",
                "The optional volumetric-analysis choice is complete, and the run is ready for review.",
            ),
        }

    if state.get("stage") == "awaiting_review":
        launch_tokens = {"launch analysis", "run analysis", "launch", "run", "go", "analyze"}
        selected_review_action = resolve_user_choice(user_message, DATA_QUALITY_REVIEW_OPTIONS)
        if not selected_review_action and normalized_lower in launch_tokens:
            selected_review_action = "Launch analysis"

        if selected_review_action == "Launch analysis":
            return await _run_data_quality_analysis_safe(state)
        if selected_review_action == "Edit table":
            state["stage"] = "awaiting_table"
            return {
                "answer": append_choice_markdown(
                    _data_quality_intro_markdown(req.clickhouse.database, table_options, len(state["available_tables"])),
                    "Table Selection",
                    "Choose the table you want to profile first.",
                    table_options,
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-edit-table",
                    "Editing table",
                    "running",
                    "The user chose to change the target table.",
                ),
            }
        if selected_review_action == "Edit columns":
            state["columns"] = []
            state["stage"] = "awaiting_columns_mode"
        elif selected_review_action == "Edit sample size":
            state["stage"] = "awaiting_sample_size"
        elif selected_review_action == "Edit row filter":
            state["stage"] = "awaiting_row_filter_mode"
        elif selected_review_action == "Edit time column":
            if state.get("date_columns"):
                state["stage"] = "awaiting_time_column"
            else:
                state["time_column"] = None
        elif selected_review_action == "Start over":
            state = _default_data_quality_state()
            state["agent_id"] = "data_quality_tables"
            state["db_type"] = "clickhouse"
            state["session_id"] = uuid.uuid4().hex
            state = await _reload_tables(state)
            state["stage"] = "awaiting_table"
            table_options = _data_quality_table_options(state)
            return {
                "answer": append_choice_markdown(
                    _data_quality_intro_markdown(req.clickhouse.database, table_options, len(state["available_tables"])),
                    "Table Selection",
                    "Choose the table you want to profile first.",
                    table_options,
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-restart",
                    "Restarted setup",
                    "running",
                    "The guided data-quality setup was reset.",
                ),
            }

        if state.get("stage") == "awaiting_columns_mode":
            return {
                "answer": build_choice_markdown(
                    "Column Scope",
                    f"I loaded `{len(state.get('available_columns') or [])}` column(s) from `{state['table']}`. Choose the profiling scope.",
                    _data_quality_column_mode_options(state),
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-columns-mode",
                    "Waiting for column scope",
                    "running",
                    "The agent is waiting for the user to redefine the profiling scope.",
                ),
            }
        if state.get("stage") == "awaiting_sample_size":
            return {
                "answer": build_choice_markdown(
                    "Sample Size",
                    "Choose how many rows should be profiled. Full scans are safety-capped.",
                    list(DATA_QUALITY_SAMPLE_OPTIONS.keys()) + [DATA_QUALITY_CUSTOM_SAMPLE_OPTION],
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-sample",
                    "Waiting for sample size",
                    "running",
                    "The agent is waiting for a new sample size.",
                ),
            }
        if state.get("stage") == "awaiting_row_filter_mode":
            return {
                "answer": build_choice_markdown(
                    "Row Filter",
                    "Choose whether to profile all rows or enter a manual row filter.",
                    [DATA_QUALITY_SKIP_ROW_FILTER_OPTION, DATA_QUALITY_ENTER_ROW_FILTER_OPTION],
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-row-filter-mode",
                    "Waiting for row filter mode",
                    "running",
                    "The agent is waiting for the optional row-filter choice.",
                ),
            }
        if state.get("stage") == "awaiting_row_filter":
            return {
                "answer": append_choice_markdown(
                    (
                        "## Row Filter\n"
                        "Type a safe boolean expression such as `region = 'FR'`."
                    ),
                    "Quick Choice",
                    "Or skip the row filter for this run.",
                    [DATA_QUALITY_SKIP_ROW_FILTER_OPTION],
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-row-filter",
                    "Waiting for row filter",
                    "running",
                    "The agent is waiting for the optional row filter.",
                ),
            }
        if state.get("stage") == "awaiting_time_column":
            return {
                "answer": build_choice_markdown(
                    "Volumetric Analysis",
                    "Choose a time column if you also want a volume-over-time analysis.",
                    [DATA_QUALITY_SKIP_TIME_OPTION] + list(state.get("date_columns") or []),
                ),
                "agent_state": state,
                "steps": _data_quality_agent_steps(
                    "dq-time-column",
                    "Waiting for time column",
                    "running",
                    "The agent is waiting for the optional volumetric-analysis choice.",
                ),
            }

        return {
            "answer": append_choice_markdown(
                _data_quality_review_markdown(state),
                "Review Actions",
                "Choose the next action for this data-quality run.",
                DATA_QUALITY_REVIEW_OPTIONS,
            ),
            "agent_state": state,
            "steps": _data_quality_agent_steps(
                "dq-review",
                "Ready to launch",
                "running",
                "The analysis parameters are ready and waiting for the final launch decision.",
            ),
        }

    if state.get("columns") and state.get("stage") == "awaiting_sample_size":
        return {
            "answer": build_choice_markdown(
                "Sample Size",
                "Choose how many rows should be profiled. Full scans are safety-capped.",
                list(DATA_QUALITY_SAMPLE_OPTIONS.keys()) + [DATA_QUALITY_CUSTOM_SAMPLE_OPTION],
            ),
            "agent_state": state,
            "steps": _data_quality_agent_steps(
                "dq-sample",
                "Waiting for sample size",
                "running",
                "The agent is waiting for the sampling strategy.",
            ),
        }

    state["stage"] = "awaiting_review"
    return {
        "answer": append_choice_markdown(
            _data_quality_review_markdown(state),
            "Review Actions",
            "Choose the next action for this data-quality run.",
            DATA_QUALITY_REVIEW_OPTIONS,
        ),
        "agent_state": state,
        "steps": _data_quality_agent_steps(
            "dq-review",
            "Ready to launch",
            "running",
            "The analysis parameters are ready and waiting for the final launch decision.",
        ),
    }


@app.post("/api/chat/data-analyst-agent")
async def chat_data_analyst_agent(req: DataAnalystAgentRequest):
    _emit_log("info", "data_analyst", "Received request", {"query": req.message})
    user_message = (req.message or "").strip()
    state = _normalize_data_analyst_state(req.agent_state.model_dump())
    max_steps = max(1, min(int(req.max_steps or DATA_ANALYST_DEFAULT_MAX_STEPS), DATA_ANALYST_MAX_STEPS))
    memory_anchor = state.get("pending_request") or user_message
    conversation_memory = _conversation_memory_markdown(
        req.history,
        current_message=memory_anchor,
        max_steps=CHAT_MEMORY_MAX_STEPS,
    )


    try:
        state["available_tables"] = state["available_tables"] or await list_clickhouse_tables(req.clickhouse)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"ClickHouse connection error: {exc}") from exc

    if not state["available_tables"]:
        raise HTTPException(status_code=400, detail="No tables were found in the configured ClickHouse database.")

    explicit_table_switch = resolve_user_choice(user_message, state["available_tables"])
    if state.get("selected_table") and explicit_table_switch and explicit_table_switch != state.get("selected_table"):
        state["selected_table"] = explicit_table_switch
        state["table_schema"] = []
        state["analysis_goal"] = ""
        state["analysis_plan"] = []
        state["success_criteria"] = []
        state["latest_review"] = ""
        state["last_sqls"] = []
        state["last_result_meta"] = []
        state["last_result_rows"] = []
        state["final_answer"] = ""
        state["last_error"] = ""
        state["last_export_path"] = ""
        state["knowledge_hits"] = []
        state["pending_request"] = ""
        state["clarification_prompt"] = ""
        state["clarification_options"] = []
        state["stage"] = "ready"

    if not user_message:
        return {
            "answer": (
                "## Data Analyst Agent\n"
                "Ask a complex business question about your ClickHouse data, and I will gather evidence through several analytical steps.\n\n"
                f"- **Database:** `{req.clickhouse.database}`\n"
                f"- **Max credited actions:** {max_steps}\n"
                "- **Capabilities:** deeper multi-step ReAct analysis, subqueries when they help, optional knowledge-base lookup, safe SQL retry, and CSV export when you explicitly request it.\n"
                f"- **Current table focus:** `{state.get('selected_table') or 'not selected yet'}`"
            ),
            "agent_state": state,
            "steps": [
                {
                    "id": "data-analyst-ready",
                    "title": "Ready for multi-step analysis",
                    "status": "success",
                    "details": "The agent is ready to inspect ClickHouse tables, run iterative queries, and synthesize a business-facing answer.",
                }
            ],
        }

    if state.get("stage") == "awaiting_table":
        table_options = state.get("clarification_options") or state["available_tables"][:DATA_ANALYST_TABLE_OPTION_LIMIT]
        selected_table = resolve_user_choice(user_message, table_options) or resolve_user_choice(user_message, state["available_tables"])
        if not selected_table:
            return {
                "answer": build_choice_markdown(
                    "Table Clarification",
                    state.get("clarification_prompt") or "Which table should I use for this analysis?",
                    table_options,
                ),
                "agent_state": state,
                "steps": [
                    {
                        "id": "data-analyst-await-table",
                        "title": "Waiting for table selection",
                        "status": "running",
                        "details": "The analysis needs one primary ClickHouse table before it can continue.",
                    }
                ],
            }
        state["selected_table"] = selected_table
        state["stage"] = "ready"
        state["clarification_prompt"] = ""
        state["clarification_options"] = []
        user_message = state.get("pending_request") or user_message

    if state.get("stage") == "awaiting_row_intent":
        row_intent_options = state.get("clarification_options") or [CLICKHOUSE_ROW_COUNT_OPTION, CLICKHOUSE_SAMPLE_ROWS_OPTION]
        row_intent_choice = resolve_user_choice(user_message, row_intent_options)
        if not row_intent_choice:
            return {
                "answer": build_choice_markdown(
                    "Row Request Clarification",
                    state.get("clarification_prompt") or "Do you want the row count or a sample list of rows?",
                    row_intent_options,
                ),
                "agent_state": state,
                "steps": [
                    {
                        "id": "data-analyst-await-row-intent",
                        "title": "Waiting for row intent clarification",
                        "status": "running",
                        "details": "The analysis needs to know whether to count rows or preview rows.",
                    }
                ],
            }

        original_request = str(state.get("pending_request") or "").strip()
        if row_intent_choice == CLICKHOUSE_ROW_COUNT_OPTION:
            state["pending_request"] = f"{original_request}\n\n[Resolved intent: row count]".strip()
        else:
            state["pending_request"] = f"{original_request}\n\n[Resolved intent: sample rows]".strip()
        state["stage"] = "ready"
        state["clarification_prompt"] = ""
        state["clarification_options"] = []
        user_message = state["pending_request"]

    current_request = state.get("pending_request") or user_message
    if not state.get("selected_table"):
        direct_table = resolve_user_choice(current_request, state["available_tables"])
        if direct_table:
            state["selected_table"] = direct_table
            state["stage"] = "ready"
        elif len(state["available_tables"]) == 1:
            state["selected_table"] = state["available_tables"][0]
            state["stage"] = "ready"
        else:
            state["pending_request"] = current_request
            table_analysis = await analyze_clickhouse_tables(
                current_request,
                state["available_tables"],
                conversation_memory,
                req.llm_base_url,
                req.llm_model,
                req.llm_provider,
                req.llm_api_key,
            )
            matched_candidates = match_available_options(
                table_analysis.get("table_candidates") or [],
                state["available_tables"],
            )
            matched_selected = match_available_options(
                [table_analysis.get("selected_table") or ""],
                state["available_tables"],
            )
            if matched_selected and not table_analysis.get("table_choice_required"):
                state["selected_table"] = matched_selected[0]
                state["stage"] = "ready"
            elif len(matched_candidates) == 1 and not table_analysis.get("table_choice_required"):
                state["selected_table"] = matched_candidates[0]
                state["stage"] = "ready"
            else:
                state["stage"] = "awaiting_table"
                state["clarification_prompt"] = (
                    str(table_analysis.get("table_choice_prompt") or "").strip()
                    or "Which table should I use for this analysis?"
                )
                state["clarification_options"] = (matched_candidates or state["available_tables"])[:DATA_ANALYST_TABLE_OPTION_LIMIT]
                return {
                    "answer": build_choice_markdown(
                        "Table Clarification",
                        state["clarification_prompt"],
                        state["clarification_options"],
                    ),
                    "agent_state": state,
                    "steps": [
                        {
                            "id": "data-analyst-table-routing",
                            "title": "Need table confirmation",
                            "status": "running",
                            "details": str(table_analysis.get("reasoning") or "Several tables remain plausible for the current analysis request."),
                        }
                    ],
                }

    current_request = user_message or state.get("pending_request") or ""
    if is_clickhouse_row_request_ambiguous(current_request):
        state["pending_request"] = current_request
        state["stage"] = "awaiting_row_intent"
        state["clarification_prompt"] = (
            "Your request can mean either **count the rows** or **show sample rows**. "
            "Choose the output you want."
        )
        state["clarification_options"] = [CLICKHOUSE_ROW_COUNT_OPTION, CLICKHOUSE_SAMPLE_ROWS_OPTION]
        return {
            "answer": build_choice_markdown(
                "Row Request Clarification",
                state["clarification_prompt"],
                state["clarification_options"],
            ),
            "agent_state": state,
            "steps": [
                {
                    "id": "data-analyst-row-intent",
                    "title": "Need row intent clarification",
                    "status": "running",
                    "details": "The request is ambiguous between a row count and a sample row listing.",
                }
            ],
        }

    if not current_request:
        return {
            "answer": (
                f"## Data Analyst Agent\nI am ready to work from table `{state.get('selected_table')}`.\n\n"
                "Ask a complex analytical question, and I will decide which evidence-gathering steps to run next."
            ),
            "agent_state": state,
            "steps": [
                {
                    "id": "data-analyst-await-request",
                    "title": "Waiting for analytical request",
                    "status": "running",
                    "details": "The primary table is selected, but the user has not asked the actual analysis question yet.",
                }
            ],
        }

    state["pending_request"] = current_request
    state["clarification_prompt"] = ""
    state["clarification_options"] = []

    if not state.get("table_schema"):
        try:
            state["table_schema"] = await describe_clickhouse_table(req.clickhouse, str(state["selected_table"]))
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Failed to inspect schema for {state['selected_table']}: {exc}") from exc

    if not state["table_schema"]:
        raise HTTPException(status_code=400, detail=f"Table '{state['selected_table']}' has no readable columns.")

    if not state.get("analysis_goal"):
        try:
            mission = await _data_analyst_define_goal(
                current_request,
                str(state["selected_table"]),
                state["table_schema"],
                conversation_memory,
                max_steps,
                req.llm_base_url,
                req.llm_model,
                req.llm_provider,
                req.llm_api_key,
            )
            state["analysis_goal"] = mission.get("analysis_goal") or "Explain the business question through a structured evidence-gathering analysis."
            state["analysis_plan"] = mission.get("analysis_plan") or []
            state["success_criteria"] = mission.get("success_criteria") or []
        except Exception:
            state["analysis_goal"] = state.get("analysis_goal") or "Explain the business question through a structured evidence-gathering analysis."
            state["analysis_plan"] = state.get("analysis_plan") or []
            state["success_criteria"] = state.get("success_criteria") or []

    persisted_state = await read_db_state()
    knowledge_enabled = _app_opensearch_config((persisted_state or {}).get("config") or {}) is not None
    export_requested = _data_analyst_export_requested(current_request)
    if not export_requested:
        state["last_export_path"] = ""
    last_result_meta = list(state.get("last_result_meta") or [])
    last_result_rows = list(state.get("last_result_rows") or [])
    knowledge_hits = list(state.get("knowledge_hits") or [])
    analysis_goal = str(state.get("analysis_goal") or "").strip()
    analysis_plan = list(state.get("analysis_plan") or [])
    success_criteria = list(state.get("success_criteria") or [])
    latest_review = str(state.get("latest_review") or "").strip()
    executed_sqls: list[str] = []
    step_log: list[dict[str, Any]] = []
    last_error = ""
    absolute_iteration_limit = max_steps * 4

    for _ in range(absolute_iteration_limit):
        used_steps = len(step_log)
        if used_steps >= max_steps:
            break

        planned = await plan_data_analyst_step(
            current_request,
            str(state["selected_table"]),
            state["table_schema"],
            conversation_memory,
            analysis_goal,
            analysis_plan,
            success_criteria,
            latest_review,
            step_log,
            max_steps,
            used_steps,
            export_requested,
            knowledge_enabled,
            last_result_rows,
            req.llm_base_url,
            req.llm_model,
            req.llm_provider,
            req.llm_api_key,
        )
        action = str(planned.get("action") or "").strip().lower()
        reasoning = str(planned.get("reasoning") or "").strip() or "The local LLM selected the next analytical action."

        if action == "finish" and _data_analyst_needs_more_evidence(current_request, step_log, max_steps, latest_review) and len(step_log) + 1 < max_steps:
            planned = await plan_data_analyst_step(
                current_request,
                str(state["selected_table"]),
                state["table_schema"],
                conversation_memory,
                analysis_goal,
                analysis_plan,
                success_criteria,
                latest_review,
                step_log,
                max_steps,
                used_steps,
                export_requested,
                knowledge_enabled,
                last_result_rows,
                req.llm_base_url,
                req.llm_model,
                req.llm_provider,
                req.llm_api_key,
                must_continue=True,
            )
            action = str(planned.get("action") or "").strip().lower()
            reasoning = str(planned.get("reasoning") or "").strip() or "The local LLM selected the next analytical action."

        if action == "finish":
            final_body = str(planned.get("final_answer") or "").strip()
            if not final_body:
                final_body = await synthesize_data_analyst_answer(
                    current_request,
                    str(state["selected_table"]),
                    analysis_goal,
                    success_criteria,
                    latest_review,
                    conversation_memory,
                    step_log,
                    last_result_meta,
                    last_result_rows,
                    knowledge_hits,
                    forced_finish=False,
                    llm_base_url=req.llm_base_url,
                    llm_model=req.llm_model,
                    llm_provider=req.llm_provider,
                    llm_api_key=req.llm_api_key,
                )
            final_body = await polish_data_analyst_answer(
                current_request,
                str(state["selected_table"]),
                analysis_goal,
                success_criteria,
                latest_review,
                final_body,
                conversation_memory,
                step_log,
                last_result_rows,
                knowledge_hits,
                forced_finish=False,
                llm_base_url=req.llm_base_url,
                llm_model=req.llm_model,
                llm_provider=req.llm_provider,
                llm_api_key=req.llm_api_key,
            )
            confidence_score, confidence_reason = _data_analyst_confidence_score(step_log, forced_finish=False)
            final_answer = build_data_analyst_response_markdown(
                final_body,
                executed_sqls or list(state.get("last_sqls") or []),
                last_result_meta,
                last_result_rows,
                knowledge_hits,
                str(state.get("last_export_path") or ""),
                confidence_score,
                confidence_reason,
                step_log=step_log,
                force_visible_table=_user_explicitly_requests_table(current_request),
            )
            state["stage"] = "ready"
            state["pending_request"] = ""
            state["last_sqls"] = executed_sqls or list(state.get("last_sqls") or [])
            state["last_result_meta"] = last_result_meta
            state["last_result_rows"] = last_result_rows[:DATA_ANALYST_MAX_RESULT_ROWS]
            state["knowledge_hits"] = knowledge_hits[:DATA_ANALYST_MAX_KNOWLEDGE_RESULTS]
            state["final_answer"] = final_answer
            state["last_error"] = last_error
            
            _emit_log(
                "info", 
                "data_analyst", 
                "Completed analysis", 
                {"total_steps": len(step_log), "confidence": confidence_score}
            )
            
            return {
                "answer": final_answer,
                "agent_state": state,
                "steps": step_log + [
                    {
                        "id": "data-analyst-finish",
                        "title": "Prepared final answer",
                        "status": "success",
                        "details": reasoning,
                        "step": len(step_log) + 1,
                        "type": "finish",
                        "reasoning": reasoning,
                        "result_summary": "The agent consolidated the evidence into the final business-facing answer.",
                        "row_count": len(last_result_rows),
                        "ok": True,
                    }
                ],
                "total_steps": len(step_log),
            }

        if action == "search_knowledge":
            knowledge_query = str(planned.get("knowledge_query") or "").strip() or current_request
            search_result = await _data_analyst_search_knowledge(knowledge_query)
            if search_result.get("ok"):
                knowledge_hits = list(search_result.get("results") or [])
            else:
                last_error = str(search_result.get("summary") or "")
            step_log.append(
                _data_analyst_format_step(
                    len(step_log) + 1,
                    "search_knowledge",
                    reasoning,
                    str(search_result.get("summary") or "Knowledge search completed."),
                    len(search_result.get("results") or []),
                    bool(search_result.get("ok")),
                )
            )
            continue

        if action == "export_csv":
            export_path = _data_analyst_unique_export_path(
                str(planned.get("suggested_path") or "").strip()
                or _data_analyst_suggest_export_path(current_request, str(state["selected_table"]))
            )
            if not export_requested:
                last_error = "CSV export is only allowed when the user explicitly asks for it."
                step_log.append(
                    _data_analyst_format_step(
                        len(step_log) + 1,
                        "export_csv",
                        reasoning,
                        last_error,
                        0,
                        False,
                        suggested_path=export_path,
                    )
                )
                continue
            if not last_result_rows:
                last_error = "There is no dataset to export yet. Run at least one successful query first."
                step_log.append(
                    _data_analyst_format_step(
                        len(step_log) + 1,
                        "export_csv",
                        reasoning,
                        last_error,
                        0,
                        False,
                        suggested_path=export_path,
                    )
                )
                continue
            headers, rows = _manager_export_headers_and_rows(
                {
                    "last_result_meta": last_result_meta,
                    "last_result_rows": last_result_rows,
                }
            )
            try:
                result = create_file_tool(
                    export_path,
                    _serialize_delimited_rows(headers, rows, "|"),
                )
            except Exception as exc:
                last_error = str(exc)
                step_log.append(
                    _data_analyst_format_step(
                        len(step_log) + 1,
                        "export_csv",
                        reasoning,
                        f"CSV export failed: {exc}",
                        len(rows),
                        False,
                        suggested_path=export_path,
                    )
                )
                continue
            state["last_export_path"] = export_path
            
            _emit_log("info", "data_analyst", "Exported CSV", {"path": export_path, "rows": len(rows)})
            
            step_log.append(
                _data_analyst_format_step(
                    len(step_log) + 1,
                    "export_csv",
                    reasoning,
                    f"{result.get('summary') or 'The CSV export completed successfully.'} The file uses a pipe (`|`) delimiter.",
                    len(rows),
                    True,
                    suggested_path=export_path,
                )
            )
            continue

        if action != "query":
            last_error = "The analytical planner returned an unsupported action."
            step_log.append(
                _data_analyst_format_step(
                    len(step_log) + 1,
                    action or "unknown",
                    reasoning,
                    last_error,
                    0,
                    False,
                )
            )
            continue

        sql = enforce_query_limit(
            clean_sql_text(str(planned.get("sql") or "")),
            _data_analyst_dynamic_query_limit(current_request, str(planned.get("sql") or ""), req.clickhouse.query_limit),
        )
        retried = False
        if not _data_analyst_sql_is_valid(sql, state["table_schema"]):
            repaired = await repair_data_analyst_sql(
                current_request,
                str(state["selected_table"]),
                state["table_schema"],
                sql,
                "The proposed SQL was unsafe, empty, or used SELECT *.",
                req.llm_base_url,
                req.llm_model,
                req.llm_provider,
                req.llm_api_key,
                _data_analyst_dynamic_query_limit(current_request, sql, req.clickhouse.query_limit),
            )
            sql = repaired["sql"]
            reasoning = repaired["reasoning"] or reasoning
            retried = True

        if not _data_analyst_sql_is_valid(sql, state["table_schema"]):
            last_error = (
                "The generated SQL was rejected because it is unsafe, uses forbidden date logic, "
                "or does not follow the BETWEEN rule for date filters."
            )
            step_log.append(
                _data_analyst_format_step(
                    len(step_log) + 1,
                    "query",
                    reasoning,
                    last_error,
                    0,
                    False,
                    sql=sql,
                    retried=retried,
                )
            )
            continue

        try:
            await execute_clickhouse_sql(req.clickhouse, f"EXPLAIN SYNTAX {sql}", readonly=False, json_format=False)
            result = await execute_clickhouse_sql(req.clickhouse, sql)
        except Exception as first_error:
            repaired = await repair_data_analyst_sql(
                current_request,
                str(state["selected_table"]),
                state["table_schema"],
                sql,
                str(first_error),
                req.llm_base_url,
                req.llm_model,
                req.llm_provider,
                req.llm_api_key,
                _data_analyst_dynamic_query_limit(current_request, sql, req.clickhouse.query_limit),
            )
            repaired_sql = repaired["sql"]
            if not _data_analyst_sql_is_valid(repaired_sql, state["table_schema"]):
                last_error = f"Query failed and the repaired SQL was still unsafe.\nOriginal error: {first_error}"
                step_log.append(
                    _data_analyst_format_step(
                        len(step_log) + 1,
                        "query",
                        reasoning,
                        last_error,
                        0,
                        False,
                        sql=sql,
                        retried=True,
                    )
                )
                continue
            try:
                await execute_clickhouse_sql(req.clickhouse, f"EXPLAIN SYNTAX {repaired_sql}", readonly=False, json_format=False)
                result = await execute_clickhouse_sql(req.clickhouse, repaired_sql)
                sql = repaired_sql
                reasoning = repaired["reasoning"] or reasoning
                retried = True
            except Exception as second_error:
                last_error = f"Query failed twice. First error: {first_error}. Retry error: {second_error}"
                step_log.append(
                    _data_analyst_format_step(
                        len(step_log) + 1,
                        "query",
                        reasoning,
                        last_error,
                        0,
                        False,
                        sql=repaired_sql,
                        retried=True,
                    )
                )
                continue

        last_result_meta = result.get("meta", [])
        last_result_rows = (result.get("data") or [])[:DATA_ANALYST_MAX_RESULT_ROWS]
        executed_sqls.append(sql)
        state["last_export_path"] = ""
        summary = _data_analyst_compact_query_summary(last_result_meta, last_result_rows)
        
        _emit_log(
            "info", 
            "data_analyst", 
            "Executed query", 
            {"sql": sql, "rows": len(last_result_rows), "retried": retried}
        )
        
        step_log.append(
            _data_analyst_format_step(
                len(step_log) + 1,
                "query",
                reasoning,
                summary,
                len(result.get("data") or []),
                True,
                sql=sql,
                retried=retried,
            )
        )

        if len(step_log) % DATA_ANALYST_REVIEW_INTERVAL == 0:
            try:
                review = await _data_analyst_review_progress(
                    current_request,
                    str(state["selected_table"]),
                    analysis_goal,
                    analysis_plan,
                    success_criteria,
                    step_log,
                    last_result_rows,
                    knowledge_hits,
                    max_steps,
                    req.llm_base_url,
                    req.llm_model,
                    req.llm_provider,
                    req.llm_api_key,
                )
                latest_review = " | ".join(
                    part for part in [
                        str(review.get("review_summary") or "").strip(),
                        f"Next focus: {str(review.get('next_focus') or '').strip()}" if str(review.get("next_focus") or "").strip() else "",
                        "Enough evidence to conclude." if review.get("enough_evidence") else "",
                    ] if part
                ).strip()
                state["latest_review"] = latest_review
            except Exception:
                pass

    final_body = await synthesize_data_analyst_answer(
        current_request,
        str(state["selected_table"]),
        analysis_goal,
        success_criteria,
        latest_review,
        conversation_memory,
        step_log,
        last_result_meta,
        last_result_rows,
        knowledge_hits,
        forced_finish=True,
        llm_base_url=req.llm_base_url,
        llm_model=req.llm_model,
        llm_provider=req.llm_provider,
        llm_api_key=req.llm_api_key,
    )
    final_body = await polish_data_analyst_answer(
        current_request,
        str(state["selected_table"]),
        analysis_goal,
        success_criteria,
        latest_review,
        final_body,
        conversation_memory,
        step_log,
        last_result_rows,
        knowledge_hits,
        forced_finish=True,
        llm_base_url=req.llm_base_url,
        llm_model=req.llm_model,
        llm_provider=req.llm_provider,
        llm_api_key=req.llm_api_key,
    )
    confidence_score, confidence_reason = _data_analyst_confidence_score(step_log, forced_finish=True)
    final_answer = build_data_analyst_response_markdown(
        final_body,
        executed_sqls or list(state.get("last_sqls") or []),
        last_result_meta,
        last_result_rows,
        knowledge_hits,
        str(state.get("last_export_path") or ""),
        confidence_score,
        confidence_reason,
        step_log=step_log,
        force_visible_table=_user_explicitly_requests_table(current_request),
    )
    state["stage"] = "ready"
    state["pending_request"] = ""
    state["last_sqls"] = executed_sqls or list(state.get("last_sqls") or [])
    state["last_result_meta"] = last_result_meta
    state["last_result_rows"] = last_result_rows[:DATA_ANALYST_MAX_RESULT_ROWS]
    state["knowledge_hits"] = knowledge_hits[:DATA_ANALYST_MAX_KNOWLEDGE_RESULTS]
    state["final_answer"] = final_answer
    state["last_error"] = last_error
    return {
        "answer": final_answer,
        "agent_state": state,
        "steps": step_log + [
            {
                "id": "data-analyst-forced-finish",
                "title": "Reached step budget",
                "status": "error",
                "details": "The agent reached the maximum step budget and generated the best final answer it could with the available evidence.",
                "step": len(step_log) + 1,
                "type": "finish",
                "reasoning": "The anti-loop limit forced the agent to stop and synthesize the answer.",
                "result_summary": "The final answer was synthesized after the step budget was exhausted.",
                "row_count": len(last_result_rows),
                "ok": False,
            }
        ],
        "total_steps": len(step_log),
    }


@app.post("/api/chat/auto-ml-agent")
async def chat_auto_ml_agent(req: AutoMlAgentRequest):
    _emit_log("info", "auto_ml", "Received request", {"query": req.message})
    user_message = (req.message or "").strip()
    state = _normalize_auto_ml_state(req.agent_state.model_dump())
    conversation_memory = _conversation_memory_markdown(
        req.history,
        current_message=state.get("pending_request") or user_message,
        max_steps=CHAT_MEMORY_MAX_STEPS,
    )

    try:
        state["available_tables"] = state["available_tables"] or await list_clickhouse_tables(req.clickhouse)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"ClickHouse connection error: {exc}") from exc

    if not state["available_tables"]:
        raise HTTPException(status_code=400, detail="No tables were found in the configured ClickHouse database.")
    if not HAVE_SKLEARN:
        raise HTTPException(status_code=400, detail="Auto-ML requires scikit-learn on the server. Install the backend ML dependencies first.")

    extracted_row_filter = _extract_auto_ml_row_filter(user_message)
    if extracted_row_filter:
        state["row_filter"] = extracted_row_filter
    validation_error = _validate_data_quality_row_filter(str(state.get("row_filter") or ""))
    if validation_error:
        raise HTTPException(status_code=400, detail=validation_error)

    extracted_sample_limit = _extract_auto_ml_sample_row_limit(user_message)
    if extracted_sample_limit:
        state["sample_row_limit"] = extracted_sample_limit
    state["sample_row_limit"] = max(100, min(10_000, int(state.get("sample_row_limit") or AUTO_ML_SAMPLE_ROWS)))

    if not user_message:
        return {
            "answer": (
                "## Auto-ML\n"
                "Ask me to benchmark predictive models on a ClickHouse table and I will compare several algorithms on the requested target.\n\n"
                "- **Model families:** linear/logistic regression, Random Forest, and XGBoost when available.\n"
                "- **Scope control:** you can narrow the dataset with a safe row filter and cap the benchmark sample size.\n"
                "- **Output:** a Markdown comparison table plus a short recommendation."
            ),
            "agent_state": state,
            "steps": [{"id": "auto-ml-ready", "title": "Ready for model benchmarking", "status": "success", "details": "The agent is ready to train and compare several machine-learning models."}],
        }

    if state.get("stage") == "awaiting_table":
        selected_table = resolve_user_choice(user_message, state.get("clarification_options") or state["available_tables"])
        if not selected_table:
            return {
                "answer": build_choice_markdown(
                    "Table Clarification",
                    state.get("clarification_prompt") or "Which table should I use for this Auto-ML benchmark?",
                    state.get("clarification_options") or state["available_tables"][:AUTO_ML_TABLE_OPTION_LIMIT],
                ),
                "agent_state": state,
                "steps": [{"id": "auto-ml-await-table", "title": "Waiting for table selection", "status": "running", "details": "The agent needs one training table before it can benchmark models."}],
            }
        state["selected_table"] = selected_table
        state["stage"] = "ready"
        state["clarification_prompt"] = ""
        state["clarification_options"] = []
        user_message = state.get("pending_request") or user_message

    if not state.get("selected_table"):
        direct_table = resolve_user_choice(user_message, state["available_tables"])
        if direct_table:
            state["selected_table"] = direct_table
        elif len(state["available_tables"]) == 1:
            state["selected_table"] = state["available_tables"][0]
        else:
            table_analysis = await analyze_clickhouse_tables(
                user_message,
                state["available_tables"],
                conversation_memory,
                req.llm_base_url,
                req.llm_model,
                req.llm_provider,
                req.llm_api_key,
            )
            matched_selected = match_available_options([table_analysis.get("selected_table") or ""], state["available_tables"])
            matched_candidates = match_available_options(table_analysis.get("table_candidates") or [], state["available_tables"])
            if matched_selected and not table_analysis.get("table_choice_required"):
                state["selected_table"] = matched_selected[0]
            elif len(matched_candidates) == 1 and not table_analysis.get("table_choice_required"):
                state["selected_table"] = matched_candidates[0]
            else:
                state["stage"] = "awaiting_table"
                state["pending_request"] = user_message
                state["clarification_prompt"] = str(table_analysis.get("table_choice_prompt") or "").strip() or "Which table should I use for this Auto-ML benchmark?"
                state["clarification_options"] = (matched_candidates or state["available_tables"])[:AUTO_ML_TABLE_OPTION_LIMIT]
                return {
                    "answer": build_choice_markdown("Table Clarification", state["clarification_prompt"], state["clarification_options"]),
                    "agent_state": state,
                    "steps": [{"id": "auto-ml-route", "title": "Need table confirmation", "status": "running", "details": str(table_analysis.get("reasoning") or "Several tables remain plausible for the requested model benchmark.")}],
                }

    if not state.get("schema_info"):
        try:
            state["schema_info"] = await describe_clickhouse_table(req.clickhouse, str(state["selected_table"]))
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Failed to inspect schema for {state['selected_table']}: {exc}") from exc

    if state.get("stage") == "awaiting_target":
        target_choice = resolve_user_choice(user_message, state.get("clarification_options") or _auto_ml_target_candidates(state["schema_info"]))
        if not target_choice:
            return {
                "answer": build_choice_markdown(
                    "Target Clarification",
                    state.get("clarification_prompt") or "Which column should be the prediction target?",
                    state.get("clarification_options") or _auto_ml_target_candidates(state["schema_info"])[:AUTO_ML_TARGET_OPTION_LIMIT],
                ),
                "agent_state": state,
                "steps": [{"id": "auto-ml-await-target", "title": "Waiting for target selection", "status": "running", "details": "The agent needs the target column before it can benchmark models."}],
            }
        state["target_column"] = target_choice
        state["stage"] = "ready"
        state["clarification_prompt"] = ""
        state["clarification_options"] = []
        user_message = state.get("pending_request") or user_message

    if not state.get("target_column"):
        inferred_target = _infer_target_column(user_message, state["schema_info"])
        if inferred_target:
            state["target_column"] = inferred_target
        else:
            candidates = _auto_ml_target_candidates(state["schema_info"])[:AUTO_ML_TARGET_OPTION_LIMIT]
            if len(candidates) == 1:
                state["target_column"] = candidates[0]
            else:
                state["stage"] = "awaiting_target"
                state["pending_request"] = user_message
                state["clarification_prompt"] = "Which column should be the prediction target for the Auto-ML benchmark?"
                state["clarification_options"] = candidates
                return {
                    "answer": build_choice_markdown("Target Clarification", state["clarification_prompt"], state["clarification_options"]),
                    "agent_state": state,
                    "steps": [{"id": "auto-ml-target", "title": "Need target confirmation", "status": "running", "details": "The benchmark needs one target column to evaluate candidate models."}],
                }

    sample_columns = [str(column.get("name") or "").strip() for column in state["schema_info"] if str(column.get("name") or "").strip()]
    _, rows, sample_sql = await _fetch_clickhouse_sample_rows(
        req.clickhouse,
        str(state["selected_table"]),
        sample_columns,
        int(state["sample_row_limit"]),
        str(state.get("row_filter") or ""),
    )
    _emit_log("sql", "auto_ml", "Fetched training sample", {"table": state["selected_table"], "rows": len(rows)})
    records, targets, feature_columns, problem_type = _prepare_automl_dataset(rows, state["schema_info"], str(state["target_column"]))
    try:
        comparison_rows, recommended_model = _run_automl_benchmark(records, targets, problem_type)
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=(
                f"Auto-ML could not benchmark models on the current scope: {exc}. "
                "Try a broader row filter, a larger sample limit, or a target column with more usable values."
            ),
        ) from exc

    try:
        narrative = await _auto_ml_narrative(
            user_message,
            str(state["selected_table"]),
            str(state["target_column"]),
            str(state.get("row_filter") or ""),
            int(state["sample_row_limit"]),
            problem_type,
            feature_columns,
            comparison_rows,
            conversation_memory,
            req.llm_base_url,
            req.llm_model,
            req.llm_provider,
            req.llm_api_key,
        )
    except Exception:
        narrative = (
            "## Executive Summary\n"
            f"I benchmarked several candidate models on `{state['selected_table']}` to predict `{state['target_column']}`.\n\n"
            "## Benchmark Scope\n"
            f"- Row filter: **{str(state.get('row_filter') or 'Full table')}**\n"
            f"- Training sample limit: **{int(state['sample_row_limit'])} rows**\n\n"
            "## Model Comparison\n"
            f"The strongest performer was **{recommended_model}** on the sampled dataset.\n\n"
            "## Recommendation\n"
            "Use the leading model as the baseline, then iterate on feature engineering and hold-out validation before production deployment."
        )

    answer = "\n\n".join(
        [
            narrative.strip(),
            "## Comparison Table",
            _markdown_table_from_rows(comparison_rows),
            "## Executed SQL",
            f"```sql\n{sample_sql}\n```",
        ]
    ).strip()
    state["feature_columns"] = feature_columns
    state["row_filter"] = str(state.get("row_filter") or "")
    state["sample_row_limit"] = int(state["sample_row_limit"])
    state["comparison_rows"] = comparison_rows
    state["problem_type"] = problem_type
    state["recommended_model"] = recommended_model
    state["final_answer"] = answer
    state["last_error"] = ""
    state["pending_request"] = ""
    state["clarification_prompt"] = ""
    state["clarification_options"] = []
    state["stage"] = "ready"
    return {
        "answer": answer,
        "agent_state": state,
        "steps": [
            {"id": "auto-ml-schema", "title": "Inspected schema", "status": "success", "details": f"Loaded {len(state['schema_info'])} columns from `{state['selected_table']}`."},
            {"id": "auto-ml-dataset", "title": "Prepared training dataset", "status": "success", "details": f"Built a {problem_type} dataset with {len(records)} rows, {len(feature_columns)} feature columns, filter `{str(state.get('row_filter') or 'full table')}`, and a sample cap of {int(state['sample_row_limit'])} rows."},
            {"id": "auto-ml-benchmark", "title": "Benchmarked candidate models", "status": "success", "details": f"Compared {len(comparison_rows)} candidate models and recommended `{recommended_model}`."},
        ],
    }


@app.post("/api/chat/data-cleaner-agent")
async def chat_data_cleaner_agent(req: DataCleanerAgentRequest):
    _emit_log("info", "data_cleaner", "Received request", {"query": req.message})
    user_message = (req.message or "").strip()
    state = _normalize_data_cleaner_state(req.agent_state.model_dump())
    extracted_row_filter = _extract_guided_row_filter(user_message)
    if extracted_row_filter:
        state["row_filter"] = extracted_row_filter
    conversation_memory = _conversation_memory_markdown(
        req.history,
        current_message=state.get("pending_request") or user_message,
        max_steps=CHAT_MEMORY_MAX_STEPS,
    )
    try:
        state["available_tables"] = state["available_tables"] or await list_clickhouse_tables(req.clickhouse)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"ClickHouse connection error: {exc}") from exc
    validation_error = _validate_data_quality_row_filter(str(state.get("row_filter") or ""))
    if validation_error:
        raise HTTPException(status_code=400, detail=f"Invalid Data Cleaner row filter: {validation_error}")

    if not user_message:
        return {
            "answer": (
                "## Data Cleaner\n"
                "Ask me to audit a ClickHouse table for duplicates, missing values, and inconsistent formats.\n\n"
                "- **Audit focus:** duplicate risk, null spikes, empty strings, and date-format inconsistencies.\n"
                "- **Output:** a business summary plus suggested SQL correction scripts."
            ),
            "agent_state": state,
            "steps": [{"id": "data-cleaner-ready", "title": "Ready for data-quality cleanup", "status": "success", "details": "The agent is ready to inspect one ClickHouse table."}],
        }

    if state.get("stage") == "awaiting_table":
        selected_table = resolve_user_choice(user_message, state.get("clarification_options") or state["available_tables"])
        if not selected_table:
            return {
                "answer": build_choice_markdown(
                    "Table Clarification",
                    state.get("clarification_prompt") or "Which table should I audit for data-quality issues?",
                    state.get("clarification_options") or state["available_tables"][:8],
                ),
                "agent_state": state,
                "steps": [{"id": "data-cleaner-await-table", "title": "Waiting for table selection", "status": "running", "details": "The cleaner needs a table before it can inspect duplicates, nulls, and inconsistent formats."}],
            }
        state["selected_table"] = selected_table
        state["stage"] = "ready"
        state["clarification_prompt"] = ""
        state["clarification_options"] = []
        user_message = state.get("pending_request") or user_message

    direct_table = resolve_user_choice(user_message, state["available_tables"])
    if direct_table and direct_table != state.get("selected_table"):
        state["selected_table"] = direct_table
        state["schema_info"] = []

    if not state.get("selected_table"):
        if len(state["available_tables"]) == 1:
            state["selected_table"] = state["available_tables"][0]
        else:
            table_analysis = await analyze_clickhouse_tables(
                user_message,
                state["available_tables"],
                conversation_memory,
                req.llm_base_url,
                req.llm_model,
                req.llm_provider,
                req.llm_api_key,
            )
            matched_selected = match_available_options([table_analysis.get("selected_table") or ""], state["available_tables"])
            matched_candidates = match_available_options(table_analysis.get("table_candidates") or [], state["available_tables"])
            if matched_selected and not table_analysis.get("table_choice_required"):
                state["selected_table"] = matched_selected[0]
            elif len(matched_candidates) == 1 and not table_analysis.get("table_choice_required"):
                state["selected_table"] = matched_candidates[0]
            else:
                state["stage"] = "awaiting_table"
                state["pending_request"] = user_message
                state["clarification_prompt"] = str(table_analysis.get("table_choice_prompt") or "").strip() or "Which table should I audit for data-quality issues?"
                state["clarification_options"] = (matched_candidates or state["available_tables"])[:8]
                return {
                    "answer": build_choice_markdown("Table Clarification", state["clarification_prompt"], state["clarification_options"]),
                    "agent_state": state,
                    "steps": [{"id": "data-cleaner-route", "title": "Need table confirmation", "status": "running", "details": str(table_analysis.get("reasoning") or "Several tables remain plausible for the requested data-quality audit.")}],
                }

    if not state.get("schema_info"):
        try:
            state["schema_info"] = await describe_clickhouse_table(req.clickhouse, str(state["selected_table"]))
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Failed to inspect schema for {state['selected_table']}: {exc}") from exc

    findings, scripts, sample_rows, audit_sqls, preview_markdown = await _data_cleaner_profile(
        req.clickhouse,
        str(state["selected_table"]),
        state["schema_info"],
        str(state.get("row_filter") or ""),
    )

    try:
        narrative = await _data_cleaner_narrative(
            user_message,
            str(state["selected_table"]),
            str(state.get("row_filter") or ""),
            findings,
            scripts,
            preview_markdown,
            conversation_memory,
            req.llm_base_url,
            req.llm_model,
            req.llm_provider,
            req.llm_api_key,
        )
    except Exception:
        narrative = (
            "## Executive Summary\n"
            + f"I reviewed `{state['selected_table']}` for duplicate risk, missing values, and inconsistent date formats"
            + (f" within the scope `{state.get('row_filter')}`.\n\n" if str(state.get("row_filter") or "").strip() else ".\n\n")
            + "## Key Quality Risks\n"
            + "\n".join(f"- **{item['title']}** — {item['detail']}" for item in findings[:4])
            + "\n\n## Recommended Fix Path\n"
            + "Prioritize the critical issues first, normalize date parsing rules, and standardize missing-value handling before downstream analytics or ML work."
        )

    script_blocks = "\n\n".join(
        f"### {script['title']}\n```sql\n{script['sql']}\n```"
        for script in scripts
    ) or "No correction script was necessary for the inspected scope."
    sql_log = "\n\n".join(
        f"```sql\n{item['sql']}\n```"
        for item in audit_sqls
        if str(item.get("sql") or "").strip()
    )
    answer = (
        f"{narrative.strip()}\n\n"
        "## Technical details\n"
        "### Suggested correction scripts\n"
        f"{script_blocks}\n\n"
        "### Sample preview\n"
        f"{preview_markdown}\n\n"
        "### Audit SQL log\n"
        f"{sql_log}"
    ).strip()

    state["findings"] = findings
    state["correction_scripts"] = scripts
    state["final_answer"] = answer
    state["pending_request"] = ""
    state["clarification_prompt"] = ""
    state["clarification_options"] = []
    state["stage"] = "ready"
    state["last_error"] = ""
    return {
        "answer": answer,
        "agent_state": state,
        "steps": [
            {"id": "data-cleaner-schema", "title": "Inspected schema", "status": "success", "details": f"Loaded {len(state['schema_info'])} columns from `{state['selected_table']}`."},
            {"id": "data-cleaner-profile", "title": "Profiled quality risks", "status": "success", "details": f"Flagged {len(findings)} finding(s) and prepared {len(scripts)} suggested SQL fix script(s)."},
        ],
    }


@app.post("/api/chat/anonymizer-agent")
async def chat_anonymizer_agent(req: AnonymizerAgentRequest):
    _emit_log("info", "anonymizer", "Received request", {"query": req.message})
    user_message = (req.message or "").strip()
    state = _normalize_anonymizer_state(req.agent_state.model_dump())
    extracted_row_filter = _extract_guided_row_filter(user_message)
    if extracted_row_filter:
        state["row_filter"] = extracted_row_filter
    conversation_memory = _conversation_memory_markdown(
        req.history,
        current_message=state.get("pending_request") or user_message,
        max_steps=CHAT_MEMORY_MAX_STEPS,
    )
    try:
        state["available_tables"] = state["available_tables"] or await list_clickhouse_tables(req.clickhouse)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"ClickHouse connection error: {exc}") from exc
    validation_error = _validate_data_quality_row_filter(str(state.get("row_filter") or ""))
    if validation_error:
        raise HTTPException(status_code=400, detail=f"Invalid Anonymizer row filter: {validation_error}")

    if not user_message:
        return {
            "answer": (
                "## Anonymizer\n"
                "Ask me to scan a ClickHouse table for PII exposure and propose masking or hashing strategies.\n\n"
                "- **Detection focus:** direct identifiers, contact fields, addresses, dates of birth, and IP-like values.\n"
                "- **Output:** a privacy-oriented summary plus suggested SQL masking patterns."
            ),
            "agent_state": state,
            "steps": [{"id": "anonymizer-ready", "title": "Ready for privacy review", "status": "success", "details": "The agent is ready to inspect one ClickHouse table for PII exposure."}],
        }

    if state.get("stage") == "awaiting_table":
        selected_table = resolve_user_choice(user_message, state.get("clarification_options") or state["available_tables"])
        if not selected_table:
            return {
                "answer": build_choice_markdown(
                    "Table Clarification",
                    state.get("clarification_prompt") or "Which table should I scan for PII exposure?",
                    state.get("clarification_options") or state["available_tables"][:8],
                ),
                "agent_state": state,
                "steps": [{"id": "anonymizer-await-table", "title": "Waiting for table selection", "status": "running", "details": "The anonymizer needs one table before it can scan for PII."}],
            }
        state["selected_table"] = selected_table
        state["stage"] = "ready"
        state["clarification_prompt"] = ""
        state["clarification_options"] = []
        user_message = state.get("pending_request") or user_message

    direct_table = resolve_user_choice(user_message, state["available_tables"])
    if direct_table and direct_table != state.get("selected_table"):
        state["selected_table"] = direct_table
        state["schema_info"] = []

    if not state.get("selected_table"):
        if len(state["available_tables"]) == 1:
            state["selected_table"] = state["available_tables"][0]
        else:
            table_analysis = await analyze_clickhouse_tables(
                user_message,
                state["available_tables"],
                conversation_memory,
                req.llm_base_url,
                req.llm_model,
                req.llm_provider,
                req.llm_api_key,
            )
            matched_selected = match_available_options([table_analysis.get("selected_table") or ""], state["available_tables"])
            matched_candidates = match_available_options(table_analysis.get("table_candidates") or [], state["available_tables"])
            if matched_selected and not table_analysis.get("table_choice_required"):
                state["selected_table"] = matched_selected[0]
            elif len(matched_candidates) == 1 and not table_analysis.get("table_choice_required"):
                state["selected_table"] = matched_candidates[0]
            else:
                state["stage"] = "awaiting_table"
                state["pending_request"] = user_message
                state["clarification_prompt"] = str(table_analysis.get("table_choice_prompt") or "").strip() or "Which table should I scan for PII exposure?"
                state["clarification_options"] = (matched_candidates or state["available_tables"])[:8]
                return {
                    "answer": build_choice_markdown("Table Clarification", state["clarification_prompt"], state["clarification_options"]),
                    "agent_state": state,
                    "steps": [{"id": "anonymizer-route", "title": "Need table confirmation", "status": "running", "details": str(table_analysis.get("reasoning") or "Several tables remain plausible for the requested privacy scan.")}],
                }

    if not state.get("schema_info"):
        try:
            state["schema_info"] = await describe_clickhouse_table(req.clickhouse, str(state["selected_table"]))
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Failed to inspect schema for {state['selected_table']}: {exc}") from exc

    pii_findings, masking_scripts, sample_rows, preview_markdown = await _anonymizer_profile(
        req.clickhouse,
        str(state["selected_table"]),
        state["schema_info"],
        str(state.get("row_filter") or ""),
    )
    try:
        narrative = await _anonymizer_narrative(
            user_message,
            str(state["selected_table"]),
            str(state.get("row_filter") or ""),
            pii_findings,
            masking_scripts,
            preview_markdown,
            conversation_memory,
            req.llm_base_url,
            req.llm_model,
            req.llm_provider,
            req.llm_api_key,
        )
    except Exception:
        narrative = (
            "## Executive Summary\n"
            + f"I reviewed `{state['selected_table']}` for personally identifiable information and privacy-sensitive fields"
            + (f" within the scope `{state.get('row_filter')}`.\n\n" if str(state.get("row_filter") or "").strip() else ".\n\n")
            + "## PII Exposure Review\n"
            + ("\n".join(f"- **{item['column']}** — {item['pii_type']} ({item['risk']} risk). {item['recommendation']}" for item in pii_findings[:6]) or "- No direct PII signal was found in the inspected sample.")
            + "\n\n## Recommended Masking Strategy\n"
            + "Hash stable identifiers, mask direct contact fields, and generalize sensitive dates before sharing or analytics outside the restricted perimeter."
        )

    script_blocks = "\n\n".join(
        f"### {script['title']}\n```sql\n{script['sql']}\n```"
        for script in masking_scripts
    ) or "No masking script was required for the inspected scope."
    answer = (
        f"{narrative.strip()}\n\n"
        "## Technical details\n"
        "### Suggested masking scripts\n"
        f"{script_blocks}\n\n"
        "### Sample preview\n"
        f"{preview_markdown}"
    ).strip()

    state["pii_findings"] = pii_findings
    state["masking_scripts"] = masking_scripts
    state["final_answer"] = answer
    state["pending_request"] = ""
    state["clarification_prompt"] = ""
    state["clarification_options"] = []
    state["stage"] = "ready"
    state["last_error"] = ""
    return {
        "answer": answer,
        "agent_state": state,
        "steps": [
            {"id": "anonymizer-schema", "title": "Inspected schema", "status": "success", "details": f"Loaded {len(state['schema_info'])} columns from `{state['selected_table']}`."},
            {"id": "anonymizer-scan", "title": "Scanned for PII exposure", "status": "success", "details": f"Flagged {len(pii_findings)} PII candidate(s) and prepared {len(masking_scripts)} masking pattern(s)."},
        ],
    }


@app.post("/api/custom-agent/analyze")
async def analyze_custom_agent(req: CustomAgentAnalysisRequest):
    python_code = (req.python_code or "").strip()
    if not python_code:
        raise HTTPException(status_code=400, detail="Python code is required to analyze a custom agent.")

    try:
        profile = await _analyze_custom_agent_code(
            python_code,
            req.title,
            req.description,
            req.llm_base_url,
            req.llm_model,
            req.llm_provider,
            req.llm_api_key,
        )
        return {
            "status": "ready",
            "profile": profile,
        }
    except Exception as exc:
        fallback = _fallback_custom_agent_profile(python_code, req.title, req.description)
        fallback["statusMessage"] = f"{fallback['statusMessage']} Root issue: {exc}"
        return {
            "status": "fallback",
            "profile": fallback,
        }


@app.post("/api/chat/custom-agent")
async def chat_custom_agent(req: CustomAgentChatRequest):
    custom_agent = {
        "id": req.custom_agent.id,
        "title": req.custom_agent.title or "Custom Agent",
        "description": req.custom_agent.description or "",
        "pythonCode": req.custom_agent.python_code or "",
        "systemPrompt": req.custom_agent.system_prompt or "",
        "managerRoutingHint": req.custom_agent.manager_routing_hint or "",
        "status": req.custom_agent.status or "draft",
        "statusMessage": req.custom_agent.status_message or "",
        "enabled": bool(req.custom_agent.enabled),
        "badgeColor": req.custom_agent.badge_color or "zinc",
    }
    if not custom_agent["pythonCode"].strip():
        raise HTTPException(status_code=400, detail="The selected custom agent has no Python code attached.")

    message = (req.message or "").strip()
    if not message:
        return {
            "answer": (
                f"## {custom_agent['title']}\n"
                f"{custom_agent['description'] or 'This custom agent is ready.'}"
            ),
            "agent_state": _normalize_custom_agent_state({"selected_agent_id": custom_agent["id"]}),
            "steps": [
                {
                    "id": "custom-agent-ready",
                    "title": "Custom agent ready",
                    "status": "success",
                    "details": "The generated custom agent is ready to work from its uploaded Python specification.",
                }
            ],
        }

    state = _normalize_custom_agent_state(req.agent_state.model_dump())
    state["selected_agent_id"] = custom_agent["id"]
    system_prompt = (
        f"You are {custom_agent['title']}. Reply in English.\n\n"
        f"User-facing description:\n{custom_agent['description'] or 'No description provided.'}\n\n"
        f"Manager routing hint:\n{custom_agent['managerRoutingHint'] or 'No routing hint provided.'}\n\n"
        f"Behavior contract:\n{custom_agent['systemPrompt'] or 'Use the Python code below as your implementation reference.'}\n\n"
        "Treat the uploaded Python code as the source specification for the agent's capabilities, workflow, constraints, and intent. "
        "Preserve its specific behavior as much as possible, but answer clearly for the user. "
        "Keep the visible answer business-friendly. Put technical notes only after the main answer inside a `## Technical details` section.\n\n"
        f"Python specification:\n```python\n{custom_agent['pythonCode']}\n```"
    )
    try:
        answer = await llm_chat(
            [{"role": "system", "content": system_prompt}, *req.history, {"role": "user", "content": message}],
            req.llm_base_url,
            req.llm_model,
            req.llm_provider,
            req.llm_api_key,
        )
    except Exception as exc:
        state["last_error"] = str(exc)
        raise HTTPException(status_code=400, detail=f"Custom agent execution failed: {exc}") from exc

    state["final_answer"] = answer
    state["last_error"] = ""
    return {
        "answer": answer,
        "agent_state": state,
        "steps": [
            {
                "id": "custom-agent-executed",
                "title": f"Executed {custom_agent['title']}",
                "status": "success",
                "details": "The custom agent answered using its uploaded Python specification.",
            }
        ],
    }


@app.post("/api/chat/manager-agent")
async def chat_manager_agent(req: ManagerAgentRequest):
    start_t = _time.time()
    _emit_log("info", "manager", "Received manager request", {"query": req.message})
    user_message = (req.message or "").strip()
    manager_state = _normalize_manager_agent_state(req.manager_state.model_dump())
    persisted_state = await read_db_state()
    app_config = (persisted_state or {}).get("config") or {}
    manager_rag_context = await _manager_functional_context_from_rag(user_message, app_config)

    clickhouse_state = dump_clickhouse_agent_state(req.clickhouse_state)
    data_analyst_state = _normalize_data_analyst_state(req.data_analyst_state.model_dump())
    auto_ml_state = _normalize_auto_ml_state(req.auto_ml_state.model_dump())
    data_cleaner_state = _normalize_data_cleaner_state(req.data_cleaner_state.model_dump())
    anonymizer_state = _normalize_anonymizer_state(req.anonymizer_state.model_dump())
    custom_agent_state = _normalize_custom_agent_state(req.custom_agent_state.model_dump())
    custom_agents = _enabled_custom_agents([agent.model_dump() for agent in req.custom_agents])
    file_manager_state = _normalize_file_manager_state(req.file_manager_state.model_dump())
    pdf_creator_state = _normalize_pdf_creator_state(req.pdf_creator_state.model_dump())
    oracle_analyst_state = _normalize_oracle_analyst_state(req.oracle_analyst_state.model_dump())
    file_manager_config = _normalize_file_manager_config(req.file_manager_config.model_dump())
    oracle_analyst_config = _normalize_oracle_analyst_config(req.oracle_analyst_config.model_dump())
    oracle_connections = req.oracle_connections or [OracleConnectionConfig(**DEFAULT_APP_CONFIG["oracleConnections"][0])]
    pending_pipeline = manager_state.get("pending_pipeline")

    def _manager_agent_state_payload() -> dict[str, Any]:
        return {
            "manager": manager_state,
            "clickhouse": clickhouse_state,
            "dataAnalyst": data_analyst_state,
            "autoMl": auto_ml_state,
            "dataCleaner": data_cleaner_state,
            "anonymizer": anonymizer_state,
            "customAgent": custom_agent_state,
            "fileManager": file_manager_state,
            "pdfCreator": pdf_creator_state,
            "oracleAnalyst": oracle_analyst_state,
        }

    if not user_message:
        manager_state["active_delegate"] = None
        return {
            "answer": (
                "## Agent Manager\n"
                "Describe the outcome you want, and I will either answer directly or route the task "
                "to Clickhouse SQL, Data Analyst, Auto-ML, Data Cleaner, Anonymizer, a custom specialist, File management, PDF creator, or Oracle SQL when a specialist is needed."
            ),
            "agent_state": _manager_agent_state_payload(),
            "steps": [
                {
                    "id": "manager-ready",
                    "title": "Manager ready",
                    "status": "success",
                    "details": "The manager can orchestrate all specialist agents currently available in RAGnarok.",
                }
            ],
        }

    async def _delegate_file_manager(message: str) -> dict[str, Any]:
        return await chat_file_manager_agent(
            FileManagerAgentRequest(
                message=message,
                history=req.history,
                llm_base_url=req.llm_base_url,
                llm_model=req.llm_model,
                llm_api_key=req.llm_api_key,
                llm_provider=req.llm_provider,
                agent_state=FileManagerAgentStateModel(**file_manager_state),
                file_manager_config=FileManagerAgentConfigModel(
                    base_path=file_manager_config["basePath"],
                    max_iterations=file_manager_config["maxIterations"],
                    system_prompt=file_manager_config["systemPrompt"],
                ),
            )
        )

    async def _delegate_data_analyst(message: str) -> dict[str, Any]:
        return await chat_data_analyst_agent(
            DataAnalystAgentRequest(
                message=message,
                history=req.history,
                clickhouse=req.clickhouse,
                llm_base_url=req.llm_base_url,
                llm_model=req.llm_model,
                llm_api_key=req.llm_api_key,
                llm_provider=req.llm_provider,
                max_steps=DATA_ANALYST_DEFAULT_MAX_STEPS,
                agent_state=DataAnalystAgentStateModel(**data_analyst_state),
            )
        )

    async def _delegate_auto_ml(message: str) -> dict[str, Any]:
        return await chat_auto_ml_agent(
            AutoMlAgentRequest(
                message=message,
                history=req.history,
                clickhouse=req.clickhouse,
                llm_base_url=req.llm_base_url,
                llm_model=req.llm_model,
                llm_api_key=req.llm_api_key,
                llm_provider=req.llm_provider,
                agent_state=AutoMlAgentStateModel(**auto_ml_state),
            )
        )

    async def _delegate_data_cleaner(message: str) -> dict[str, Any]:
        return await chat_data_cleaner_agent(
            DataCleanerAgentRequest(
                message=message,
                history=req.history,
                clickhouse=req.clickhouse,
                llm_base_url=req.llm_base_url,
                llm_model=req.llm_model,
                llm_api_key=req.llm_api_key,
                llm_provider=req.llm_provider,
                agent_state=DataCleanerAgentStateModel(**data_cleaner_state),
            )
        )

    async def _delegate_anonymizer(message: str) -> dict[str, Any]:
        return await chat_anonymizer_agent(
            AnonymizerAgentRequest(
                message=message,
                history=req.history,
                clickhouse=req.clickhouse,
                llm_base_url=req.llm_base_url,
                llm_model=req.llm_model,
                llm_api_key=req.llm_api_key,
                llm_provider=req.llm_provider,
                agent_state=AnonymizerAgentStateModel(**anonymizer_state),
            )
        )

    async def _delegate_custom_agent(message: str, custom_agent_id: str) -> dict[str, Any]:
        custom_agent = _find_custom_agent(custom_agents, custom_agent_id)
        if not custom_agent:
            raise HTTPException(status_code=400, detail=f"Unknown custom agent `{custom_agent_id}`.")
        return await chat_custom_agent(
            CustomAgentChatRequest(
                message=message,
                history=req.history,
                custom_agent=CustomAgentConfigModel(
                    id=custom_agent["id"],
                    title=custom_agent["title"],
                    description=custom_agent["description"],
                    python_code=custom_agent["pythonCode"],
                    system_prompt=custom_agent["systemPrompt"],
                    manager_routing_hint=custom_agent["managerRoutingHint"],
                    status=custom_agent["status"],
                    status_message=custom_agent["statusMessage"],
                    enabled=custom_agent["enabled"],
                    badge_color=custom_agent["badgeColor"],
                ),
                agent_state=CustomAgentStateModel(**custom_agent_state),
                llm_base_url=req.llm_base_url,
                llm_model=req.llm_model,
                llm_api_key=req.llm_api_key,
                llm_provider=req.llm_provider,
            )
        )

    async def _delegate_pdf_creator(message: str) -> dict[str, Any]:
        return await chat_pdf_creator_agent(
            PdfCreatorAgentRequest(
                message=message,
                history=req.history,
                llm_base_url=req.llm_base_url,
                llm_model=req.llm_model,
                llm_api_key=req.llm_api_key,
                llm_provider=req.llm_provider,
                agent_state=PdfCreatorAgentStateModel(**pdf_creator_state),
                file_manager_config=FileManagerAgentConfigModel(
                    base_path=file_manager_config["basePath"],
                    max_iterations=file_manager_config["maxIterations"],
                    system_prompt=file_manager_config["systemPrompt"],
                ),
            )
        )

    async def _delegate_oracle_analyst(message: str) -> dict[str, Any]:
        return await chat_oracle_analyst_agent(
            OracleAnalystAgentRequest(
                message=message,
                history=req.history,
                oracle_connections=oracle_connections,
                oracle_analyst_config=OracleAnalystConfigModel(
                    connection_id=oracle_analyst_config["connectionId"],
                    row_limit=oracle_analyst_config["rowLimit"],
                    max_retries=oracle_analyst_config["maxRetries"],
                    max_iterations=oracle_analyst_config["maxIterations"],
                    toolkit_id=oracle_analyst_config["toolkitId"],
                    system_prompt=oracle_analyst_config["systemPrompt"],
                ),
                llm_base_url=req.llm_base_url,
                llm_model=req.llm_model,
                llm_api_key=req.llm_api_key,
                llm_provider=req.llm_provider,
                agent_state=OracleAnalystAgentStateModel(**oracle_analyst_state),
            )
        )

    if (
        isinstance(pending_pipeline, dict)
        and pending_pipeline.get("kind") == "clickhouse_to_file"
        and pending_pipeline.get("stage") == "awaiting_export_details"
    ):
        if is_negative_response(user_message) or normalize_choice(user_message).lower() in {"cancel", "stop", "never mind"}:
            manager_state["pending_pipeline"] = None
            manager_state["active_delegate"] = None
            return {
                "answer": "## Agent Manager\nThe pending file export was cancelled.",
                "agent_state": _manager_agent_state_payload(),
                "steps": [
                    {
                        "id": "manager-export-cancelled",
                        "title": "Cancelled pending export",
                        "status": "success",
                        "details": "The user cancelled the second step of the ClickHouse-to-file workflow.",
                    }
                ],
            }

        pending_pipeline = _manager_update_export_pipeline_from_reply(pending_pipeline, user_message)
        manager_state["pending_pipeline"] = pending_pipeline
        if _manager_pending_pipeline_requires_details(pending_pipeline):
            manager_state["active_delegate"] = None
            return {
                "answer": _manager_export_details_prompt(pending_pipeline),
                "agent_state": _manager_agent_state_payload(),
                "steps": [
                    {
                        "id": "manager-export-details",
                        "title": "Waiting for export details",
                        "status": "running",
                        "details": "The manager still needs the target format or path before delegating to File management.",
                    }
                ],
            }

        delegate = "file_management"
        routing = {
            "delegate": "file_management",
            "reasoning": "The manager now has the remaining export details and can continue with File management.",
            "handoff_message": json.dumps(
                _build_file_export_payload_from_clickhouse(pending_pipeline, clickhouse_state),
                ensure_ascii=False,
            ),
        }
    else:
        routing = await analyze_manager_routing(
            user_message,
            req.history,
            manager_state,
            clickhouse_state,
            data_analyst_state,
            auto_ml_state,
            data_cleaner_state,
            anonymizer_state,
            custom_agents,
            custom_agent_state,
            file_manager_state,
            pdf_creator_state,
            oracle_analyst_state,
            req.llm_base_url,
            req.llm_model,
            req.llm_provider,
            req.llm_api_key,
            manager_rag_context.get("context") or "",
        )
        if routing["delegate"] == "clickhouse_query":
            export_pipeline = _extract_clickhouse_file_export_pipeline(user_message)
            pdf_pipeline = _extract_clickhouse_pdf_export_pipeline(user_message)
            if export_pipeline and not manager_state.get("pending_pipeline"):
                manager_state["pending_pipeline"] = export_pipeline
            elif pdf_pipeline and not manager_state.get("pending_pipeline"):
                manager_state["pending_pipeline"] = pdf_pipeline
        elif routing["delegate"] != "file_management":
            manager_state["pending_pipeline"] = None

    delegate = routing["delegate"]
    manager_state["last_routing_reason"] = routing["reasoning"]
    manager_state["last_delegate_label"] = _manager_specialist_label(delegate)
    if manager_rag_context.get("used") and delegate in {"clickhouse_query", "data_analyst", "auto_ml", "data_cleaner", "anonymizer", "custom_agent", "manager"}:
        routing["handoff_message"] = (
            f"{routing['handoff_message']}\n\n"
            "Functional context from RAG:\n"
            f"{manager_rag_context.get('context')}"
        ).strip()
    
    _emit_log(
        "decision",
        "manager",
        f"Delegating to {delegate}",
        {"agent": delegate, "rationale": routing["reasoning"]}
    )

    base_steps = []
    if manager_rag_context.get("enabled"):
        base_steps.append(
            {
                "id": "manager-functional-context",
                "title": "Loaded functional context",
                "status": "success" if manager_rag_context.get("used") else "running",
                "details": str(manager_rag_context.get("summary") or "No matching functional context was found in the knowledge base."),
            }
        )
    base_steps.append(
        {
            "id": "manager-analyze",
            "title": "Analyzed request",
            "status": "success",
            "details": routing["reasoning"],
        }
    )

    if delegate == "manager":
        manager_state["active_delegate"] = None
        manager_state["pending_pipeline"] = None
        
        _emit_log("llm", "manager", "Answering directly as manager", {"query": req.message})
        answer = await _run_manager_direct_response(
            req.history,
            req.llm_base_url,
            req.llm_model,
            req.llm_provider,
            req.llm_api_key,
            req.system_prompt,
            manager_rag_context.get("context") or "",
        )
        return {
            "answer": answer,
            "agent_state": _manager_agent_state_payload(),
            "steps": base_steps + [
                {
                    "id": "manager-direct",
                    "title": "Answered directly",
                    "status": "success",
                    "details": "No specialist tool was needed for this turn.",
                }
            ],
        }

    try:
        if delegate == "clickhouse_query":
            delegated = await chat_clickhouse_agent(
                ClickHouseAgentRequest(
                    message=(
                        routing["handoff_message"] + "\n\nDo not start a chart flow for this turn. Return the text result only."
                        if manager_state.get("pending_pipeline")
                        else routing["handoff_message"]
                    ),
                    history=req.history,
                    clickhouse=req.clickhouse,
                    llm_base_url=req.llm_base_url,
                    llm_model=req.llm_model,
                    llm_api_key=req.llm_api_key,
                    llm_provider=req.llm_provider,
                    agent_state=ClickHouseAgentState(**clickhouse_state),
                )
            )
            clickhouse_state = (
                delegated.get("agent_state")
                if isinstance(delegated.get("agent_state"), dict)
                else clickhouse_state
            )
            manager_state["active_delegate"] = (
                "clickhouse_query" if _clickhouse_state_needs_followup(clickhouse_state) else None
            )
        elif delegate == "data_analyst":
            delegated = await _delegate_data_analyst(routing["handoff_message"])
            data_analyst_state = _normalize_data_analyst_state(delegated.get("agent_state"))
            manager_state["active_delegate"] = (
                "data_analyst" if _data_analyst_state_needs_followup(data_analyst_state) else None
            )
        elif delegate == "auto_ml":
            delegated = await _delegate_auto_ml(routing["handoff_message"])
            auto_ml_state = _normalize_auto_ml_state(delegated.get("agent_state"))
            manager_state["active_delegate"] = (
                "auto_ml" if _auto_ml_state_needs_followup(auto_ml_state) else None
            )
        elif delegate == "data_cleaner":
            delegated = await _delegate_data_cleaner(routing["handoff_message"])
            data_cleaner_state = _normalize_data_cleaner_state(delegated.get("agent_state"))
            manager_state["active_delegate"] = (
                "data_cleaner" if _data_cleaner_state_needs_followup(data_cleaner_state) else None
            )
        elif delegate == "anonymizer":
            delegated = await _delegate_anonymizer(routing["handoff_message"])
            anonymizer_state = _normalize_anonymizer_state(delegated.get("agent_state"))
            manager_state["active_delegate"] = (
                "anonymizer" if _anonymizer_state_needs_followup(anonymizer_state) else None
            )
        elif delegate == "custom_agent":
            if not custom_agents:
                raise HTTPException(status_code=400, detail="No enabled custom agent is available for delegation.")
            delegated = await _delegate_custom_agent(
                routing["handoff_message"],
                str(routing.get("custom_agent_id") or custom_agents[0]["id"]),
            )
            custom_agent_state = _normalize_custom_agent_state(delegated.get("agent_state"))
            manager_state["active_delegate"] = None
        elif delegate == "file_management":
            delegated = await _delegate_file_manager(routing["handoff_message"])
            file_manager_state = _normalize_file_manager_state(delegated.get("agent_state"))
            manager_state["active_delegate"] = (
                "file_management" if _file_manager_state_needs_followup(file_manager_state) else None
            )
            manager_state["pending_pipeline"] = None
        elif delegate == "pdf_creator":
            delegated = await _delegate_pdf_creator(routing["handoff_message"])
            pdf_creator_state = _normalize_pdf_creator_state(delegated.get("agent_state"))
            manager_state["active_delegate"] = (
                "pdf_creator" if _pdf_creator_state_needs_followup(pdf_creator_state) else None
            )
            manager_state["pending_pipeline"] = None
        elif delegate == "oracle_analyst":
            delegated = await _delegate_oracle_analyst(routing["handoff_message"])
            oracle_analyst_state = _normalize_oracle_analyst_state(delegated.get("agent_state"))
            manager_state["active_delegate"] = (
                "oracle_analyst" if _oracle_state_needs_followup(oracle_analyst_state) else None
            )
            manager_state["pending_pipeline"] = None
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported manager delegate: {delegate}")
    except HTTPException as exc:
        raise HTTPException(
            status_code=exc.status_code,
            detail=f"Manager delegation to {_manager_specialist_label(delegate)} failed: {exc.detail}",
        ) from exc
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=f"Manager delegation to {_manager_specialist_label(delegate)} failed: {exc}",
        ) from exc

    specialist_label = _manager_specialist_label(delegate)
    specialist_steps = _prefix_agent_steps(delegated.get("steps") or [], delegate)
    manager_steps = base_steps + [
        {
            "id": "manager-route",
            "title": f"Delegated to {specialist_label}",
            "status": "running" if manager_state.get("active_delegate") else "success",
            "details": (
                "The specialist needs more user input to continue."
                if manager_state.get("active_delegate")
                else "The specialist completed its turn and returned the result."
            ),
        }
    ]

    if (
        delegate == "clickhouse_query"
        and isinstance(manager_state.get("pending_pipeline"), dict)
        and not manager_state.get("active_delegate")
    ):
        pending_pipeline = dict(manager_state["pending_pipeline"])
        pipeline_kind = pending_pipeline.get("kind")

        if pipeline_kind == "clickhouse_to_file":
            if _manager_pending_pipeline_requires_details(pending_pipeline):
                pending_pipeline["stage"] = "awaiting_export_details"
                manager_state["pending_pipeline"] = pending_pipeline
                manager_steps.append(
                    {
                        "id": "manager-await-export-details",
                        "title": "Waiting for export details",
                        "status": "running",
                        "details": "The ClickHouse query is complete, but the manager still needs the export format or target path.",
                    }
                )
                return {
                    "answer": _manager_compose_chained_answer(
                        delegated.get("answer") or "",
                        _manager_export_details_prompt(pending_pipeline),
                        "File Export",
                    ),
                    "actions": delegated.get("actions"),
                    "chart": delegated.get("chart"),
                    "agent_state": _manager_agent_state_payload(),
                    "steps": manager_steps + specialist_steps,
                }

            try:
                chained = await _delegate_file_manager(
                    json.dumps(
                        _build_file_export_payload_from_clickhouse(pending_pipeline, clickhouse_state),
                        ensure_ascii=False,
                    )
                )
                file_manager_state = _normalize_file_manager_state(chained.get("agent_state"))
                manager_state["active_delegate"] = (
                    "file_management" if _file_manager_state_needs_followup(file_manager_state) else None
                )
                manager_state["pending_pipeline"] = None
                chained_steps = _prefix_agent_steps(chained.get("steps") or [], "file_management")
                manager_steps.append(
                    {
                        "id": "manager-chain-file-export",
                        "title": "Continued to File management",
                        "status": "running" if manager_state.get("active_delegate") else "success",
                        "details": "The manager continued the same request by exporting the ClickHouse result through File management.",
                    }
                )
                return {
                    "answer": _manager_compose_chained_answer(
                        delegated.get("answer") or "",
                        chained.get("answer") or "",
                        "File Export",
                    ),
                    "actions": chained.get("actions") or delegated.get("actions"),
                    "chart": delegated.get("chart") or chained.get("chart"),
                    "agent_state": _manager_agent_state_payload(),
                    "steps": manager_steps + specialist_steps + chained_steps,
                }
            except HTTPException as exc:
                manager_state["pending_pipeline"] = None
                manager_state["active_delegate"] = None
                manager_steps.append(
                    {
                        "id": "manager-chain-file-export-failed",
                        "title": "File export failed",
                        "status": "error",
                        "details": str(exc.detail),
                    }
                )
                return {
                    "answer": _manager_compose_chained_answer(
                        delegated.get("answer") or "",
                        f"## File Export\nI could not complete the export step.\n\n```text\n{exc.detail}\n```",
                        "File Export",
                    ),
                    "actions": delegated.get("actions"),
                    "chart": delegated.get("chart"),
                    "agent_state": _manager_agent_state_payload(),
                    "steps": manager_steps + specialist_steps,
                }

        if pipeline_kind == "clickhouse_to_pdf":
            try:
                chained = await _delegate_pdf_creator(
                    json.dumps(
                        _build_pdf_export_payload_from_clickhouse(
                            pending_pipeline,
                            delegated.get("answer") or "",
                        ),
                        ensure_ascii=False,
                    )
                )
                pdf_creator_state = _normalize_pdf_creator_state(chained.get("agent_state"))
                manager_state["active_delegate"] = (
                    "pdf_creator" if _pdf_creator_state_needs_followup(pdf_creator_state) else None
                )
                manager_state["pending_pipeline"] = None
                chained_steps = _prefix_agent_steps(chained.get("steps") or [], "pdf_creator")
                manager_steps.append(
                    {
                        "id": "manager-chain-pdf-export",
                        "title": "Continued to PDF creator",
                        "status": "running" if manager_state.get("active_delegate") else "success",
                        "details": "The manager continued the same request by exporting the ClickHouse result through PDF creator.",
                    }
                )
                return {
                    "answer": _manager_compose_chained_answer(
                        delegated.get("answer") or "",
                        chained.get("answer") or "",
                        "PDF Export",
                    ),
                    "actions": chained.get("actions") or delegated.get("actions"),
                    "chart": delegated.get("chart") or chained.get("chart"),
                    "agent_state": _manager_agent_state_payload(),
                    "steps": manager_steps + specialist_steps + chained_steps,
                }
            except HTTPException as exc:
                manager_state["pending_pipeline"] = None
                manager_state["active_delegate"] = None
                manager_steps.append(
                    {
                        "id": "manager-chain-pdf-export-failed",
                        "title": "PDF export failed",
                        "status": "error",
                        "details": str(exc.detail),
                    }
                )
                return {
                    "answer": _manager_compose_chained_answer(
                        delegated.get("answer") or "",
                        f"## PDF Export\nI could not complete the PDF step.\n\n```text\n{exc.detail}\n```",
                        "PDF Export",
                    ),
                    "actions": delegated.get("actions"),
                    "chart": delegated.get("chart"),
                    "agent_state": _manager_agent_state_payload(),
                    "steps": manager_steps + specialist_steps,
                }

    return {
        "answer": delegated.get("answer") or "## Answer\nThe delegated agent completed its turn.",
        "actions": delegated.get("actions"),
        "chart": delegated.get("chart"),
        "agent_state": _manager_agent_state_payload(),
        "steps": manager_steps + specialist_steps,
    }


# ── ClickHouse endpoints ──────────────────────────────────────────────────────

@app.post("/api/clickhouse/test")
async def test_clickhouse_connection(req: ClickHouseTestRequest):
    try:
        info = await execute_clickhouse_sql(
            req.clickhouse,
            "SELECT currentDatabase() AS database_name, version() AS version",
        )
        tables = await list_clickhouse_tables(req.clickhouse)
        first_row = info.get("data", [{}])[0] if info.get("data") else {}
        return {
            "status": "ok",
            "database": first_row.get("database_name", req.clickhouse.database),
            "version": first_row.get("version", "unknown"),
            "tables": tables[:20],
            "table_count": len(tables),
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/oracle/test")
async def test_oracle_connection_endpoint(req: OracleTestRequest):
    try:
        return await test_oracle_connection(req.connection)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.post("/api/sql-draft/preview")
async def preview_sql_draft(req: SqlDraftPreviewRequest):
    row_limit = max(1, min(int(req.row_limit or 1000), 10_000))
    engine = str(req.engine or "clickhouse").strip().lower()

    if engine == "oracle":
        connection = _resolve_oracle_connection(
            req.oracle_connections or [OracleConnectionConfig(**DEFAULT_APP_CONFIG["oracleConnections"][0])],
            req.oracle_analyst_config,
        )
        sql = clean_sql_text(req.sql)
        if not is_safe_read_only_oracle_sql(sql):
            raise HTTPException(status_code=400, detail="Only read-only Oracle SELECT queries are allowed.")
        try:
            result = await execute_oracle_query(connection, sql, row_limit)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

        return {
            "engine": "oracle",
            "executed_sql": result["sql"],
            "columns": list(result.get("columns") or []),
            "rows": list(result.get("rows") or []),
            "row_count": int(result.get("row_count") or 0),
            "shown_rows": len(result.get("rows") or []),
            "row_limit": row_limit,
        }

    sql = enforce_clickhouse_preview_limit(req.sql, row_limit)
    if not is_safe_read_only_sql(sql):
        raise HTTPException(status_code=400, detail="Only read-only ClickHouse queries are allowed.")

    try:
        result = await execute_clickhouse_sql(
            req.clickhouse,
            sql,
            max_result_rows_override=row_limit,
        )
    except Exception as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    rows = list(result.get("data") or [])
    columns = [str(item.get("name") or "") for item in (result.get("meta") or []) if str(item.get("name") or "").strip()]
    if not columns and rows:
        columns = [str(key) for key in rows[0].keys()]

    return {
        "engine": "clickhouse",
        "executed_sql": sql,
        "columns": columns,
        "rows": rows,
        "row_count": len(rows),
        "shown_rows": len(rows),
        "row_limit": row_limit,
    }


@app.post("/api/file-manager/test-path")
async def test_file_manager_path(req: FileManagerPathTestRequest):
    try:
        config = _normalize_file_manager_config(req.file_manager_config.model_dump())
        target = _resolve_agent_path(".", config["basePath"])
        if not target.exists():
            raise ValueError(f"The configured folder does not exist: {target}")
        if not target.is_dir():
            raise ValueError(f"The configured path is not a directory: {target}")

        entries = sorted(target.iterdir(), key=lambda item: item.name.lower())
        preview = [
            {
                "name": item.name,
                "kind": "directory" if item.is_dir() else "file",
            }
            for item in entries[:8]
        ]

        return {
            "status": "ok",
            "resolvedPath": str(target),
            "restricted": bool(config["basePath"]),
            "entryCount": len(entries),
            "preview": preview,
        }
    except Exception as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


def build_choice_markdown(title: str, prompt: str, options: list[str]) -> str:
    bullet_list = "\n".join(f"- [ ] {option}" for option in options)
    return f"## {title}\n{prompt}\n\n{bullet_list}"


def append_choice_markdown(base_answer: str, title: str, prompt: str, options: list[str]) -> str:
    choice = build_choice_markdown(title, prompt, options)
    return f"{base_answer}\n\n---\n\n{choice}" if base_answer else choice


def reset_clickhouse_clarification(state: ClickHouseAgentState) -> None:
    state.clarification_prompt = ""
    state.clarification_options = []


CLICKHOUSE_ROW_COUNT_OPTION = "Return the row count"
CLICKHOUSE_SAMPLE_ROWS_OPTION = "Return sample rows"


@app.post("/api/chat/clickhouse-agent")
async def chat_clickhouse_agent(req: ClickHouseAgentRequest):
    # This agent intentionally mixes two behaviors:
    # 1. fast-path handling for obvious utility requests such as schema listing
    #    or simple row previews
    # 2. LLM-guided SQL planning when the request is analytical or ambiguous
    # The goal is to avoid unnecessary clarification for routine database tasks.
    _emit_log("info", "clickhouse", "Received request", {"query": req.message})
    state = req.agent_state

    user_message = (req.message or "").strip()
    memory_anchor = state.pending_request or user_message
    conversation_memory = _conversation_memory_markdown(
        req.history,
        current_message=memory_anchor,
        max_steps=CHAT_MEMORY_MAX_STEPS,
    )

    try:
        state.available_tables = state.available_tables or await list_clickhouse_tables(req.clickhouse)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"ClickHouse connection error: {e}")

    if not state.available_tables:
        raise HTTPException(status_code=400, detail="No tables were found in the configured ClickHouse database.")

    explicit_table_switch = resolve_user_choice(user_message, state.available_tables)
    if state.selected_table and explicit_table_switch and explicit_table_switch != state.selected_table:
        state.selected_table = explicit_table_switch
        state.table_schema = []
        state.pending_request = ""
        reset_clickhouse_query_resolution(state)
        state.stage = "ready"

    if state.stage == "ready" and not state.pending_request and state.last_result_rows and is_chart_followup_request(user_message):
        state.chart_requested = True
        state.stage = "awaiting_chart_x"

    chart_flow_cancelled = False

    if state.last_result_rows and state.last_result_meta and (
        state.chart_requested
        or state.stage in {"awaiting_chart_offer", "awaiting_chart_x", "awaiting_chart_y", "awaiting_chart_type"}
    ):
        chart_context = infer_chart_options(state.last_result_meta, state.last_result_rows)
        if not chart_context["can_chart"]:
            reset_clickhouse_chart_state(state)
            state.stage = "ready"
        else:
            state.chart_x_options = state.chart_x_options or chart_context["x_options"]
            state.chart_y_options = state.chart_y_options or chart_context["y_options"]
            state.chart_type_options = state.chart_type_options or chart_context["type_options"]

            if state.stage == "awaiting_chart_offer":
                chart_offer_choice = resolve_user_choice(
                    user_message,
                    state.chart_offer_options or [CHART_CREATE_OPTION, CHART_SKIP_OPTION],
                )
                if not chart_offer_choice and is_affirmative_response(user_message):
                    chart_offer_choice = CHART_CREATE_OPTION
                if not chart_offer_choice and is_negative_response(user_message):
                    chart_offer_choice = CHART_SKIP_OPTION
                if chart_offer_choice == CHART_SKIP_OPTION:
                    reset_clickhouse_chart_state(state)
                    state.stage = "ready"
                    return {
                        "answer": build_clickhouse_response_markdown(
                            "I kept the latest result in **text format only**. No chart was created.",
                            [state.last_sql],
                        ),
                        "agent_state": dump_clickhouse_agent_state(state),
                        "steps": [
                            {
                                "id": "ch-chart-skip",
                                "title": "Skipped chart generation",
                                "status": "success",
                                "details": "The user chose to keep the tabular/text answer only.",
                            }
                        ],
                    }
                if chart_offer_choice == CHART_CREATE_OPTION:
                    state.chart_requested = True
                    state.stage = "awaiting_chart_x"
                elif user_message:
                    reset_clickhouse_chart_state(state)
                    state.stage = "ready"

            if state.chart_requested or state.stage in {"awaiting_chart_x", "awaiting_chart_y", "awaiting_chart_type"}:
                requested_chart_type = detect_requested_chart_type(user_message)
                x_options = state.chart_x_options
                y_options = [
                    option for option in state.chart_y_options
                    if option != state.selected_chart_x
                ] or state.chart_y_options

                if not state.selected_chart_x:
                    x_choice = resolve_user_choice(user_message, x_options) if state.stage == "awaiting_chart_x" else None
                    if len(x_options) == 1:
                        state.selected_chart_x = x_options[0]
                    elif x_choice:
                        state.selected_chart_x = x_choice
                    elif state.stage == "awaiting_chart_x" and user_message:
                        reset_clickhouse_chart_state(state)
                        state.stage = "ready"
                        chart_flow_cancelled = True
                    else:
                        state.stage = "awaiting_chart_x"
                        return {
                            "answer": build_choice_markdown(
                                "Chart X Axis",
                                "Choose the field to use on the X axis.",
                                x_options,
                            ),
                            "agent_state": dump_clickhouse_agent_state(state),
                            "steps": [
                                {
                                    "id": "ch-chart-x",
                                    "title": "Waiting for X axis selection",
                                    "status": "running",
                                    "details": "The user must choose which field should drive the horizontal axis.",
                                }
                            ],
                        }

                if chart_flow_cancelled:
                    pass
                elif not state.selected_chart_y:
                    y_choice = resolve_user_choice(user_message, y_options) if state.stage == "awaiting_chart_y" else None
                    if len(y_options) == 1:
                        state.selected_chart_y = y_options[0]
                    elif y_choice:
                        state.selected_chart_y = y_choice
                    elif state.stage == "awaiting_chart_y" and user_message:
                        reset_clickhouse_chart_state(state)
                        state.stage = "ready"
                        chart_flow_cancelled = True
                    else:
                        state.stage = "awaiting_chart_y"
                        return {
                            "answer": build_choice_markdown(
                                "Chart Y Axis",
                                "Choose the metric to use on the Y axis.",
                                y_options,
                            ),
                            "agent_state": dump_clickhouse_agent_state(state),
                            "steps": [
                                {
                                    "id": "ch-chart-y",
                                    "title": "Waiting for Y axis selection",
                                    "status": "running",
                                    "details": "The user must choose the metric to visualize.",
                                }
                            ],
                        }

                if chart_flow_cancelled:
                    pass
                elif not state.selected_chart_type:
                    chart_type_choice = None
                    if requested_chart_type and CHART_TYPE_LABELS.get(requested_chart_type) in state.chart_type_options:
                        chart_type_choice = requested_chart_type
                    elif state.stage == "awaiting_chart_type":
                        resolved_type_label = resolve_user_choice(user_message, state.chart_type_options)
                        if resolved_type_label:
                            chart_type_choice = CHART_TYPE_BY_LABEL.get(resolved_type_label.lower())

                    if len(state.chart_type_options) == 1:
                        chart_type_choice = CHART_TYPE_BY_LABEL.get(state.chart_type_options[0].lower())

                    if chart_type_choice:
                        state.selected_chart_type = chart_type_choice
                    elif state.stage == "awaiting_chart_type" and user_message:
                        reset_clickhouse_chart_state(state)
                        state.stage = "ready"
                        chart_flow_cancelled = True
                    else:
                        state.stage = "awaiting_chart_type"
                        return {
                            "answer": build_choice_markdown(
                                "Chart Type",
                                "Choose the chart type.",
                                state.chart_type_options,
                            ),
                            "agent_state": dump_clickhouse_agent_state(state),
                            "steps": [
                                {
                                    "id": "ch-chart-type",
                                    "title": "Waiting for chart type",
                                    "status": "running",
                                    "details": "The user must choose how to visualize the selected axes.",
                                }
                            ],
                        }

                if not chart_flow_cancelled:
                    chart = build_chart(
                        state.last_result_rows,
                        state.selected_chart_x,
                        state.selected_chart_y,
                        state.selected_chart_type,
                    )
                    if not chart:
                        reset_clickhouse_chart_state(state)
                        state.stage = "ready"
                        return {
                            "answer": build_clickhouse_response_markdown(
                                (
                                    "I could not build a usable chart from the latest result because the selected "
                                    "axes do not produce enough valid numeric data points."
                                ),
                                [state.last_sql],
                            ),
                            "agent_state": dump_clickhouse_agent_state(state),
                            "steps": [
                                {
                                    "id": "ch-chart-failed",
                                    "title": "Chart generation failed",
                                    "status": "error",
                                    "details": "Not enough valid data points remained after filtering null or non-numeric values.",
                                }
                            ],
                        }

                    chart_type_label = CHART_TYPE_LABELS.get(state.selected_chart_type, "Chart")
                    answer = build_clickhouse_response_markdown(
                        (
                            f"I created a **{chart_type_label.lower()}** based on the latest query result.\n\n"
                            f"- **X axis:** `{state.selected_chart_x}`\n"
                            f"- **Y axis:** `{state.selected_chart_y}`\n"
                            f"- **Purpose:** make the result easier to read and compare visually"
                        ),
                        [state.last_sql],
                    )
                    reset_clickhouse_chart_state(state)
                    state.stage = "ready"
                    return {
                        "answer": answer,
                        "chart": chart,
                        "agent_state": dump_clickhouse_agent_state(state),
                        "steps": [
                            {
                                "id": "ch-chart-built",
                                "title": "Generated chart",
                                "status": "success",
                                "details": f"Built a {chart_type_label.lower()} from the latest ClickHouse query result.",
                            }
                        ],
                    }

    if state.stage == "awaiting_table":
        table_options = state.clarification_options or state.available_tables
        selected_table = resolve_user_choice(user_message, table_options) or resolve_user_choice(user_message, state.available_tables)
        if not selected_table:
            return {
                "answer": build_choice_markdown(
                    "Table Clarification",
                    state.clarification_prompt or "Which table should I use for this request?",
                    table_options,
                ),
                "agent_state": dump_clickhouse_agent_state(state),
                "steps": [
                    {
                        "id": "ch-await-table",
                        "title": "Waiting for table selection",
                        "status": "running",
                        "details": "Multiple tables remain plausible for the current request.",
                    }
                ],
            }
        state.selected_table = selected_table
        state.stage = "ready"
        reset_clickhouse_clarification(state)

    if state.stage == "ready" and user_message and not explicit_table_switch:
        state.pending_request = user_message
        reset_clickhouse_query_resolution(state)

    if is_clickhouse_table_list_request(state.pending_request or user_message):
        table_rows = [
            {
                "Table": table_name,
                "Database": req.clickhouse.database,
            }
            for table_name in state.available_tables
        ]
        tables_sql = (
            "SELECT name\n"
            "FROM system.tables\n"
            f"WHERE database = {quote_clickhouse_literal(req.clickhouse.database)}\n"
            "ORDER BY name"
        )
        answer = build_clickhouse_response_markdown(
            (
                f"Here is the current table inventory for **{req.clickhouse.database}**.\n\n"
                f"{build_clickhouse_markdown_table(table_rows, headers=['Table', 'Database'], max_rows=500)}"
            ),
            [tables_sql],
        )
        state.last_sql = tables_sql
        state.last_result_meta = [
            {"name": "Table", "type": "String"},
            {"name": "Database", "type": "String"},
        ]
        state.last_result_rows = table_rows
        state.pending_request = ""
        reset_clickhouse_clarification(state)
        reset_clickhouse_chart_state(state)
        state.stage = "ready"
        return {
            "answer": answer,
            "agent_state": dump_clickhouse_agent_state(state),
            "sql": tables_sql,
            "steps": [
                {
                    "id": "ch-list-tables",
                    "title": "Listed available tables",
                    "status": "success",
                    "details": f"Returned {len(table_rows)} table(s) from `{req.clickhouse.database}` as a Markdown table.",
                }
            ],
        }

    if not state.selected_table:
        selected_table = resolve_user_choice(user_message, state.available_tables)
        if selected_table:
            state.selected_table = selected_table
            state.stage = "ready"
            reset_clickhouse_clarification(state)
            if not state.pending_request:
                state.pending_request = ""
        elif len(state.available_tables) == 1:
            state.selected_table = state.available_tables[0]
            state.stage = "ready"
            reset_clickhouse_clarification(state)
        else:
            if user_message and not state.pending_request:
                state.pending_request = user_message
            if not state.pending_request:
                return {
                    "answer": (
                        "## Clickhouse SQL Agent\n"
                        "Ask your question directly. I will infer the best table when the intent is clear, "
                        "and I will only ask you to choose if the request stays ambiguous."
                    ),
                    "agent_state": dump_clickhouse_agent_state(state),
                    "steps": [
                        {
                            "id": "ch-ready",
                            "title": "Ready for direct query",
                            "status": "success",
                            "details": f"Loaded {len(state.available_tables)} table(s) from `{req.clickhouse.database}`.",
                        }
                    ],
                }

            table_analysis = await analyze_clickhouse_tables(
                state.pending_request,
                state.available_tables,
                conversation_memory,
                req.llm_base_url,
                req.llm_model,
                req.llm_provider,
                req.llm_api_key,
            )
            matched_candidates = match_available_options(
                table_analysis["table_candidates"],
                state.available_tables,
            )
            matched_selected_table = match_available_options(
                [table_analysis["selected_table"]],
                state.available_tables,
            )

            if matched_selected_table and not table_analysis["table_choice_required"]:
                state.selected_table = matched_selected_table[0]
                state.stage = "ready"
                reset_clickhouse_clarification(state)
            elif len(matched_candidates) == 1 and not table_analysis["table_choice_required"]:
                state.selected_table = matched_candidates[0]
                state.stage = "ready"
                reset_clickhouse_clarification(state)
            else:
                state.stage = "awaiting_table"
                state.clarification_options = matched_candidates or state.available_tables
                state.clarification_prompt = (
                    table_analysis["table_choice_prompt"]
                    or "Which table should I use for this request?"
                )
                return {
                    "answer": build_choice_markdown(
                        "Table Clarification",
                        state.clarification_prompt,
                        state.clarification_options,
                    ),
                    "agent_state": dump_clickhouse_agent_state(state),
                    "steps": [
                        {
                            "id": "ch-tables",
                            "title": "Loaded ClickHouse tables",
                            "status": "success",
                            "details": f"Found {len(state.available_tables)} table(s) in database `{req.clickhouse.database}`.",
                        },
                        {
                            "id": "ch-table-routing",
                            "title": "Need table confirmation",
                            "status": "running",
                            "details": table_analysis["reasoning"] or "Several tables look plausible for the current request.",
                        },
                    ],
                }

    if state.selected_table and not state.pending_request and not explicit_table_switch:
        state.pending_request = user_message

    if not state.pending_request:
        return {
            "answer": (
                f"## Clickhouse SQL Agent\nI am focused on table `{state.selected_table}`.\n\n"
                "Please tell me what you want to know, and I will inspect the schema before writing SQL."
            ),
            "agent_state": dump_clickhouse_agent_state(state),
            "steps": [
                {
                    "id": "ch-await-request",
                    "title": "Waiting for analytical request",
                    "status": "running",
                    "details": "The table is selected, but the user has not asked a data question yet.",
                }
            ],
        }

    if not state.table_schema:
        try:
            state.table_schema = await describe_clickhouse_table(req.clickhouse, state.selected_table)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to inspect schema for {state.selected_table}: {e}")

    if not state.table_schema:
        raise HTTPException(status_code=400, detail=f"Table '{state.selected_table}' has no readable columns.")

    if state.stage == "awaiting_row_intent":
        row_intent_choice = resolve_user_choice(
            user_message,
            state.clarification_options or [CLICKHOUSE_ROW_COUNT_OPTION, CLICKHOUSE_SAMPLE_ROWS_OPTION],
        )
        if row_intent_choice == CLICKHOUSE_ROW_COUNT_OPTION:
            state.pending_request = f"{state.pending_request}\n\n[Resolved intent: row count]".strip()
            state.stage = "ready"
            reset_clickhouse_clarification(state)
        elif row_intent_choice == CLICKHOUSE_SAMPLE_ROWS_OPTION:
            state.pending_request = f"{state.pending_request}\n\n[Resolved intent: sample rows]".strip()
            state.stage = "ready"
            reset_clickhouse_clarification(state)
        else:
            return {
                "answer": build_choice_markdown(
                    "Row Request Clarification",
                    state.clarification_prompt or "Do you want the row count or a sample list of rows?",
                    state.clarification_options or [CLICKHOUSE_ROW_COUNT_OPTION, CLICKHOUSE_SAMPLE_ROWS_OPTION],
                ),
                "agent_state": dump_clickhouse_agent_state(state),
                "steps": [
                    {
                        "id": "ch-row-intent",
                        "title": "Waiting for row intent clarification",
                        "status": "running",
                        "details": "The request could mean either a row count or a sample of rows.",
                    }
                ],
            }

    # Direct schema requests should not go through field/date disambiguation.
    if state.pending_request and is_clickhouse_schema_request(state.pending_request):
        schema_sql = (
            "SELECT name, type, default_kind, default_expression "
            "FROM system.columns "
            f"WHERE database = {quote_clickhouse_literal(req.clickhouse.database)} "
            f"AND table = {quote_clickhouse_literal(state.selected_table)} "
            "ORDER BY position"
        )
        schema_section = build_clickhouse_schema_section(state.selected_table, state.table_schema)
        answer = build_clickhouse_response_markdown(
            f"I found **{len(state.table_schema)} columns** in `{state.selected_table}`.",
            [schema_sql],
            [schema_section],
        )
        state.last_sql = schema_sql
        state.last_result_meta = [
            {"name": "Column", "type": "String"},
            {"name": "Type", "type": "String"},
            {"name": "Default", "type": "String"},
        ]
        state.last_result_rows = [
            {
                "Column": column.get("name", ""),
                "Type": column.get("type", ""),
                "Default": column.get("default_expression", "") or column.get("default_kind", "") or "",
            }
            for column in state.table_schema
        ][:200]
        state.stage = "ready"
        state.pending_request = ""
        reset_clickhouse_clarification(state)
        reset_clickhouse_chart_state(state)
        return {
            "answer": answer,
            "agent_state": dump_clickhouse_agent_state(state),
            "sql": schema_sql,
            "steps": [
                {
                    "id": "ch-selected-table",
                    "title": "Selected ClickHouse table",
                    "status": "success",
                    "details": f"Using table `{state.selected_table}`.",
                },
                {
                    "id": "ch-schema-direct",
                    "title": "Returned table schema directly",
                    "status": "success",
                    "details": f"Listed {len(state.table_schema)} column(s) without requesting any extra clarification.",
                },
            ],
        }

    if state.pending_request and is_clickhouse_row_request_ambiguous(state.pending_request):
        state.stage = "awaiting_row_intent"
        state.clarification_prompt = (
            "Your request can mean either **count the rows** or **show sample rows**. "
            "Choose the output you want."
        )
        state.clarification_options = [CLICKHOUSE_ROW_COUNT_OPTION, CLICKHOUSE_SAMPLE_ROWS_OPTION]
        return {
            "answer": build_choice_markdown(
                "Row Request Clarification",
                state.clarification_prompt,
                state.clarification_options,
            ),
            "agent_state": dump_clickhouse_agent_state(state),
            "steps": [
                {
                    "id": "ch-row-intent",
                    "title": "Clarified row request",
                    "status": "running",
                    "details": "The request is ambiguous between row count and row listing.",
                }
            ],
        }

    if state.pending_request and is_clickhouse_row_count_request(state.pending_request):
        count_sql = (
            f"SELECT count() AS row_count\n"
            f"FROM {quote_clickhouse_identifier(state.selected_table)}"
        )
        try:
            start_exec = _time.time()
            result = await execute_clickhouse_sql(req.clickhouse, count_sql)
            elapsed = _time.time() - start_exec
            _emit_log(
                "sql",
                "clickhouse",
                f"Executed direct count query successfully ({elapsed:.2f}s)",
                {"sql": count_sql, "rows": len(result.get('data', []))},
            )
        except Exception as exc:
            _emit_log("error", "clickhouse", f"Direct count query failed: {exc}", {"sql": count_sql})
            raise

        count_value = 0
        if result.get("data"):
            try:
                count_value = int((result.get("data") or [{}])[0].get("row_count") or 0)
            except Exception:
                count_value = 0

        state.last_sql = count_sql
        state.last_result_meta = result.get("meta", [])
        state.last_result_rows = result.get("data", [])[:1]
        answer = build_clickhouse_response_markdown(
            f"`{state.selected_table}` contains **{count_value:,} row(s)**.",
            [count_sql],
        )
        state.stage = "ready"
        state.pending_request = ""
        reset_clickhouse_clarification(state)
        reset_clickhouse_chart_state(state)
        return {
            "answer": answer,
            "agent_state": dump_clickhouse_agent_state(state),
            "sql": count_sql,
            "steps": [
                {
                    "id": "ch-selected-table",
                    "title": "Selected ClickHouse table",
                    "status": "success",
                    "details": f"Using table `{state.selected_table}`.",
                },
                {
                    "id": "ch-count-direct",
                    "title": "Returned row count directly",
                    "status": "success",
                    "details": f"Counted **{count_value:,}** row(s) without switching to a row preview.",
                },
            ],
        }

    # Direct raw-row requests are another fast path. When the user explicitly
    # asks for rows with all fields, return a bounded preview immediately.
    if state.pending_request and is_clickhouse_sample_rows_request(state.pending_request):
        requested_limit = extract_clickhouse_requested_row_limit(
            state.pending_request,
            default=10,
            max_limit=max(1, min(int(req.clickhouse.query_limit or 100), 100)),
        )
        sample_sql = (
            f"SELECT * FROM {quote_clickhouse_identifier(state.selected_table)}\n"
            f"LIMIT {requested_limit}"
        )
        try:
            start_exec = _time.time()
            result = await execute_clickhouse_sql(req.clickhouse, sample_sql)
            elapsed = _time.time() - start_exec
            _emit_log(
                "sql",
                "clickhouse",
                f"Executed direct sample query successfully ({elapsed:.2f}s)",
                {"sql": sample_sql, "rows": len(result.get('data', []))},
            )
        except Exception as exc:
            _emit_log("error", "clickhouse", f"Direct sample query failed: {exc}", {"sql": sample_sql})
            raise

        state.last_sql = sample_sql
        state.last_result_meta = result.get("meta", [])
        export_row_cap = max(1, min(int(req.clickhouse.query_limit or 200), 2000))
        state.last_result_rows = result.get("data", [])[:export_row_cap]
        preview_headers = [str(item.get("name", "")) for item in state.last_result_meta if str(item.get("name", ""))]
        preview_table = build_clickhouse_markdown_table(
            result.get("data", []),
            headers=preview_headers or None,
            max_rows=requested_limit,
        )
        answer = build_clickhouse_response_markdown(
            (
                f"Here are **{len(result.get('data', []))} row(s)** from `{state.selected_table}` "
                f"with **all available columns**."
            ),
            [sample_sql],
            [f"## Data Preview\n{preview_table}" if preview_table else ""],
        )
        state.stage = "ready"
        state.pending_request = ""
        reset_clickhouse_clarification(state)
        reset_clickhouse_chart_state(state)
        return {
            "answer": answer,
            "agent_state": dump_clickhouse_agent_state(state),
            "sql": sample_sql,
            "steps": [
                {
                    "id": "ch-selected-table",
                    "title": "Selected ClickHouse table",
                    "status": "success",
                    "details": f"Using table `{state.selected_table}`.",
                },
                {
                    "id": "ch-sample-direct",
                    "title": "Returned sample rows directly",
                    "status": "success",
                    "details": f"Fetched {len(result.get('data', []))} row(s) with all columns and formatted them as a table.",
                },
            ],
        }

    if state.stage == "awaiting_field":
        selected_field = resolve_user_choice(user_message, state.clarification_options)
        if not selected_field:
            return {
                "answer": build_choice_markdown(
                    "Field Clarification",
                    state.clarification_prompt or "Please choose the field I should use.",
                    state.clarification_options,
                ),
                "agent_state": dump_clickhouse_agent_state(state),
                "steps": [
                    {
                        "id": "ch-await-field",
                        "title": "Waiting for field clarification",
                        "status": "running",
                        "details": "The request can map to multiple plausible columns.",
                    }
                ],
            }
        state.selected_field = selected_field
        state.stage = "ready"
        reset_clickhouse_clarification(state)

    if state.stage == "awaiting_date":
        selected_date = resolve_user_choice(user_message, state.clarification_options)
        if not selected_date:
            return {
                "answer": build_choice_markdown(
                    "Date Clarification",
                    state.clarification_prompt or "Please choose the date column I should use.",
                    state.clarification_options,
                ),
                "agent_state": dump_clickhouse_agent_state(state),
                "steps": [
                    {
                        "id": "ch-await-date",
                        "title": "Waiting for date clarification",
                        "status": "running",
                        "details": "The request needs a date column and multiple date-like columns are available.",
                    }
                ],
            }
        state.selected_date_field = selected_date
        state.stage = "ready"
        reset_clickhouse_clarification(state)

    analysis = await analyze_clickhouse_schema(
        state.pending_request,
        state.selected_table,
        state.table_schema,
        conversation_memory,
        req.llm_base_url,
        req.llm_model,
        req.llm_provider,
        req.llm_api_key,
    )

    state.candidate_fields = match_schema_columns(analysis["field_candidates"], state.table_schema)
    detected_date_fields = match_schema_columns(analysis["date_candidates"], state.table_schema)
    heuristic_dates = find_date_columns(state.table_schema)
    state.date_fields = detected_date_fields or heuristic_dates

    if analysis["field_choice_required"] and len(state.candidate_fields) > 1 and not state.selected_field:
        state.stage = "awaiting_field"
        state.clarification_prompt = analysis["field_choice_prompt"] or "Which field should I use for this request?"
        state.clarification_options = state.candidate_fields
        return {
            "answer": build_choice_markdown(
                "Field Clarification",
                state.clarification_prompt,
                state.clarification_options,
            ),
            "agent_state": dump_clickhouse_agent_state(state),
            "steps": [
                {
                    "id": "ch-schema",
                    "title": "Inspected table schema",
                    "status": "success",
                    "details": f"Loaded {len(state.table_schema)} columns from `{state.selected_table}`.",
                },
                {
                    "id": "ch-field-clarification",
                    "title": "Need field confirmation",
                    "status": "running",
                    "details": analysis["reasoning"] or "Several columns look plausible for the request.",
                },
            ],
        }

    if not state.selected_field and len(state.candidate_fields) == 1:
        state.selected_field = state.candidate_fields[0]

    if analysis["needs_date_column"] and len(state.date_fields) > 1 and not state.selected_date_field:
        state.stage = "awaiting_date"
        state.clarification_prompt = analysis["date_choice_prompt"] or "Which date column should I use for this request?"
        state.clarification_options = state.date_fields
        return {
            "answer": build_choice_markdown(
                "Date Clarification",
                state.clarification_prompt,
                state.clarification_options,
            ),
            "agent_state": dump_clickhouse_agent_state(state),
            "steps": [
                {
                    "id": "ch-schema",
                    "title": "Inspected table schema",
                    "status": "success",
                    "details": f"Loaded {len(state.table_schema)} columns from `{state.selected_table}`.",
                },
                {
                    "id": "ch-date-clarification",
                    "title": "Need date confirmation",
                    "status": "running",
                    "details": "More than one date-like column can satisfy the request.",
                },
            ],
        }

    if not state.selected_date_field and len(state.date_fields) == 1:
        state.selected_date_field = state.date_fields[0]

    generated = await generate_clickhouse_sql(
        state.pending_request,
        state.selected_table,
        state.table_schema,
        state.selected_field,
        state.selected_date_field,
        conversation_memory,
        req.llm_base_url,
        req.llm_model,
        req.llm_provider,
        req.llm_api_key,
        req.clickhouse.query_limit,
    )

    sql = generated["sql"]
    if not is_safe_read_only_sql(sql):
        raise HTTPException(status_code=400, detail="The generated SQL was rejected because it is not a safe read-only query.")

    try:
        start_exec = _time.time()
        await execute_clickhouse_sql(req.clickhouse, f"EXPLAIN SYNTAX {sql}", readonly=False, json_format=False)
        result = await execute_clickhouse_sql(req.clickhouse, sql)
        elapsed = _time.time() - start_exec
        _emit_log("sql", "clickhouse", f"Executed query successfully ({elapsed:.2f}s)", {"sql": sql, "rows": len(result.get("data", []))})
    except Exception as first_error:
        _emit_log("warning", "clickhouse", f"First execution failed: {first_error}", {"error": str(first_error), "sql": sql})
        repaired = await generate_clickhouse_sql(

            state.pending_request,
            state.selected_table,
            state.table_schema,
            state.selected_field,
            state.selected_date_field,
            conversation_memory,
            req.llm_base_url,
            req.llm_model,
            req.llm_provider,
            req.llm_api_key,
            req.clickhouse.query_limit,
            error_feedback=str(first_error),
        )
        sql = repaired["sql"]
        if not is_safe_read_only_sql(sql):
            _emit_log("error", "clickhouse", "Repaired SQL rejected (not read-only)", {"sql": sql})
            raise HTTPException(status_code=400, detail="The repaired SQL was rejected because it is not read-only.")
        generated["reasoning"] = repaired["reasoning"] or generated["reasoning"]
        try:
            start_exec = _time.time()
            result = await execute_clickhouse_sql(req.clickhouse, sql)
            elapsed = _time.time() - start_exec
            _emit_log("sql", "clickhouse", f"Executed repaired query successfully ({elapsed:.2f}s)", {"sql": sql, "rows": len(result.get("data", []))})
        except Exception as e:
            _emit_log("error", "clickhouse", f"Repaired execution failed: {e}", {"error": str(e), "sql": sql})
            raise

    original_request = state.pending_request

    result_summary = await summarize_clickhouse_result(
        original_request,
        sql,
        generated["reasoning"],
        result.get("data", []),
        conversation_memory,
        req.llm_base_url,
        req.llm_model,
        req.llm_provider,
        req.llm_api_key,
    )
    tabular_section = ""
    if _user_explicitly_requests_table(original_request):
        tabular_section = build_clickhouse_markdown_table(
            result.get("data", []),
            headers=[str(item.get("name") or "") for item in (result.get("meta") or []) if str(item.get("name") or "").strip()] or None,
            max_rows=min(max(5, int(req.clickhouse.query_limit or 20)), 20),
        )
    answer = build_clickhouse_response_markdown(
        result_summary,
        [sql],
        [f"## Result Table\n{tabular_section}" if tabular_section else ""],
    )

    state.last_sql = sql
    state.last_result_meta = result.get("meta", [])
    export_row_cap = max(1, min(int(req.clickhouse.query_limit or 200), 2000))
    state.last_result_rows = result.get("data", [])[:export_row_cap]

    chart_context = infer_chart_options(state.last_result_meta, state.last_result_rows)
    requested_chart = detect_chart_request(original_request)
    if chart_context["can_chart"] and (requested_chart or chart_context["recommended"]):
        reset_clickhouse_chart_state(state)
        if requested_chart:
            initialize_chart_selection(
                state,
                chart_context["x_options"],
                chart_context["y_options"],
                chart_context["type_options"],
                requested_chart_type=detect_requested_chart_type(original_request),
            )
            prompt = next_chart_prompt(state)
            if prompt:
                answer = build_clickhouse_response_markdown(
                    result_summary,
                    [sql],
                    [
                        build_choice_markdown(
                            prompt["title"],
                            prompt["prompt"],
                            prompt["options"],
                        )
                    ],
                )
            else:
                chart = build_chart(
                    state.last_result_rows,
                    state.selected_chart_x,
                    state.selected_chart_y,
                    state.selected_chart_type,
                )
                if chart:
                    state.stage = "ready"
                    answer = build_clickhouse_response_markdown(
                        result_summary,
                        [sql],
                        [
                            (
                                "## Visualization\n"
                                f"I also generated a **{CHART_TYPE_LABELS.get(state.selected_chart_type, 'chart').lower()}** "
                                f"for `{state.selected_chart_y}` by `{state.selected_chart_x}` to make the result easier to interpret."
                            )
                        ],
                    )
                    reset_clickhouse_chart_state(state)
                    return {
                        "answer": answer,
                        "chart": chart,
                        "agent_state": dump_clickhouse_agent_state(state),
                        "sql": sql,
                        "steps": [
                            {
                                "id": "ch-selected-table",
                                "title": "Selected ClickHouse table",
                                "status": "success",
                                "details": f"Using table `{state.selected_table}`.",
                            },
                            {
                                "id": "ch-inspect-schema",
                                "title": "Inspected schema",
                                "status": "success",
                                "details": f"Loaded {len(state.table_schema)} columns to map the request safely.",
                            },
                            {
                                "id": "ch-generate-sql",
                                "title": "Generated safe SQL",
                                "status": "success",
                                "details": generated["reasoning"] or "Built a read-only ClickHouse query with the configured local LLM.",
                            },
                            {
                                "id": "ch-execute",
                                "title": "Executed query",
                                "status": "success",
                                "details": f"Returned {len(result.get('data', []))} row(s).",
                            },
                            {
                                "id": "ch-chart-auto",
                                "title": "Generated chart",
                                "status": "success",
                                "details": "The chart request had only one clear X/Y/type combination, so it was generated directly.",
                            },
                        ],
                    }
        else:
            state.chart_suggested = True
            state.chart_offer_options = [CHART_CREATE_OPTION, CHART_SKIP_OPTION]
            state.chart_x_options = chart_context["x_options"]
            state.chart_y_options = chart_context["y_options"]
            state.chart_type_options = chart_context["type_options"]
            state.stage = "awaiting_chart_offer"
            answer = build_clickhouse_response_markdown(
                result_summary,
                [sql],
                [
                    build_choice_markdown(
                        "Visualization",
                        "This result would work well as a chart. If you want, I can let you choose X, Y, and the chart type.",
                        state.chart_offer_options,
                    )
                ],
            )

    if state.stage not in {"awaiting_chart_offer", "awaiting_chart_x", "awaiting_chart_y", "awaiting_chart_type"}:
        state.stage = "ready"
    state.pending_request = ""
    reset_clickhouse_clarification(state)

    return {
        "answer": answer,
        "agent_state": dump_clickhouse_agent_state(state),
        "sql": sql,
        "steps": [
            {
                "id": "ch-selected-table",
                "title": "Selected ClickHouse table",
                "status": "success",
                "details": f"Using table `{state.selected_table}`.",
            },
            {
                "id": "ch-inspect-schema",
                "title": "Inspected schema",
                "status": "success",
                "details": f"Loaded {len(state.table_schema)} columns to map the request safely.",
            },
            {
                "id": "ch-generate-sql",
                "title": "Generated safe SQL",
                "status": "success",
                "details": generated["reasoning"] or "Built a read-only ClickHouse query with the configured local LLM.",
            },
            {
                "id": "ch-execute",
                "title": "Executed query",
                "status": "success",
                "details": f"Returned {len(result.get('data', []))} row(s).",
            },
        ],
    }


# ── MCP endpoints ─────────────────────────────────────────────────────────────

@app.post("/api/llm/models")
async def list_llm_models(req: LlmModelsRequest):
    normalized_base_url = _normalize_local_service_url(req.base_url)
    provider = (req.provider or "ollama").strip().lower()
    headers = {"Content-Type": "application/json"}
    if req.api_key:
        headers["Authorization"] = f"Bearer {req.api_key}"

    if provider == "ollama":
        endpoint = normalized_base_url.rstrip("/") + "/api/tags"
    else:
        endpoint = normalized_base_url.rstrip("/") + "/models"

    try:
        async with httpx.AsyncClient(
            **_httpx_async_client_kwargs(
                endpoint,
                timeout=60.0,
                verify=False if req.disable_ssl_verification else None,
            )
        ) as client:
            response = await client.get(endpoint, headers=headers)
            response.raise_for_status()
        data = response.json()
        models = (
            [str(model.get("name") or "").strip() for model in data.get("models", []) if str(model.get("name") or "").strip()]
            if provider == "ollama"
            else [str(model.get("id") or "").strip() for model in data.get("data", []) if str(model.get("id") or "").strip()]
        )
        return {"status": "ok", "models": models, "model_count": len(models)}
    except Exception as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.post("/api/embedding/models")
async def list_embedding_models(req: EmbeddingModelsRequest):
    endpoint, resolution_message = _derive_embedding_models_endpoint(req.base_url)
    headers = {"Content-Type": "application/json"}
    if req.api_key:
        headers["Authorization"] = f"Bearer {req.api_key}"

    try:
        async with httpx.AsyncClient(
            **_httpx_async_client_kwargs(
                endpoint,
                timeout=60.0,
                verify=False if req.disable_ssl_verification else None,
            )
        ) as client:
            response = await client.get(endpoint, headers=headers)
            response.raise_for_status()
        data = _parse_http_json_response(
            response,
            service_label="Embedding model discovery endpoint",
            endpoint=endpoint,
        )
        models = [str(model.get("id") or "").strip() for model in data.get("data", []) if str(model.get("id") or "").strip()]
        if not models and isinstance(data.get("models"), list):
            models = [str(model.get("name") or "").strip() for model in data.get("models", []) if str(model.get("name") or "").strip()]
        message = f"Model discovery used `{endpoint}` ({resolution_message})."
        if not models:
            message += " The endpoint responded, but no models were listed."
        return {"status": "ok", "models": models, "model_count": len(models), "message": message, "models_endpoint": endpoint}
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=f"Embedding model discovery failed via `{endpoint}` ({resolution_message}): {exc}",
        ) from exc

@app.post("/api/mcp/test")
async def test_mcp_connection(req: MCPTestRequest):
    """Connect to an MCP server via the Python MCP SDK and return its available tools."""
    try:
        async with _mcp_client_session(
            req.url,
            disable_ssl_verification=req.disable_ssl_verification,
        ) as session:
            tool_definitions = await session.list_tools()
            resource_definitions = await session.list_resources()
            tools = [
                {"name": tool.name, "description": tool.description or ""}
                for tool in (tool_definitions.tools or [])
            ]
        return {
            "status": "ok",
            "tools": tools,
            "tool_count": len(tools),
            "resource_count": len(resource_definitions.resources or []),
        }
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"MCP connection failed: {_format_mcp_exception(e, req.url)}",
        ) from e


def _mcp_tool_to_openai(tool) -> dict:
    """Convert an MCP tool definition to OpenAI function-calling format."""
    schema = {}
    if hasattr(tool, "inputSchema") and tool.inputSchema:
        schema = tool.inputSchema if isinstance(tool.inputSchema, dict) else {}
    return {
        "type": "function",
        "function": {
            "name": tool.name,
            "description": tool.description or "",
            "parameters": schema or {"type": "object", "properties": {}},
        },
    }


@app.post("/api/chat/mcp")
async def chat_mcp(req: MCPChatRequest):
    """Agentic loop: connects to an MCP server via the Python MCP SDK, lets the LLM call tools, returns the final answer."""
    try:
        async with _mcp_client_session(
            req.mcp_url,
            disable_ssl_verification=req.disable_ssl_verification,
        ) as session:
            tool_result = await session.list_tools()
            tool_definitions = list(tool_result.tools or [])
            openai_tools = [_mcp_tool_to_openai(tool) for tool in tool_definitions]
            memory_history = _normalized_history_messages(
                req.history,
                current_message=req.message,
                max_steps=CHAT_MEMORY_MAX_STEPS,
            )

            # Build initial messages
            messages: list[dict] = []
            if req.system_prompt:
                messages.append({"role": "system", "content": req.system_prompt})
            for m in memory_history:
                role = "user" if m.get("role") == "user" else "assistant"
                messages.append({"role": role, "content": m.get("content", "")})
            messages.append({"role": "user", "content": req.message})

            tool_calls_log: list[dict] = []
            MAX_TURNS = 5

            for _ in range(MAX_TURNS):
                # Call LLM with tools
                headers = {"Content-Type": "application/json"}
                if req.llm_api_key:
                    headers["Authorization"] = f"Bearer {req.llm_api_key}"
                normalized_base_url = _normalize_local_service_url(req.llm_base_url)

                if req.llm_provider == "ollama":
                    endpoint = normalized_base_url.rstrip("/") + "/api/chat"
                    payload: dict = {
                        "model": req.llm_model,
                        "messages": messages,
                        "stream": False,
                        "tools": openai_tools,
                    }
                    async with httpx.AsyncClient(
                        **_httpx_async_client_kwargs(
                            endpoint,
                            timeout=120.0,
                            verify=False if req.disable_ssl_verification else None,
                        )
                    ) as llm_client:
                        resp = await llm_client.post(endpoint, json=payload, headers=headers)
                        resp.raise_for_status()
                    data = resp.json()
                    llm_msg = data.get("message", {})
                    content = llm_msg.get("content", "")
                    raw_tool_calls = llm_msg.get("tool_calls", [])
                else:
                    endpoint = normalized_base_url.rstrip("/") + "/chat/completions"
                    payload = {
                        "model": req.llm_model,
                        "messages": messages,
                        "tools": openai_tools,
                    }
                    async with httpx.AsyncClient(
                        **_httpx_async_client_kwargs(
                            endpoint,
                            timeout=120.0,
                            verify=False if req.disable_ssl_verification else None,
                        )
                    ) as llm_client:
                        resp = await llm_client.post(endpoint, json=payload, headers=headers)
                        resp.raise_for_status()
                    data = resp.json()
                    choice = data["choices"][0]["message"]
                    content = choice.get("content") or ""
                    raw_tool_calls = choice.get("tool_calls", [])

                # No tool calls → final answer
                if not raw_tool_calls:
                    return {
                        "answer": content,
                        "tool_calls": tool_calls_log,
                    }

                # Append assistant message with tool calls
                messages.append({
                    "role": "assistant",
                    "content": content,
                    "tool_calls": raw_tool_calls,
                })

                # Execute each tool call via the Python MCP SDK
                for tc in raw_tool_calls:
                    # Normalise across Ollama / OpenAI formats
                    if isinstance(tc, dict) and "function" in tc:
                        fn = tc["function"]
                        tool_name = fn.get("name", "")
                        raw_args = fn.get("arguments", "{}")
                        tool_id = tc.get("id", tool_name)
                    else:
                        # Ollama sometimes returns {"name":..., "arguments":...} directly
                        tool_name = tc.get("name", "")
                        raw_args = tc.get("arguments", "{}")
                        tool_id = tool_name

                    tool_args = json.loads(raw_args) if isinstance(raw_args, str) else raw_args

                    try:
                        result = await session.call_tool(tool_name, tool_args or {})
                        tool_output = _format_mcp_tool_result(result)
                    except Exception as e:
                        tool_output = f"[Tool error] {e}"

                    tool_calls_log.append({
                        "tool": tool_name,
                        "args": tool_args,
                        "result": tool_output,
                    })

                    messages.append({
                        "role": "tool",
                        "tool_call_id": tool_id,
                        "name": tool_name,
                        "content": tool_output,
                    })

            # Safety net: ask LLM for a final answer without tools
            messages.append({
                "role": "user",
                "content": "Please summarise the results above as a final answer.",
            })
            final = await llm_chat(
                messages,
                req.llm_base_url,
                req.llm_model,
                req.llm_provider,
                req.llm_api_key,
                disable_ssl_verification=req.disable_ssl_verification,
            )
            return {"answer": final, "tool_calls": tool_calls_log}

    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"MCP chat failed: {_format_mcp_exception(e, req.mcp_url)}",
        ) from e


# ── OpenSearch management endpoints ──────────────────────────────────────────

@app.post("/api/opensearch/test")
async def test_opensearch(req: TestConnectionRequest):
    """Test connectivity to an OpenSearch cluster."""
    def _ping():
        client = get_os_client(req.url, req.username, req.password)
        return client.info()

    try:
        info = await asyncio.to_thread(_ping)
        cluster_name = info.get("cluster_name", "OpenSearch")
        version = info.get("version", {}).get("number", "unknown")
        return {"status": "ok", "cluster_name": cluster_name, "version": version}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/opensearch/setup-index")
async def setup_index(req: SetupIndexRequest):
    """Create a kNN-enabled index if it does not already exist."""
    def _setup():
        client = get_os_client(
            req.opensearch.url, req.opensearch.username, req.opensearch.password
        )
        index = req.opensearch.index
        if client.indices.exists(index=index):
            return {"status": "exists", "index": index}

        mapping = {
            "settings": {"index.knn": True},
            "mappings": {
                "properties": {
                    "chunk_id":  {"type": "keyword"},
                    "doc_id":    {"type": "keyword"},
                    "doc_name":  {"type": "keyword"},
                    "text":      {"type": "text", "analyzer": "standard"},
                    "embedding": {
                        "type":      "knn_vector",
                        "dimension": req.embedding_dimension,
                        "method": {
                            "name":       "hnsw",
                            "space_type": "cosinesimil",
                            "engine":     "nmslib",
                            "parameters": {"ef_construction": 128, "m": 24},
                        },
                    },
                }
            },
        }
        client.indices.create(index=index, body=mapping)
        return {"status": "created", "index": index}

    try:
        result = await asyncio.to_thread(_setup)
        return result
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


# ── Embedding test endpoint ───────────────────────────────────────────────────

@app.post("/api/embedding/test")
async def test_embedding(req: EmbeddingTestRequest):
    """Generate a real test embedding and optionally verify dimension compatibility with the OpenSearch index."""
    try:
        vector = await get_embedding(
            "embedding connectivity test", req.embedding_base_url, req.embedding_model,
            req.embedding_api_key, verify_ssl=req.embedding_verify_ssl
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Embedding error: {e}")

    dimension = len(vector)
    result: dict = {"status": "ok", "dimension": dimension, "model": req.embedding_model}

    if req.opensearch:
        def _check_mapping():
            client = get_os_client(req.opensearch.url, req.opensearch.username, req.opensearch.password)
            index = req.opensearch.index
            if not client.indices.exists(index=index):
                return None
            mapping = client.indices.get_mapping(index=index)
            props = mapping.get(index, {}).get("mappings", {}).get("properties", {})
            return props.get("embedding", {}).get("dimension")

        try:
            index_dim = await asyncio.to_thread(_check_mapping)
            if index_dim is None:
                result["opensearch"] = {
                    "status": "no_index",
                    "message": f"Index '{req.opensearch.index}' does not exist yet — use Setup Index.",
                }
            elif index_dim == dimension:
                result["opensearch"] = {"status": "compatible", "index_dimension": index_dim}
            else:
                result["opensearch"] = {
                    "status": "incompatible",
                    "index_dimension": index_dim,
                    "message": (
                        f"Dimension mismatch: model produces {dimension}‑D vectors "
                        f"but index expects {index_dim}‑D. Delete and re-create the index."
                    ),
                }
        except Exception as e:
            result["opensearch"] = {"status": "error", "message": str(e)}

    return result


# ── Document ingest endpoint ──────────────────────────────────────────────────

@app.post("/api/documents/ingest")
async def ingest_document(req: IngestRequest):
    """Chunk a document, embed each chunk, and index into OpenSearch."""
    chunks = chunk_text(req.text, max_words=req.chunk_size, overlap_sentences=req.chunk_overlap)
    doc_id = str(uuid.uuid4())

    try:
        embeddings = []
        for chunk in chunks:
            vec = await get_embedding(
                chunk, req.embedding_base_url, req.embedding_model,
                req.embedding_api_key, verify_ssl=req.embedding_verify_ssl
            )
            embeddings.append(vec)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Embedding error: {e}")

    def _index():
        client = get_os_client(
            req.opensearch.url, req.opensearch.username, req.opensearch.password
        )
        index = req.opensearch.index

        if not client.indices.exists(index=index):
            client.indices.create(index=index, body={
                "settings": {"index.knn": True},
                "mappings": {
                    "properties": {
                        "chunk_id":  {"type": "keyword"},
                        "doc_id":    {"type": "keyword"},
                        "doc_name":  {"type": "keyword"},
                        "text":      {"type": "text", "analyzer": "standard"},
                        "embedding": {
                            "type":      "knn_vector",
                            "dimension": req.embedding_dimension,
                        },
                    }
                },
            })

        indexed = 0
        for i, (chunk, vec) in enumerate(zip(chunks, embeddings)):
            client.index(index=index, body={
                "chunk_id":  f"{doc_id}_{i}",
                "doc_id":    doc_id,
                "doc_name":  req.doc_name,
                "text":      chunk,
                "embedding": vec,
            })
            indexed += 1
        return indexed

    try:
        count = await asyncio.to_thread(_index)
        return {"status": "ok", "doc_id": doc_id, "chunks_indexed": count}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


# ── RAG chat endpoint ─────────────────────────────────────────────────────────

@app.post("/api/chat/rag")
async def chat_rag(req: ChatRequest):
    _emit_log("info", "rag", "Received RAG request", {"query": req.message})
    message = req.message

    history = _normalized_history_messages(
        req.history,
        current_message=message,
        max_steps=CHAT_MEMORY_MAX_STEPS,
    )

    # 1. HyDE — generate a hypothetical answer to improve semantic recall
    try:
        start_hyde = _time.time()
        hyde_answer = await llm_chat(
            [{"role": "user", "content": (
                "Write a concise factual answer for semantic search. "
                "No filler, just key facts:\n\n" + message
            )}],
            req.llm_base_url, req.llm_model, req.llm_provider, req.llm_api_key,
        )
        hyde_elapsed = _time.time() - start_hyde
        _emit_log("llm", "rag", f"HyDE generated ({hyde_elapsed:.2f}s)", {"hyde_reply": hyde_answer})
        search_text = hyde_answer or message
    except Exception as e:
        _emit_log("warning", "rag", f"HyDE failed", {"error": str(e)})
        search_text = message

    # 2. Embed the query
    try:
        start_embed = _time.time()
        query_vector = await get_embedding(
            search_text, req.embedding_base_url, req.embedding_model,
            req.embedding_api_key, verify_ssl=req.embedding_verify_ssl
        )
        embed_elapsed = _time.time() - start_embed
        _emit_log(
            "decision",
            "rag",
            f"Query embedded ({embed_elapsed:.2f}s)",
            {
                "vector_len": len(query_vector),
                "target_index": req.opensearch.index,
                "knn_neighbors": req.knn_neighbors,
            },
        )
    except Exception as e:
        _emit_log("error", "rag", f"Embedding failed", {"error": str(e)})
        raise HTTPException(status_code=500, detail=f"Embedding error: {e}")


    # 3. kNN search in OpenSearch
    def _search():
        client = get_os_client(
            req.opensearch.url, req.opensearch.username, req.opensearch.password
        )
        index = req.opensearch.index
        if not client.indices.exists(index=index):
            return []

        response = client.search(
            index=index,
            body={
                "size": req.knn_neighbors,
                "query": {
                    "knn": {
                        "embedding": {
                            "vector": query_vector,
                            "k": req.knn_neighbors,
                        }
                    }
                },
                "_source": ["chunk_id", "doc_id", "doc_name", "text"],
            },
        )
        return [
            {
                "id":       h["_id"],
                "chunk_id": h["_source"].get("chunk_id", h["_id"]),
                "doc_id":   h["_source"].get("doc_id", ""),
                "doc_name": h["_source"].get("doc_name", "document"),
                "text":     h["_source"].get("text", ""),
                "score":    h.get("_score", 0.0),
            }
            for h in response.get("hits", {}).get("hits", [])
        ]

    try:
        results = await asyncio.to_thread(_search)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"OpenSearch error: {e}")

    # Fallback — index empty or unreachable
    if not results:
        try:
            answer = await llm_chat(
                [
                    {"role": "system", "content": (
                        "You are a helpful assistant. No documents were found in the knowledge base. "
                        "Answer from general knowledge and mention that no documents are indexed."
                    )},
                    {"role": "user", "content": message},
                ],
                req.llm_base_url, req.llm_model, req.llm_provider, req.llm_api_key,
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"LLM error: {e}")
        return {"answer": answer, "sources": [], "confidence": 0.0}

    # 4. Keyword boost (Python-side)
    for r in results:
        kw = keyword_score(message, r["text"])
        r["score"] = r["score"] * 0.7 + kw * 0.3
    results.sort(key=lambda x: x["score"], reverse=True)
    top = results[:10]

    # 5. LLM reranking
    try:
        rerank_prompt = (
            f"Score each chunk's relevance to the query 0–10.\n"
            f"Return JSON array only: [{{\"index\": 0, \"relevanceScore\": 8}}, ...]\n\n"
            f"Query: {message}\n\nChunks:\n"
            + "\n\n".join(f"[Chunk {i}]\n{c['text']}" for i, c in enumerate(top))
        )
        rerank_text = await llm_chat(
            [{"role": "user", "content": rerank_prompt}],
            req.llm_base_url, req.llm_model, req.llm_provider, req.llm_api_key,
            response_format="json",
        )
        json_match = re.search(r"\[.*\]", rerank_text, re.DOTALL)
        if json_match:
            for s in json.loads(json_match.group()):
                idx = s.get("index", -1)
                if 0 <= idx < len(top):
                    llm_score = s.get("relevanceScore", 0) / 10.0
                    top[idx]["score"] = top[idx]["score"] * 0.3 + llm_score * 0.7
        top.sort(key=lambda x: x["score"], reverse=True)
        top = [c for c in top if c["score"] > 0.3][:5]
    except Exception as e:
        print(f"Reranking skipped: {e}")
        top = [c for c in top if c["score"] > 0.2][:5]

    # 6. Generate answer with citations
    context = "\n\n".join(
        f"[Source {i + 1}: {c['doc_name']}]\n{c['text']}"
        for i, c in enumerate(top)
    )
    messages_payload = [
        {"role": "system", "content": (
            "You are a helpful assistant. Use the retrieved context to answer.\n"
            "Cite sources using [1], [2], etc. If the answer is not in the context, say so.\n\n"
            f"Context:\n{context}"
        )}
    ]
    for m in history:
        role = "user" if m.get("role") == "user" else "assistant"
        messages_payload.append({"role": role, "content": m.get("content", "")})
    messages_payload.append({"role": "user", "content": message})

    try:
        answer = await llm_chat(
            messages_payload,
            req.llm_base_url, req.llm_model, req.llm_provider, req.llm_api_key,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLM error: {e}")

    return {
        "answer": answer,
        "sources": [
            {
                "id":      c["chunk_id"],
                "docName": c["doc_name"],
                "text":    c["text"],
                "score":   c["score"],
            }
            for c in top
        ],
        "confidence": top[0]["score"] if top else 0.0,
    }


# ── Static file serving ───────────────────────────────────────────────────────

dist_path = Path(__file__).parent / "dist"
if dist_path.exists():
    app.mount("/assets", StaticFiles(directory=str(dist_path / "assets")), name="assets")

    @app.get("/{full_path:path}")
    async def serve_spa(full_path: str):
        file = dist_path / full_path
        if file.is_file():
            return FileResponse(str(file))
        return FileResponse(str(dist_path / "index.html"))


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", "8000"))
    print(f"RAGnarok backend · http://localhost:{port}")
    uvicorn.run("server:app", host="0.0.0.0", port=port, reload=True)
